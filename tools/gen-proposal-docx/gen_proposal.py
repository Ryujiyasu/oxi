#!/usr/bin/env python3
"""Generate the Oxi proposal as a .docx file using python-docx."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, size=10.5, bold=False, color=None, font_name='Yu Gothic'):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_heading_styled(doc, text, level):
    """Add a heading with custom styling."""
    if level == 0:
        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color='2E4057')
        return p
    elif level == 1:
        # Section header (① ② etc.)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2E4057')
        pBdr.append(bottom)
        pPr.append(pBdr)
        run = p.add_run(text)
        set_run_font(run, size=14, bold=True, color='2E4057')
        return p
    elif level == 2:
        # Sub-section (1-1, 1-2 etc.)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, color='2E4057')
        return p
    elif level == 3:
        # Sub-sub-section
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        set_run_font(run, size=11, bold=True)
        return p

def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    run = p.add_run('・' + text)
    set_run_font(run, size=10.5)
    return p

def add_table(doc, headers, rows):
    """Add a table with blue header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '808080')
        borders.append(border)
    tblPr.append(borders)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '2E4057')
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color='FFFFFF')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run, size=10)
            # Bold the last row if it looks like a total
            if cell_text.startswith('合計') or cell_text.startswith('実施費用'):
                run.font.bold = True

    return table

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # Set margins (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ============================================
    # Title page
    # ============================================
    add_heading_styled(doc, '脱OfficeのためのOSS完全互換スイートOxiの開発', 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('一人で世界標準を作る時代が来た')
    set_run_font(run, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('提案者：安河内 竜二　　2026年3月')
    set_run_font(run, size=11)

    # ============================================
    # Section 1
    # ============================================
    add_heading_styled(doc, '① プロジェクトの背景、目的、目標', 1)

    add_heading_styled(doc, '1-1. 背景', 2)

    add_heading_styled(doc, '空白の30年と、逆転の機会', 3)
    add_para(doc, '空白の30年で、日本は周りの国に抜かれていった。IT産業においてその差は特に顕著であり、ソフトウェア開発の主導権を海外企業に譲り続けた結果、日本の官公庁・企業は文書処理の基盤すらMicrosoft（米国）に依存する構造が固定化された。')
    add_para(doc, '財政難の自治体は古いソフトを使い続けるしかない。生産性は上がらず、職員の負担は増え、人手不足が加速する悪循環に陥っている。LibreOfficeへの移行を試みた自治体も多いが、日本語レイアウトが崩れるという根本的な問題により断念した事例が全国に存在する。しかし今、この状況を逆転できる歴史上初めてのタイミングが来ている。')

    add_heading_styled(doc, '大企業は会議をしている間に、一人が世界を作る', 3)
    add_para(doc, '大企業はリスク管理と組織の意思決定コストという負債を抱えている。新技術の採用には会議・承認・法務確認が必要で、身動きが取れない。日本の大企業はレガシーシステムという技術的負債も抱え、過去の投資を守るために新しいスタックへの移行が遅れる。LibreOfficeなど既存OSSもC++による巨大なコードベースという負債を持ち、日本語組版への対応が後回しにされ続けてきた。')
    add_para(doc, '身軽な個人開発者にはこれらの負債がすべてない。意思決定はゼロ秒、方向転換は即時、AIを100%全力で使い倒せる。実際、提案者はこの提案書の作成と並行して開発を開始し、わずか2日で.docx/.xlsx/.pptxの3形式すべてに対応したブラウザ動作デモを完成させた。これは理論ではなく、すでに証明された事実である。')

    add_heading_styled(doc, 'AIコーディングのウィンドウは今しか開いていない', 3)
    add_para(doc, 'Claude Codeをはじめとするエージェント型AI開発環境の登場により、かつては大規模チームにしか不可能だった複雑な実装が個人で現実的になった。このウィンドウは永遠に開いていない。大企業がAIコーディングを本格採用した瞬間、リソース量で逆転される。今動いた者だけが標準を作れる。')

    add_heading_styled(doc, '世界を日本に合わせる', 3)
    add_para(doc, 'これまで日本はWordに合わせて文書を作り、英語UIを日本語化し、海外OSSに日本語対応パッチを当ててきた。常に「世界標準に合わせる」側だった。Oxiは逆の発想で作る。判子・縦書き・禁則処理・ルビ・公文書書式を「例外処理」ではなく「コア機能」として設計する。日本の文書文化を一等市民として扱い、世界がそれに合わせる基盤を作る。同じ問題を抱えるCJK文字圏（中国・韓国・台湾）のデファクトスタンダードになり得る。')

    add_heading_styled(doc, '1-2. 目的', 2)
    add_para(doc, 'すでに動作するOSSオフィススイート「Oxi」（.docx/.xlsx/.pptx/.pdf対応・Rust+Wasm実装）をベースに、Microsoft Officeとの完全互換性を確立し、日本の官公庁・自治体・教育機関への実導入を実現する。ブラウザさえあればWindows・Mac・Linux・Chromebook どこでも動作し、インストール不要・ライセンス費用ゼロ。MITライセンスによる完全無償公開で、財政難の自治体が古いソフトを使い続けざるを得ない悪循環を断ち切る。')

    add_heading_styled(doc, '1-3. 目標（事業期間：2026年7月〜2027年2月）', 2)
    add_bullet(doc, '.docx/.xlsx/.pptx全形式のレイアウト精度向上（Microsoft Officeとのピクセル単位一致を目標とする自動テストスイート構築）')
    add_bullet(doc, '日本語組版の完全実装（JIS X 4051準拠：禁則処理・ルビ・縦書き・均等割り付け）')
    add_bullet(doc, '国内自治体・教育機関との実証実験（最低1団体での試験導入）')
    add_bullet(doc, 'GitHubスター500以上・OSSコミュニティの形成')
    add_bullet(doc, 'CJK文字圏（中国語・韓国語）対応の基盤整備')

    add_heading_styled(doc, '1-4. 競合との差別化', 2)
    add_table(doc,
        ['比較項目', 'Oxi（本提案）', 'LibreOffice', 'Google Docs', 'OnlyOffice'],
        [
            ['実装言語', 'Rust + Wasm', 'C++ / Java', 'JavaScript', 'JavaScript'],
            ['Word互換精度', '完全一致を目標', '差異あり', '差異あり', '差異あり'],
            ['日本語組版', '◎ コア設計', '△ 後付け対応', '△ 部分対応', '△ 部分対応'],
            ['ブラウザ・OS非依存', '◎ Wasmのみ', '× Windows依存強', '◎', '△'],
            ['判子・電子署名', '◎ PAdES対応', '×', '×', '×'],
            ['ライセンス', 'MIT（完全無償）', 'LGPL', 'プロプライエタリ', '一部商用制限'],
        ])

    add_page_break(doc)

    # ============================================
    # Section 2
    # ============================================
    add_heading_styled(doc, '② プロジェクトの内容', 1)

    add_heading_styled(doc, '2-1. 現在の実装状況（提案書提出時点）', 2)
    add_para(doc, '提案書提出時点で以下の機能がすべて実装・動作済みである。')
    add_bullet(doc, '.docx / .xlsx / .pptx パーサー、言語非依存IR、レイアウトエンジン')
    add_bullet(doc, '段落・表・画像・ヘッダー/フッター・ページ罫線の描画')
    add_bullet(doc, '日本語禁則処理（JIS X 4051）、13フォント・約55KBのフォントメトリクス')
    add_bullet(doc, '3形式のラウンドトリップ編集（元のXMLを保持したままパッチ方式で編集・ダウンロード）')
    add_bullet(doc, 'PDF 1.7 パーサー・テキスト抽出・PDF生成（oxipdf-core）')
    add_bullet(doc, 'デジタル判子生成 + PAdES PDF電子署名（oxihanko）')
    add_bullet(doc, 'Wasmビルド + ブラウザWebデモ（https://ryujiyasu.github.io/oxi/）')
    add_bullet(doc, 'E2Eテストスイート・ゴールデンテスト基盤')

    add_heading_styled(doc, '2-2. アーキテクチャ', 2)
    add_bullet(doc, 'oxi-common：共通OOXMLユーティリティ（ZIP・XML・relationships）')
    add_bullet(doc, 'oxidocs-core：.docxエンジン（パーサー・IR・レイアウト・フォントメトリクス・エディタ）')
    add_bullet(doc, 'oxicells-core：.xlsxエンジン（パーサー・IR・エディタ）')
    add_bullet(doc, 'oxislides-core：.pptxエンジン（パーサー・IR・エディタ）')
    add_bullet(doc, 'oxipdf-core：PDF 1.7エンジン（パーサー・テキスト抽出・PDF生成）')
    add_bullet(doc, 'oxihanko：デジタル判子生成 + PAdES電子署名')
    add_bullet(doc, 'oxi-wasm：WebAssemblyバインディング（wasm-bindgen）')

    add_heading_styled(doc, '2-3. 事業期間中に実施する内容', 2)
    add_bullet(doc, '【精度向上】Word/Excel/PowerPointとのピクセル単位レイアウト一致を目指す自動テストスイートの構築・公開。OSSコミュニティと協力して差分を継続的に解消する。')
    add_bullet(doc, '【日本語組版完全実装】縦書き・ルビ（振り仮名）・均等割り付けの実装。JIS X 4051に基づく完全対応。')
    add_bullet(doc, '【自治体実証実験】LibreOfficeへの移行を断念した自治体・教育機関へのアプローチと試験導入。PMネットワーク・Code for Japan等を活用。')
    add_bullet(doc, '【CJK展開基盤】中国語・韓国語の組版対応基盤を整備し、アジア全体のOSSコミュニティへのリーチを確立する。')

    add_page_break(doc)

    # ============================================
    # Section 3
    # ============================================
    add_heading_styled(doc, '③ プロジェクトの計画', 1)

    add_heading_styled(doc, '3-1. スケジュール', 2)
    add_table(doc,
        ['期間', 'フェーズ', '主な作業内容'],
        [
            ['7月〜8月', '精度基盤', '自動テストスイート構築、差分計測システム確立、禁則処理の精度向上'],
            ['9月〜10月', '日本語組版', '縦書き・ルビ・均等割り付けの完全実装、公文書書式への対応強化'],
            ['11月〜12月', '社会実験', '自治体・教育機関への実証実験開始、フィードバックによる改善、コミュニティ形成'],
            ['1月〜2月', '拡張・総括', 'CJK展開基盤整備、実証実験結果報告、成果報告書作成'],
        ])

    add_heading_styled(doc, '3-2. 克服すべき課題と解決策', 2)
    add_bullet(doc, '【課題1】Wordのレイアウトアルゴリズムは非公開：既存のWord文書を大量にサンプリングしレンダリング結果を比較する自動テストスイートを構築・公開。OSSコミュニティと協力して差分を継続的に解消する。')
    add_bullet(doc, '【課題2】フォントメトリクスの取得：GitHub ActionsのWindows runnerを用いてフォントをレンダリング・計測しテーブル化する独自システムをすでに構築済み（13フォント・約55KB JSON）。')
    add_bullet(doc, '【課題3】自治体との接点づくり：採択後のPMネットワーク活用、Code for Japan経由のアプローチ、公開済みWebデモを活用した直接アプローチを並行実施。')
    add_bullet(doc, '【課題4】普及の壁：提案書提出時点ですでに動くデモが公開済み。「LibreOfficeで日本語が崩れた」という体験を持つ担当者への具体的な代替として提示できる。')

    add_page_break(doc)

    # ============================================
    # Section 4
    # ============================================
    add_heading_styled(doc, '④ 提案者の実力・能力を示す実績', 1)

    add_heading_styled(doc, '4-1. 提案者プロフィール', 2)
    add_para(doc, '氏名：安河内 竜二（ヤスコウチ リュウジ）')
    add_para(doc, '所属：株式会社エムスクエア・ラボ　取締役CTO')

    add_heading_styled(doc, '4-2. 学歴・職歴', 2)
    add_table(doc,
        ['年月', '内容'],
        [
            ['2013年3月', '京都大学農学部 卒業'],
            ['2015年3月', '京都大学大学院理学研究科 生物科学専攻 修了'],
            ['2016年4月', '月桂冠株式会社 入社（システム開発・IT業務）'],
            ['2020年〜', '株式会社エムスクエア・ラボ 入社、取締役CTO就任（現任）'],
        ])

    add_heading_styled(doc, '4-3. 技術スキル・実績', 2)
    add_bullet(doc, 'Rust + Wasm：提案書提出から2日で.docx/.xlsx/.pptx/.pdf対応・判子・PAdES署名・E2Eテスト基盤まで完成（https://ryujiyasu.github.io/oxi/）')
    add_bullet(doc, 'OOXMLフォーマット解析：Word/Excel/PowerPointの内部仕様を独自に解析・実装。PDF 1.7パーサー・PAdES電子署名も実装済み')
    add_bullet(doc, '農地ナビ＋：農地情報・衛星データを統合した農業向けナビゲーションシステムの開発・運用')
    add_bullet(doc, 'Cloudflare Workers/D1/KV：複数プロダクトの本番運用実績')
    add_bullet(doc, 'React/TypeScript：フロントエンド開発の実務経験')
    add_bullet(doc, 'AI活用開発：Claude APIを用いたプロダクト開発（SaaS・チャットボット等）の実績')

    add_heading_styled(doc, '4-4. プロジェクト実現性の根拠', 2)
    add_para(doc, '提案書提出から2日で全形式対応・判子機能・PAdES電子署名・E2Eテスト基盤まで含むブラウザ動作デモを完成させた。これは理論ではなく実績である。大企業が仕様検討をしている間に、個人がAIコーディングを全力投球することでここまで実装できることの証明でもある。')
    add_para(doc, 'GitHubリポジトリ：https://github.com/Ryujiyasu/oxi（MITライセンス・公開済み）')
    add_para(doc, 'Webデモ：https://ryujiyasu.github.io/oxi/')

    add_page_break(doc)

    # ============================================
    # Section 5
    # ============================================
    add_heading_styled(doc, '⑤ 事業化・社会実装の新規性・優位性、想定するターゲットと規模', 1)

    add_heading_styled(doc, '5-1. 市場規模', 2)
    add_para(doc, '国内のMicrosoft Officeライセンス市場は年間推計3,000〜5,000億円規模と言われる。官公庁・自治体（約1,800団体）および大学・教育機関（約1,200校）だけでも、年間数百億円のライセンス費用が発生している。OxiはMITライセンスで完全無償提供するため、財政難の自治体が古いソフトを使い続けざるを得ない悪循環を解決できる唯一のOSS基盤となり得る。Wasm＋ブラウザネイティブ設計によりWindowsへの依存も不要となるため、OS調達コストの削減も同時に実現できる。')

    add_heading_styled(doc, '5-2. ターゲットユーザー', 2)
    add_bullet(doc, '一次ターゲット：脱Office・脱Windows推進中の地方自治体・官公庁')
    add_bullet(doc, '二次ターゲット：教育機関（大学・高専・高校）')
    add_bullet(doc, '三次ターゲット：文書管理SaaSを開発するスタートアップ・IT企業（APIとして利用）')
    add_bullet(doc, 'グローバル：CJK文字圏（中国・韓国・台湾）のOSSコミュニティ')

    add_heading_styled(doc, '5-3. 競合サービスとの比較', 2)
    add_para(doc, 'LibreOffice：歴史あるOSSだがC++実装・デスクトップ前提・日本語レイアウト精度に根本的な問題がある。実際に多くの自治体がLibreOfficeへの移行を試みて断念した。Oxiはブラウザネイティブ・インストール不要・日本語を一等市民として設計している点で全方位的に差別化される。')
    add_para(doc, 'Google Docs / OnlyOffice：独自形式への変換が主体で.docxの完全互換を目指していない。また前者はプロプライエタリ、後者は地政学的リスクを抱える。')

    add_page_break(doc)

    # ============================================
    # Section 6
    # ============================================
    add_heading_styled(doc, '⑥ 事業化・社会実装の具体的な進め方', 1)

    add_heading_styled(doc, '6-1. ビジネスモデル', 2)
    add_para(doc, 'コアエンジンはMITライセンスで永続的に無償公開する。事業期間終了後は以下のモデルを組み合わせる。')
    add_bullet(doc, '【モデル1】OSSコア＋エンタープライズサポート（RedHatモデル）：基本機能を無償提供し、サポート契約・カスタマイズ開発を有償化。自治体1団体年間数百万円 × 全国1,800団体がポテンシャル。')
    add_bullet(doc, '【モデル2】クラウドSaaS：Oxiエンジンを用いたOffice互換文書エディタをSaaSとして提供。月額課金モデル。')
    add_bullet(doc, '【モデル3】APIライセンス：.docx変換・レンダリング・PDF生成APIを他社SaaSへBtoBで提供。')

    add_heading_styled(doc, '6-2. 事業期間中の事業化に向けた作業', 2)
    add_bullet(doc, 'GitHubリポジトリのスター獲得・コミュニティ形成（目標：500スター）')
    add_bullet(doc, '公開済みWebデモを活用した自治体・教育機関へのアプローチ')
    add_bullet(doc, '採択後のPMネットワークを活用した自治体コネクション獲得')
    add_bullet(doc, 'NLnet Foundation等の欧州OSSファンディングへの申請準備')

    add_page_break(doc)

    # ============================================
    # Section 7
    # ============================================
    add_heading_styled(doc, '⑦ 事業期間終了後の事業化・社会実装に関する計画', 1)

    add_heading_styled(doc, '7-1. アウトプット形態', 2)
    add_para(doc, 'すでにGitHubでMITライセンスにて公開済み（https://github.com/Ryujiyasu/oxi）。事業期間を通じてコミュニティからのバグ報告・PRを取り込みながら精度を高める。完全互換はコミュニティと協力しながら育てていく。')

    add_heading_styled(doc, '7-2. ロードマップ（事業期間終了後）', 2)
    add_table(doc,
        ['フェーズ', '時期', '内容'],
        [
            ['v2', '〜2027年末', 'CRDTによるリアルタイム共同編集、AIアシスト機能、エンドツーエンド暗号化、PWA/オフライン対応'],
            ['v3', '〜2028年末', 'SaaS正式開始、CJK市場（中国・台湾）展開、プラグインシステム、Tauriデスクトップ・モバイルアプリ'],
            ['v4', '〜2029年末', 'エンタープライズ向けコンプライアンス・監査証跡、業界特化（法務・医療・行政）、LibreOfficeに代わる次世代OSSスイートとして世界標準化'],
        ])

    add_page_break(doc)

    # ============================================
    # Section 8
    # ============================================
    add_heading_styled(doc, '⑧ 予算内訳のまとめ', 1)

    add_heading_styled(doc, '8-1. 稼働計画', 2)
    add_table(doc,
        ['メンバー', '月平均稼働時間', '事業期間（月数）', '総稼働時間', '金額（税抜）'],
        [
            ['安河内 竜二（課税事業者）', '80時間/月', '8ヶ月', '640時間', '3,200,000円'],
            ['合計', '', '', '640時間', '3,200,000円'],
        ])

    add_heading_styled(doc, '8-2. 費用総額', 2)
    add_table(doc,
        ['項目', '金額'],
        [
            ['作業費（税抜）：640時間 × 5,000円/時間', '3,200,000円'],
            ['消費税（10%）', '320,000円'],
            ['実施費用総額（税込・申請金額）', '3,520,000円'],
        ])

    add_para(doc, '※時間単価：5,000円/時間（公募要領4.(4)①に基づく一律単価）')
    add_para(doc, '※1名プロジェクトのため上限800万円の範囲内')

    add_page_break(doc)

    # ============================================
    # Section 9
    # ============================================
    add_heading_styled(doc, '⑨ 提案プロジェクトと所属企業のビジネスモデル・技術開発との差異', 1)

    add_heading_styled(doc, '9-1. 所属組織について', 2)
    add_para(doc, '提案者は株式会社エムスクエア・ラボの取締役CTOとして在籍している。同社は農業ロボット・スマート農業システムの研究開発・販売を主事業としており、農業IoTセンサー・農機自動化システム・農業データ分析プラットフォームの開発に特化している。')

    add_heading_styled(doc, '9-2. 本提案プロジェクトとの差異', 2)
    add_para(doc, '本提案プロジェクトは文書処理・Officeフォーマット互換性というテーマであり、農業ロボット・スマート農業という所属組織の事業内容と技術的・ビジネス的に全く異なる領域である。所属組織の技術開発はROS・センサーフュージョン・農業機械制御を中心としており、RustやWebAssemblyを用いたブラウザ上の文書処理エンジンの開発とは接点がない。')
    add_para(doc, '本プロジェクトは提案者が個人として取り組む新規OSS開発であり、所属組織のビジネスモデルや技術開発と類似・重複する内容は一切含まない。')

    add_heading_styled(doc, '9-3. 所属組織からの了解', 2)
    add_para(doc, '所属組織（株式会社エムスクエア・ラボ）からは、本事業による支援措置を受けること、および開発成果がイノベータ個人に帰属することについて了解を得ている。契約時には書面による承諾書を提出する。')

    # Footer
    add_para(doc, '')
    p = add_para(doc, '本提案書のPDFはOxi（oxipdf-core）で生成しました。', size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    output = sys.argv[1] if len(sys.argv) > 1 else 'proposal.docx'
    doc.save(output)
    print(f'Written to {output}')

if __name__ == '__main__':
    main()
