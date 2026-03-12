"""Generate comprehensive test docx files for Oxi rendering verification.

Each file tests specific OOXML features that need to be rendered correctly.
Run: python tools/gen_test_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import docx.oxml

OUT = os.path.join(os.path.dirname(__file__), '..', 'tests', 'fixtures')
os.makedirs(OUT, exist_ok=True)

# ============================================================
# 1. Text Formatting
# ============================================================
def gen_text_formatting():
    doc = Document()
    doc.add_heading('Text Formatting Test', level=1)

    p = doc.add_paragraph()
    r = p.add_run('Normal text. ')
    r = p.add_run('Bold text. ')
    r.bold = True
    r = p.add_run('Italic text. ')
    r.italic = True
    r = p.add_run('Bold+Italic. ')
    r.bold = True
    r.italic = True

    p = doc.add_paragraph()
    r = p.add_run('Single underline. ')
    r.underline = WD_UNDERLINE.SINGLE
    r = p.add_run('Double underline. ')
    r.underline = WD_UNDERLINE.DOUBLE
    r = p.add_run('Dotted underline. ')
    r.underline = WD_UNDERLINE.DOTTED
    r = p.add_run('Wave underline. ')
    r.underline = WD_UNDERLINE.WAVY

    p = doc.add_paragraph()
    r = p.add_run('Strikethrough. ')
    r.font.strike = True
    r = p.add_run('Double strikethrough. ')
    r.font.double_strike = True

    p = doc.add_paragraph()
    r = p.add_run('Superscript')
    r.font.superscript = True
    r = p.add_run(' Normal ')
    r = p.add_run('Subscript')
    r.font.subscript = True

    p = doc.add_paragraph()
    r = p.add_run('Red text ')
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r = p.add_run('Green text ')
    r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    r = p.add_run('Blue text ')
    r.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    p = doc.add_paragraph()
    for size in [8, 10, 12, 14, 16, 20, 24]:
        r = p.add_run(f'{size}pt ')
        r.font.size = Pt(size)

    p = doc.add_paragraph()
    r = p.add_run('Small Caps Text ')
    r.font.small_caps = True
    r = p.add_run('ALL CAPS text ')
    r.font.all_caps = True

    # Highlight
    p = doc.add_paragraph()
    r = p.add_run('Yellow highlight ')
    r.font.highlight_color = 7  # YELLOW
    r = p.add_run('Cyan highlight ')
    r.font.highlight_color = 3  # TURQUOISE

    doc.save(os.path.join(OUT, 'test_text_formatting.docx'))
    print('  -> test_text_formatting.docx')


# ============================================================
# 2. Tab Stops (Left, Center, Right, Decimal)
# ============================================================
def gen_tab_stops():
    doc = Document()
    doc.add_heading('Tab Stop Test', level=1)

    # Left tab
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(1.5), WD_TAB_ALIGNMENT.LEFT)
    p.add_run('Left:\tAligned at 1.5"')

    # Center tab
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.CENTER)
    p.add_run('Center:\tCentered at 3"')

    # Right tab
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(5.0), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run('Right:\tRight-aligned at 5"')

    # Decimal tab
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.DECIMAL)
    p.add_run('Decimal:\t123.45')

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.DECIMAL)
    p.add_run('Decimal:\t1,234.5')

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.DECIMAL)
    p.add_run('Decimal:\t67.890')

    # Multiple tabs on one line
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
    pf.tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run('Col1\tCol2\tCol3')

    # Tab with leader dots
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(5.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.add_run('Chapter 1\t10')

    doc.save(os.path.join(OUT, 'test_tab_stops.docx'))
    print('  -> test_tab_stops.docx')


# ============================================================
# 3. Page Break & Multi-page
# ============================================================
def gen_page_break():
    doc = Document()
    doc.add_heading('Page 1 Content', level=1)
    doc.add_paragraph('This is the first page with some text content.')
    doc.add_paragraph('More text on page 1.')

    # Explicit page break
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.add_heading('Page 2 Content', level=1)
    doc.add_paragraph('This text should appear on the second page.')

    # Another page break
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.add_heading('Page 3 Content', level=1)
    doc.add_paragraph('This text should appear on the third page.')

    doc.save(os.path.join(OUT, 'test_page_break.docx'))
    print('  -> test_page_break.docx')


# ============================================================
# 4. Hyperlinks & Bookmarks
# ============================================================
def gen_hyperlinks():
    doc = Document()
    doc.add_heading('Hyperlink & Bookmark Test', level=1)

    doc.add_paragraph('Normal text before hyperlink.')

    # Add hyperlink via XML manipulation (python-docx doesn't natively support hyperlinks well)
    p = doc.add_paragraph()
    p.add_run('Visit ')
    _add_hyperlink(p, 'https://example.com', 'Example Website')
    p.add_run(' for more info.')

    # Internal bookmark
    p = doc.add_paragraph('Text with bookmark target below.')

    # Bookmark (add via XML)
    p2 = doc.add_paragraph()
    _add_bookmark(p2, 'my_anchor', 'Bookmarked Text Here')

    doc.save(os.path.join(OUT, 'test_hyperlinks.docx'))
    print('  -> test_hyperlinks.docx')


def _add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'  <w:r>'
        f'    <w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        f'    <w:t>{text}</w:t>'
        f'  </w:r>'
        f'</w:hyperlink>'
    )
    paragraph._element.append(hyperlink)


def _add_bookmark(paragraph, name, text):
    """Add a bookmark anchor."""
    bm_start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="1" w:name="{name}"/>'
    )
    bm_end = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="1"/>'
    )
    paragraph._element.append(bm_start)
    r = paragraph.add_run(text)
    r.bold = True
    paragraph._element.append(bm_end)


# ============================================================
# 5. Table: Cell Merge, Row Height, vAlign
# ============================================================
def gen_table_advanced():
    doc = Document()
    doc.add_heading('Advanced Table Test', level=1)

    # 4x4 table with merges
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # Horizontal merge (gridSpan): merge A1+B1
    a1 = table.cell(0, 0)
    b1 = table.cell(0, 1)
    a1.merge(b1)
    a1.text = 'Merged A1+B1'

    table.cell(0, 2).text = 'C1'
    table.cell(0, 3).text = 'D1'

    # Vertical merge: merge A2+A3
    a2 = table.cell(1, 0)
    a3 = table.cell(2, 0)
    a2.merge(a3)
    a2.text = 'Merged A2+A3'

    table.cell(1, 1).text = 'B2'
    table.cell(1, 2).text = 'C2'
    table.cell(1, 3).text = 'D2'

    table.cell(2, 1).text = 'B3'
    table.cell(2, 2).text = 'C3'
    table.cell(2, 3).text = 'D3'

    # Row 4 with cell shading and vAlign
    for i in range(4):
        cell = table.cell(3, i)
        cell.text = f'Row4-{chr(65+i)}'

    # Cell shading
    cell_shd = table.cell(3, 0)._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFD700" w:val="clear"/>')
    cell_shd.append(shd)

    # vAlign center
    table.cell(3, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # vAlign bottom
    table.cell(3, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

    # Row height
    row = table.rows[3]
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="720" w:hRule="atLeast"/>')
    trPr.append(trHeight)

    doc.add_paragraph()  # spacer

    # Table with alignment (centered)
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(0, 0).text = 'Centered'
    table2.cell(0, 1).text = 'Table'
    table2.cell(1, 0).text = 'Row 2'
    table2.cell(1, 1).text = 'Col 2'

    doc.save(os.path.join(OUT, 'test_table_advanced.docx'))
    print('  -> test_table_advanced.docx')


# ============================================================
# 6. Indent, Spacing, Borders
# ============================================================
def gen_indent_spacing():
    doc = Document()
    doc.add_heading('Indent & Spacing Test', level=1)

    # Left indent
    p = doc.add_paragraph('Left indent 1 inch')
    p.paragraph_format.left_indent = Inches(1.0)

    # Right indent
    p = doc.add_paragraph('Right indent 1 inch')
    p.paragraph_format.right_indent = Inches(1.0)

    # First line indent
    p = doc.add_paragraph('First line indent 0.5 inch. This is a longer paragraph to show that only the first line is indented while subsequent lines wrap normally at the left margin.')
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Hanging indent
    p = doc.add_paragraph('Hanging indent: first line at 0, rest at 0.5 inch. This paragraph demonstrates a hanging indent where subsequent lines are indented more than the first.')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

    # Space before/after
    p = doc.add_paragraph('Space before 24pt, after 12pt')
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph('Normal paragraph after spacing.')

    # Line spacing
    p = doc.add_paragraph('Line spacing 2.0 (double). This paragraph has double line spacing to test that the layout engine correctly applies the multiplier to the base line height.')
    p.paragraph_format.line_spacing = 2.0

    p = doc.add_paragraph('Exact line spacing 24pt. This paragraph has exact line spacing set.')
    p.paragraph_format.line_spacing = Pt(24)
    from docx.enum.text import WD_LINE_SPACING
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # Paragraph border
    p = doc.add_paragraph('Paragraph with bottom border')
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Paragraph shading
    p = doc.add_paragraph('Paragraph with light gray background')
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E0E0E0" w:val="clear"/>')
    pPr.append(shd)

    doc.save(os.path.join(OUT, 'test_indent_spacing.docx'))
    print('  -> test_indent_spacing.docx')


# ============================================================
# 7. Fields (PAGE, NUMPAGES) in header/footer
# ============================================================
def gen_fields():
    doc = Document()

    # Add header with PAGE field
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.add_run('Header - Page ')
    _add_field(hp, 'PAGE')
    hp.add_run(' of ')
    _add_field(hp, 'NUMPAGES')

    # Add footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.add_run('Footer text')

    doc.add_heading('Page 1 - Fields Test', level=1)
    doc.add_paragraph('This document tests PAGE and NUMPAGES field codes in headers.')

    # Add enough content for 2 pages
    for i in range(30):
        doc.add_paragraph(f'Line {i+1}: Lorem ipsum dolor sit amet, consectetur adipiscing elit.')

    doc.save(os.path.join(OUT, 'test_fields.docx'))
    print('  -> test_fields.docx')


def _add_field(paragraph, field_code):
    """Add a simple field code (e.g. PAGE, NUMPAGES) to a paragraph."""
    run = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {field_code} </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_sep)

    run4 = paragraph.add_run('#')  # placeholder

    run5 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar_end)


# ============================================================
# 8. Numbered Lists (multilevel)
# ============================================================
def gen_lists():
    doc = Document()
    doc.add_heading('List Test', level=1)

    # Bullet list
    doc.add_paragraph('First bullet item', style='List Bullet')
    doc.add_paragraph('Second bullet item', style='List Bullet')
    doc.add_paragraph('Third bullet item', style='List Bullet')

    doc.add_paragraph()  # spacer

    # Numbered list
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    doc.add_paragraph('Third numbered item', style='List Number')

    doc.add_paragraph()  # spacer

    # Multilevel (using List Bullet 2, 3)
    doc.add_paragraph('Level 1 item', style='List Bullet')
    doc.add_paragraph('Level 2 item', style='List Bullet 2')
    doc.add_paragraph('Level 2 item', style='List Bullet 2')
    doc.add_paragraph('Level 3 item', style='List Bullet 3')
    doc.add_paragraph('Level 1 item', style='List Bullet')

    doc.save(os.path.join(OUT, 'test_lists.docx'))
    print('  -> test_lists.docx')


# ============================================================
# 9. Japanese text (mixed CJK/Latin, kinsoku)
# ============================================================
def gen_japanese():
    doc = Document()
    doc.add_heading('日本語テスト', level=1)

    doc.add_paragraph('これは日本語のテスト文書です。禁則処理が正しく動作するか確認します。')

    doc.add_paragraph('句読点の禁則：「これは括弧の中です。」句読点（、。）が行頭に来ないことを確認。')

    # Mixed CJK + Latin
    doc.add_paragraph('English text mixed with 日本語テキスト in the same paragraph. ABCDEあいうえお12345。')

    # Long CJK paragraph (for line break testing)
    doc.add_paragraph(
        'これは長い日本語の段落です。複数行にわたるテキストで、禁則処理のテストを行います。'
        '句読点（、。）や括弧（「」『』）が行頭・行末で正しく処理されるか確認します。'
        'また、全角数字１２３４５や全角アルファベットＡＢＣも含みます。'
        'さらに、カタカナのテストも行います：アイウエオカキクケコ。'
    )

    # Different fonts
    p = doc.add_paragraph()
    r = p.add_run('ＭＳ ゴシック: ')
    r.font.name = 'MS Gothic'
    r = p.add_run('テスト文字列')
    r.font.name = 'MS Gothic'

    p = doc.add_paragraph()
    r = p.add_run('ＭＳ 明朝: ')
    r.font.name = 'MS Mincho'
    r = p.add_run('テスト文字列')
    r.font.name = 'MS Mincho'

    doc.save(os.path.join(OUT, 'test_japanese.docx'))
    print('  -> test_japanese.docx')


# ============================================================
# 10. SDT (Structured Document Tags / Content Controls)
# ============================================================
def gen_sdt():
    doc = Document()
    doc.add_heading('SDT / Content Control Test', level=1)

    doc.add_paragraph('Text before SDT.')

    # Add block-level SDT via XML
    body = doc.element.body
    sdt = parse_xml(
        f'<w:sdt {nsdecls("w")}>'
        f'  <w:sdtPr><w:alias w:val="TestControl"/></w:sdtPr>'
        f'  <w:sdtContent>'
        f'    <w:p><w:r><w:t>Content inside SDT block</w:t></w:r></w:p>'
        f'  </w:sdtContent>'
        f'</w:sdt>'
    )
    body.append(sdt)

    doc.add_paragraph('Text after SDT.')

    doc.save(os.path.join(OUT, 'test_sdt.docx'))
    print('  -> test_sdt.docx')


# ============================================================
# 11. All-in-one comprehensive test
# ============================================================
def gen_all_in_one():
    doc = Document()

    # === Section 1: Text Formatting ===
    doc.add_heading('Section 1: Text Formatting', level=1)
    p = doc.add_paragraph()
    p.add_run('Normal ')
    r = p.add_run('Bold ')
    r.bold = True
    r = p.add_run('Italic ')
    r.italic = True
    r = p.add_run('Underline ')
    r.underline = WD_UNDERLINE.SINGLE
    r = p.add_run('Strike ')
    r.font.strike = True
    r = p.add_run('Red ')
    r.font.color.rgb = RGBColor(0xFF, 0, 0)

    # === Section 2: Alignment ===
    doc.add_heading('Section 2: Alignment', level=2)
    doc.add_paragraph('Left aligned (default)')
    p = doc.add_paragraph('Center aligned')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Right aligned')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('Justified text. ' * 10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # === Section 3: Lists ===
    doc.add_heading('Section 3: Lists', level=2)
    doc.add_paragraph('Bullet 1', style='List Bullet')
    doc.add_paragraph('Bullet 2', style='List Bullet')
    doc.add_paragraph('Number 1', style='List Number')
    doc.add_paragraph('Number 2', style='List Number')

    # === Section 4: Table with merge ===
    doc.add_heading('Section 4: Table', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = 'Merged'
    table.cell(0, 2).text = 'C1'
    for r in range(1, 3):
        for c in range(3):
            table.cell(r, c).text = f'R{r+1}C{c+1}'

    # === Section 5: Indent ===
    doc.add_heading('Section 5: Indentation', level=2)
    p = doc.add_paragraph('Left indent 1"')
    p.paragraph_format.left_indent = Inches(1.0)
    p = doc.add_paragraph('First line indent 0.5" with enough text to wrap to the next line and show the hanging behavior clearly.')
    p.paragraph_format.first_line_indent = Inches(0.5)

    # === Section 6: Japanese ===
    doc.add_heading('Section 6: 日本語', level=2)
    doc.add_paragraph('日本語テキストのテスト。禁則処理：「括弧」や、句読点。')
    doc.add_paragraph('Mixed: Hello世界！ABCDEあいうえお12345。')

    # === Page break to page 2 ===
    p = doc.add_paragraph()
    p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.add_heading('Page 2: More Content', level=1)
    doc.add_paragraph('This content is on the second page after an explicit page break.')

    # Tab stops
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(2.0), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(5.0), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run('Left\tCenter\tRight')

    doc.save(os.path.join(OUT, 'test_all_in_one.docx'))
    print('  -> test_all_in_one.docx')


# ============================================================
# Main
# ============================================================
if __name__ == '__main__':
    print('Generating test docx files...')
    gen_text_formatting()
    gen_tab_stops()
    gen_page_break()
    try:
        gen_hyperlinks()
    except Exception as e:
        print(f'  -> test_hyperlinks.docx SKIPPED: {e}')
    gen_table_advanced()
    gen_indent_spacing()
    gen_fields()
    gen_lists()
    gen_japanese()
    gen_sdt()
    gen_all_in_one()
    print('Done!')
