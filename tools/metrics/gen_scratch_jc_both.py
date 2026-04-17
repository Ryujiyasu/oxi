"""Scratch minimal docx with ONLY jc=both — test if R3 finding holds independently.

If this reproduces ・=9.5pt, then jc=both alone is the trigger.
If it shows ・=12pt, then R3's trigger requires additional d77a properties
(sectPr/settings/fontTable not replicated here).
"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data",
                 "scratch_jc_both.docx")
)

TEST_TEXT = "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"


def set_font(run, family="ＭＳ ゴシック", size_pt=12.0):
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    rFonts.set(qn("w:hint"), "eastAsia")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)


def main():
    doc = Document()
    # compressPunctuation on
    settings = doc.settings.element
    for e in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(e)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)

    # Paragraph with explicit jc=both
    p = doc.add_paragraph()
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr"); p._element.insert(0, pPr)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "both")
    pPr.append(jc)

    r = p.add_run(TEST_TEXT)
    set_font(r)

    doc.save(OUT)
    print(f"[scratch_jc_both] {OUT}")


if __name__ == "__main__":
    main()
