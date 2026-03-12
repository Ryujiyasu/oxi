"""
Verify whether "Exactly" line spacing mode enforces a minimum height.
Tests both CJK and Western fonts with Exactly values below natural line height.
Also tests: does the minimum equal Single spacing height?
"""
import win32com.client
import os, time, tempfile

def create_test_docx(path):
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.oxml.ns import qn

    doc = Document()
    section = doc.sections[0]
    sectPr = section._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)
    grid_el = sectPr.makeelement(qn('w:docGrid'), {qn('w:type'): 'default', qn('w:linePitch'): '360'})
    sectPr.append(grid_el)

    fonts_and_sizes = [
        ("Calibri", 10.5),
        ("Calibri", 14),
        ("Times New Roman", 12),
        ("Arial", 11),
        ("Yu Gothic", 10.5),
        ("Yu Gothic", 14),
        ("MS Gothic", 10.5),
        ("MS Mincho", 10.5),
    ]

    tests = []

    for font, size in fonts_and_sizes:
        # First: Single spacing (baseline)
        tests.append((f"{font} {size}pt Single", font, size, None, None))
        # Then: Exactly = various values from very small to natural
        for exact_pt in [5, 8, 10, 12, 14, 16, 18, 20, 24]:
            exact_twips = int(exact_pt * 20)
            tests.append((f"{font} {size}pt Exact={exact_pt}", font, size, exact_twips, "exact"))

    for desc, font, size, line_val, rule in tests:
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()

        # Disable grid snap
        snap = pPr.makeelement(qn('w:snapToGrid'), {qn('w:val'): '0'})
        pPr.append(snap)

        # Spacing
        attrs = {qn('w:before'): '0', qn('w:after'): '0'}
        if line_val is not None:
            attrs[qn('w:line')] = str(line_val)
            attrs[qn('w:lineRule')] = rule
        else:
            attrs[qn('w:line')] = '240'
            attrs[qn('w:lineRule')] = 'auto'
        spacing = pPr.makeelement(qn('w:spacing'), attrs)
        pPr.append(spacing)

        run = p.add_run("ABCabc test text")
        run.font.name = font
        run.font.size = Pt(size)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font)

    doc.save(path)
    return tests

def measure(word, doc_path, tests):
    doc = word.Documents.Open(doc_path)
    time.sleep(1)
    doc.Repaginate()
    time.sleep(0.5)

    results = []
    n = len(tests)
    for i in range(n):
        para = doc.Paragraphs(i + 1)
        rng = para.Range
        rng.Collapse(1)
        y = rng.Information(6)

        if i + 1 < n:
            next_rng = doc.Paragraphs(i + 2).Range
            next_rng.Collapse(1)
            y_next = next_rng.Information(6)
            h = y_next - y
        else:
            h = None

        results.append((tests[i][0], h))

    doc.Close(0)
    return results

def main():
    tmp = os.path.join(tempfile.gettempdir(), "test_exact_minimum.docx")
    print("Creating test document...")
    tests = create_test_docx(tmp)

    print("Opening Word COM...")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        results = measure(word, tmp, tests)

        # Group by font+size
        current_font = None
        single_h = None

        print("\n=== Exactly Mode Minimum Test ===")
        print(f"{'Description':<40} {'Measured':>10} {'vs Single':>12}")
        print("-" * 65)

        for desc, h in results:
            if "Single" in desc:
                single_h = h
                current_font = desc.split(" Single")[0]
                print(f"\n--- {current_font} ---")
                h_str = f"{h:.2f}pt" if h else "N/A"
                print(f"  {'Single spacing':<36} {h_str:>10} {'(baseline)':>12}")
            elif h is not None and single_h is not None:
                exact_val = desc.split("Exact=")[1]
                exact_pt = float(exact_val)
                diff = h - single_h
                clamp = "CLAMPED" if abs(h - single_h) < 0.05 and exact_pt < single_h else ""
                follows = "EXACT" if abs(h - exact_pt) < 0.05 else ""
                note = clamp or follows or f"diff={diff:+.2f}"
                print(f"  Exactly {exact_val:>4}pt: {h:10.2f}pt {note:>12}")

    finally:
        word.Quit()

    os.unlink(tmp)

if __name__ == "__main__":
    main()
