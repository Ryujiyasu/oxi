"""Author minimal repro docx files for investigating Word's LINE-LEVEL
row-split behavior — i.e. where inside a table cell paragraph Word places
the page break when the wrapped lines don't all fit on one page.

d77a mimic: MS Gothic 10.5pt, A4 2.5cm margins, single-column body. Push
the table to the bottom of p.1 so a controlled number of lines overflow
to p.2.

Output dir: tools/metrics/row_split_content_repro/
"""
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import sys
    print("python-docx required: pip install python-docx")
    sys.exit(1)


OUT_DIR = Path(__file__).parent / "row_split_content_repro"
OUT_DIR.mkdir(exist_ok=True)

MS_GOTHIC = "ＭＳ ゴシック"


def set_table_borders_all(table, sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_tr_cant_split(row):
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    old = trPr.find(qn('w:cantSplit'))
    if old is None:
        trPr.append(OxmlElement('w:cantSplit'))


def set_para_flag(p, tag: str, val: str | None = None):
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn(f'w:{tag}'))
    if old is not None:
        pPr.remove(old)
    el = OxmlElement(f'w:{tag}')
    if val is not None:
        el.set(qn('w:val'), val)
    pPr.append(el)


def set_run_gothic(run, pt: float = 10.5):
    r = run._r
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    # Remove existing fonts
    for child_tag in ('w:rFonts', 'w:sz', 'w:szCs'):
        old = rPr.find(qn(child_tag))
        if old is not None:
            rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), MS_GOTHIC)
    rFonts.set(qn('w:eastAsia'), MS_GOTHIC)
    rFonts.set(qn('w:hAnsi'), MS_GOTHIC)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(pt * 2)))
    rPr.append(szCs)


def set_section_docgrid(section, line_pitch_tw: int = 400):
    # line_pitch in twips — 400 twips = 20pt grid (d77a default)
    sectPr = section._sectPr
    old = sectPr.find(qn('w:docGrid'))
    if old is not None:
        sectPr.remove(old)
    dg = OxmlElement('w:docGrid')
    dg.set(qn('w:type'), 'linesAndChars')
    dg.set(qn('w:linePitch'), str(line_pitch_tw))
    dg.set(qn('w:charSpace'), '0')
    sectPr.append(dg)


def build_base_doc(filler_paragraphs: int = 20) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    set_section_docgrid(section, line_pitch_tw=400)

    for i in range(filler_paragraphs):
        p = doc.add_paragraph()
        run = p.add_run(f"Filler line {i+1:02d}. " + "あ" * 30)
        set_run_gothic(run, pt=10.5)
    return doc


def add_text_cell_para(cell, text: str, pt: float = 10.5,
                        keep_lines: bool = False,
                        widow_off: bool = False,
                        first_para: bool = False):
    if first_para:
        p = cell.paragraphs[0]
        # Clear existing content
        for r in list(p._p.findall(qn('w:r'))):
            p._p.remove(r)
    else:
        p = cell.add_paragraph()
    if keep_lines:
        set_para_flag(p, 'keepLines', None)
    if widow_off:
        set_para_flag(p, 'widowControl', '0')
    run = p.add_run(text)
    set_run_gothic(run, pt=pt)
    return p


def make_rsA(filler: int = 18):
    """Baseline: 1-row 1-cell, ONE long paragraph wrapping to ~6-10 lines,
    placed so middle overflows to p.2. cantSplit off, default everything."""
    doc = build_base_doc(filler_paragraphs=filler)
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    cell = table.cell(0, 0)
    # Long single paragraph — 40 chars × 10 lines = 400 chars
    text = "".join(f"{i+1:02d}文字目アイウエオカキクケコサシスセソタチツテト" for i in range(20))
    add_text_cell_para(cell, text, first_para=True)
    path = OUT_DIR / "RS_A_baseline.docx"
    doc.save(path)
    return path


def make_rsB(filler: int = 18):
    """cantSplit ON: whole row must move to p.2."""
    doc = build_base_doc(filler_paragraphs=filler)
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    set_tr_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    text = "".join(f"{i+1:02d}文字目アイウエオカキクケコサシスセソタチツテト" for i in range(20))
    add_text_cell_para(cell, text, first_para=True)
    path = OUT_DIR / "RS_B_cantSplit.docx"
    doc.save(path)
    return path


def make_rsC(filler: int = 18):
    """keepLines ON on the cell paragraph: para should stay together."""
    doc = build_base_doc(filler_paragraphs=filler)
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    cell = table.cell(0, 0)
    text = "".join(f"{i+1:02d}文字目アイウエオカキクケコサシスセソタチツテト" for i in range(20))
    add_text_cell_para(cell, text, keep_lines=True, first_para=True)
    path = OUT_DIR / "RS_C_keepLines.docx"
    doc.save(path)
    return path


def make_rsD(filler: int = 18):
    """1-row 2-cell: each cell has a long paragraph. Verify split happens
    per cell independently or together."""
    doc = build_base_doc(filler_paragraphs=filler)
    table = doc.add_table(rows=1, cols=2)
    set_table_borders_all(table)
    text_l = "".join(f"L{i+1:02d}アイウエオカキクケコサシスセソタチ" for i in range(10))
    text_r = "".join(f"R{i+1:02d}あいうえおかきくけこさしすせそたち" for i in range(10))
    add_text_cell_para(table.cell(0, 0), text_l, first_para=True)
    add_text_cell_para(table.cell(0, 1), text_r, first_para=True)
    path = OUT_DIR / "RS_D_2cell.docx"
    doc.save(path)
    return path


def make_rsE(filler: int = 18):
    """3-row table, each row 1 cell 1 short paragraph; third row overflows.
    Verify Word splits inside row 3 or moves row 3 entirely."""
    doc = build_base_doc(filler_paragraphs=filler)
    table = doc.add_table(rows=3, cols=1)
    set_table_borders_all(table)
    # Rows 1 and 2 short content
    for i in range(2):
        short = f"Row {i+1}: 短い内容、一行で収まるはず。"
        add_text_cell_para(table.cell(i, 0), short, first_para=True)
    # Row 3 long — wraps and overflows
    text = "".join(f"{i+1:02d}文字目アイウエオカキクケコサシスセソタチツテト" for i in range(8))
    add_text_cell_para(table.cell(2, 0), text, first_para=True)
    path = OUT_DIR / "RS_E_multirow.docx"
    doc.save(path)
    return path


def make_rsF(filler: int = 18):
    """1-row 1-cell, THREE paragraphs each 3-4 lines. Check whether Word
    breaks at paragraph boundary or line boundary inside a paragraph."""
    doc = build_base_doc(filler_paragraphs=filler)
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    cell = table.cell(0, 0)
    para_text = [
        "段落1: " + "アイウエオカキクケコサシスセソタチツテト" * 4,
        "段落2: " + "あいうえおかきくけこさしすせそたちつてと" * 4,
        "段落3: " + "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnop" * 3,
    ]
    for i, t in enumerate(para_text):
        add_text_cell_para(cell, t, first_para=(i == 0))
    path = OUT_DIR / "RS_F_3paras.docx"
    doc.save(path)
    return path


def make_rsG(filler: int = 18):
    """Baseline + widowControl=0: does disabling widow/orphan change the split?"""
    doc = build_base_doc(filler_paragraphs=filler)
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    cell = table.cell(0, 0)
    text = "".join(f"{i+1:02d}文字目アイウエオカキクケコサシスセソタチツテト" for i in range(20))
    add_text_cell_para(cell, text, widow_off=True, first_para=True)
    path = OUT_DIR / "RS_G_widowOff.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    results = []
    for fn in (make_rsA, make_rsB, make_rsC, make_rsD, make_rsE, make_rsF, make_rsG):
        path = fn()
        results.append(path)
        print(f"  built: {path}")
    print(f"\nTotal: {len(results)} repros in {OUT_DIR}")
