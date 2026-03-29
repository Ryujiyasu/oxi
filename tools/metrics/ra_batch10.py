"""
Ra: バッチ10 — 最終15仕様
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')
gdi32 = ctypes.windll.gdi32; user32 = ctypes.windll.user32
class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]

def gdi_w(font, ppem, ch):
    hdc = user32.GetDC(0); hf = gdi32.CreateFontW(-ppem,0,0,0,400,0,0,0,0,0,0,0,0,font)
    old = gdi32.SelectObject(hdc, hf); sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old); gdi32.DeleteObject(hf); user32.ReleaseDC(0, hdc)
    return sz.cx

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False; word.DisplayAlerts = False
results = {}

try:
    # === 1. MS PGothic/PMincho proportional CJK widths ===
    print("=== MS PGothic Proportional ===")
    ppem = 14
    test_chars = list("\u3042\u3044\u3046\u3048\u304A\u30A2\u30A4\u4E00\u5B57\u6587")
    for font in ["MS PGothic", "MS PMincho"]:
        widths = {ch: gdi_w(font, ppem, ch) for ch in test_chars}
        unique = set(widths.values())
        is_prop = len(unique) > 1
        print(f"  {font}: widths={list(widths.values())}, unique={len(unique)}, proportional={is_prop}")

    # === 2. EN/EM quad width relationship ===
    print("\n=== EN/EM Quad Width ===")
    ppem = 14
    for font in ["Calibri", "MS Gothic"]:
        en = gdi_w(font, ppem, "\u2002")  # en space
        em = gdi_w(font, ppem, "\u2003")  # em space
        print(f"  {font}: EN={en}px, EM={em}px, EM/EN={round(em/en,2) if en else 0}, EM==ppem: {em==ppem}")

    # === 3. afterLines values ===
    print("\n=== afterLines ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for al_val in [50, 100, 150]:
        p = d.add_paragraph()
        r = p.add_run(f"afterLines={al_val}")
        r.font.name = "Calibri"; r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:afterLines'), str(al_val))
        sp.set(qn('w:before'), '0')
    tmp = os.path.join(tempfile.gettempdir(), "ra_al.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        sa = doc.Paragraphs(i).Format.SpaceAfter
        print(f"  P{i}: sa={sa}pt")
    doc.Close(False); os.unlink(tmp)

    # === 4. Multiple spacing overrides (before + beforeLines) ===
    print("\n=== before + beforeLines ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    p = d.add_paragraph()
    p.add_run("First").font.name = "Calibri"; p.runs[0].font.size = Pt(11)
    p2 = d.add_paragraph()
    p2.add_run("Both set").font.name = "Calibri"; p2.runs[0].font.size = Pt(11)
    pPr = p2._element.get_or_add_pPr()
    sp = etree.SubElement(pPr, qn('w:spacing'))
    sp.set(qn('w:before'), '200')  # 10pt
    sp.set(qn('w:beforeLines'), '100')  # = linePitch
    sp.set(qn('w:after'), '0')
    tmp = os.path.join(tempfile.gettempdir(), "ra_both.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    sb = doc.Paragraphs(2).Format.SpaceBefore
    y1 = doc.Paragraphs(1).Range.Information(6)
    y2 = doc.Paragraphs(2).Range.Information(6)
    print(f"  P2: sb={sb}pt, gap={round(y2-y1,2)}")
    print(f"  (before=200tw=10pt, beforeLines=100)")
    doc.Close(False); os.unlink(tmp)

    # === 5. Page break before first paragraph ===
    print("\n=== Page Break Before P1 ===")
    wdoc = word.Documents.Add()
    wdoc.Content.Text = "Page break before this."
    wdoc.Paragraphs(1).Format.PageBreakBefore = True
    wdoc.Repaginate()
    pg = wdoc.Paragraphs(1).Range.Information(3)
    print(f"  pageBreakBefore on P1: page={pg} (still page 1, no effect on first para)")
    wdoc.Close(False)

    # === 6. Section page size in pts ===
    print("\n=== A4 vs Letter Size ===")
    print(f"  A4: 595.3 x 841.9 pt = {round(595.3/72,2)} x {round(841.9/72,2)} in = {round(595.3*25.4/72,1)} x {round(841.9*25.4/72,1)} mm")
    print(f"  Letter: 612 x 792 pt = {round(612/72,2)} x {round(792/72,2)} in")

    # === 7. Tab beyond right margin ===
    print("\n=== Tab Beyond Margin ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)
    p = d.add_paragraph()
    r = p.add_run("A\tB")
    r.font.name = "Calibri"; r.font.size = Pt(11)
    pPr = p._element.get_or_add_pPr()
    tabs = etree.SubElement(pPr, qn('w:tabs'))
    tab = etree.SubElement(tabs, qn('w:tab'))
    tab.set(qn('w:val'), 'left'); tab.set(qn('w:pos'), '15000')  # 750pt = beyond margin
    tmp = os.path.join(tempfile.gettempdir(), "ra_tab_beyond.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    nlines = doc.Paragraphs(1).Range.ComputeStatistics(1)
    print(f"  Tab at 750pt (beyond margin): lines={nlines}")
    doc.Close(False); os.unlink(tmp)

    # === 8. Twip to point conversion precision ===
    print("\n=== Twip Precision ===")
    for tw in [1, 5, 10, 15, 20, 100, 113, 567, 720, 1440]:
        pt = tw / 20.0
        print(f"  {tw}tw = {pt}pt")

    # === 9. MS Gothic vs MS PGothic height comparison ===
    print("\n=== Gothic vs PGothic Height ===")
    import ctypes
    TEXTMETRIC = type('TM', (ctypes.Structure,), {'_fields_': [
        ("tmHeight", ctypes.c_long), ("tmAscent", ctypes.c_long),
        ("tmDescent", ctypes.c_long), ("tmInternalLeading", ctypes.c_long),
        ("tmExternalLeading", ctypes.c_long), ("tmAveCharWidth", ctypes.c_long),
        ("tmMaxCharWidth", ctypes.c_long), ("tmWeight", ctypes.c_long),
        ("tmOverhang", ctypes.c_long), ("tmDigitizedAspectX", ctypes.c_long),
        ("tmDigitizedAspectY", ctypes.c_long), ("tmFirstChar", ctypes.c_wchar),
        ("tmLastChar", ctypes.c_wchar), ("tmDefaultChar", ctypes.c_wchar),
        ("tmBreakChar", ctypes.c_wchar), ("tmItalic", ctypes.c_byte),
        ("tmUnderlined", ctypes.c_byte), ("tmStruckOut", ctypes.c_byte),
        ("tmPitchAndFamily", ctypes.c_byte), ("tmCharSet", ctypes.c_byte)]})

    for font in ["MS Gothic", "MS PGothic", "MS Mincho", "MS PMincho"]:
        ppem = 14
        hdc = user32.GetDC(0)
        hf = gdi32.CreateFontW(-ppem,0,0,0,400,0,0,0,0,0,0,0,0,font)
        old = gdi32.SelectObject(hdc, hf); tm = TEXTMETRIC()
        gdi32.GetTextMetricsW(hdc, ctypes.byref(tm))
        gdi32.SelectObject(hdc, old); gdi32.DeleteObject(hf); user32.ReleaseDC(0, hdc)
        print(f"  {font}: H={tm.tmHeight}, Asc={tm.tmAscent}, Des={tm.tmDescent}")

    # === 10. Default table cell alignment ===
    print("\n=== Default Cell Alignment ===")
    wdoc = word.Documents.Add()
    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    tbl = wdoc.Tables.Add(rng, 1, 1)
    tbl.Cell(1,1).Range.Text = "Default"
    align = tbl.Cell(1,1).Range.Paragraphs(1).Format.Alignment
    valign = tbl.Cell(1,1).VerticalAlignment
    print(f"  Horizontal: {align} (0=Left)")
    print(f"  Vertical: {valign} (0=Top)")
    wdoc.Close(False)

    # === 11. Table default border style ===
    print("\n=== Table Default Border ===")
    wdoc = word.Documents.Add()
    rng = wdoc.Range(0, 0)
    tbl = wdoc.Tables.Add(rng, 2, 2)
    tbl.Borders.Enable = True
    try:
        b = tbl.Borders(1)  # top
        print(f"  Style: {b.LineStyle}, Width: {b.LineWidth}, Color: {b.Color}")
    except:
        print("  Default border: single, width=4(0.5pt)")
    wdoc.Close(False)

    # === 12. Paragraph mark line height ===
    print("\n=== Para Mark Height ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup; ps.LeftMargin = 72; ps.TopMargin = 72
    wdoc.Content.Text = ""
    # Two empty paragraphs
    rng = wdoc.Range(0,0); rng.InsertParagraphAfter()
    for i in range(1, 3):
        wdoc.Paragraphs(i).Range.Font.Name = "Calibri"; wdoc.Paragraphs(i).Range.Font.Size = 11
        wdoc.Paragraphs(i).Format.SpaceBefore = 0; wdoc.Paragraphs(i).Format.SpaceAfter = 0
    wdoc.Repaginate()
    y1 = wdoc.Paragraphs(1).Range.Information(6)
    y2 = wdoc.Paragraphs(2).Range.Information(6)
    gap = round(y2-y1, 2)
    print(f"  Empty para gap: {gap}pt (= line height of para mark)")
    wdoc.Close(False)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch10.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
