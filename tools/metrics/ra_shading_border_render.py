"""
Ra: 段落シェーディング・ボーダーのレンダリング精度
- シェーディングの範囲 (テキスト幅? カラム幅?)
- ボーダーのspace属性 (テキストからの距離)
- between ボーダー (段落間)
- テーブル内の段落ボーダー
"""
import win32com.client, json, os, tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def test_shading_extent():
    """Paragraph shading: does it span text width or full column?"""
    wdoc = word.Documents.Add()
    try:
        ps = wdoc.Sections(1).PageSetup
        ps.LeftMargin = 72
        ps.RightMargin = 72

        wdoc.Content.Text = ""

        # P1: short text with shading
        rng = wdoc.Range(0, 0)
        rng.InsertAfter("Short shaded text")
        p1 = wdoc.Paragraphs(1)
        p1.Range.Font.Name = "Calibri"
        p1.Range.Font.Size = 11
        p1.Shading.BackgroundPatternColor = 0xC0C0C0  # light gray

        # P2: long text with shading + indent
        rng2 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
        rng2.InsertParagraphAfter()
        p2 = wdoc.Paragraphs(2)
        p2.Range.Text = "Indented shaded text with longer content here."
        p2.Range.Font.Name = "Calibri"
        p2.Range.Font.Size = 11
        p2.Format.LeftIndent = 36
        p2.Format.RightIndent = 36
        p2.Shading.BackgroundPatternColor = 0xFFFF00  # yellow

        # P3: no shading (baseline)
        rng3 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
        rng3.InsertParagraphAfter()
        p3 = wdoc.Paragraphs(3)
        p3.Range.Text = "No shading baseline."
        p3.Range.Font.Name = "Calibri"
        p3.Range.Font.Size = 11

        wdoc.Repaginate()

        data = {"scenario": "shading_extent", "paragraphs": []}
        for i in range(1, 4):
            para = wdoc.Paragraphs(i)
            prng = para.Range
            shading_color = para.Shading.BackgroundPatternColor
            data["paragraphs"].append({
                "index": i,
                "x": round(prng.Information(5), 4),
                "y": round(prng.Information(6), 4),
                "left_indent": round(para.Format.LeftIndent, 4),
                "right_indent": round(para.Format.RightIndent, 4),
                "shading_color": shading_color,
                "text": prng.Text.strip()[:40],
            })

        return data
    finally:
        wdoc.Close(False)


def test_border_space():
    """Paragraph border with space attribute."""
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # P1: border with space=0
    p1 = d.add_paragraph()
    p1.add_run("Border space=0").font.name = "Calibri"
    p1.runs[0].font.size = Pt(11)
    pPr1 = p1._element.get_or_add_pPr()
    pBdr1 = etree.SubElement(pPr1, qn('w:pBdr'))
    for side in ['top', 'bottom', 'left', 'right']:
        el = etree.SubElement(pBdr1, qn(f'w:{side}'))
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')

    # P2: border with space=4
    p2 = d.add_paragraph()
    p2.add_run("Border space=4").font.name = "Calibri"
    p2.runs[0].font.size = Pt(11)
    pPr2 = p2._element.get_or_add_pPr()
    pBdr2 = etree.SubElement(pPr2, qn('w:pBdr'))
    for side in ['top', 'bottom', 'left', 'right']:
        el = etree.SubElement(pBdr2, qn(f'w:{side}'))
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '000000')

    # P3: border with space=12
    p3 = d.add_paragraph()
    p3.add_run("Border space=12").font.name = "Calibri"
    p3.runs[0].font.size = Pt(11)
    pPr3 = p3._element.get_or_add_pPr()
    pBdr3 = etree.SubElement(pPr3, qn('w:pBdr'))
    for side in ['top', 'bottom', 'left', 'right']:
        el = etree.SubElement(pBdr3, qn(f'w:{side}'))
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '12')
        el.set(qn('w:color'), '000000')

    return d


def test_between_border():
    """Between border (between consecutive paragraphs)."""
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    for i in range(3):
        p = d.add_paragraph()
        p.add_run(f"Paragraph {i+1} with between border").font.name = "Calibri"
        p.runs[0].font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
        between = etree.SubElement(pBdr, qn('w:between'))
        between.set(qn('w:val'), 'single')
        between.set(qn('w:sz'), '4')
        between.set(qn('w:space'), '1')
        between.set(qn('w:color'), '000000')

    return d


try:
    d1 = test_shading_extent()
    results.append(d1)
    print("=== shading_extent ===")
    for p in d1["paragraphs"]:
        print(f"  P{p['index']}: x={p['x']}, li={p['left_indent']}, ri={p['right_indent']}, "
              f"shading={p['shading_color']}  [{p['text']}]")

    d2 = test_border_space()
    tmp = os.path.join(tempfile.gettempdir(), "ra_border_space.docx")
    d2.save(tmp)
    doc = word.Documents.Open(tmp)
    print(f"\n=== border_space ===")
    border_data = {"scenario": "border_space", "paragraphs": []}
    for i in range(1, doc.Paragraphs.Count + 1):
        para = doc.Paragraphs(i)
        rng = para.Range
        y = rng.Information(6)
        x = rng.Information(5)
        border_data["paragraphs"].append({
            "index": i, "x": round(x, 4), "y": round(y, 4),
            "text": rng.Text.strip()[:30],
        })
        print(f"  P{i}: x={round(x,2)}, y={round(y,2)}  [{rng.Text.strip()[:30]}]")
    results.append(border_data)
    doc.Close(False)
    os.unlink(tmp)

    d3 = test_between_border()
    tmp = os.path.join(tempfile.gettempdir(), "ra_between_border.docx")
    d3.save(tmp)
    doc = word.Documents.Open(tmp)
    print(f"\n=== between_border ===")
    between_data = {"scenario": "between_border", "paragraphs": []}
    for i in range(1, doc.Paragraphs.Count + 1):
        para = doc.Paragraphs(i)
        rng = para.Range
        y = rng.Information(6)
        between_data["paragraphs"].append({
            "index": i, "y": round(y, 4), "text": rng.Text.strip()[:30],
        })
    # Compute gaps
    for i in range(1, len(between_data["paragraphs"])):
        gap = between_data["paragraphs"][i]["y"] - between_data["paragraphs"][i-1]["y"]
        between_data["paragraphs"][i]["gap"] = round(gap, 4)
        print(f"  P{i+1}: y={between_data['paragraphs'][i]['y']}, gap={round(gap,2)}")
    results.append(between_data)
    doc.Close(False)
    os.unlink(tmp)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_shading_border.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
