"""
Build default-tab fixtures for §6.4 verification.

Test: paragraph with N consecutive TAB characters (no custom tabs defined),
varying:
  - DefaultTabStop in section settings: 360tw (18pt), 720tw (36pt = default),
    1440tw (72pt)
  - Pre-tab text width (so first tab lands at different positions)

Each fixture has paragraphs:
  P1: "" + TAB → measure where text after tab lands (= first default tab pos)
  P2: "X" + TAB → measure (X width matters)
  P3: "X" + TAB + "Y" + TAB → measure both tab landings

Output: output/default_tab_fixtures/
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output", "default_tab_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


def set_default_tab(doc, tw):
    """Set <w:defaultTabStop w:val="N"/> in settings.xml."""
    # python-docx exposes settings via doc.settings
    settings = doc.settings.element
    dts = settings.find(qn("w:defaultTabStop"))
    if dts is None:
        dts = OxmlElement("w:defaultTabStop")
        settings.append(dts)
    dts.set(qn("w:val"), str(tw))


def add_run_with_tab(p, text="", tabs_after=1):
    """Add a run with text followed by N TAB characters."""
    if text:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    for _ in range(tabs_after):
        tab_run = p.add_run()
        tab_elem = OxmlElement("w:tab")
        tab_run._r.append(tab_elem)


def build_fixture(path, *, default_tab_tw):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    set_default_tab(doc, default_tab_tw)

    # P1: just a tab + marker
    p1 = doc.add_paragraph()
    add_run_with_tab(p1, text="", tabs_after=1)
    r = p1.add_run("M1")
    r.font.name = "Calibri"
    r.font.size = Pt(11)

    # P2: short text + tab + marker
    p2 = doc.add_paragraph()
    add_run_with_tab(p2, text="A", tabs_after=1)
    r = p2.add_run("M2")
    r.font.name = "Calibri"
    r.font.size = Pt(11)

    # P3: longer text + tab + marker
    p3 = doc.add_paragraph()
    add_run_with_tab(p3, text="ABCDE", tabs_after=1)
    r = p3.add_run("M3")
    r.font.name = "Calibri"
    r.font.size = Pt(11)

    # P4: text + 2 consecutive tabs + marker
    p4 = doc.add_paragraph()
    add_run_with_tab(p4, text="A", tabs_after=2)
    r = p4.add_run("M4")
    r.font.name = "Calibri"
    r.font.size = Pt(11)

    # P5: text + tab + text + tab + marker
    p5 = doc.add_paragraph()
    add_run_with_tab(p5, text="X", tabs_after=1)
    r = p5.add_run("Y")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    tab_run = p5.add_run()
    tab_elem = OxmlElement("w:tab")
    tab_run._r.append(tab_elem)
    r = p5.add_run("M5")
    r.font.name = "Calibri"
    r.font.size = Pt(11)

    for p in (p1, p2, p3, p4, p5):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.save(path)


def main():
    cases = [
        ("DT_default720tw.docx", 720),  # Word default
        ("DT_360tw.docx",        360),
        ("DT_1440tw.docx",       1440),
    ]
    for name, tw in cases:
        path = os.path.join(OUT_DIR, name)
        try:
            build_fixture(path, default_tab_tw=tw)
            print(f"  built {name}")
        except Exception as e:
            print(f"  ERR {name}: {e}")
    print(f"\nBuilt {len(cases)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
