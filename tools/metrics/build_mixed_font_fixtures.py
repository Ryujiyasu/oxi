"""
Build §1.7 mixed-font line height fixtures.

Each fixture has 3 paragraphs:
  P1: pure font A at size A
  P2: pure font B at size B
  P3: mixed line "<font A text> + <font B text>" inline

If §1.7 max rule holds: P3.line_h = max(P1.line_h, P2.line_h).

Sweep:
  - 4 font pairs (Calibri/MS Mincho, TNR/MS Gothic, Yu Mincho/Yu Gothic, Calibri/Yu Mincho)
  - 4 size combinations: {8/8, 11/14, 14/11, 18/24}

Output: output/mixed_font_fixtures/
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output", "mixed_font_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


def add_doc_grid(section, line_pitch_tw):
    sectPr = section._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        docGrid = OxmlElement("w:docGrid")
        sectPr.append(docGrid)
    docGrid.set(qn("w:type"), "lines")
    docGrid.set(qn("w:linePitch"), str(line_pitch_tw))


def add_run(p, text, font, size_pt):
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    return run


def build_fixture(path, *, font_a, size_a, font_b, size_b,
                   line_pitch_tw=0):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    if line_pitch_tw:
        add_doc_grid(sec, line_pitch_tw)

    # P1: pure font A
    p1 = doc.add_paragraph()
    add_run(p1, "AAA pure", font_a, size_a)
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)

    # P2: pure font B
    p2 = doc.add_paragraph()
    add_run(p2, "BBB pure", font_b, size_b)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    # P3: mixed line — font A then font B then font A
    p3 = doc.add_paragraph()
    add_run(p3, "MIX-A ", font_a, size_a)
    add_run(p3, "MIX-B ", font_b, size_b)
    add_run(p3, "back-A", font_a, size_a)
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)

    # P4: control — pure font A again to measure P3→P4 gap = P3 line height
    p4 = doc.add_paragraph()
    add_run(p4, "AAA control", font_a, size_a)
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(0)

    doc.save(path)


def main():
    font_pairs = [
        ("Calibri",         "MS Mincho"),
        ("Times New Roman", "MS Gothic"),
        ("Yu Mincho",       "Yu Gothic"),
        ("Calibri",         "Yu Mincho"),
    ]
    size_combos = [(8, 8), (11, 14), (14, 11), (18, 24)]
    grid_options = [0, 360]  # noGrid + grid 18pt

    cases = []
    for font_a, font_b in font_pairs:
        for size_a, size_b in size_combos:
            for pitch in grid_options:
                fname_safe_a = font_a.replace(" ", "")
                fname_safe_b = font_b.replace(" ", "")
                grid_label = "noGrid" if pitch == 0 else f"grid{pitch}tw"
                name = f"MF_{fname_safe_a}{size_a}_{fname_safe_b}{size_b}_{grid_label}.docx"
                path = os.path.join(OUT_DIR, name)
                try:
                    build_fixture(path, font_a=font_a, size_a=size_a,
                                  font_b=font_b, size_b=size_b,
                                  line_pitch_tw=pitch)
                    cases.append({"path": name, "font_a": font_a, "size_a": size_a,
                                  "font_b": font_b, "size_b": size_b, "pitch_tw": pitch})
                except Exception as e:
                    print(f"  ERR {name}: {e}")
    print(f"\nBuilt {len(cases)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
