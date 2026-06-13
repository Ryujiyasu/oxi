# -*- coding: utf-8 -*-
"""S560 minimal repro: continuous section break that CHANGES column count
(1-col -> 2-col). Establish Word's layout model for kyotei36spec root cause
(Oxi merges all continuous sections into 1 IR Page with the LAST section's
column count, mis-laying-out the 1-col content as 2-col).

repro_short : Sec A (1-col, short) | continuous | Sec B (2-col, B1..B10)
repro_tall  : Sec A (1-col, fills page) | continuous | Sec B (2-col, B1..B10)
Measure each paragraph x/y/page in Word to confirm:
 (1) Sec A content is full-width (1-col), Sec B flows in 2 balanced columns.
 (2) Sec B begins at the CURRENT Y on the same page (short) OR overflows to a
     new page when Sec A fills the page (tall) -> kyotei case.
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

def set_cols(sectPr, num):
    # remove existing cols
    for c in sectPr.findall(qn('w:cols')):
        sectPr.remove(c)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), '425')
    sectPr.append(cols)

def set_type(sectPr, t):
    for c in sectPr.findall(qn('w:type')):
        sectPr.remove(c)
    typ = OxmlElement('w:type')
    typ.set(qn('w:val'), t)
    sectPr.insert(0, typ)

def build(path, tall):
    doc = Document()
    # Section A: 1-col. The default last-section sectPr is on body; we add a
    # continuous section break to start Section B.
    nA = 25 if tall else 2
    for i in range(nA):
        doc.add_paragraph('A%02d この段落はセクションA(一段組)の本文です。' % i)
    # add a continuous section break -> this creates a new section; the break's
    # sectPr (Section A's properties) goes into the LAST paragraph of A.
    secA = doc.add_section(start_type=2)  # 2 = CONTINUOUS (WD_SECTION_START.CONTINUOUS)
    set_type(secA._sectPr, 'continuous')
    set_cols(secA._sectPr, 1)
    # Section B paragraphs (2-col), set on body-final sectPr
    for i in range(10):
        doc.add_paragraph('B%02d これはセクションB(二段組)の本文テキストであり折り返します。' % i)
    # body-final sectPr -> Section B props: continuous + 2 cols
    body_sectPr = doc.sections[-1]._sectPr
    set_type(body_sectPr, 'continuous')
    set_cols(body_sectPr, 2)
    doc.save(path)
    print('wrote', path, 'tall=%s'%tall)

import os
out='tools/golden-test/repros/s560_cont_cols'
os.makedirs(out, exist_ok=True)
build(out+'/s560short_repro.docx', tall=False)
build(out+'/s560tall_repro.docx', tall=True)
