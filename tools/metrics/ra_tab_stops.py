"""
Ra: タブストップの正確な位置計算をCOM計測で確定
- left / center / right / decimal タブの位置
- デフォルトタブ間隔 (明示タブなし)
- カスタムタブ位置の基準点 (マージン起点? ページ起点?)
- タブ文字の幅計算
- indent との相互作用
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt, Twips, Inches
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def make_doc(scenario):
    d = Document(TEMPLATE)
    sec = d.sections[0]

    # Remove grid
    sectPr = sec._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)

    # Remove default paragraphs
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    if scenario == "default_tabs":
        # Default tab stops (no custom tabs)
        p = d.add_paragraph()
        r = p.add_run("A\tB\tC\tD\tE")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    elif scenario == "custom_left_tabs":
        # Custom left tab stops at specific positions
        p = d.add_paragraph()
        r = p.add_run("Col1\tCol2\tCol3\tCol4")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # Add tab stops via XML
        pPr = p._element.get_or_add_pPr()
        tabs = etree.SubElement(pPr, qn('w:tabs'))
        for pos in [1440, 2880, 4320, 5760]:  # 1in, 2in, 3in, 4in in twips
            tab = etree.SubElement(tabs, qn('w:tab'))
            tab.set(qn('w:val'), 'left')
            tab.set(qn('w:pos'), str(pos))

    elif scenario == "center_right_tabs":
        # Center and right tab stops
        p = d.add_paragraph()
        r = p.add_run("Left\tCenter\tRight")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pf = p.paragraph_format
        pPr = p._element.get_or_add_pPr()
        tabs = etree.SubElement(pPr, qn('w:tabs'))
        # Center tab at 3in
        tab_c = etree.SubElement(tabs, qn('w:tab'))
        tab_c.set(qn('w:val'), 'center')
        tab_c.set(qn('w:pos'), str(4320))
        # Right tab at 6in
        tab_r = etree.SubElement(tabs, qn('w:tab'))
        tab_r.set(qn('w:val'), 'right')
        tab_r.set(qn('w:pos'), str(8640))

    elif scenario == "tabs_with_indent":
        # Tab stops + left indent
        p = d.add_paragraph()
        r = p.add_run("Indented\tTabbed")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pf = p.paragraph_format
        pf.left_indent = Pt(36)  # 0.5 inch
        pPr = p._element.get_or_add_pPr()
        tabs = etree.SubElement(pPr, qn('w:tabs'))
        tab = etree.SubElement(tabs, qn('w:tab'))
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), str(2880))  # 2in

    elif scenario == "decimal_tab":
        # Decimal tab stop
        p = d.add_paragraph()
        r = p.add_run("Price:\t123.45")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        tabs = etree.SubElement(pPr, qn('w:tabs'))
        tab = etree.SubElement(tabs, qn('w:tab'))
        tab.set(qn('w:val'), 'decimal')
        tab.set(qn('w:pos'), str(4320))

    elif scenario == "tab_leader":
        # Tab with leader (dot, dash, etc.)
        p = d.add_paragraph()
        r = p.add_run("Chapter 1\t10")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        tabs = etree.SubElement(pPr, qn('w:tabs'))
        tab = etree.SubElement(tabs, qn('w:tab'))
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), str(8640))
        tab.set(qn('w:leader'), 'dot')

    elif scenario == "multiple_lines_tabs":
        # Multiple lines with tabs (for gap measurement)
        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"Row{i+1}\tData{i+1}\tEnd{i+1}")
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pPr = p._element.get_or_add_pPr()
            tabs = etree.SubElement(pPr, qn('w:tabs'))
            for pos in [2880, 5760]:
                tab = etree.SubElement(tabs, qn('w:tab'))
                tab.set(qn('w:val'), 'left')
                tab.set(qn('w:pos'), str(pos))

    return d


def measure_tab_positions(doc_path, scenario):
    """Measure character positions to determine tab stop behavior."""
    doc = word.Documents.Open(doc_path)
    try:
        data = {"scenario": scenario, "paragraphs": []}

        for pi in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(pi)
            rng = para.Range
            para_data = {
                "index": pi,
                "y_pt": round(rng.Information(6), 4),
                "text": rng.Text.strip()[:60],
                "char_positions": []
            }

            # Measure x position of each character
            for ci in range(rng.Start, rng.End):
                char_rng = doc.Range(ci, ci + 1)
                char = char_rng.Text
                x = char_rng.Information(5)  # wdHorizontalPositionRelativeToPage
                y = char_rng.Information(6)  # wdVerticalPositionRelativeToPage
                para_data["char_positions"].append({
                    "char": char,
                    "x_pt": round(x, 4),
                    "y_pt": round(y, 4),
                    "char_code": ord(char) if len(char) == 1 else -1
                })

            data["paragraphs"].append(para_data)

        # Also get default tab stop interval
        data["default_tab_stop"] = round(doc.DefaultTabStop, 4)

        # Get page margins
        sec = doc.Sections(1)
        data["margin_left"] = round(sec.PageSetup.LeftMargin, 4)
        data["margin_right"] = round(sec.PageSetup.RightMargin, 4)
        data["page_width"] = round(sec.PageSetup.PageWidth, 4)

        return data
    finally:
        doc.Close(False)


scenarios = [
    "default_tabs",
    "custom_left_tabs",
    "center_right_tabs",
    "tabs_with_indent",
    "decimal_tab",
    "tab_leader",
    "multiple_lines_tabs",
]

try:
    for sc in scenarios:
        d = make_doc(sc)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_tab_{sc}.docx")
        d.save(tmp)
        data = measure_tab_positions(tmp, sc)
        results.append(data)

        print(f"\n=== {sc} ===")
        print(f"  DefaultTabStop: {data['default_tab_stop']}pt")
        print(f"  MarginLeft: {data['margin_left']}pt, PageWidth: {data['page_width']}pt")

        for p in data["paragraphs"]:
            print(f"  P{p['index']} (y={p['y_pt']}): {p['text']}")
            # Show tab-separated segment positions
            tab_segments = []
            current_segment = {"start_x": None, "chars": ""}
            for cp in p["char_positions"]:
                if cp["char_code"] == 9:  # tab character
                    if current_segment["start_x"] is not None:
                        tab_segments.append(current_segment)
                    current_segment = {"start_x": None, "chars": ""}
                else:
                    if current_segment["start_x"] is None:
                        current_segment["start_x"] = cp["x_pt"]
                    current_segment["chars"] += cp["char"]

            if current_segment["start_x"] is not None:
                tab_segments.append(current_segment)

            for si, seg in enumerate(tab_segments):
                print(f"    Segment {si}: x={seg['start_x']}pt  \"{seg['chars']}\"")

        os.unlink(tmp)

finally:
    word.Quit()

# Save
out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_tab_stops.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")

# Analysis
print("\n=== TAB STOP ANALYSIS ===")
for data in results:
    sc = data["scenario"]
    ml = data["margin_left"]
    if data["paragraphs"]:
        p = data["paragraphs"][0]
        segs = []
        current_x = None
        for cp in p["char_positions"]:
            if cp["char_code"] == 9:
                current_x = None
            elif current_x is None:
                current_x = cp["x_pt"]
                segs.append(current_x)

        print(f"\n{sc}:")
        print(f"  Segment X positions (page-relative): {segs}")
        print(f"  Segment X positions (margin-relative): {[round(x - ml, 2) for x in segs]}")
