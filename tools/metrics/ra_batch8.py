"""
Ra: バッチ8 — compat14/15最終確認、文字分類詳細、テーブル幅計算、
    spaceBeforeLines grid snap、段落最終行justify、CJK compress幅
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')
gdi32 = ctypes.windll.gdi32; user32 = ctypes.windll.user32
class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]

def gdi_w(font, ppem, ch):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return sz.cx

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False; word.DisplayAlerts = False
results = {}

try:
    # === 1. compat=14 LS sweep FIX: ensure style override works ===
    print("=== compat=14 LS Sweep FIXED ===")
    for line_val in [200, 240, 276, 360, 480]:
        wdoc = word.Documents.Add()
        # Save as compat=14
        wdoc.SetCompatibilityMode(14)
        ps = wdoc.Sections(1).PageSetup
        ps.LeftMargin = 72; ps.TopMargin = 72
        wdoc.Content.Text = ""
        for i in range(3):
            if i > 0:
                r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
                r.InsertParagraphAfter()
            p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
            p.Range.Text = f"Line {i+1}"
            p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
            p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0
            p.Format.LineSpacingRule = 4  # wdLineSpaceMultiple
            p.Format.LineSpacing = line_val / 240.0 * 12  # set as multiple
        wdoc.Repaginate()
        y1 = wdoc.Paragraphs(1).Range.Information(6)
        y2 = wdoc.Paragraphs(2).Range.Information(6)
        gap = round(y2 - y1, 2)
        factor = line_val / 240.0
        expected = round(13.5 * factor, 2)  # gdi_h based
        print(f"  line={line_val}(x{factor:.3f}): gap={gap}, expected_gdi={expected}")
        results[f"c14_ls_{line_val}"] = gap
        wdoc.Close(False)

    # === 2. compat=15 LS sweep for comparison ===
    print("\n=== compat=15 LS Sweep ===")
    for line_val in [200, 240, 276, 360, 480]:
        wdoc = word.Documents.Add()
        ps = wdoc.Sections(1).PageSetup
        ps.LeftMargin = 72; ps.TopMargin = 72
        wdoc.Content.Text = ""
        for i in range(3):
            if i > 0:
                r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
                r.InsertParagraphAfter()
            p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
            p.Range.Text = f"Line {i+1}"
            p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
            p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0
            p.Format.LineSpacingRule = 4
            p.Format.LineSpacing = line_val / 240.0 * 12
        wdoc.Repaginate()
        y1 = wdoc.Paragraphs(1).Range.Information(6)
        y2 = wdoc.Paragraphs(2).Range.Information(6)
        gap = round(y2 - y1, 2)
        factor = line_val / 240.0
        expected = round(13.5 * factor, 2)
        print(f"  line={line_val}(x{factor:.3f}): gap={gap}, expected_gdi={expected}")
        results[f"c15_ls_{line_val}"] = gap
        wdoc.Close(False)

    # === 3. CJK punctuation compression in justify ===
    print("\n=== CJK Punct Compression ===")
    ppem = 14
    for font in ["MS Gothic", "Yu Gothic"]:
        fw = gdi_w(font, ppem, "\u3042")  # fullwidth reference
        for ch, name in [("\u3001", "comma"), ("\u3002", "period"),
                          ("\u300C", "left-bracket"), ("\u300D", "right-bracket"),
                          ("\uFF08", "fw-lparen"), ("\uFF09", "fw-rparen")]:
            w = gdi_w(font, ppem, ch)
            half = fw // 2
            print(f"  {font} {name}(U+{ord(ch):04X}): {w}px, fw={fw}px, "
                  f"50%={half}px, is_compressible={w==fw}")

    # === 4. tcW type=pct (percentage cell width) ===
    print("\n=== tcW percentage ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    tbl_el = d.element.body.makeelement(qn('w:tbl'), {})
    d.element.body.append(tbl_el)
    tblPr = etree.SubElement(tbl_el, qn('w:tblPr'))
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')  # 100%

    tblGrid = etree.SubElement(tbl_el, qn('w:tblGrid'))
    for gw in [2500, 2500]:  # equal
        gc = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc.set(qn('w:w'), str(gw))

    tr = etree.SubElement(tbl_el, qn('w:tr'))
    for ci, pct in enumerate([2500, 2500]):  # 50% each
        tc = etree.SubElement(tr, qn('w:tc'))
        tcPr = etree.SubElement(tc, qn('w:tcPr'))
        tcW = etree.SubElement(tcPr, qn('w:tcW'))
        tcW.set(qn('w:w'), str(pct)); tcW.set(qn('w:type'), 'pct')
        p = etree.SubElement(tc, qn('w:p'))
        run = etree.SubElement(p, qn('w:r'))
        t = etree.SubElement(run, qn('w:t'))
        t.text = f"C{ci+1}"

    tmp = os.path.join(tempfile.gettempdir(), "ra_tcw_pct.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    tbl = doc.Tables(1)
    content_w = doc.Sections(1).PageSetup.PageWidth - doc.Sections(1).PageSetup.LeftMargin - doc.Sections(1).PageSetup.RightMargin
    for c in range(1, 3):
        w = round(tbl.Columns(c).Width, 2)
        pct = round(w / content_w * 100, 1)
        print(f"  Col {c}: {w}pt ({pct}% of {round(content_w,1)})")
    doc.Close(False); os.unlink(tmp)

    # === 5. Multiple font sizes on same line: line height ===
    print("\n=== Mixed Size Line Height ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup; ps.LeftMargin = 72
    wdoc.Content.Text = ""
    # Line 1: 11pt only
    rng = wdoc.Range(0, 0); rng.InsertAfter("11pt only line here")
    wdoc.Paragraphs(1).Range.Font.Name = "Calibri"; wdoc.Paragraphs(1).Range.Font.Size = 11
    wdoc.Paragraphs(1).Format.SpaceBefore = 0; wdoc.Paragraphs(1).Format.SpaceAfter = 0

    # Line 2: mixed 11pt + 20pt
    r2 = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); r2.InsertParagraphAfter()
    p2 = wdoc.Paragraphs(2)
    p2.Range.Text = ""; r2a = wdoc.Range(p2.Range.Start, p2.Range.Start)
    r2a.InsertAfter("Small "); wdoc.Range(p2.Range.Start, p2.Range.Start+6).Font.Size = 11
    r2b = wdoc.Range(p2.Range.End-1, p2.Range.End-1)
    r2b.InsertAfter("BIG"); wdoc.Range(p2.Range.End-4, p2.Range.End-1).Font.Size = 20
    wdoc.Range(p2.Range.Start, p2.Range.End).Font.Name = "Calibri"
    p2.Format.SpaceBefore = 0; p2.Format.SpaceAfter = 0

    # Line 3: 11pt only (for gap measurement)
    r3 = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); r3.InsertParagraphAfter()
    p3 = wdoc.Paragraphs(3)
    p3.Range.Text = "11pt again"; p3.Range.Font.Name = "Calibri"; p3.Range.Font.Size = 11
    p3.Format.SpaceBefore = 0; p3.Format.SpaceAfter = 0

    wdoc.Repaginate()
    for i in range(1, 4):
        y = wdoc.Paragraphs(i).Range.Information(6)
        if i > 1:
            prev = wdoc.Paragraphs(i-1).Range.Information(6)
            print(f"  P{i}: y={round(y,2)}, gap={round(y-prev,2)}")
        else:
            print(f"  P{i}: y={round(y,2)}")
    wdoc.Close(False)

    # === 6. Page number count ===
    print("\n=== Page Dimensions ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    print(f"  A4 default: {round(ps.PageWidth,1)} x {round(ps.PageHeight,1)} pt")
    print(f"  Margins: T={round(ps.TopMargin,1)} B={round(ps.BottomMargin,1)} "
          f"L={round(ps.LeftMargin,1)} R={round(ps.RightMargin,1)}")
    print(f"  HeaderDist={round(ps.HeaderDistance,1)}, FooterDist={round(ps.FooterDistance,1)}")
    print(f"  Orientation={ps.Orientation}")  # 0=portrait, 1=landscape
    results["page_dims"] = {
        "w": round(ps.PageWidth, 4), "h": round(ps.PageHeight, 4),
        "tM": round(ps.TopMargin, 4), "bM": round(ps.BottomMargin, 4),
        "lM": round(ps.LeftMargin, 4), "rM": round(ps.RightMargin, 4),
        "hDist": round(ps.HeaderDistance, 4), "fDist": round(ps.FooterDistance, 4),
    }
    wdoc.Close(False)

    # === 7. Hanging punct (overflowPunct) ===
    print("\n=== Overflow Punctuation ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    d.sections[0].page_width = Pt(250)
    d.sections[0].left_margin = Pt(36); d.sections[0].right_margin = Pt(36)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    p = d.add_paragraph()
    r = p.add_run("ABCDEFGHIJ" * 3 + "\u3001")  # ends with 、
    r.font.name = "MS Gothic"; r.font.size = Pt(10.5)
    pPr = p._element.get_or_add_pPr()
    sp = etree.SubElement(pPr, qn('w:spacing'))
    sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')

    tmp = os.path.join(tempfile.gettempdir(), "ra_overflow.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    nlines = doc.Paragraphs(1).Range.ComputeStatistics(1)
    print(f"  Text with trailing comma: {nlines} lines")
    results["overflow_punct"] = nlines
    doc.Close(False); os.unlink(tmp)

    # === 8. Tab stop clear behavior ===
    print("\n=== Tab Stop Clear ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    # Style has tabs, paragraph clears them
    p1 = d.add_paragraph()
    r1 = p1.add_run("A\tB")
    r1.font.name = "Calibri"; r1.font.size = Pt(11)
    pPr = p1._element.get_or_add_pPr()
    tabs = etree.SubElement(pPr, qn('w:tabs'))
    # Clear tab at default positions
    tab_clear = etree.SubElement(tabs, qn('w:tab'))
    tab_clear.set(qn('w:val'), 'clear')
    tab_clear.set(qn('w:pos'), '840')  # 42pt default
    # Set new tab at 200pt
    tab_new = etree.SubElement(tabs, qn('w:tab'))
    tab_new.set(qn('w:val'), 'left')
    tab_new.set(qn('w:pos'), '4000')  # 200pt

    tmp = os.path.join(tempfile.gettempdir(), "ra_tabclear.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    prng = doc.Paragraphs(1).Range
    for ci in range(prng.Start, min(prng.End, prng.Start + 10)):
        cr = doc.Range(ci, ci+1)
        ch = cr.Text
        if ord(ch) not in (13, 7):
            x = cr.Information(5)
            print(f"  '{ch}': x={round(x, 2)}")
    doc.Close(False); os.unlink(tmp)

    # === 9. Compat mode distribution in pipeline docs ===
    print("\n=== Pipeline Doc Compat Modes ===")
    import glob
    docx_dir = os.path.join('tools', 'golden-test', 'documents', 'docx')
    if os.path.isdir(docx_dir):
        compat_counts = {}
        files = glob.glob(os.path.join(docx_dir, '*.docx'))[:20]
        for f in files:
            try:
                doc = word.Documents.Open(os.path.abspath(f))
                cm = doc.CompatibilityMode
                compat_counts[cm] = compat_counts.get(cm, 0) + 1
                doc.Close(False)
            except:
                pass
        print(f"  Compat distribution: {compat_counts}")
        results["compat_distribution"] = compat_counts

    # === 10. Default font for CJK in fresh doc ===
    print("\n=== Default CJK Font ===")
    wdoc = word.Documents.Add()
    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    rng.InsertAfter("Test")
    p = wdoc.Paragraphs(1)
    rng_p = p.Range
    print(f"  Default font: {rng_p.Font.Name}")
    print(f"  NameAscii: {rng_p.Font.NameAscii}")
    print(f"  NameFarEast: {rng_p.Font.NameFarEast}")
    print(f"  NameOther: {rng_p.Font.NameOther}")
    results["default_fonts"] = {
        "name": rng_p.Font.Name,
        "ascii": rng_p.Font.NameAscii,
        "fareast": rng_p.Font.NameFarEast,
    }
    wdoc.Close(False)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch8.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
