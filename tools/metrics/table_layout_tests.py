#!/usr/bin/env python3
"""
Measure Word table layout behaviors via Word COM + PDF export + PyMuPDF.
Resolves unknowns about column widths, row heights, cell padding, borders.
"""
import os, tempfile, time, json
import pythoncom
pythoncom.CoInitialize()
import fitz
from docx import Document
from docx.shared import Pt, Twips, Cm, Inches, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
import win32com.client

TMPDIR = tempfile.gettempdir()

def docx_to_pdf(docx_path, pdf_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path).replace("/", "\\"))
        time.sleep(1)
        doc.SaveAs(os.path.abspath(pdf_path).replace("/", "\\"), FileFormat=17)
        time.sleep(0.5)
        doc.Close(False)
    finally:
        word.Quit()

def get_pdf_details(pdf_path, page_idx=0):
    """Extract text positions, lines/rects from PDF."""
    doc = fitz.open(pdf_path)
    if len(doc) <= page_idx:
        doc.close()
        return {"texts": [], "lines": [], "rects": [], "page_size": (0,0)}
    page = doc[page_idx]

    # Text
    texts = []
    blocks = page.get_text("dict")["blocks"]
    for b in blocks:
        if "lines" in b:
            for line in b["lines"]:
                text = "".join(s["text"] for s in line["spans"]).strip()
                if text:
                    texts.append({
                        "x": line["bbox"][0], "y": line["bbox"][1],
                        "x2": line["bbox"][2], "y2": line["bbox"][3],
                        "text": text[:60],
                        "font_size": line["spans"][0]["size"] if line["spans"] else 0
                    })

    # Lines and rects (table borders)
    drawings = page.get_drawings()
    lines = []
    rects = []
    for d in drawings:
        if d["type"] == "l":  # line
            lines.append({"x0": d["rect"][0], "y0": d["rect"][1],
                          "x1": d["rect"][2], "y1": d["rect"][3],
                          "color": d.get("color"), "width": d.get("width", 1)})
        elif d["type"] == "r":  # rectangle
            rects.append({"x0": d["rect"][0], "y0": d["rect"][1],
                          "x1": d["rect"][2], "y1": d["rect"][3],
                          "color": d.get("color"), "fill": d.get("fill")})

    page_size = (page.rect.width, page.rect.height)
    doc.close()
    return {"texts": texts, "lines": lines, "rects": rects, "page_size": page_size}

def set_no_spacing(doc):
    style_el = doc.styles['Normal'].element
    pPr = style_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        style_el.append(pPr)
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:line="240" w:lineRule="auto"/>')
    pPr.append(sp)

def run_test(name, doc):
    docx_path = os.path.join(TMPDIR, f"tbl_{name}.docx")
    pdf_path = os.path.join(TMPDIR, f"tbl_{name}.pdf")
    doc.save(docx_path)
    docx_to_pdf(docx_path, pdf_path)
    details = get_pdf_details(pdf_path)
    os.unlink(docx_path)
    os.unlink(pdf_path)
    return details


# ============================================================
# TEST 1: Default cell margins (padding)
# ============================================================
def test1_default_margins():
    print("\n=== TEST 1: Default cell margins ===")
    doc = Document()
    set_no_spacing(doc)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f"R{i}C{j}"

    details = run_test("margins", doc)

    # Find border lines
    print(f"  Page: {details['page_size']}")

    # Extract horizontal border Y positions and vertical border X positions
    h_borders = set()
    v_borders = set()
    for d in details.get("lines", []):
        pass  # lines might be empty, check drawings

    # Use text positions to infer cell structure
    print("  Text positions:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f} x2={t['x2']:.2f} y2={t['y2']:.2f}  \"{t['text']}\"")

    # Key question: what is the left padding from cell edge to text?
    # Table Grid style default padding = 0.08" = 5.76pt on left/right
    # And 0pt on top/bottom (or is it?)


# ============================================================
# TEST 2: Cell margin values - explicit test
# ============================================================
def test2_cell_margins_explicit():
    print("\n=== TEST 2: Cell margins with known widths ===")
    doc = Document()
    set_no_spacing(doc)

    # Create table with specific column widths
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # Set explicit column widths via tblGrid
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))

    # Set table width to 400pt (8000 twips)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="8000" w:type="dxa"/>')
    tblPr.append(tblW)

    # Set tblGrid with 200pt + 200pt
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = parse_xml(f'<w:tblGrid {nsdecls("w")}><w:gridCol w:w="4000"/><w:gridCol w:w="4000"/></w:tblGrid>')
    tbl.insert(tbl.index(tblPr) + 1, grid)

    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "LLLL"
    table.cell(1, 1).text = "RRRR"

    details = run_test("margins_explicit", doc)

    print("  Table: 400pt wide, 2 cols x 200pt each")
    print("  Text positions:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")

    # Calculate: text_x - table_left_edge = left_margin + left_padding
    # Default tblCellMar left = 0.08" = 5.76pt

# ============================================================
# TEST 3: Row height - empty vs content
# ============================================================
def test3_row_heights():
    print("\n=== TEST 3: Row heights ===")
    doc = Document()
    set_no_spacing(doc)

    table = doc.add_table(rows=5, cols=1)
    table.style = 'Table Grid'

    table.cell(0, 0).text = "Row with text"
    # Row 1: empty
    table.cell(2, 0).text = "Row with text again"
    # Row 3: set explicit trHeight = 40pt (800tw), atLeast
    row3 = table.rows[3]
    trPr = row3._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
        row3._tr.insert(0, trPr)
    trH = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="800" w:hRule="atLeast"/>')
    trPr.append(trH)
    table.cell(3, 0).text = "Row with trHeight=40pt atLeast"

    # Row 4: trHeight exact = 20pt
    row4 = table.rows[4]
    trPr4 = row4._tr.find(qn('w:trPr'))
    if trPr4 is None:
        trPr4 = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
        row4._tr.insert(0, trPr4)
    trH4 = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="400" w:hRule="exact"/>')
    trPr4.append(trH4)
    table.cell(4, 0).text = "Row exact=20pt"

    doc.add_paragraph("After table")

    details = run_test("row_heights", doc)

    print("  Text positions:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")

# ============================================================
# TEST 4: Table borders - thickness and position
# ============================================================
def test4_borders():
    print("\n=== TEST 4: Table borders ===")
    doc = Document()
    set_no_spacing(doc)

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"

    doc.add_paragraph("After table")

    details = run_test("borders", doc)

    print("  Texts:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")

    # Get drawings to find border lines
    pdf_path = os.path.join(TMPDIR, "tbl_borders2.pdf")
    docx_path = os.path.join(TMPDIR, "tbl_borders2.docx")
    doc.save(docx_path)
    docx_to_pdf(docx_path, pdf_path)

    pdf_doc = fitz.open(pdf_path)
    page = pdf_doc[0]
    drawings = page.get_drawings()

    print(f"\n  Drawings ({len(drawings)}):")
    for i, d in enumerate(drawings):
        items_str = str(d.get("items", []))[:100]
        rect = d.get("rect", (0,0,0,0))
        color = d.get("color")
        fill = d.get("fill")
        width = d.get("width", 0)
        print(f"    [{i}] rect=({rect[0]:.1f},{rect[1]:.1f},{rect[2]:.1f},{rect[3]:.1f}) color={color} fill={fill} w={width}")
        for item in d.get("items", []):
            print(f"         {item[0]} {[f'{v:.2f}' if isinstance(v, float) else v for v in item[1:]]}")

    pdf_doc.close()
    os.unlink(docx_path)
    os.unlink(pdf_path)

# ============================================================
# TEST 5: Table alignment (left, center, right) and indent
# ============================================================
def test5_alignment():
    print("\n=== TEST 5: Table alignment ===")

    for align in ["left", "center", "right"]:
        doc = Document()
        set_no_spacing(doc)

        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.cell(0, 0).text = f"Align={align}"

        # Set alignment
        tblPr = table._tbl.find(qn('w:tblPr'))
        jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="{align}"/>')
        tblPr.append(jc)

        # Set table width to 200pt (4000tw)
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="4000" w:type="dxa"/>')
        tblPr.append(tblW)

        details = run_test(f"align_{align}", doc)

        for t in details["texts"]:
            print(f"  {align}: text x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")

# ============================================================
# TEST 6: Cell vertical alignment
# ============================================================
def test6_valign():
    print("\n=== TEST 6: Cell vertical alignment ===")
    doc = Document()
    set_no_spacing(doc)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Set row height = 60pt
    row = table.rows[0]
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
        row._tr.insert(0, trPr)
    trH = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="1200" w:hRule="exact"/>')
    trPr.append(trH)

    # Cell 0: top (default)
    table.cell(0, 0).text = "top"

    # Cell 1: center
    cell1 = table.cell(0, 1)
    cell1.text = "center"
    tc1 = cell1._tc
    tcPr1 = tc1.find(qn('w:tcPr'))
    if tcPr1 is None:
        tcPr1 = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc1.insert(0, tcPr1)
    vAlign1 = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr1.append(vAlign1)

    # Cell 2: bottom
    cell2 = table.cell(0, 2)
    cell2.text = "bottom"
    tc2 = cell2._tc
    tcPr2 = tc2.find(qn('w:tcPr'))
    if tcPr2 is None:
        tcPr2 = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc2.insert(0, tcPr2)
    vAlign2 = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="bottom"/>')
    tcPr2.append(vAlign2)

    doc.add_paragraph("After table")

    details = run_test("valign", doc)

    print("  Row height = 60pt exact")
    for t in details["texts"]:
        print(f"  x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")

# ============================================================
# TEST 7: vMerge (vertical cell merging)
# ============================================================
def test7_vmerge():
    print("\n=== TEST 7: vMerge ===")
    doc = Document()
    set_no_spacing(doc)

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    # Col 0: merged across 3 rows
    cell_00 = table.cell(0, 0)
    cell_00.text = "Merged 3 rows"
    tc00 = cell_00._tc
    tcPr00 = tc00.find(qn('w:tcPr'))
    if tcPr00 is None:
        tcPr00 = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc00.insert(0, tcPr00)
    vm_restart = parse_xml(f'<w:vMerge {nsdecls("w")} w:val="restart"/>')
    tcPr00.append(vm_restart)

    for row_idx in [1, 2]:
        cell = table.cell(row_idx, 0)
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            tc.insert(0, tcPr)
        vm_cont = parse_xml(f'<w:vMerge {nsdecls("w")}/>')
        tcPr.append(vm_cont)

    # Col 1: separate cells
    table.cell(0, 1).text = "R0C1"
    table.cell(1, 1).text = "R1C1"
    table.cell(2, 1).text = "R2C1"

    details = run_test("vmerge", doc)

    print("  Text positions:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")

# ============================================================
# TEST 8: gridSpan (horizontal cell merging)
# ============================================================
def test8_gridspan():
    print("\n=== TEST 8: gridSpan ===")
    doc = Document()
    set_no_spacing(doc)

    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    # Row 0: first cell spans 2 columns
    cell_00 = table.cell(0, 0)
    cell_00.text = "Span 2 cols"
    tc00 = cell_00._tc
    tcPr00 = tc00.find(qn('w:tcPr'))
    if tcPr00 is None:
        tcPr00 = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc00.insert(0, tcPr00)
    gs = parse_xml(f'<w:gridSpan {nsdecls("w")} w:val="2"/>')
    tcPr00.append(gs)

    # Remove 2nd cell in row 0 (merged into first)
    # python-docx handles this differently, let's use merge
    cell_00_merged = table.cell(0, 0).merge(table.cell(0, 1))
    cell_00_merged.text = "Span 2 cols"

    table.cell(0, 2).text = "C2"
    table.cell(1, 0).text = "R1C0"
    table.cell(1, 1).text = "R1C1"
    table.cell(1, 2).text = "R1C2"

    details = run_test("gridspan", doc)

    print("  Text positions:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f} x2={t['x2']:.2f}  \"{t['text']}\"")

# ============================================================
# TEST 9: Table with tblInd (indent)
# ============================================================
def test9_indent():
    print("\n=== TEST 9: Table indent ===")
    doc = Document()
    set_no_spacing(doc)

    doc.add_paragraph("Reference line")

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Indented table"

    # Set tblInd = 72pt (1 inch, 1440tw)
    tblPr = table._tbl.find(qn('w:tblPr'))
    ind = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="1440" w:type="dxa"/>')
    tblPr.append(ind)

    doc.add_paragraph("After indented table")

    details = run_test("indent", doc)

    print("  Text positions:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")

# ============================================================
# TEST 10: Default cell padding values
# ============================================================
def test10_default_padding():
    print("\n=== TEST 10: Default cell padding ===")
    doc = Document()
    set_no_spacing(doc)

    # Table with explicit cell margins
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'

    # Set explicit cell margins: top=5pt, bottom=5pt, left=10pt, right=10pt
    tblPr = table._tbl.find(qn('w:tblPr'))
    cellMar = parse_xml(f'''<w:tblCellMar {nsdecls("w")}>
        <w:top w:w="100" w:type="dxa"/>
        <w:bottom w:w="100" w:type="dxa"/>
        <w:left w:w="200" w:type="dxa"/>
        <w:right w:w="200" w:type="dxa"/>
    </w:tblCellMar>''')
    tblPr.append(cellMar)

    table.cell(0, 0).text = "Custom margins (T=5,B=5,L=10,R=10)"
    table.cell(1, 0).text = "Row 2"

    # Also a table with NO margins (0)
    table2 = doc.add_table(rows=1, cols=1)
    table2.style = 'Table Grid'
    tblPr2 = table2._tbl.find(qn('w:tblPr'))
    cellMar2 = parse_xml(f'''<w:tblCellMar {nsdecls("w")}>
        <w:top w:w="0" w:type="dxa"/>
        <w:bottom w:w="0" w:type="dxa"/>
        <w:left w:w="0" w:type="dxa"/>
        <w:right w:w="0" w:type="dxa"/>
    </w:tblCellMar>''')
    tblPr2.append(cellMar2)
    table2.cell(0, 0).text = "Zero margins"

    details = run_test("padding", doc)

    print("  Text positions:")
    for t in details["texts"]:
        print(f"    x={t['x']:.2f} y={t['y']:.2f}  \"{t['text']}\"")


if __name__ == "__main__":
    test1_default_margins()
    test2_cell_margins_explicit()
    test3_row_heights()
    test4_borders()
    test5_alignment()
    test6_valign()
    test7_vmerge()
    test8_gridspan()
    test9_indent()
    test10_default_padding()
    print("\n=== ALL TABLE TESTS COMPLETE ===")
