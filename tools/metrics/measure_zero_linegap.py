"""
Measure actual line heights for lineGap=0 fonts (MS Mincho, MS Gothic, etc.)
to understand Word's "zero-line-gap adjustment" behavior.

Two approaches:
1. GDI TEXTMETRIC: Direct Win32 API call to get tmHeight/tmExternalLeading
   - For lineGap=0 fonts, Word adds an artificial external leading
2. Word verification: Create docx, convert to PDF, measure Y positions in PDF

The key question: when a font has lineGap=0 (hhea) and tmExternalLeading=0 (GDI),
does Word add extra spacing? If so, how much?
"""

import ctypes
from ctypes import wintypes
import json
import os
import tempfile
import time

# ─── GDI structures ───

gdi32 = ctypes.windll.gdi32
user32 = ctypes.windll.user32

class TEXTMETRICW(ctypes.Structure):
    _fields_ = [
        ("tmHeight", wintypes.LONG),
        ("tmAscent", wintypes.LONG),
        ("tmDescent", wintypes.LONG),
        ("tmInternalLeading", wintypes.LONG),
        ("tmExternalLeading", wintypes.LONG),
        ("tmAveCharWidth", wintypes.LONG),
        ("tmMaxCharWidth", wintypes.LONG),
        ("tmWeight", wintypes.LONG),
        ("tmOverhang", wintypes.LONG),
        ("tmDigitizedAspectX", wintypes.LONG),
        ("tmDigitizedAspectY", wintypes.LONG),
        ("tmFirstChar", wintypes.WORD),
        ("tmLastChar", wintypes.WORD),
        ("tmDefaultChar", wintypes.WORD),
        ("tmBreakChar", wintypes.WORD),
        ("tmItalic", ctypes.c_byte),
        ("tmUnderlined", ctypes.c_byte),
        ("tmStruckOut", ctypes.c_byte),
        ("tmPitchAndFamily", ctypes.c_byte),
        ("tmCharSet", ctypes.c_byte),
    ]

class LOGFONTW(ctypes.Structure):
    _fields_ = [
        ("lfHeight", wintypes.LONG),
        ("lfWidth", wintypes.LONG),
        ("lfEscapement", wintypes.LONG),
        ("lfOrientation", wintypes.LONG),
        ("lfWeight", wintypes.LONG),
        ("lfItalic", ctypes.c_byte),
        ("lfUnderline", ctypes.c_byte),
        ("lfStrikeOut", ctypes.c_byte),
        ("lfCharSet", ctypes.c_byte),
        ("lfOutPrecision", ctypes.c_byte),
        ("lfClipPrecision", ctypes.c_byte),
        ("lfQuality", ctypes.c_byte),
        ("lfPitchAndFamily", ctypes.c_byte),
        ("lfFaceName", ctypes.c_wchar * 32),
    ]

class OUTLINETEXTMETRICW(ctypes.Structure):
    _fields_ = [
        ("otmSize", wintypes.UINT),
        ("otmTextMetrics", TEXTMETRICW),
        ("otmFiller", ctypes.c_byte),
        ("otmPanoseNumber", ctypes.c_byte * 10),
        ("otmfsSelection", wintypes.UINT),
        ("otmfsType", wintypes.UINT),
        ("otmsCharSlopeRise", wintypes.INT),
        ("otmsCharSlopeRun", wintypes.INT),
        ("otmItalicAngle", wintypes.INT),
        ("otmEMSquare", wintypes.UINT),
        ("otmAscent", wintypes.INT),
        ("otmDescent", wintypes.INT),
        ("otmLineGap", wintypes.UINT),
        ("otmsCapEmHeight", wintypes.UINT),
        ("otmsXHeight", wintypes.UINT),
        ("otmrcFontBox_left", wintypes.LONG),
        ("otmrcFontBox_top", wintypes.LONG),
        ("otmrcFontBox_right", wintypes.LONG),
        ("otmrcFontBox_bottom", wintypes.LONG),
        ("otmMacAscent", wintypes.INT),
        ("otmMacDescent", wintypes.INT),
        ("otmMacLineGap", wintypes.UINT),
        ("otmusMinimumPPEM", wintypes.UINT),
        ("otmptSubscriptSize_x", wintypes.LONG),
        ("otmptSubscriptSize_y", wintypes.LONG),
        ("otmptSubscriptOffset_x", wintypes.LONG),
        ("otmptSubscriptOffset_y", wintypes.LONG),
        ("otmptSuperscriptSize_x", wintypes.LONG),
        ("otmptSuperscriptSize_y", wintypes.LONG),
        ("otmptSuperscriptOffset_x", wintypes.LONG),
        ("otmptSuperscriptOffset_y", wintypes.LONG),
        ("otmsStrikeoutSize", wintypes.UINT),
        ("otmsStrikeoutPosition", wintypes.INT),
        ("otmsUnderscoreSize", wintypes.INT),
        ("otmsUnderscorePosition", wintypes.INT),
        ("otmpFamilyName", wintypes.LPWSTR),
        ("otmpFaceName", wintypes.LPWSTR),
        ("otmpStyleName", wintypes.LPWSTR),
        ("otmpFullName", wintypes.LPWSTR),
    ]


def MulDiv(a, b, c):
    """Windows MulDiv: (a*b + c/2) / c with integer rounding."""
    return (a * b + c // 2) // c


def get_gdi_metrics(font_name, size_pt, dpi=96, weight=400):
    """Get GDI TEXTMETRIC and OUTLINETEXTMETRIC for a font."""
    hdc = user32.GetDC(0)
    gdi32.SetMapMode(hdc, 1)  # MM_TEXT

    size_twips = int(size_pt * 20)
    lf_height = -MulDiv(size_twips, dpi, 1440)

    lf = LOGFONTW()
    lf.lfHeight = lf_height
    lf.lfWeight = weight
    lf.lfCharSet = 1  # DEFAULT_CHARSET
    lf.lfFaceName = font_name
    lf.lfOutPrecision = 7  # OUT_TT_ONLY_PRECIS
    lf.lfQuality = 5  # CLEARTYPE_QUALITY

    hfont = gdi32.CreateFontIndirectW(ctypes.byref(lf))
    old_font = gdi32.SelectObject(hdc, hfont)

    tm = TEXTMETRICW()
    gdi32.GetTextMetricsW(hdc, ctypes.byref(tm))

    # Also get OUTLINETEXTMETRIC for lineGap info
    otm_size = gdi32.GetOutlineTextMetricsW(hdc, 0, None)
    otm = None
    otm_data = {}
    if otm_size > 0:
        buf = ctypes.create_string_buffer(otm_size)
        gdi32.GetOutlineTextMetricsW(hdc, otm_size, buf)
        otm = OUTLINETEXTMETRICW.from_buffer_copy(buf)
        otm_data = {
            "otmEMSquare": otm.otmEMSquare,
            "otmAscent": otm.otmAscent,
            "otmDescent": otm.otmDescent,
            "otmLineGap": otm.otmLineGap,
            "otmMacAscent": otm.otmMacAscent,
            "otmMacDescent": otm.otmMacDescent,
            "otmMacLineGap": otm.otmMacLineGap,
        }

    gdi32.SelectObject(hdc, old_font)
    gdi32.DeleteObject(hfont)
    user32.ReleaseDC(0, hdc)

    return {
        "font": font_name,
        "size_pt": size_pt,
        "dpi": dpi,
        "lfHeight": lf_height,
        "ppem": abs(lf_height),
        "tmHeight": tm.tmHeight,
        "tmAscent": tm.tmAscent,
        "tmDescent": tm.tmDescent,
        "tmInternalLeading": tm.tmInternalLeading,
        "tmExternalLeading": tm.tmExternalLeading,
        "height_pt": round(tm.tmHeight * 72.0 / dpi, 4),
        "extlead_pt": round(tm.tmExternalLeading * 72.0 / dpi, 4),
        "gdi_line_height_pt": round((tm.tmHeight + tm.tmExternalLeading) * 72.0 / dpi, 4),
        **otm_data,
    }


def measure_gdi_all():
    """Measure GDI metrics for all fonts at all sizes."""
    FONTS = [
        ("MS Mincho", 400),
        ("MS Gothic", 400),
        ("MS PMincho", 400),
        ("MS PGothic", 400),
        ("Yu Mincho", 400),
        ("Yu Gothic", 400),
        ("Calibri", 400),
        ("Times New Roman", 400),
        ("Arial", 400),
        ("Century", 400),
        ("Cambria", 400),
    ]

    SIZES = [8, 9, 10, 10.5, 11, 12, 14, 16, 18, 20, 24, 28, 36, 48, 72]

    results = {}

    # Measure at 96 DPI (standard) and 150 DPI (Word internal)
    for dpi in [96, 150]:
        print(f"\n{'='*90}")
        print(f"DPI = {dpi}")
        print(f"{'='*90}")
        print(f"{'Font':<20} {'Size':>5} {'ppem':>5} {'H':>4} {'A':>4} {'D':>4} {'IL':>4} {'EL':>4} {'H_pt':>7} {'EL_pt':>7} {'GDI_lh':>8} {'otmLG':>6} {'macLG':>6}")
        print("-" * 110)

        for font_name, weight in FONTS:
            for size in SIZES:
                m = get_gdi_metrics(font_name, size, dpi, weight)
                key = f"{font_name}_{weight}_dpi{dpi}"
                if key not in results:
                    results[key] = []
                results[key].append(m)

                otm_lg = m.get("otmLineGap", "?")
                mac_lg = m.get("otmMacLineGap", "?")
                print(f"{font_name:<20} {size:>5.1f} {m['ppem']:>5} {m['tmHeight']:>4} {m['tmAscent']:>4} {m['tmDescent']:>4} {m['tmInternalLeading']:>4} {m['tmExternalLeading']:>4} {m['height_pt']:>7.2f} {m['extlead_pt']:>7.2f} {m['gdi_line_height_pt']:>8.2f} {otm_lg:>6} {mac_lg:>6}")

    return results


def verify_with_word():
    """Create docx files and convert to PDF to measure actual Word line heights."""
    try:
        import fitz
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        import pythoncom
        pythoncom.CoInitialize()
        import win32com.client
    except ImportError as e:
        print(f"\nSkipping Word verification (missing dependency: {e})")
        return []

    TMPDIR = tempfile.gettempdir()

    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False
    word_app.DisplayAlerts = False

    def docx_to_pdf(docx_path, pdf_path):
        doc = word_app.Documents.Open(os.path.abspath(docx_path).replace("/", "\\"))
        time.sleep(0.3)
        doc.SaveAs(os.path.abspath(pdf_path).replace("/", "\\"), FileFormat=17)
        time.sleep(0.2)
        doc.Close(False)

    def set_no_spacing(doc):
        style_el = doc.styles['Normal'].element
        pPr = style_el.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            style_el.append(pPr)
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        sp = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:line="240" w:lineRule="auto"/>')
        pPr.append(sp)

    def set_snap_false(para):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            para._element.insert(0, pPr)
        snap = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
        pPr.append(snap)

    def get_y_positions(pdf_path):
        doc = fitz.open(pdf_path)
        page = doc[0]
        blocks = page.get_text("dict")["blocks"]
        positions = []
        for b in blocks:
            if "lines" in b:
                for line in b["lines"]:
                    text = "".join(s["text"] for s in line["spans"])
                    positions.append((line["bbox"][1], text.strip()))
        doc.close()
        return positions

    FONTS_SIZES = [
        ("MS Mincho", 10.5),
        ("MS Mincho", 11),
        ("MS Mincho", 12),
        ("MS Mincho", 14),
        ("MS Mincho", 16),
        ("MS Mincho", 20),
        ("MS Mincho", 24),
        ("MS Gothic", 10.5),
        ("MS Gothic", 11),
        ("MS Gothic", 12),
        ("MS Gothic", 14),
        ("MS Gothic", 16),
        ("MS Gothic", 20),
        ("Yu Mincho", 10.5),
        ("Yu Mincho", 12),
        ("Yu Mincho", 14),
        ("Yu Gothic Regular", 10.5),
        ("Yu Gothic Regular", 12),
        ("Yu Gothic Regular", 14),
        ("Calibri", 10.5),
        ("Calibri", 11),
        ("Calibri", 12),
        ("Calibri", 14),
        ("Calibri", 16),
        ("Calibri", 20),
        ("Times New Roman", 10.5),
        ("Times New Roman", 12),
        ("Times New Roman", 14),
    ]

    results = []
    print(f"\n{'='*80}")
    print("Word PDF verification (snap=false, single spacing)")
    print(f"{'='*80}")
    print(f"{'Font':<25} {'Size':>5}  {'AvgGap':>8}  {'Gaps'}")
    print("-" * 80)

    for font_name, font_size in FONTS_SIZES:
        doc = Document()
        set_no_spacing(doc)

        for i in range(10):
            p = doc.add_paragraph()
            r = p.add_run(f"Line {i+1} {font_name}")
            r.font.name = font_name
            r.font.size = Pt(font_size)
            set_snap_false(p)

        docx_path = os.path.join(TMPDIR, "zero_linegap_test.docx")
        pdf_path = os.path.join(TMPDIR, "zero_linegap_test.pdf")
        doc.save(docx_path)
        docx_to_pdf(docx_path, pdf_path)

        positions = get_y_positions(pdf_path)
        gaps = []
        for i in range(1, len(positions)):
            gaps.append(round(positions[i][0] - positions[i-1][0], 2))

        avg_gap = sum(gaps) / len(gaps) if gaps else 0
        gap_str = " ".join(f"{g:.2f}" for g in gaps[:5])
        print(f"{font_name:<25} {font_size:>5.1f}  {avg_gap:>8.2f}  {gap_str}")

        results.append({
            "font": font_name,
            "size_pt": font_size,
            "avg_gap_pt": round(avg_gap, 4),
            "gaps": gaps,
            "num_lines": len(positions),
        })

        try:
            os.unlink(docx_path)
            os.unlink(pdf_path)
        except:
            pass

    word_app.Quit()
    return results


def analyze_zero_linegap(gdi_results):
    """Analyze the zero-line-gap adjustment pattern."""
    print(f"\n\n{'='*90}")
    print("ANALYSIS: Zero lineGap fonts vs normal fonts")
    print(f"{'='*90}")

    # At 150 DPI (Word's internal DPI)
    zero_lg_fonts = ["MS Mincho", "MS Gothic", "MS PMincho", "MS PGothic"]
    nonzero_lg_fonts = ["Calibri", "Times New Roman", "Arial", "Century", "Cambria"]
    yu_fonts = ["Yu Mincho", "Yu Gothic"]

    print("\n--- External Leading comparison at 150 DPI ---")
    print(f"{'Font':<20} {'Size':>5} {'EL':>4} {'tmH':>4} {'H_pt':>7} {'EL_pt':>7} {'GDI_lh':>8} {'Ratio(lh/sz)':>12}")
    print("-" * 80)

    for font_name in zero_lg_fonts + ["---"] + yu_fonts + ["---"] + nonzero_lg_fonts:
        if font_name == "---":
            print()
            continue
        key = f"{font_name}_400_dpi150"
        if key not in gdi_results:
            continue
        for m in gdi_results[key]:
            ratio = m['gdi_line_height_pt'] / m['size_pt']
            print(f"{font_name:<20} {m['size_pt']:>5.1f} {m['tmExternalLeading']:>4} {m['tmHeight']:>4} {m['height_pt']:>7.2f} {m['extlead_pt']:>7.2f} {m['gdi_line_height_pt']:>8.2f} {ratio:>12.4f}")

    # Check if EL is always 0 for zero-linegap fonts
    print("\n\n--- Is tmExternalLeading always 0 for lineGap=0 fonts? ---")
    for font_name in zero_lg_fonts + yu_fonts:
        key = f"{font_name}_400_dpi150"
        if key not in gdi_results:
            continue
        els = [m['tmExternalLeading'] for m in gdi_results[key]]
        otm_lg = gdi_results[key][0].get('otmLineGap', '?')
        mac_lg = gdi_results[key][0].get('otmMacLineGap', '?')
        print(f"  {font_name:<20} otmLineGap={otm_lg}, macLineGap={mac_lg}, EL values: {els}")

    # Word's line height formula: (tmHeight + tmExternalLeading) * lineSpacing / 240 * (pt conversion)
    # But for zero-EL fonts, Word seems to add something. Let's see what.
    print("\n\n--- Word single-line formula check ---")
    print("Word formula: gdi_height = (tmHeight + tmExternalLeading) in device units")
    print("              line_height_pt = gdi_height * 72 / dpi")
    print("For lineGap=0 fonts with EL=0, gdi_height = tmHeight only")
    print("But Word might add artificial EL. Let's compute what EL *should* be:")
    print()

    for font_name in zero_lg_fonts:
        key96 = f"{font_name}_400_dpi96"
        key150 = f"{font_name}_400_dpi150"
        if key96 not in gdi_results:
            continue

        # At both DPIs, check consistency
        print(f"  {font_name}:")
        for m in gdi_results.get(key150, []):
            if m['size_pt'] in [10.5, 12, 14]:
                # Without fix: line height = tmHeight only
                no_fix_pt = m['height_pt']
                # With fix: what would Word need to add?
                # Typical ratio for fonts with lineGap is ~1.2x fontSize
                # So artificial_el = desired_height - tmHeight
                print(f"    {m['size_pt']:>5.1f}pt: tmH={m['tmHeight']}, EL={m['tmExternalLeading']}, "
                      f"H_pt={m['height_pt']:.2f}, GDI_lh={m['gdi_line_height_pt']:.2f}")


def main():
    OUTPUT_PATH = r"c:\Users\ryuji\oxi-1\tools\metrics\output\zero_linegap_measurements.json"

    print("=" * 90)
    print("Measuring GDI TEXTMETRIC for lineGap=0 fonts")
    print("=" * 90)

    # Step 1: Pure GDI measurements
    gdi_results = measure_gdi_all()

    # Step 2: Analyze
    analyze_zero_linegap(gdi_results)

    # Step 3: Save GDI results first (in case Word verification crashes)
    output = {
        "description": "Zero lineGap font measurements - GDI TEXTMETRIC + Word verification",
        "date": "2026-03-16",
        "gdi_measurements": gdi_results,
        "word_pdf_verification": [],
    }
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    print(f"\nGDI results saved to {OUTPUT_PATH}")

    # Step 4: Word PDF verification
    word_results = verify_with_word()

    # Step 5: Save combined results
    output["word_pdf_verification"] = word_results
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    print(f"\nFull results saved to {OUTPUT_PATH}")

    # Final comparison if Word results are available
    if word_results:
        print(f"\n\n{'='*90}")
        print("FINAL COMPARISON: GDI prediction vs Word actual")
        print(f"{'='*90}")
        print(f"{'Font':<25} {'Size':>5} {'GDI_lh(96)':>10} {'GDI_lh(150)':>11} {'Word_actual':>11} {'Delta(W-G96)':>12}")
        print("-" * 80)

        for wr in word_results:
            fn = wr['font']
            gdi_fn = fn.replace(" Regular", "")
            sz = wr['size_pt']

            gdi96 = None
            gdi150 = None
            for m in gdi_results.get(f"{gdi_fn}_400_dpi96", []):
                if abs(m['size_pt'] - sz) < 0.01:
                    gdi96 = m['gdi_line_height_pt']
            for m in gdi_results.get(f"{gdi_fn}_400_dpi150", []):
                if abs(m['size_pt'] - sz) < 0.01:
                    gdi150 = m['gdi_line_height_pt']

            g96_str = f"{gdi96:.2f}" if gdi96 else "?"
            g150_str = f"{gdi150:.2f}" if gdi150 else "?"
            word_val = wr['avg_gap_pt']
            delta = f"{word_val - (gdi96 or 0):.2f}" if gdi96 else "?"
            print(f"{fn:<25} {sz:>5.1f} {g96_str:>10} {g150_str:>11} {word_val:>11.2f} {delta:>12}")


if __name__ == "__main__":
    main()
