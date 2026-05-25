"""
Minimal repros isolating cursor_y after row-split.

Target hypothesis: after a 1-cell row-split pushes content onto p.2,
Oxi sets cursor_y = page_top + (row_bottom - split_y) which can be
LESS than the actual bottom of the continuation box. The body paragraph
that follows the table then overlaps with the continuation box.

Repros vary:
  CR_1 — 1 line overflow (Oxi should place body para below y=89 on p.2)
  CR_2 — 3 line overflow
  CR_3 — 1 line overflow + empty paragraph after table
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent / "rowsplit_cursor_repro"
OUT_DIR.mkdir(exist_ok=True)


def set_run_gothic(run, pt: float = 10.5):
    rPr = run._r.get_or_add_rPr()
    for tag in ['rFonts']:
        for e in rPr.findall(qn('w:' + tag)):
            rPr.remove(e)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'MS Gothic')
    rf.set(qn('w:eastAsia'), 'MS Gothic')
    rf.set(qn('w:hAnsi'), 'MS Gothic')
    rPr.append(rf)
    run.font.size = Pt(pt)


def set_table_borders_all(tbl):
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)


def set_section_docgrid(section, line_pitch_tw=360):
    sectPr = section._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)
    dg = OxmlElement('w:docGrid')
    dg.set(qn('w:type'), 'linesAndChars')
    dg.set(qn('w:linePitch'), str(line_pitch_tw))
    dg.set(qn('w:charSpace'), '0')
    sectPr.append(dg)


def add_adjust_line_height_in_table(doc):
    """Insert <w:adjustLineHeightInTable/> into settings.xml.
    d77a + most real docs have this flag; python-docx omits by default.
    Without it, Oxi (and Word, per Day 28 V70 measurement) renders cell
    paragraphs at text_render_height instead of snapping to docGrid linePitch."""
    settings = doc.settings.element
    # Remove any existing one to avoid duplicates
    for e in settings.findall(qn('w:adjustLineHeightInTable')):
        settings.remove(e)
    el = OxmlElement('w:adjustLineHeightInTable')
    # Insert near the top of settings (after zoom/view if present)
    settings.insert(0, el)


def build_base(filler_n: int = 40):
    """Base doc: many fillers to push content near bottom."""
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    set_section_docgrid(section, line_pitch_tw=360)
    for i in range(filler_n):
        p = doc.add_paragraph()
        run = p.add_run(f"Filler {i+1:02d}")
        set_run_gothic(run, pt=10.5)
    return doc


def add_body_para(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_gothic(run, pt=10.5)
    return p


def add_table_with_long_para(doc, text: str, cells_first_para_text: str = None):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    for r in list(p._p.findall(qn('w:r'))):
        p._p.remove(r)
    run = p.add_run(text)
    set_run_gothic(run, pt=10.5)
    return table


def make_cr1():
    """1-line overflow: table pushed to bottom, 1 line goes to next page."""
    doc = build_base(filler_n=52)
    text = "あ" * 90 + "終"
    add_table_with_long_para(doc, text)
    add_body_para(doc, "AFTER_TABLE_BODY_PARA")
    doc.save(OUT_DIR / "CR_1_one_line_overflow.docx")


def make_cr2():
    """3-line overflow."""
    doc = build_base(filler_n=52)
    text = "あ" * 200 + "終"
    add_table_with_long_para(doc, text)
    add_body_para(doc, "AFTER_TABLE_BODY_PARA")
    doc.save(OUT_DIR / "CR_2_three_line_overflow.docx")


def make_cr3():
    """1-line overflow + empty paragraph between table and next content."""
    doc = build_base(filler_n=52)
    text = "あ" * 90 + "終"
    add_table_with_long_para(doc, text)
    doc.add_paragraph()
    add_body_para(doc, "AFTER_TABLE_BODY_PARA")
    doc.save(OUT_DIR / "CR_3_one_line_plus_empty.docx")


def make_cr4():
    """S269 add: content sized to force overflow in Oxi's CURRENT geometric
    accounting (lines emitted at text_render_height=13.5pt for MS Gothic 10.5pt).
    Need > 698/13.5 ≈ 51 lines × 43 chars/line ≈ 2200 chars worth in cell.
    3500 chars ≈ 81 lines × 13.5 = 1094pt → forces multi-page row split."""
    doc = build_base(filler_n=52)
    text = "あ" * 3500 + "終"
    add_table_with_long_para(doc, text)
    add_body_para(doc, "AFTER_TABLE_BODY_PARA")
    doc.save(OUT_DIR / "CR_4_multipage_content.docx")


def make_cr5():
    """S269 add: small filler so table starts mid-pg1 with enough free
    space to bypass widow gate, content sized to overflow ~2 lines."""
    doc = build_base(filler_n=30)
    text = "あ" * 700 + "終"
    add_table_with_long_para(doc, text)
    add_body_para(doc, "AFTER_TABLE_BODY_PARA")
    doc.save(OUT_DIR / "CR_5_no_widow_overflow.docx")


def make_cr6():
    """S269 add: CR_4 + <w:adjustLineHeightInTable/> flag (matches d77a).
    With the flag, Oxi snaps cell lines to docGrid linePitch=18pt (per Day 28
    V70). This is the proper d77a-equivalent minimal repro for Pattern A:
    Bug #2 (mod.rs:8794 cursor.set geometric vs structural) without the
    flag-absent cell-line-spacing divergence that confused CR_4 measurement."""
    doc = build_base(filler_n=52)
    add_adjust_line_height_in_table(doc)
    # 2500 chars / 43 per line ≈ 58 lines × 18pt = 1044pt
    # pg3 content area ≈ 698pt → ~38 lines fit, 20 overflow → multi-page split
    text = "あ" * 2500 + "終"
    add_table_with_long_para(doc, text)
    add_body_para(doc, "AFTER_TABLE_BODY_PARA")
    doc.save(OUT_DIR / "CR_6_adjustLineHeight_flag.docx")


def make_cr7():
    """S269 add: CR_6 with TRAILING EMPTY paragraph in cell. Mirrors
    d77a t8/t10 + e3c545 t2 (3 of 4 real-doc splits). Word formula
    predicts body = last_cont_top + lh × (1 + 1 trailing_empty) = +2×lh.
    Without trailing-empty handling, Oxi would still place body at
    last_cont_top + lh × 1 → -lh deficit."""
    doc = build_base(filler_n=52)
    add_adjust_line_height_in_table(doc)
    text = "あ" * 2500 + "終"
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    cell = table.cell(0, 0)
    # First para: long content
    p = cell.paragraphs[0]
    for r in list(p._p.findall(qn('w:r'))):
        p._p.remove(r)
    run = p.add_run(text)
    set_run_gothic(run, pt=10.5)
    # Trailing empty paragraph inside the same cell. Use an empty run rather
    # than no run so Oxi's parser captures the Paragraph block. The runs are
    # all-empty-text → trailing_empty_count detects it.
    trail_p = cell.add_paragraph()
    trail_run = trail_p.add_run("")
    set_run_gothic(trail_run, pt=10.5)
    add_body_para(doc, "AFTER_TABLE_BODY_PARA")
    doc.save(OUT_DIR / "CR_7_trailing_empty.docx")


if __name__ == "__main__":
    make_cr1()
    make_cr2()
    make_cr3()
    make_cr4()
    make_cr5()
    make_cr6()
    make_cr7()
    print("Built:")
    for f in sorted(OUT_DIR.glob("*.docx")):
        print(" ", f.name)
