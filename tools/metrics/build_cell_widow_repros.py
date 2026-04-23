"""Repros to determine Word's page-break rule in table cells
when widowControl=0 (explicitly disabled).

d77a Normal style has `<w:widowControl w:val="0"/>`. Yet Word moves para 95
(last content para of a cell, multi-line, only 1 line fits on split page)
entirely to next page. Oxi respects widowControl=0 → allows orphan → +1 line
on split page → pixel mismatch.

Hypothesis candidates:
  H1: Word ignores widowControl=0 for ALL table cell paragraphs.
  H2: Word applies special protection to the LAST paragraph of a cell.
  H3: Word ignores widowControl=0 only for paragraphs whose FIRST line
      would be orphaned (not widow).

Repros (A4, 2.5cm margins, default font):
  W1: widowControl=0 globally, cell with 3 paragraphs, last para multi-line,
      arranged so only 1 line of last para fits. Expected: Word push if H2.
  W2: widowControl=1 globally, same structure. Baseline: Word pushes.
  W3: widowControl=0 globally, cell with 3 paragraphs, MIDDLE para multi-line
      with potential orphan. (Not last.) Expected: Word push if H1,
      Word does NOT push if H2/H3.
  W4: widowControl=0 globally, BODY paragraph (no cell), orphan line.
      Expected: Word does NOT push (respects widow=0 in body).
"""
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Twips, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import sys
    print("python-docx required")
    sys.exit(1)

OUT = Path(__file__).parent / "cell_widow_repro"
OUT.mkdir(exist_ok=True)


def set_docdefault_widow(doc, val: str):
    """Set docDefaults pPr widowControl."""
    styles_part = doc.part.part_related_by(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles"
    )
    styles_xml = styles_part.blob.decode('utf-8')
    # Not modifying styles.xml from python-docx easily — use style instead.
    # Fallback: set widowControl on each paragraph individually.
    pass


def set_para_widow(p, val: str):
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn('w:widowControl'))
    if old is not None: pPr.remove(old)
    el = OxmlElement('w:widowControl')
    el.set(qn('w:val'), val)
    # Insert widowControl as FIRST child (OOXML order constraint)
    pPr.insert(0, el)


def set_table_borders_all(table, sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz)); b.set(qn('w:color'),'auto')
        tblBorders.append(b)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None: tblPr.remove(old)
    tblPr.append(tblBorders)


def make_page(filler_n=22):
    doc = Document()
    s = doc.sections[0]
    s.page_height = Cm(29.7); s.page_width = Cm(21.0)
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    for i in range(filler_n):
        p = doc.add_paragraph(f"F{i+1} " + "あ"*30)
        set_para_widow(p, "0")
    return doc


def w1_cell_last_para_orphan(widow: str) -> Path:
    """Cell with 3 paragraphs, last multi-line, only 1 line fits on split."""
    doc = make_page(filler_n=22)
    t = doc.add_table(rows=1, cols=1)
    set_table_borders_all(t)
    cell = t.cell(0, 0)
    cell.text = ""
    # Para 1-2: short paragraphs
    p1 = cell.paragraphs[0]
    p1.add_run("Cell P1: " + "い"*20)
    set_para_widow(p1, widow)
    p2 = cell.add_paragraph()
    p2.add_run("Cell P2: " + "う"*20)
    set_para_widow(p2, widow)
    # Last para: multi-line (5 lines worth)
    p3 = cell.add_paragraph()
    p3.add_run("Cell P3 (last, multi-line): " + "え"*160)
    set_para_widow(p3, widow)
    name = f"W1_cell_last_orphan_widow{widow}.docx"
    p = OUT / name
    doc.save(p)
    return p


def w3_cell_middle_para_orphan(widow: str) -> Path:
    """Cell with 3 paragraphs, MIDDLE para multi-line."""
    doc = make_page(filler_n=22)
    t = doc.add_table(rows=1, cols=1)
    set_table_borders_all(t)
    cell = t.cell(0, 0)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.add_run("Cell P1: " + "い"*20)
    set_para_widow(p1, widow)
    # Middle para: multi-line
    p2 = cell.add_paragraph()
    p2.add_run("Cell P2 (middle, multi-line): " + "う"*160)
    set_para_widow(p2, widow)
    # Last para: short
    p3 = cell.add_paragraph()
    p3.add_run("Cell P3 (last, short): " + "え"*15)
    set_para_widow(p3, widow)
    name = f"W3_cell_middle_orphan_widow{widow}.docx"
    p = OUT / name
    doc.save(p)
    return p


def w4_body_orphan(widow: str) -> Path:
    """Body paragraph (no table), multi-line, orphan on split."""
    doc = make_page(filler_n=22)
    # One large body paragraph spanning pages
    p = doc.add_paragraph()
    p.add_run("Body orphan test: " + "お"*200)
    set_para_widow(p, widow)
    # Trailing short para to force split visibility
    p2 = doc.add_paragraph()
    p2.add_run("Trailing.")
    set_para_widow(p2, widow)
    name = f"W4_body_orphan_widow{widow}.docx"
    p_ = OUT / name
    doc.save(p_)
    return p_


if __name__ == "__main__":
    created = []
    for w in ("0", "1"):
        created.append(w1_cell_last_para_orphan(w))
        created.append(w3_cell_middle_para_orphan(w))
        created.append(w4_body_orphan(w))
    for p in created:
        print("created:", p)
