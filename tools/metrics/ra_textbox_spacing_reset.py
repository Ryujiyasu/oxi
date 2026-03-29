"""
Ra: TextBox内の spacing リセット条件をCOM計測で確定
- TextBox内段落の lineSpacing, spaceBefore, spaceAfter はリセットされるか？
- Normal スタイルから継承した値はどうなるか？
- テーブルセルと同じ挙動か？
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt, Twips, Emu, Inches
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def add_textbox_with_text(doc, text_lines, anchor_para=None):
    """Add a textbox via XML manipulation."""
    # We'll use Word COM to insert textbox after saving, since python-docx can't do it natively
    pass


def make_doc_body(scenario):
    """Create doc with body paragraphs that have known spacing from Normal style."""
    d = Document(TEMPLATE)

    # Set Normal style to have sa=8pt, sb=0, ls=1.15x
    # (This is typical Word default)
    style = d.styles['Normal']
    pf = style.paragraph_format
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)

    # Remove default paragraphs
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Add 3 body paragraphs as baseline
    for i in range(3):
        p = d.add_paragraph()
        r = p.add_run(f"Body paragraph {i+1}")
        r.font.name = "Calibri"
        r.font.size = Pt(11)

    return d


def measure_textbox_via_com(doc_path):
    """Open in Word, insert textbox via COM, measure paragraph positions inside."""
    doc = word.Documents.Open(doc_path)
    try:
        data = {"body_paragraphs": [], "textbox_paragraphs": []}

        # Measure body paragraphs first
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            rng = para.Range
            y = rng.Information(6)
            data["body_paragraphs"].append({
                "index": i,
                "y_pt": round(y, 4),
                "text": rng.Text.strip()[:40],
                "line_spacing": round(para.Format.LineSpacing, 4),
                "line_spacing_rule": para.Format.LineSpacingRule,
                "space_before": round(para.Format.SpaceBefore, 4),
                "space_after": round(para.Format.SpaceAfter, 4),
            })

        # Compute body gaps
        for i in range(1, len(data["body_paragraphs"])):
            gap = data["body_paragraphs"][i]["y_pt"] - data["body_paragraphs"][i-1]["y_pt"]
            data["body_paragraphs"][i]["gap_from_prev"] = round(gap, 4)

        # Insert textbox with 3 paragraphs via COM
        rng_insert = doc.Range(0, 0)
        # AddTextbox(Orientation, Left, Top, Width, Height, Anchor)
        tb = doc.Shapes.AddTextbox(
            1,  # msoTextOrientationHorizontal
            100,  # Left (pt)
            300,  # Top (pt)
            200,  # Width (pt)
            150,  # Height (pt)
            rng_insert
        )
        tf = tb.TextFrame

        # Set internal margins
        tf.MarginLeft = 7.2  # ~0.1 inch default
        tf.MarginRight = 7.2
        tf.MarginTop = 3.6
        tf.MarginBottom = 3.6

        # Add text to textbox
        tf.TextRange.Text = "TB Line 1\rTB Line 2\rTB Line 3"

        # Set font properties directly (don't set style - causes error in textbox)
        for i in range(1, tf.TextRange.Paragraphs.Count + 1):
            para = tf.TextRange.Paragraphs(i)
            para.Range.Font.Name = "Calibri"
            para.Range.Font.Size = 11

        # Now measure textbox paragraph properties
        for i in range(1, tf.TextRange.Paragraphs.Count + 1):
            para = tf.TextRange.Paragraphs(i)
            rng = para.Range
            data["textbox_paragraphs"].append({
                "index": i,
                "text": rng.Text.strip()[:40],
                "line_spacing": round(para.Format.LineSpacing, 4),
                "line_spacing_rule": para.Format.LineSpacingRule,
                "space_before": round(para.Format.SpaceBefore, 4),
                "space_after": round(para.Format.SpaceAfter, 4),
            })

        return data
    finally:
        doc.Close(False)


def measure_table_cell_via_com(doc_path):
    """For comparison: measure table cell paragraph properties."""
    doc = word.Documents.Open(doc_path)
    try:
        data = {"table_cell_paragraphs": []}

        # Insert a 1x1 table via COM
        rng_insert = doc.Range(doc.Content.End - 1, doc.Content.End - 1)
        tbl = doc.Tables.Add(rng_insert, 1, 1)

        cell = tbl.Cell(1, 1)
        cell.Range.Text = "Cell Line 1\rCell Line 2\rCell Line 3"

        # Set font properties directly
        for i in range(1, cell.Range.Paragraphs.Count + 1):
            para = cell.Range.Paragraphs(i)
            para.Range.Font.Name = "Calibri"
            para.Range.Font.Size = 11

        # Measure
        for i in range(1, cell.Range.Paragraphs.Count + 1):
            para = cell.Range.Paragraphs(i)
            data["table_cell_paragraphs"].append({
                "index": i,
                "text": para.Range.Text.strip()[:40],
                "line_spacing": round(para.Format.LineSpacing, 4),
                "line_spacing_rule": para.Format.LineSpacingRule,
                "space_before": round(para.Format.SpaceBefore, 4),
                "space_after": round(para.Format.SpaceAfter, 4),
            })

        return data
    finally:
        doc.Close(False)


try:
    # Test 1: Body paragraphs + TextBox paragraphs
    d = make_doc_body("textbox_spacing")
    tmp = os.path.join(tempfile.gettempdir(), "ra_textbox_sp.docx")
    d.save(tmp)

    print("=== Measuring body + textbox ===")
    data1 = measure_textbox_via_com(tmp)
    results.append({"test": "textbox_vs_body", **data1})

    print("\nBody paragraphs:")
    for p in data1["body_paragraphs"]:
        gap_str = f"  gap={p.get('gap_from_prev', '-')}" if 'gap_from_prev' in p else ""
        print(f"  P{p['index']}: y={p['y_pt']}pt, ls={p['line_spacing']}pt(rule={p['line_spacing_rule']}), "
              f"sb={p['space_before']}pt, sa={p['space_after']}pt{gap_str}")

    print("\nTextBox paragraphs:")
    for p in data1["textbox_paragraphs"]:
        print(f"  P{p['index']}: ls={p['line_spacing']}pt(rule={p['line_spacing_rule']}), "
              f"sb={p['space_before']}pt, sa={p['space_after']}pt  [{p['text']}]")

    # Test 2: Table cell paragraphs for comparison
    d2 = make_doc_body("table_cell_comparison")
    tmp2 = os.path.join(tempfile.gettempdir(), "ra_tablecell_sp.docx")
    d2.save(tmp2)

    print("\n=== Measuring table cell ===")
    data2 = measure_table_cell_via_com(tmp2)
    results.append({"test": "table_cell", **data2})

    print("\nTable cell paragraphs:")
    for p in data2["table_cell_paragraphs"]:
        print(f"  P{p['index']}: ls={p['line_spacing']}pt(rule={p['line_spacing_rule']}), "
              f"sb={p['space_before']}pt, sa={p['space_after']}pt  [{p['text']}]")

    os.unlink(tmp)
    os.unlink(tmp2)

finally:
    word.Quit()

# Save results
out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_textbox_spacing_reset.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)

print(f"\nResults saved to {out_path}")

# Analysis
print("\n=== ANALYSIS ===")
body = data1["body_paragraphs"]
tb = data1["textbox_paragraphs"]
tc = data2["table_cell_paragraphs"]

print(f"Body P1: ls={body[0]['line_spacing']}, sa={body[0]['space_after']}, sb={body[0]['space_before']}")
print(f"TextBox P1: ls={tb[0]['line_spacing']}, sa={tb[0]['space_after']}, sb={tb[0]['space_before']}")
print(f"TableCell P1: ls={tc[0]['line_spacing']}, sa={tc[0]['space_after']}, sb={tc[0]['space_before']}")

if tb[0]['line_spacing'] == tc[0]['line_spacing'] and tb[0]['space_after'] == tc[0]['space_after']:
    print("=> TextBox and TableCell have SAME spacing reset behavior")
else:
    print("=> TextBox and TableCell have DIFFERENT spacing behavior!")
    print(f"   TextBox: ls={tb[0]['line_spacing']}, sa={tb[0]['space_after']}")
    print(f"   TableCell: ls={tc[0]['line_spacing']}, sa={tc[0]['space_after']}")
