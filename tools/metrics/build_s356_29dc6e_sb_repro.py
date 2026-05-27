"""S356 — minimal repro: isolate 29dc6e i=130's sb-suppression discriminator.

Build a series of 2-row tables where row 2's first paragraph has different
pPr signatures. Pair each variant with a "control" (same except sb=0) so we
can measure the actual sb applied by Word via dy difference.

Variants:
  V_E0_*: just `<w:before>` (Day 33 part 17 condition without lineRule)
  V_E1_*: `<w:before>` + `<w:lineRule="exact">`
  V_E2_*: V_E1 + `<w:beforeLines>` (29dc6e signature without pStyle)
  V_E3_*: V_E2 + `<w:pStyle w:val="ac">` (full 29dc6e signature)
  V_E4_*: V_E3 + docGrid linesAndChars+charSpace=1453 (family doc-level)

Each variant comes in two forms: _A (sb=0 control) and _B (sb=146 = 7.3pt).

Measurement: take row2_first_para_y - row1_first_para_y.
  Δy(B) - Δy(A) = sb actually applied by Word.
  - If Δ = 7.3pt → Word applies sb (current S136 rule correct for this case)
  - If Δ = 0 → Word suppresses sb (S136 rule wrong for this case)
"""
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.abspath("tools/golden-test/repros/s356_29dc6e_sb")
os.makedirs(OUT_DIR, exist_ok=True)


def set_doc_grid_linesAndChars(doc, line_pitch=292, char_space=1453):
    sectPr = doc.sections[0]._sectPr
    for ex in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(ex)
    dg = OxmlElement("w:docGrid")
    dg.set(qn("w:type"), "linesAndChars")
    dg.set(qn("w:linePitch"), str(line_pitch))
    dg.set(qn("w:charSpace"), str(char_space))
    sectPr.append(dg)


def add_para_with_signature(cell, text, sb_twips, *, has_lineRule=False,
                             has_beforeLines=False, has_pStyle_ac=False,
                             pmark_spacing0=False):
    """Add a paragraph to cell with controlled pPr."""
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    # Clear existing
    for name in ["widowControl", "wordWrap", "autoSpaceDE", "autoSpaceDN",
                 "adjustRightInd", "jc", "spacing", "rPr", "pStyle"]:
        ex = pPr.find(qn(f"w:{name}"))
        if ex is not None:
            pPr.remove(ex)

    if has_pStyle_ac:
        # ac style is custom — we'll define it below at doc level
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), "ac")
        pPr.append(ps)
        # ac inherits these:
        for tag in ["widowControl", "wordWrap", "autoSpaceDE", "autoSpaceDN", "adjustRightInd"]:
            e = OxmlElement(f"w:{tag}")
            e.set(qn("w:val"), "0")
            pPr.append(e)

    sp = OxmlElement("w:spacing")
    if has_beforeLines and sb_twips > 0:
        # beforeLines=50 = 50% × line. For linePitch=292 (14.6pt), 50% = 7.3pt = 146 dxa
        sp.set(qn("w:beforeLines"), "50")
    if sb_twips > 0:
        sp.set(qn("w:before"), str(sb_twips))
    if has_lineRule:
        sp.set(qn("w:line"), "259")
        sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)

    if pmark_spacing0:
        pmark_rpr = OxmlElement("w:rPr")
        csp = OxmlElement("w:spacing")
        csp.set(qn("w:val"), "0")
        pmark_rpr.append(csp)
        pPr.append(pmark_rpr)

    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "ＭＳ 明朝"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hAnsi"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    return p


def define_ac_style(doc):
    """Define a custom paragraph style 'ac' matching 29dc6e."""
    styles_part = doc.styles.element
    # Check if style 'ac' already exists
    for s in styles_part.findall(qn("w:style")):
        if s.get(qn("w:styleId")) == "ac":
            return
    style_el = OxmlElement("w:style")
    style_el.set(qn("w:type"), "paragraph")
    style_el.set(qn("w:customStyle"), "1")
    style_el.set(qn("w:styleId"), "ac")
    name = OxmlElement("w:name")
    name.set(qn("w:val"), "ac")
    style_el.append(name)
    pPr = OxmlElement("w:pPr")
    for tag in ["widowControl", "wordWrap", "autoSpaceDE", "autoSpaceDN", "adjustRightInd"]:
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:val"), "0")
        pPr.append(e)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "210")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    style_el.append(pPr)
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hAnsi"), "ＭＳ 明朝")
    rFonts.set(qn("w:cs"), "ＭＳ 明朝")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")
    rPr.append(sz)
    style_el.append(rPr)
    styles_part.append(style_el)


def make_doc(name, *, sb_twips, has_lineRule, has_beforeLines, has_pStyle_ac,
             use_docGrid_lac):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"

    if has_pStyle_ac:
        define_ac_style(doc)
    if use_docGrid_lac:
        set_doc_grid_linesAndChars(doc)

    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Table Grid"

    # Row 1: simple paragraph, no sb (reference)
    cell1 = tbl.rows[0].cells[0]
    # Remove default empty paragraph
    default_p = cell1.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    add_para_with_signature(cell1, "ROW1_REF", sb_twips=0,
                            has_lineRule=False, has_beforeLines=False,
                            has_pStyle_ac=False)

    # Row 2: target paragraph with controlled signature
    cell2 = tbl.rows[1].cells[0]
    default_p2 = cell2.paragraphs[0]
    default_p2._element.getparent().remove(default_p2._element)
    add_para_with_signature(cell2, "ROW2_TARGET", sb_twips=sb_twips,
                            has_lineRule=has_lineRule,
                            has_beforeLines=has_beforeLines,
                            has_pStyle_ac=has_pStyle_ac)

    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    # 5 signatures × 2 (sb=0 vs sb=146)
    configs = [
        ("E0_before_only",          dict(has_lineRule=False, has_beforeLines=False, has_pStyle_ac=False, use_docGrid_lac=False)),
        ("E1_before_lineRule",      dict(has_lineRule=True,  has_beforeLines=False, has_pStyle_ac=False, use_docGrid_lac=False)),
        ("E2_before_lineRule_bl",   dict(has_lineRule=True,  has_beforeLines=True,  has_pStyle_ac=False, use_docGrid_lac=False)),
        ("E3_full_29dc6e_no_grid",  dict(has_lineRule=True,  has_beforeLines=True,  has_pStyle_ac=True,  use_docGrid_lac=False)),
        ("E4_full_29dc6e_w_grid",   dict(has_lineRule=True,  has_beforeLines=True,  has_pStyle_ac=True,  use_docGrid_lac=True)),
    ]
    for name, cfg in configs:
        make_doc(f"{name}_A_sb0",   sb_twips=0,   **cfg)
        make_doc(f"{name}_B_sb146", sb_twips=146, **cfg)
