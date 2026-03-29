"""
Ra: バッチ7 — テーブル詳細、TextBox詳細、フォント解決詳細、ページ設定詳細
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')
gdi32 = ctypes.windll.gdi32
user32 = ctypes.windll.user32
class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]

def gdi_w(font, ppem, ch):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return sz.cx

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False
results = {}

try:
    # === 1. Table cell vertical merge height ===
    print("=== Vertical Merge Height ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72
    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    tbl = wdoc.Tables.Add(rng, 4, 3)
    tbl.Borders.Enable = True
    # Vertical merge R1C1-R3C1
    tbl.Cell(1,1).Merge(tbl.Cell(3,1))
    # Fill
    for ci in range(1, tbl.Range.Cells.Count + 1):
        cell = tbl.Range.Cells(ci)
        cell.Range.Text = f"Cell{ci}"
        cell.Range.Font.Name = "Calibri"; cell.Range.Font.Size = 11
    wdoc.Repaginate()
    merged = tbl.Range.Cells(1)
    print(f"  Merged cell(R1-3,C1): w={round(merged.Width,2)}, h={round(merged.Height,2)}")
    print(f"  RowIndex={merged.RowIndex}, ColIndex={merged.ColumnIndex}")
    results["vmerge_height"] = round(merged.Height, 4)
    wdoc.Close(False)

    # === 2. Table row height: atLeast with padding ===
    print("\n=== Row Height atLeast + padding ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup; ps.LeftMargin = 72
    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    tbl = wdoc.Tables.Add(rng, 3, 2)
    tbl.Borders.Enable = True
    tbl.TopPadding = 5; tbl.BottomPadding = 5
    for r in range(1, 4):
        tbl.Rows(r).Height = 40; tbl.Rows(r).HeightRule = 1  # atLeast
        for c in range(1, 3):
            tbl.Cell(r,c).Range.Text = f"R{r}C{c}"
            tbl.Cell(r,c).Range.Font.Name = "Calibri"; tbl.Cell(r,c).Range.Font.Size = 11
    wdoc.Repaginate()
    for r in range(1, 4):
        y = tbl.Cell(r,1).Range.Paragraphs(1).Range.Information(6)
        print(f"  R{r}: y={round(y,2)}, h={round(tbl.Rows(r).Height,2)}")
        if r > 1:
            prev = tbl.Cell(r-1,1).Range.Paragraphs(1).Range.Information(6)
            print(f"    gap={round(y-prev,2)}")
    wdoc.Close(False)

    # === 3. TextBox text wrapping width ===
    print("\n=== TextBox Text Wrap Width ===")
    wdoc = word.Documents.Add()
    wdoc.Content.Text = "Body."
    tb = wdoc.Shapes.AddTextbox(1, 72, 200, 150, 100, wdoc.Range(0, 0))
    tf = tb.TextFrame
    tf.TextRange.Text = "Short text that might wrap inside this narrow textbox."
    tf.TextRange.Font.Name = "Calibri"; tf.TextRange.Font.Size = 11
    wdoc.Repaginate()
    nlines = tf.TextRange.Paragraphs(1).Range.ComputeStatistics(1)
    print(f"  TB w={tb.Width}, marginL={round(tf.MarginLeft,2)}, marginR={round(tf.MarginRight,2)}")
    print(f"  Content width: {round(tb.Width - tf.MarginLeft - tf.MarginRight, 2)}")
    print(f"  Lines: {nlines}")
    results["tb_wrap_width"] = round(tb.Width - tf.MarginLeft - tf.MarginRight, 4)
    wdoc.Close(False)

    # === 4. TextBox AutoSize (shrink to fit) ===
    print("\n=== TextBox AutoSize ===")
    wdoc = word.Documents.Add()
    wdoc.Content.Text = "Body."
    tb = wdoc.Shapes.AddTextbox(1, 72, 200, 200, 200, wdoc.Range(0, 0))
    tf = tb.TextFrame
    tf.TextRange.Text = "Short"
    tf.TextRange.Font.Name = "Calibri"; tf.TextRange.Font.Size = 11
    print(f"  Before AutoSize: w={tb.Width}, h={tb.Height}")
    try:
        tf.AutoSize = 1  # msoAutoSizeShapeToFitText
        wdoc.Repaginate()
        print(f"  After AutoSize: w={tb.Width}, h={round(tb.Height,2)}")
    except Exception as e:
        print(f"  AutoSize error: {e}")
    results["tb_autosize"] = round(tb.Height, 4)
    wdoc.Close(False)

    # === 5. Font name aliases ===
    print("\n=== Font Aliases ===")
    aliases = [
        ("\uff2d\uff33 \u30b4\u30b7\u30c3\u30af", "MS Gothic"),  # ＭＳ ゴシック
        ("\uff2d\uff33 \u660e\u671d", "MS Mincho"),  # ＭＳ 明朝
        ("\u6e38\u30b4\u30b7\u30c3\u30af", "Yu Gothic"),  # 游ゴシック
    ]
    ppem = 14
    for ja_name, en_name in aliases:
        try:
            ja_w = gdi_w(ja_name, ppem, "A")
            en_w = gdi_w(en_name, ppem, "A")
            match = ja_w == en_w
            print(f"  {en_name}: ja={ja_w}, en={en_w}, match={match}")
        except:
            print(f"  {en_name}: error")

    # === 6. Paragraph alignment + indent interaction ===
    print("\n=== Center/Right + Indent ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup; ps.LeftMargin = 72; ps.RightMargin = 72
    wdoc.Content.Text = ""
    for align, ali_name in [(1, "Center"), (2, "Right")]:
        for li in [0, 36]:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
            r.InsertParagraphAfter()
            p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
            p.Range.Text = f"{ali_name} li={li}"
            p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
            p.Format.Alignment = align
            p.Format.LeftIndent = li
    wdoc.Repaginate()
    # Measure (Information(5) unreliable for center/right, note this)
    for i in range(2, wdoc.Paragraphs.Count + 1):
        p = wdoc.Paragraphs(i)
        x = p.Range.Information(5)
        text = p.Range.Text.strip()[:20]
        print(f"  P{i}: x={round(x,2)}, align={p.Format.Alignment}, li={p.Format.LeftIndent} [{text}]")
    wdoc.Close(False)

    # === 7. Default paragraph spacing in real Japanese documents ===
    print("\n=== Template Paragraph Defaults ===")
    doc = word.Documents.Open(os.path.abspath(TEMPLATE))
    # Add a fresh paragraph
    rng = doc.Range(doc.Content.End - 1, doc.Content.End - 1)
    rng.InsertParagraphAfter()
    p = doc.Paragraphs(doc.Paragraphs.Count)
    p.Range.Text = "Template default"
    p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
    print(f"  SA: {p.Format.SpaceAfter}, SB: {p.Format.SpaceBefore}")
    print(f"  LS: {p.Format.LineSpacing}, Rule: {p.Format.LineSpacingRule}")
    print(f"  LI: {p.Format.LeftIndent}, Align: {p.Format.Alignment}")
    results["template_defaults"] = {
        "sa": round(p.Format.SpaceAfter, 2), "sb": round(p.Format.SpaceBefore, 2),
        "ls": round(p.Format.LineSpacing, 2), "rule": p.Format.LineSpacingRule,
    }
    doc.Close(False)

    # === 8. Character width: Latin-1 accented chars ===
    print("\n=== Latin-1 Accented Chars ===")
    accented = "AaEeOoUu\u00C0\u00E0\u00C9\u00E9\u00D6\u00F6\u00DC\u00FC"
    ppem = 14
    for font in ["Calibri", "Arial"]:
        print(f"  {font} ppem={ppem}:")
        for ch in accented:
            w = gdi_w(font, ppem, ch)
            if ord(ch) >= 0xC0:
                base_ch = chr(ord(ch) & ~0x20 & ~0x10)  # rough base
                base_w = gdi_w(font, ppem, "A")  # just compare to A
                diff = w - base_w
                print(f"    U+{ord(ch):04X}'{ch}': {w}px (vs A={base_w}, diff={diff})")

    # === 9. Multiple paragraph borders (top+bottom) ===
    print("\n=== Para Borders Top+Bottom ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    # P1: no border
    p1 = d.add_paragraph()
    p1.add_run("No border").font.name = "Calibri"
    p1.runs[0].font.size = Pt(11)
    # P2: top+bottom border
    p2 = d.add_paragraph()
    p2.add_run("Top+Bottom border").font.name = "Calibri"
    p2.runs[0].font.size = Pt(11)
    pPr = p2._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    for side in ['top', 'bottom']:
        el = etree.SubElement(pBdr, qn(f'w:{side}'))
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '1'); el.set(qn('w:color'), '000000')
    # P3: no border
    p3 = d.add_paragraph()
    p3.add_run("After border").font.name = "Calibri"
    p3.runs[0].font.size = Pt(11)

    for p in [p1, p2, p3]:
        pPr_sp = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr_sp, qn('w:spacing'))
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')

    tmp = os.path.join(tempfile.gettempdir(), "ra_topbot_bdr.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        y = doc.Paragraphs(i).Range.Information(6)
        print(f"  P{i}: y={round(y,2)}", end="")
        if i > 1:
            prev = doc.Paragraphs(i-1).Range.Information(6)
            print(f", gap={round(y-prev,2)}", end="")
        print()
    doc.Close(False); os.unlink(tmp)

    # === 10. Font size: half-point precision ===
    print("\n=== Half-Point Font Sizes ===")
    for fs_half in [17, 19, 21, 23, 25]:  # 8.5, 9.5, 10.5, 11.5, 12.5pt
        fs = fs_half / 2.0
        ppem = round(fs * 96 / 72)
        print(f"  {fs}pt: ppem={ppem}")

    # === 11. Line height: multiple ls values without grid ===
    print("\n=== LS Values Sweep ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    for line_val in [200, 220, 240, 260, 276, 300, 360, 480]:
        p = d.add_paragraph()
        r = p.add_run(f"ls={line_val}")
        r.font.name = "Calibri"; r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:line'), str(line_val)); sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')

    tmp = os.path.join(tempfile.gettempdir(), "ra_ls_sweep.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        y = doc.Paragraphs(i).Range.Information(6)
        ls = doc.Paragraphs(i).Format.LineSpacing
        if i > 1:
            prev = doc.Paragraphs(i-1).Range.Information(6)
            gap = round(y - prev, 2)
            factor = round(gap / 13.5, 4)  # relative to Calibri 11pt gdi_h
            print(f"  ls={ls}: gap={gap}pt, factor={factor}")
    doc.Close(False); os.unlink(tmp)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch7.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
