"""S467 minimal repro: isolate the title-block w:pBdr bottom-border vertical
contribution. Build a 26pt heading paragraph (single, sa=15) WITH vs WITHOUT a
bottom border (sz=8=1.0pt, space=4pt) followed by a body paragraph; measure the
heading->body gap in Word(COM) and Oxi(GDI dump). (with - without) = the border's
vertical contribution. Compare Word vs Oxi to localize the gen2 title -0.75 drift."""
import os, io, json, subprocess, statistics
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import win32com.client as win32

REPO = r"C:\Users\ryuji\oxi-main"
OUT = os.path.join(REPO, "tools", "golden-test", "repros", "grid_snap")
RENDERER = os.path.join(REPO, "tools", "oxi-gdi-renderer", "target", "release", "oxi-gdi-renderer.exe")
VPOS, PAGE = 6, 3


def add_bottom_border(para, sz=8, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), "4F81BD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build(fn, with_border):
    d = docx.Document()
    sec = d.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.25)
    st = d.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    st.paragraph_format.line_spacing = 1.0
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    st.paragraph_format.space_after = Pt(0)
    # heading paragraph: Calibri 26pt, single, sa=15
    h = d.add_paragraph()
    r = h.add_run("Title Heading Text")
    r.font.name = "Calibri"; r.font.size = Pt(26)
    h.paragraph_format.line_spacing = 1.0
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h.paragraph_format.space_after = Pt(15)
    h.paragraph_format.space_before = Pt(0)
    if with_border:
        add_bottom_border(h, sz=8, space=4)
    for i in range(4):
        d.add_paragraph("Body line %d the quick brown fox" % i)
    p = os.path.join(OUT, fn); d.save(p); return p


def oxi_tops(path):
    dump = r"C:/Users/ryuji/AppData/Local/Temp/s467_pbdr.json"
    subprocess.run([RENDERER, path, r"C:/Users/ryuji/AppData/Local/Temp/s467pbdr", "150",
                    "--dump-layout=" + dump], capture_output=True, text=True)
    d = json.load(io.open(dump, encoding="utf-8"))
    P = {}
    for e in d["pages"][0]["elements"]:
        if e.get("type") != "text" or e.get("para_idx") is None:
            continue
        P[e["para_idx"]] = min(P.get(e["para_idx"], 1e9), e["y"])
    return [P[k] for k in sorted(P)]


def word_tops(word, path):
    doc = word.Documents.Open(path, ReadOnly=True)
    ys = []
    for para in doc.Paragraphs:
        st = doc.Range(para.Range.Start, para.Range.Start)
        if st.Information(PAGE) != 1:
            continue
        if not para.Range.Text.strip():
            continue
        ys.append(round(st.Information(VPOS), 3))
    doc.Close(False)
    return ys


def main():
    pa = build("pbdr_OFF.docx", False)
    pb = build("pbdr_ON.docx", True)
    word = win32.gencache.EnsureDispatch("Word.Application"); word.Visible = False
    wa = word_tops(word, pa); wb = word_tops(word, pb)
    word.Quit()
    oa = oxi_tops(pa); ob = oxi_tops(pb)
    # gap heading(idx0) -> first body(idx1)
    wgap_off = wa[1] - wa[0]; wgap_on = wb[1] - wb[0]
    ogap_off = oa[1] - oa[0]; ogap_on = ob[1] - ob[0]
    print("heading->body gap (Title 26pt single sa=15; border sz=8(1.0pt) space=4)")
    print("           WORD      OXI")
    print("no-border  %7.3f  %7.3f" % (wgap_off, ogap_off))
    print("border     %7.3f  %7.3f" % (wgap_on, ogap_on))
    print("border contrib: WORD %+.3f   OXI %+.3f   (Oxi deficit %+.3f)" % (
        wgap_on - wgap_off, ogap_on - ogap_off, (ogap_on - ogap_off) - (wgap_on - wgap_off)))
    print("ON gap: WORD %.3f vs OXI %.3f -> Oxi deficit %+.3f" % (wgap_on, ogap_on, ogap_on - wgap_on))


if __name__ == "__main__":
    main()
