"""
Ra2: Section break Y handoff (§11) — robust 2-stage measurement.

Stage 1: Author each fixture .docx via python-docx (no Word COM needed).
Stage 2: Open each fixture in Word COM, measure paragraph Ys, save JSON.

This avoids the fragile Documents.Add → InsertBreak → Sections(2) pattern
that intermittently fails with RPC_E_CALL_REJECTED.

Cases:
  S11.1.1  continuous break, single column, S1=N paras, body Calibri 11pt
  S11.1.2  same, body Calibri 14pt
  S11.1.3  same, body MS Mincho 10.5pt (CJK + grid interaction)
  S11.3.1  continuous break + s2 has 2 columns, S1=N paras
  S11.3.2  continuous break + s2 has 3 columns
  S11.3.3  margin change across continuous break

Spec questions answered:
  Q1.1 Section 2 first-line Y = Section 1 last-line Y + line_h_grid?
  Q3.1 Section 2 column-1 first-line Y vs section 1 last-line Y?
  Q3.2 Section 2 column-2 first-line Y — top of column block (= column-1 start)
       or aligned with column-1 last filled position?
  Q3.3 Where do column-2 paragraphs sit when col-1 doesn't fill (Word balances)?
"""
import os
import json
import time

import win32com.client
import pythoncom

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUT_DIR, exist_ok=True)
FIX_DIR = os.path.join(OUT_DIR, "section_break_fixtures")
os.makedirs(FIX_DIR, exist_ok=True)
OUT_JSON = os.path.join(OUT_DIR, "ra2_section_break_y_handoff.json")


def make_section_break_doc(path, *, n_s1, body_font, body_size_pt,
                           s2_columns=1, s2_left_margin_in=1.0,
                           s1_left_margin_in=1.0, s2_n_paragraphs=20):
    """Author a docx with a continuous section break in the middle."""
    doc = Document()
    # Section 1 setup
    sec1 = doc.sections[0]
    sec1.top_margin = Pt(72)
    sec1.bottom_margin = Pt(72)
    sec1.left_margin = Inches(s1_left_margin_in)
    sec1.right_margin = Inches(s1_left_margin_in)
    sec1.header_distance = Pt(36)
    sec1.footer_distance = Pt(36)

    # Section 1 body
    for i in range(n_s1):
        p = doc.add_paragraph(f"S1_B{i+1}")
        for run in p.runs:
            run.font.name = body_font
            run.font.size = Pt(body_size_pt)
            # Set CJK font name too
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), body_font)
            rfonts.set(qn("w:ascii"), body_font)
            rfonts.set(qn("w:hAnsi"), body_font)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Add a continuous section break
    sec2 = doc.add_section(WD_SECTION.CONTINUOUS)
    sec2.top_margin = Pt(72)
    sec2.bottom_margin = Pt(72)
    sec2.left_margin = Inches(s2_left_margin_in)
    sec2.right_margin = Inches(s2_left_margin_in)
    sec2.header_distance = Pt(36)
    sec2.footer_distance = Pt(36)

    # Set column count for section 2 via XML (python-docx doesn't expose this directly)
    sectPr = sec2._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(s2_columns))
    cols.set(qn("w:space"), "720")  # 720 twips = 36pt = 0.5in default

    # Section 2 body
    for i in range(s2_n_paragraphs):
        p = doc.add_paragraph(f"S2_B{i+1}")
        for run in p.runs:
            run.font.name = body_font
            run.font.size = Pt(body_size_pt)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), body_font)
            rfonts.set(qn("w:ascii"), body_font)
            rfonts.set(qn("w:hAnsi"), body_font)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.save(path)


def measure_doc(word, path):
    """Open path in Word, measure all paragraphs, close. Returns dict."""
    wdoc = word.Documents.Open(path)
    try:
        wdoc.Repaginate()
        time.sleep(0.1)
        n_paras = wdoc.Paragraphs.Count
        n_secs = wdoc.Sections.Count
        paras = []
        for i in range(1, n_paras + 1):
            p = wdoc.Paragraphs(i)
            try:
                paras.append({
                    "i": i,
                    "y": round(p.Range.Information(6), 4),
                    "x": round(p.Range.Information(5), 4),
                    "page": p.Range.Information(3),
                    "text": p.Range.Text.strip()[:20],
                })
            except Exception as e:
                paras.append({"i": i, "error": str(e)})
        sec_info = []
        for s in range(1, n_secs + 1):
            sec = wdoc.Sections(s)
            sp = {"i": s}
            try:
                sp["cols"] = sec.PageSetup.TextColumns.Count
            except Exception:
                pass
            try:
                sp["leftMargin"] = round(sec.PageSetup.LeftMargin, 4)
                sp["rightMargin"] = round(sec.PageSetup.RightMargin, 4)
            except Exception:
                pass
            sec_info.append(sp)
        return {
            "path": os.path.basename(path),
            "n_sections": n_secs,
            "n_paragraphs": n_paras,
            "sections": sec_info,
            "paragraphs": paras,
        }
    finally:
        wdoc.Close(False)


def main():
    # Stage 1: author all fixtures
    fixtures = []
    cases = [
        # (name, kwargs)
        ("S11_1_1col_n1_calibri11", dict(n_s1=1, body_font="Calibri", body_size_pt=11, s2_columns=1)),
        ("S11_1_1col_n2_calibri11", dict(n_s1=2, body_font="Calibri", body_size_pt=11, s2_columns=1)),
        ("S11_1_1col_n5_calibri11", dict(n_s1=5, body_font="Calibri", body_size_pt=11, s2_columns=1)),
        ("S11_1_1col_n5_calibri14", dict(n_s1=5, body_font="Calibri", body_size_pt=14, s2_columns=1)),
        ("S11_1_1col_n5_msmincho10p5", dict(n_s1=5, body_font="MS Mincho", body_size_pt=10.5, s2_columns=1)),
        ("S11_3_2col_n1_calibri11", dict(n_s1=1, body_font="Calibri", body_size_pt=11, s2_columns=2, s2_n_paragraphs=40)),
        ("S11_3_2col_n5_calibri11", dict(n_s1=5, body_font="Calibri", body_size_pt=11, s2_columns=2, s2_n_paragraphs=40)),
        ("S11_3_3col_n5_calibri11", dict(n_s1=5, body_font="Calibri", body_size_pt=11, s2_columns=3, s2_n_paragraphs=40)),
        ("S11_3_2col_n5_msmincho10p5", dict(n_s1=5, body_font="MS Mincho", body_size_pt=10.5, s2_columns=2, s2_n_paragraphs=40)),
        ("S11_3_margin_change_n2", dict(n_s1=2, body_font="Calibri", body_size_pt=11, s2_columns=1,
                                         s1_left_margin_in=1.0, s2_left_margin_in=2.0)),
    ]
    for name, kwargs in cases:
        path = os.path.join(FIX_DIR, f"{name}.docx")
        make_section_break_doc(path, **kwargs)
        fixtures.append((name, path, kwargs))

    print(f"Authored {len(fixtures)} fixtures.")

    # Stage 2: open in Word and measure
    word = win32com.client.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    time.sleep(2.0)

    results = []
    try:
        for name, path, kwargs in fixtures:
            print(f"\n=== {name} ===")
            try:
                m = measure_doc(word, path)
            except Exception as e:
                print(f"  ERR: {e}")
                continue
            m["case_name"] = name
            m["case_kwargs"] = kwargs
            results.append(m)
            print(f"  sections={m['n_sections']} paras={m['n_paragraphs']}")
            for s in m["sections"]:
                print(f"    sec{s['i']}: cols={s.get('cols')} lm={s.get('leftMargin')}")
            for p in m["paragraphs"][:18]:
                print(f"    P{p['i']:3}@p{p.get('page')} y={p.get('y'):>7} x={p.get('x'):>6} '{p.get('text', '')}'")
            if len(m["paragraphs"]) > 18:
                print(f"    ... ({len(m['paragraphs']) - 18} more)")
    finally:
        # Save data first, then attempt to quit Word (may fail if Word died).
        with open(OUT_JSON, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"\nSaved {len(results)} records to {OUT_JSON}")
        try:
            word.Quit()
        except Exception as e:
            print(f"  (word.Quit failed, ignoring: {e})")


if __name__ == "__main__":
    main()
