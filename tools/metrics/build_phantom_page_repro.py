"""
Minimal repros for phantom-page suppression spec.

Target hypothesis: when consecutive empty paragraphs cross a page boundary,
Word keeps them on the current page (even if cursor_y approaches the bottom
margin) rather than pushing them alone onto a fresh page.

Repros:
  RPH_1: content + 1 empty at bottom + content (control — no phantom risk)
  RPH_2: content + 3 empty at bottom + content (phantom candidate)
  RPH_3: content + 2 empty + form-feed para + content (matches d77a p.10 end)
  RPH_4: content JUST filling page + 1 empty + content (empty must push)
  RPH_5: content + 3 empty + form-feed (d77a-like) starting near overflow

All 10.5pt MS Gothic, A4, 2.5cm margins, linePitch=360tw (Word default).
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent / "phantom_page_repro"
OUT_DIR.mkdir(exist_ok=True)


def set_run_gothic(run, pt: float = 10.5):
    rPr = run._r.get_or_add_rPr()
    for e in rPr.findall(qn('w:rFonts')):
        rPr.remove(e)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'MS Gothic')
    rf.set(qn('w:eastAsia'), 'MS Gothic')
    rf.set(qn('w:hAnsi'), 'MS Gothic')
    rPr.append(rf)
    run.font.size = Pt(pt)


def set_section_docgrid(section, line_pitch_tw=360):
    sectPr = section._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)
    dg = OxmlElement('w:docGrid')
    dg.set(qn('w:type'), 'linesAndChars')
    dg.set(qn('w:linePitch'), str(line_pitch_tw))
    dg.set(qn('w:charSpace'), '0')
    sectPr.append(dg)


def base_doc(filler_n: int = 0):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    set_section_docgrid(section, line_pitch_tw=360)
    for i in range(filler_n):
        p = doc.add_paragraph()
        run = p.add_run(f"Filler {i+1:02d} 本文用のフィラーです。")
        set_run_gothic(run, pt=10.5)
    return doc


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_gothic(run, pt=10.5)
    return p


def add_empty(doc):
    return doc.add_paragraph()


def add_page_break_para(doc):
    """Empty paragraph with <w:br w:type='page'/>."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    set_run_gothic(run, pt=10.5)
    return p


# ------------------------------------------------------------
# RPH_1: 1 empty at bottom, then content. Baseline.
def make_rph1():
    doc = base_doc(filler_n=35)
    add_body(doc, "RPH1_BEFORE_EMPTY")
    add_empty(doc)
    add_body(doc, "RPH1_AFTER_EMPTY")
    doc.save(OUT_DIR / "RPH_1_one_empty.docx")


# RPH_2: 3 empties at bottom, then content.
def make_rph2():
    doc = base_doc(filler_n=35)
    add_body(doc, "RPH2_BEFORE_EMPTIES")
    for _ in range(3):
        add_empty(doc)
    add_body(doc, "RPH2_AFTER_EMPTIES")
    doc.save(OUT_DIR / "RPH_2_three_empties.docx")


# RPH_3: 2 empties + form-feed para + content. d77a-like structure.
def make_rph3():
    doc = base_doc(filler_n=35)
    add_body(doc, "RPH3_BEFORE_EMPTIES")
    for _ in range(2):
        add_empty(doc)
    add_page_break_para(doc)
    add_body(doc, "RPH3_AFTER_PAGEBREAK")
    doc.save(OUT_DIR / "RPH_3_empties_plus_pagebreak.docx")


# RPH_4: content JUST filling page + 1 empty. Empty will push to next page.
def make_rph4():
    doc = base_doc(filler_n=37)  # more fillers
    add_empty(doc)
    add_body(doc, "RPH4_AFTER_EMPTY")
    doc.save(OUT_DIR / "RPH_4_single_empty_overflow.docx")


# RPH_5: heavier flow: many fillers + 3 empties + form-feed + content.
def make_rph5():
    doc = base_doc(filler_n=36)
    for _ in range(3):
        add_empty(doc)
    add_page_break_para(doc)
    add_body(doc, "RPH5_AFTER_PAGEBREAK")
    doc.save(OUT_DIR / "RPH_5_three_empties_plus_pagebreak.docx")


# RPH_6: 5 consecutive empties in middle of page (no overflow). Control.
def make_rph6():
    doc = base_doc(filler_n=5)
    add_body(doc, "RPH6_BEFORE")
    for _ in range(5):
        add_empty(doc)
    add_body(doc, "RPH6_AFTER")
    doc.save(OUT_DIR / "RPH_6_middle_empties.docx")


if __name__ == "__main__":
    make_rph1()
    make_rph2()
    make_rph3()
    make_rph4()
    make_rph5()
    make_rph6()
    print("Built:")
    for f in sorted(OUT_DIR.glob("*.docx")):
        print(" ", f.name)
