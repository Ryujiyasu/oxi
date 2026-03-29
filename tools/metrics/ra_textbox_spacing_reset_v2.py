"""
Ra: TextBox/TableCell spacing リセット — docDefaults経由の継承値 vs 明示値
- docDefaultsでsa=8pt設定 → テーブルセル/TextBoxで保持 or リセット?
- Normalスタイルで明示sa=8pt → 同上
- 段落XMLで明示sa=8pt → 同上
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def make_doc_docdefaults_sa():
    """Set spaceAfter=8pt via docDefaults only, NOT in Normal style or paragraph."""
    d = Document(TEMPLATE)

    # Remove all default paragraphs
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Set docDefaults pPrDefault to have spaceAfter=160twips(8pt)
    styles_el = d.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = etree.SubElement(styles_el, qn('w:docDefaults'))
    pPrDefault = docDefaults.find(qn('w:pPrDefault'))
    if pPrDefault is None:
        pPrDefault = etree.SubElement(docDefaults, qn('w:pPrDefault'))
    pPr = pPrDefault.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(pPrDefault, qn('w:pPr'))
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:after'), '160')  # 8pt = 160 twips
    spacing.set(qn('w:line'), '276')   # 1.15x = 276/240
    spacing.set(qn('w:lineRule'), 'auto')

    # Clear Normal style spacing to ensure only docDefaults provides values
    normal_style = d.styles['Normal']
    style_el = normal_style.element
    pPr_style = style_el.find(qn('w:pPr'))
    if pPr_style is not None:
        sp = pPr_style.find(qn('w:spacing'))
        if sp is not None:
            pPr_style.remove(sp)

    # Add 3 body paragraphs (no explicit spacing in XML)
    for i in range(3):
        p = d.add_paragraph()
        r = p.add_run(f"Body {i+1} - spacing from docDefaults only")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        # Remove any spacing from paragraph XML
        pPr_p = p._element.find(qn('w:pPr'))
        if pPr_p is not None:
            sp = pPr_p.find(qn('w:spacing'))
            if sp is not None:
                pPr_p.remove(sp)

    return d


def make_doc_normalstyle_sa():
    """Set spaceAfter=8pt via Normal style, NOT in docDefaults or paragraph."""
    d = Document(TEMPLATE)

    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Clear docDefaults spacing
    styles_el = d.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is not None:
        pPrDefault = docDefaults.find(qn('w:pPrDefault'))
        if pPrDefault is not None:
            pPr = pPrDefault.find(qn('w:pPr'))
            if pPr is not None:
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    pPr.remove(sp)

    # Set Normal style spacing
    normal_style = d.styles['Normal']
    pf = normal_style.paragraph_format
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)

    # Add body paragraphs (no explicit spacing)
    for i in range(3):
        p = d.add_paragraph()
        r = p.add_run(f"Body {i+1} - spacing from Normal style")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pPr_p = p._element.find(qn('w:pPr'))
        if pPr_p is not None:
            sp = pPr_p.find(qn('w:spacing'))
            if sp is not None:
                pPr_p.remove(sp)

    return d


def measure_all(doc_path, label):
    """Open in Word, add TextBox + Table via COM, measure spacing properties."""
    doc = word.Documents.Open(doc_path)
    try:
        data = {"label": label, "body": [], "textbox": [], "table_cell": []}

        # Measure body
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            data["body"].append({
                "index": i,
                "y_pt": round(para.Range.Information(6), 4),
                "line_spacing": round(para.Format.LineSpacing, 4),
                "ls_rule": para.Format.LineSpacingRule,
                "sb": round(para.Format.SpaceBefore, 4),
                "sa": round(para.Format.SpaceAfter, 4),
            })

        # Add TextBox
        tb = doc.Shapes.AddTextbox(1, 100, 400, 200, 100, doc.Range(0, 0))
        tf = tb.TextFrame
        tf.TextRange.Text = "TB1\rTB2\rTB3"
        for i in range(1, tf.TextRange.Paragraphs.Count + 1):
            para = tf.TextRange.Paragraphs(i)
            para.Range.Font.Name = "Calibri"
            para.Range.Font.Size = 11
        # Don't set any spacing - let it inherit
        for i in range(1, tf.TextRange.Paragraphs.Count + 1):
            para = tf.TextRange.Paragraphs(i)
            data["textbox"].append({
                "index": i,
                "line_spacing": round(para.Format.LineSpacing, 4),
                "ls_rule": para.Format.LineSpacingRule,
                "sb": round(para.Format.SpaceBefore, 4),
                "sa": round(para.Format.SpaceAfter, 4),
            })

        # Add Table (1x1)
        rng_end = doc.Range(doc.Content.End - 1, doc.Content.End - 1)
        tbl = doc.Tables.Add(rng_end, 1, 1)
        cell = tbl.Cell(1, 1)
        cell.Range.Text = "TC1\rTC2\rTC3"
        for i in range(1, cell.Range.Paragraphs.Count + 1):
            para = cell.Range.Paragraphs(i)
            para.Range.Font.Name = "Calibri"
            para.Range.Font.Size = 11
        for i in range(1, cell.Range.Paragraphs.Count + 1):
            para = cell.Range.Paragraphs(i)
            data["table_cell"].append({
                "index": i,
                "line_spacing": round(para.Format.LineSpacing, 4),
                "ls_rule": para.Format.LineSpacingRule,
                "sb": round(para.Format.SpaceBefore, 4),
                "sa": round(para.Format.SpaceAfter, 4),
            })

        return data
    finally:
        doc.Close(False)


try:
    # Test 1: docDefaults-only spacing
    d1 = make_doc_docdefaults_sa()
    tmp1 = os.path.join(tempfile.gettempdir(), "ra_sp_docdefaults.docx")
    d1.save(tmp1)
    data1 = measure_all(tmp1, "docDefaults_sa=8pt")
    results.append(data1)
    os.unlink(tmp1)

    # Test 2: Normal style spacing
    d2 = make_doc_normalstyle_sa()
    tmp2 = os.path.join(tempfile.gettempdir(), "ra_sp_normalstyle.docx")
    d2.save(tmp2)
    data2 = measure_all(tmp2, "NormalStyle_sa=8pt")
    results.append(data2)
    os.unlink(tmp2)

finally:
    word.Quit()

# Save
out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_textbox_spacing_reset_v2.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)

# Print analysis
for data in results:
    print(f"\n=== {data['label']} ===")
    for ctx_name in ["body", "textbox", "table_cell"]:
        items = data[ctx_name]
        if items:
            p1 = items[0]
            print(f"  {ctx_name}: ls={p1['line_spacing']}(rule={p1['ls_rule']}), sb={p1['sb']}, sa={p1['sa']}")

print("\n=== COMPARISON ===")
for data in results:
    body_sa = data["body"][0]["sa"]
    tb_sa = data["textbox"][0]["sa"] if data["textbox"] else "N/A"
    tc_sa = data["table_cell"][0]["sa"] if data["table_cell"] else "N/A"
    body_ls = data["body"][0]["line_spacing"]
    tb_ls = data["textbox"][0]["line_spacing"] if data["textbox"] else "N/A"
    tc_ls = data["table_cell"][0]["line_spacing"] if data["table_cell"] else "N/A"
    print(f"{data['label']}:")
    print(f"  sa: body={body_sa}, textbox={tb_sa}, table={tc_sa}")
    print(f"  ls: body={body_ls}, textbox={tb_ls}, table={tc_ls}")
    if tb_sa != body_sa:
        print(f"  => TextBox RESETS sa (body={body_sa} -> textbox={tb_sa})")
    if tc_sa != body_sa:
        print(f"  => TableCell RESETS sa (body={body_sa} -> table={tc_sa})")
