"""
Build tab leader test fixtures via python-docx.

Tests:
  - Leader types: none, dot, hyphen, underscore, middleDot
  - Tab alignments: left, center, right, decimal
  - Mixed: leading text + tab + trailing text

Each fixture has 5 paragraphs each with one leader type, all using a single
right-aligned tab at 432pt (typical TOC layout: "Chapter Name......Page #").

Saved to output/tab_leader_fixtures/
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output", "tab_leader_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


def add_tab_to_pPr(p, *, pos_tw, val="right", leader=None):
    """Add a single <w:tab/> to paragraph properties."""
    pPr = p._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), val)
    tab.set(qn("w:pos"), str(pos_tw))
    if leader:
        tab.set(qn("w:leader"), leader)
    tabs.append(tab)


def build_fixture(path, *, leader_types, tab_pos_pt=288, tab_val="right",
                   font="Calibri", size_pt=11):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Pt(36)
    sec.footer_distance = Pt(36)

    pos_tw = int(tab_pos_pt * 20)

    for leader in leader_types:
        # Each para: "Chapter X" + TAB + "Page Y"
        # The tab goes between two runs.
        p = doc.add_paragraph()
        add_tab_to_pPr(p, pos_tw=pos_tw, val=tab_val, leader=leader)

        leader_label = leader if leader else "none"
        run1 = p.add_run(f"Item leader={leader_label}")
        run1.font.name = font
        run1.font.size = Pt(size_pt)
        # Insert a tab character
        tab_run = p.add_run()
        tab_elem = OxmlElement("w:tab")
        tab_run._r.append(tab_elem)
        run3 = p.add_run("99")
        run3.font.name = font
        run3.font.size = Pt(size_pt)

        for r in (run1, run3):
            rpr = r._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), font)
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.save(path)


def main():
    leader_types = [None, "dot", "hyphen", "underscore", "middleDot"]
    cases = [
        ("TL_right_tab.docx", {"tab_val": "right", "tab_pos_pt": 432}),
        ("TL_left_tab.docx", {"tab_val": "left", "tab_pos_pt": 288}),
        ("TL_center_tab.docx", {"tab_val": "center", "tab_pos_pt": 288}),
        ("TL_decimal_tab.docx", {"tab_val": "decimal", "tab_pos_pt": 288}),
    ]
    for name, kwargs in cases:
        path = os.path.join(OUT_DIR, name)
        build_fixture(path, leader_types=leader_types, **kwargs)
        print(f"  built {name}")
    print(f"\nBuilt {len(cases)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
