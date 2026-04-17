"""Test: does <w:kern w:val='2'/> on run rPr trigger ・ compression?

d77a's Normal style has kerning=2 (threshold 2pt). All rFonts inherit this.
"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data")
)
TEST_TEXT = "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"


def make(label, with_kern):
    doc = Document()
    settings = doc.settings.element
    for e in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(e)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)

    p = doc.add_paragraph()
    r = p.add_run(TEST_TEXT)
    rel = r._element
    rPr = rel.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); rel.insert(0, rPr)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "ＭＳ ゴシック")
    rFonts.set(qn("w:eastAsia"), "ＭＳ ゴシック")
    rFonts.set(qn("w:hAnsi"), "ＭＳ ゴシック")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)
    if with_kern:
        kern = OxmlElement("w:kern"); kern.set(qn("w:val"), "2")
        rPr.append(kern)

    out = os.path.join(OUT_DIR, f"yakumono_kern_{label}.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    print(make("K1_nokern", False))
    print(make("K2_kern2", True))
    print(make("K3_kern20", True))  # will set val=20 via edit
    # Modify K3 to have val=20
    import zipfile, shutil, re
    p = os.path.join(OUT_DIR, "yakumono_kern_K3_kern20.docx")
    tmp = p + ".tmp"
    with zipfile.ZipFile(p, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == "word/document.xml":
                    xml = data.decode("utf-8")
                    xml = xml.replace('<w:kern w:val="2"/>', '<w:kern w:val="20"/>')
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
    os.replace(tmp, p)
