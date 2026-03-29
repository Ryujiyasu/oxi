"""
Ra: atLeast lineSpacingRule の詳細挙動をCOM計測
- natural height > specified value のケース
- natural height < specified value のケース
- CJK 83/64 乗数との相互作用
- grid snap との相互作用
"""
import win32com.client, json, os, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def make_doc(font_name, font_size, atleast_twips, num_lines=5, grid=False, grid_pitch=360):
    d = Document(TEMPLATE)
    sec = d.sections[0]
    sectPr = sec._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)
    if grid:
        dg = etree.SubElement(sectPr, qn('w:docGrid'))
        dg.set(qn('w:type'), 'lines')
        dg.set(qn('w:linePitch'), str(grid_pitch))

    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    for i in range(num_lines):
        p = d.add_paragraph()
        r = p.add_run(f"Line {i+1} atLeast test text ABCDEFG")
        r.font.name = font_name
        r.font.size = Pt(font_size)
        pPr = p._element.get_or_add_pPr()
        spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:line'), str(atleast_twips))
        spacing.set(qn('w:lineRule'), 'atLeast')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')

    return d


def measure(doc_path, label):
    doc = word.Documents.Open(doc_path)
    try:
        data = {"label": label, "paragraphs": []}
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            rng = para.Range
            data["paragraphs"].append({
                "index": i,
                "y_pt": round(rng.Information(6), 4),
                "line_spacing": round(para.Format.LineSpacing, 4),
                "ls_rule": para.Format.LineSpacingRule,
            })
        for i in range(1, len(data["paragraphs"])):
            data["paragraphs"][i]["gap"] = round(
                data["paragraphs"][i]["y_pt"] - data["paragraphs"][i-1]["y_pt"], 4)
        return data
    finally:
        doc.Close(False)


try:
    # Calibri 11pt natural line height ~ 13.8pt
    # MS Gothic 10.5pt natural ~ 14.6pt (with 83/64)
    tests = [
        # (font, size, atLeast_twips, description, grid)
        ("Calibri", 11, 200, "Calibri 11pt, atLeast=10pt (natural>spec)", False),
        ("Calibri", 11, 280, "Calibri 11pt, atLeast=14pt (natural~spec)", False),
        ("Calibri", 11, 400, "Calibri 11pt, atLeast=20pt (natural<spec)", False),
        ("Calibri", 11, 240, "Calibri 11pt, atLeast=12pt (natural>spec)", False),

        ("MS Gothic", 10.5, 200, "MSGothic 10.5pt, atLeast=10pt (natural>spec, CJK)", False),
        ("MS Gothic", 10.5, 300, "MSGothic 10.5pt, atLeast=15pt (natural~spec, CJK)", False),
        ("MS Gothic", 10.5, 400, "MSGothic 10.5pt, atLeast=20pt (natural<spec, CJK)", False),

        # With grid
        ("Calibri", 11, 200, "Calibri 11pt, atLeast=10pt + grid(18pt)", True),
        ("Calibri", 11, 400, "Calibri 11pt, atLeast=20pt + grid(18pt)", True),
        ("MS Gothic", 10.5, 200, "MSGothic 10.5pt, atLeast=10pt + grid(18pt)", True),
    ]

    for font, fs, tw, desc, grid in tests:
        d = make_doc(font, fs, tw, grid=grid)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_al_{font.replace(' ','_')}_{fs}_{tw}{'_g' if grid else ''}.docx")
        d.save(tmp)
        data = measure(tmp, desc)
        results.append(data)
        os.unlink(tmp)

        gap = data["paragraphs"][1].get("gap", 0) if len(data["paragraphs"]) > 1 else 0
        reported_ls = data["paragraphs"][0]["line_spacing"]
        specified = tw / 20.0
        print(f"  {desc}:")
        print(f"    specified={specified}pt, gap={gap}pt, reported_ls={reported_ls}pt")
        if gap > specified + 0.1:
            print(f"    → natural({gap}) > specified({specified}) → uses natural")
        elif abs(gap - specified) < 0.5:
            print(f"    → uses specified({specified})")
        else:
            print(f"    → uses grid-snapped value")

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_atleast_linespacing.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
