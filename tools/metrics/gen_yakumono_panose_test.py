"""Test: does real PANOSE on fontTable MS Gothic trigger ・ compression?

d77a's fontTable has panose1="020B0609070205080204". My previous repros
have panose1="00000000000000000000" and <w:notTrueType/>.
"""
import os, shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import tempfile

OUT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data")
)

TEST_TEXT = "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"


def make_base_doc(path):
    doc = Document()
    # compressPunctuation
    settings = doc.settings.element
    for e in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(e)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)

    p = doc.add_paragraph()
    r = p.add_run(TEST_TEXT)
    # MS Gothic 12pt
    rel = r._element
    rPr = rel.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); rel.insert(0, rPr)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "ＭＳ ゴシック")
    rFonts.set(qn("w:eastAsia"), "ＭＳ ゴシック")
    rFonts.set(qn("w:hAnsi"), "ＭＳ ゴシック")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)

    doc.save(path)


def rewrite_fonttable(docx_path, new_font_xml):
    """Replace MS Gothic font entry in fontTable.xml with new XML."""
    # Create temp copy
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == "word/fontTable.xml":
                    xml = data.decode("utf-8")
                    # Replace MS Gothic font block
                    import re
                    # Find the existing MS Gothic font (python-docx emits notTrueType)
                    pattern = r'<w:font w:name="ＭＳ ゴシック">.*?</w:font>'
                    new_full = f'<w:font w:name="ＭＳ ゴシック">{new_font_xml}</w:font>'
                    xml_new = re.sub(pattern, new_full, xml, count=1, flags=re.DOTALL)
                    data = xml_new.encode("utf-8")
                zout.writestr(item, data)
    os.replace(tmp, docx_path)


def main():
    # Variant P1: d77a-like PANOSE with altName
    P1 = os.path.join(OUT_DIR, "yakumono_panose_P1.docx")
    make_base_doc(P1)
    d77a_msgothic = (
        '<w:altName w:val="MS Gothic"/>'
        '<w:panose1 w:val="020B0609070205080204"/>'
        '<w:charset w:val="80"/>'
        '<w:family w:val="modern"/>'
        '<w:pitch w:val="fixed"/>'
        '<w:sig w:usb0="E00002FF" w:usb1="6AC7FDFB" w:usb2="08000012" w:usb3="00000000" w:csb0="0002009F" w:csb1="00000000"/>'
    )
    rewrite_fonttable(P1, d77a_msgothic)
    print(f"[P1] {P1} (d77a-like fontTable)")

    # Variant P2: baseline (no rewrite)
    P2 = os.path.join(OUT_DIR, "yakumono_panose_P2.docx")
    make_base_doc(P2)
    print(f"[P2] {P2} (python-docx default fontTable)")

    # Verify
    for p in [P1, P2]:
        with zipfile.ZipFile(p) as zf:
            ft = zf.read("word/fontTable.xml").decode("utf-8")
        import re
        m = re.search(r'<w:font w:name="ＭＳ ゴシック">.*?</w:font>', ft, re.DOTALL)
        if m:
            print(f"  {os.path.basename(p)} MS Gothic:")
            print(f"    {m.group(0)[:200]}")


if __name__ == "__main__":
    main()
