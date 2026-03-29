"""
Ra: contextualSpacing の挙動をCOM計測で確定
- 同一スタイル隣接段落で sa/sb が 0 になるか
- 異なるスタイルでは効果なしか
- contextualSpacing=True と False の比較
- spaceBefore/spaceAfter の具体的な抑制ルール
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def make_doc(scenario):
    """Create test doc for contextualSpacing scenarios."""
    d = Document(TEMPLATE)
    sec = d.sections[0]

    # Remove grid to isolate spacing behavior
    sectPr = sec._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)

    # Remove default paragraphs
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    if scenario == "same_style_ctx_on":
        # 3 paragraphs, same style (Normal), contextualSpacing=True, sa=12pt, sb=12pt
        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"Paragraph {i+1} - Normal style contextualSpacing=True")
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            # Set contextualSpacing
            pPr = p._element.get_or_add_pPr()
            cs = etree.SubElement(pPr, qn('w:contextualSpacing'))

    elif scenario == "same_style_ctx_off":
        # 3 paragraphs, same style (Normal), contextualSpacing=False (default), sa=12pt, sb=12pt
        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"Paragraph {i+1} - Normal style contextualSpacing=False")
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)

    elif scenario == "diff_style_ctx_on":
        # 3 paragraphs, alternating styles (Normal, Heading1, Normal), contextualSpacing=True
        styles_fonts = [("Calibri", Pt(11)), ("Arial", Pt(16)), ("Calibri", Pt(11))]
        style_names = ["Normal", "Heading1", "Normal"]
        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"Paragraph {i+1} - {style_names[i]} contextualSpacing=True")
            r.font.name = styles_fonts[i][0]
            r.font.size = styles_fonts[i][1]
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            pPr = p._element.get_or_add_pPr()
            cs = etree.SubElement(pPr, qn('w:contextualSpacing'))
            # Set different pStyle for middle paragraph
            if i == 1:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is None:
                    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
                pStyle.set(qn('w:val'), 'Heading1')

    elif scenario == "same_style_asymmetric":
        # Same style, contextualSpacing=True, different sa/sb
        # P1: sa=0, sb=20pt, ctx=True
        # P2: sa=10pt, sb=0, ctx=True
        # P3: sa=6pt, sb=6pt, ctx=True
        configs = [(0, 20), (10, 0), (6, 6)]
        for i, (sb_val, sa_val) in enumerate(configs):
            p = d.add_paragraph()
            r = p.add_run(f"Paragraph {i+1} - sb={sb_val} sa={sa_val}")
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            pf = p.paragraph_format
            pf.space_before = Pt(sb_val)
            pf.space_after = Pt(sa_val)
            pPr = p._element.get_or_add_pPr()
            cs = etree.SubElement(pPr, qn('w:contextualSpacing'))

    elif scenario == "ctx_one_side_only":
        # P1 has contextualSpacing, P2 does not. Same style.
        for i in range(2):
            p = d.add_paragraph()
            r = p.add_run(f"Paragraph {i+1}")
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            if i == 0:
                pPr = p._element.get_or_add_pPr()
                cs = etree.SubElement(pPr, qn('w:contextualSpacing'))

    elif scenario == "ctx_both_sides":
        # Both P1 and P2 have contextualSpacing. Same style.
        for i in range(2):
            p = d.add_paragraph()
            r = p.add_run(f"Paragraph {i+1}")
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            pPr = p._element.get_or_add_pPr()
            cs = etree.SubElement(pPr, qn('w:contextualSpacing'))

    return d


def measure(doc_path, scenario):
    """Open doc in Word and measure Y positions of all paragraphs."""
    doc = word.Documents.Open(doc_path)
    try:
        data = {"scenario": scenario, "paragraphs": []}
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            rng = para.Range
            y = rng.Information(6)  # wdVerticalPositionRelativeToPage
            x = rng.Information(5)  # wdHorizontalPositionRelativeToPage
            text = rng.Text.strip()[:40]
            data["paragraphs"].append({
                "index": i,
                "y_pt": round(y, 4),
                "x_pt": round(x, 4),
                "text": text
            })
        # Compute gaps
        for i in range(1, len(data["paragraphs"])):
            gap = data["paragraphs"][i]["y_pt"] - data["paragraphs"][i-1]["y_pt"]
            data["paragraphs"][i]["gap_from_prev"] = round(gap, 4)
        return data
    finally:
        doc.Close(False)


scenarios = [
    "same_style_ctx_on",
    "same_style_ctx_off",
    "diff_style_ctx_on",
    "same_style_asymmetric",
    "ctx_one_side_only",
    "ctx_both_sides",
]

try:
    for sc in scenarios:
        d = make_doc(sc)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_ctx_{sc}.docx")
        d.save(tmp)
        data = measure(tmp, sc)
        results.append(data)
        print(f"\n=== {sc} ===")
        for p in data["paragraphs"]:
            gap_str = f"  gap={p.get('gap_from_prev', '-')}" if 'gap_from_prev' in p else ""
            print(f"  P{p['index']}: y={p['y_pt']}pt{gap_str}  [{p['text']}]")
        os.unlink(tmp)
finally:
    word.Quit()

# Save results
out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_contextual_spacing.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)

print(f"\nResults saved to {out_path}")

# Analysis
print("\n=== ANALYSIS ===")
ctx_on = next(r for r in results if r["scenario"] == "same_style_ctx_on")
ctx_off = next(r for r in results if r["scenario"] == "same_style_ctx_off")
print(f"Same style, ctx ON:  P1→P2 gap = {ctx_on['paragraphs'][1].get('gap_from_prev', 'N/A')}")
print(f"Same style, ctx OFF: P1→P2 gap = {ctx_off['paragraphs'][1].get('gap_from_prev', 'N/A')}")

diff = next(r for r in results if r["scenario"] == "diff_style_ctx_on")
print(f"Diff style, ctx ON:  P1→P2 gap = {diff['paragraphs'][1].get('gap_from_prev', 'N/A')}")

one_side = next(r for r in results if r["scenario"] == "ctx_one_side_only")
both_sides = next(r for r in results if r["scenario"] == "ctx_both_sides")
print(f"One-side ctx:  P1→P2 gap = {one_side['paragraphs'][1].get('gap_from_prev', 'N/A')}")
print(f"Both-sides ctx: P1→P2 gap = {both_sides['paragraphs'][1].get('gap_from_prev', 'N/A')}")
