"""
Ra: スタイル継承チェーンの完全な優先順位をCOM確定
- docDefaults → Normal → basedOn chain → paragraph direct
- font/size/indent/spacing 各プロパティの継承
- 未指定プロパティの解決ルール
"""
import win32com.client, json, os, tempfile
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def test_font_inheritance():
    """Font family inheritance chain."""
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Set docDefaults rPrDefault font to Arial
    styles_el = d.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = etree.SubElement(styles_el, qn('w:docDefaults'))
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = etree.SubElement(docDefaults, qn('w:rPrDefault'))
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(rPrDefault, qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')

    # Set Normal style font to Times New Roman
    normal = d.styles['Normal']
    normal_rPr = normal.element.find(qn('w:rPr'))
    if normal_rPr is None:
        normal_rPr = etree.SubElement(normal.element, qn('w:rPr'))
    norm_fonts = normal_rPr.find(qn('w:rFonts'))
    if norm_fonts is None:
        norm_fonts = etree.SubElement(normal_rPr, qn('w:rFonts'))
    norm_fonts.set(qn('w:ascii'), 'Times New Roman')
    norm_fonts.set(qn('w:hAnsi'), 'Times New Roman')

    # P1: Normal style, no direct font → should use TNR (Normal overrides docDefaults)
    p1 = d.add_paragraph()
    p1.add_run("P1: Normal style, no direct font")

    # P2: Normal style, direct font Calibri
    p2 = d.add_paragraph()
    r2 = p2.add_run("P2: direct Calibri")
    r2.font.name = "Calibri"

    # P3: No explicit formatting at all
    p3 = d.add_paragraph()
    p3.add_run("P3: bare run")

    return d


def test_size_inheritance():
    """Font size inheritance chain."""
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # docDefaults size = 22 halfpoints = 11pt
    styles_el = d.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = etree.SubElement(styles_el, qn('w:docDefaults'))
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = etree.SubElement(docDefaults, qn('w:rPrDefault'))
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(rPrDefault, qn('w:rPr'))
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '22')  # 11pt

    # Normal style size = 24 halfpoints = 12pt
    normal = d.styles['Normal']
    normal_rPr = normal.element.find(qn('w:rPr'))
    if normal_rPr is None:
        normal_rPr = etree.SubElement(normal.element, qn('w:rPr'))
    norm_sz = normal_rPr.find(qn('w:sz'))
    if norm_sz is None:
        norm_sz = etree.SubElement(normal_rPr, qn('w:sz'))
    norm_sz.set(qn('w:val'), '24')  # 12pt

    # P1: Normal, no direct size → 12pt (Normal)
    p1 = d.add_paragraph()
    p1.add_run("P1: Normal 12pt")

    # P2: direct size 9pt
    p2 = d.add_paragraph()
    r2 = p2.add_run("P2: direct 9pt")
    r2.font.size = Pt(9)

    # P3: clear Normal size, should fall back to docDefaults 11pt
    # (Can't easily clear Normal's size from python-docx, so skip)

    return d


def test_spacing_inheritance():
    """Spacing (sa/sb/ls) inheritance."""
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # docDefaults: sa=8pt
    styles_el = d.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = etree.SubElement(styles_el, qn('w:docDefaults'))
    pPrDefault = docDefaults.find(qn('w:pPrDefault'))
    if pPrDefault is None:
        pPrDefault = etree.SubElement(docDefaults, qn('w:pPrDefault'))
    pPr_def = pPrDefault.find(qn('w:pPr'))
    if pPr_def is None:
        pPr_def = etree.SubElement(pPrDefault, qn('w:pPr'))
    sp_def = etree.SubElement(pPr_def, qn('w:spacing'))
    sp_def.set(qn('w:after'), '160')  # 8pt

    # Normal style: sa=10pt
    normal = d.styles['Normal']
    pf = normal.paragraph_format
    pf.space_after = Pt(10)

    # P1: Normal, no direct → sa=10pt (Normal overrides docDefaults)
    p1 = d.add_paragraph()
    p1.add_run("P1: Normal sa=10pt").font.name = "Calibri"
    p1.runs[0].font.size = Pt(11)

    # P2: direct sa=4pt
    p2 = d.add_paragraph()
    p2.add_run("P2: direct sa=4pt").font.name = "Calibri"
    p2.runs[0].font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(4)

    # P3: direct sa=0pt
    p3 = d.add_paragraph()
    p3.add_run("P3: direct sa=0pt").font.name = "Calibri"
    p3.runs[0].font.size = Pt(11)
    p3.paragraph_format.space_after = Pt(0)

    return d


def measure_styles(doc_path, label):
    doc = word.Documents.Open(doc_path)
    try:
        data = {"label": label, "paragraphs": []}
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            rng = para.Range
            data["paragraphs"].append({
                "index": i,
                "font_name": rng.Font.Name,
                "font_size": round(rng.Font.Size, 4),
                "space_after": round(para.Format.SpaceAfter, 4),
                "space_before": round(para.Format.SpaceBefore, 4),
                "line_spacing": round(para.Format.LineSpacing, 4),
                "left_indent": round(para.Format.LeftIndent, 4),
                "text": rng.Text.strip()[:40],
            })
        return data
    finally:
        doc.Close(False)


try:
    # Font inheritance
    d1 = test_font_inheritance()
    tmp = os.path.join(tempfile.gettempdir(), "ra_style_font.docx")
    d1.save(tmp)
    data1 = measure_styles(tmp, "font_inheritance")
    results.append(data1)
    os.unlink(tmp)
    print("=== font_inheritance (docDefaults=Arial, Normal=TNR) ===")
    for p in data1["paragraphs"]:
        print(f"  P{p['index']}: font={p['font_name']}, size={p['font_size']}  [{p['text']}]")

    # Size inheritance
    d2 = test_size_inheritance()
    tmp = os.path.join(tempfile.gettempdir(), "ra_style_size.docx")
    d2.save(tmp)
    data2 = measure_styles(tmp, "size_inheritance")
    results.append(data2)
    os.unlink(tmp)
    print(f"\n=== size_inheritance (docDefaults=11pt, Normal=12pt) ===")
    for p in data2["paragraphs"]:
        print(f"  P{p['index']}: font={p['font_name']}, size={p['font_size']}  [{p['text']}]")

    # Spacing inheritance
    d3 = test_spacing_inheritance()
    tmp = os.path.join(tempfile.gettempdir(), "ra_style_spacing.docx")
    d3.save(tmp)
    data3 = measure_styles(tmp, "spacing_inheritance")
    results.append(data3)
    os.unlink(tmp)
    print(f"\n=== spacing_inheritance (docDefaults sa=8pt, Normal sa=10pt) ===")
    for p in data3["paragraphs"]:
        print(f"  P{p['index']}: sa={p['space_after']}, sb={p['space_before']}, "
              f"ls={p['line_spacing']}  [{p['text']}]")

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_style_priority.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
