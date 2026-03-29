"""
Ra: exact lineSpacingRule の精度をCOM計測で確定
- twip→pt変換の正確な丸め方法
- 小フォントサイズ(7-9pt) + exact の行位置
- exactでのテキスト配置(上寄せ? 下寄せ? 中央?)
- exact + character spacing の累積誤差
"""
import win32com.client, json, os, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def make_exact_test(line_twips, font_size=11, font_name="Calibri", num_lines=5):
    """Create test doc with exact line spacing."""
    d = Document(TEMPLATE)
    sec = d.sections[0]
    # Remove grid
    sectPr = sec._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)

    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    for i in range(num_lines):
        p = d.add_paragraph()
        r = p.add_run(f"Line {i+1} exact spacing test ABCDEFG abcdefg 123")
        r.font.name = font_name
        r.font.size = Pt(font_size)
        pPr = p._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')

    return d


def make_exact_cs_test(line_twips, cs_twips_list, font_size=9, font_name="MS Gothic"):
    """Create test doc with exact spacing + character spacing."""
    d = Document(TEMPLATE)
    sec = d.sections[0]
    sectPr = sec._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)

    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    for i, cs_tw in enumerate(cs_twips_list):
        p = d.add_paragraph()
        r = p.add_run(f"cs={cs_tw}tw あいうえおかきくけこ ABCDEFG")
        r.font.name = font_name
        r.font.size = Pt(font_size)
        # Set character spacing
        rPr = r._element.get_or_add_rPr()
        sp_el = etree.SubElement(rPr, qn('w:spacing'))
        sp_el.set(qn('w:val'), str(cs_tw))

        pPr = p._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')

    return d


def measure(doc_path, label):
    doc = word.Documents.Open(doc_path)
    try:
        data = {"label": label, "paragraphs": []}
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            rng = para.Range
            data["paragraphs"].append({
                "index": i,
                "y_pt": round(rng.Information(6), 4),
                "x_pt": round(rng.Information(5), 4),
                "line_spacing": round(para.Format.LineSpacing, 4),
                "ls_rule": para.Format.LineSpacingRule,
            })
        # Compute gaps
        for i in range(1, len(data["paragraphs"])):
            gap = data["paragraphs"][i]["y_pt"] - data["paragraphs"][i-1]["y_pt"]
            data["paragraphs"][i]["gap"] = round(gap, 4)
        return data
    finally:
        doc.Close(False)


try:
    # Test 1: Various exact line heights (twips → pt precision)
    exact_values = [
        (180, "9pt"),    # 180/20 = 9.0pt exactly
        (200, "10pt"),   # 200/20 = 10.0pt
        (240, "12pt"),   # 240/20 = 12.0pt
        (210, "10.5pt"), # 210/20 = 10.5pt
        (195, "9.75pt"), # 195/20 = 9.75pt
        (183, "9.15pt"), # 183/20 = 9.15pt (non-round)
        (187, "9.35pt"), # 187/20 = 9.35pt (non-round)
        (173, "8.65pt"), # 173/20 = 8.65pt (non-round)
        (160, "8pt"),    # 160/20 = 8.0pt
        (280, "14pt"),   # 280/20 = 14.0pt
    ]

    print("=== EXACT LINE SPACING PRECISION ===\n")
    for twips, label in exact_values:
        expected_pt = twips / 20.0
        d = make_exact_test(twips)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_exact_{twips}.docx")
        d.save(tmp)
        data = measure(tmp, f"exact_{twips}tw")
        results.append(data)
        os.unlink(tmp)

        if len(data["paragraphs"]) >= 2:
            gap = data["paragraphs"][1].get("gap", 0)
            reported_ls = data["paragraphs"][0]["line_spacing"]
            diff = abs(gap - expected_pt)
            status = "OK" if diff < 0.1 else f"DIFF={diff:.4f}"
            print(f"  {twips}tw ({label}): expected={expected_pt:.4f}pt, "
                  f"gap={gap:.4f}pt, reported_ls={reported_ls:.4f}pt  [{status}]")

    # Test 2: Small font sizes with exact spacing
    print("\n=== SMALL FONTS + EXACT ===\n")
    small_tests = [
        (7, 140, "7pt font, 7pt exact"),
        (7, 160, "7pt font, 8pt exact"),
        (8, 160, "8pt font, 8pt exact"),
        (9, 180, "9pt font, 9pt exact"),
        (9, 200, "9pt font, 10pt exact"),
    ]
    for fs, twips, label in small_tests:
        d = make_exact_test(twips, font_size=fs)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_exact_small_{fs}_{twips}.docx")
        d.save(tmp)
        data = measure(tmp, label)
        results.append(data)
        os.unlink(tmp)

        if len(data["paragraphs"]) >= 2:
            gap = data["paragraphs"][1].get("gap", 0)
            expected = twips / 20.0
            print(f"  {label}: gap={gap:.4f}pt (expected={expected:.4f}pt)")

    # Test 3: Character spacing precision with exact line spacing
    print("\n=== CHARACTER SPACING + EXACT ===\n")
    cs_values = [-9, -6, -3, 0, 3, 6, 9, 12, -12, -15, 20, -20]
    d = make_exact_cs_test(200, cs_values, font_size=9, font_name="MS Gothic")
    tmp = os.path.join(tempfile.gettempdir(), "ra_exact_cs.docx")
    d.save(tmp)

    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        para = doc.Paragraphs(i)
        rng = para.Range
        y = rng.Information(6)

        # Measure first 20 char positions to check spacing
        chars = []
        for ci in range(rng.Start, min(rng.End, rng.Start + 30)):
            cr = doc.Range(ci, ci + 1)
            ch = cr.Text
            x = cr.Information(5)
            if ord(ch) not in (13, 7):
                chars.append({"ch": ch, "x": round(x, 4)})

        cs_tw = cs_values[i-1] if i-1 < len(cs_values) else 0
        cs_px = (cs_tw * 96 + 720) // 1440 if cs_tw >= 0 else -(-cs_tw * 96 + 720) // 1440
        cs_pt = cs_px * 72.0 / 96.0

        # Check actual spacing between CJK characters
        cjk_gaps = []
        for j in range(1, len(chars)):
            if ord(chars[j]["ch"]) > 0x3000 and ord(chars[j-1]["ch"]) > 0x3000:
                gap = chars[j]["x"] - chars[j-1]["x"]
                cjk_gaps.append(round(gap, 2))

        print(f"  cs={cs_tw}tw → GDI: {cs_px}px = {cs_pt:.4f}pt | "
              f"CJK char gaps: {cjk_gaps[:5]}")

        results.append({
            "cs_twips": cs_tw, "cs_px_expected": cs_px, "cs_pt_expected": round(cs_pt, 4),
            "cjk_char_gaps": cjk_gaps,
        })

    doc.Close(False)
    os.unlink(tmp)

    # Test 4: Exact spacing text vertical placement
    print("\n=== EXACT SPACING: TEXT VERTICAL PLACEMENT ===\n")
    # Create doc with very large exact spacing to see where text sits
    for twips, fs in [(400, 11), (300, 11), (200, 11)]:
        d = make_exact_test(twips, font_size=fs, num_lines=3)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_exact_vpos_{twips}.docx")
        d.save(tmp)
        doc = word.Documents.Open(tmp)

        p1 = doc.Paragraphs(1)
        p2 = doc.Paragraphs(2)
        y1 = p1.Range.Information(6)
        y2 = p2.Range.Information(6)
        gap = y2 - y1
        expected = twips / 20.0

        print(f"  exact={twips}tw({expected}pt), font={fs}pt: "
              f"P1 y={round(y1, 2)}, P2 y={round(y2, 2)}, gap={round(gap, 2)}")

        doc.Close(False)
        os.unlink(tmp)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_exact_linespacing.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
