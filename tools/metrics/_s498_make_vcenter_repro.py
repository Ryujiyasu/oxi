# -*- coding: utf-8 -*-
"""S498 vAlign=center repro: single-cell tables, vAlign=center, row height atLeast in
{30,40,50,60}pt with one CJK line (MS Mincho 10.5pt). Isolates Word's vertical-centering
of cell content vs Oxi. cp932-safe."""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

OUT = os.path.join(os.path.dirname(__file__), '..', 'golden-test', 'repros', 'vcenter_cellY')
os.makedirs(OUT, exist_ok=True)
CJK = 'あいうえお'

def build(atleast_pt):
    doc = Document()
    doc.add_paragraph('REF')
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(CJK)
    r.font.size = Pt(10.5)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'ＭＳ 明朝', qn('w:ascii'): 'ＭＳ 明朝', qn('w:hAnsi'): 'ＭＳ 明朝'})
    rpr.insert(0, rf)
    row = t.rows[0]
    row.height = Pt(atleast_pt)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    doc.add_paragraph('AFTER')
    path = os.path.join(OUT, 'vcenter_%d.docx' % atleast_pt)
    doc.save(path)
    return path

for a in [30, 40, 50, 60]:
    print('built', build(a))
print('DONE')
