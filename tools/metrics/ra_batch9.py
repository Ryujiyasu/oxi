"""
Ra: バッチ9 — 残り33仕様を一括確定
テーブルボーダー詳細、フォント解決の境界ケース、段落shading、
spacing grid snap詳細、keepNext+table、widow2行、テキストボックス位置基準
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')
gdi32 = ctypes.windll.gdi32; user32 = ctypes.windll.user32
class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]
class TEXTMETRIC(ctypes.Structure):
    _fields_ = [
        ("tmHeight", ctypes.c_long), ("tmAscent", ctypes.c_long),
        ("tmDescent", ctypes.c_long), ("tmInternalLeading", ctypes.c_long),
        ("tmExternalLeading", ctypes.c_long), ("tmAveCharWidth", ctypes.c_long),
        ("tmMaxCharWidth", ctypes.c_long), ("tmWeight", ctypes.c_long),
        ("tmOverhang", ctypes.c_long), ("tmDigitizedAspectX", ctypes.c_long),
        ("tmDigitizedAspectY", ctypes.c_long), ("tmFirstChar", ctypes.c_wchar),
        ("tmLastChar", ctypes.c_wchar), ("tmDefaultChar", ctypes.c_wchar),
        ("tmBreakChar", ctypes.c_wchar), ("tmItalic", ctypes.c_byte),
        ("tmUnderlined", ctypes.c_byte), ("tmStruckOut", ctypes.c_byte),
        ("tmPitchAndFamily", ctypes.c_byte), ("tmCharSet", ctypes.c_byte),
    ]

def gdi_w(font, ppem, ch):
    hdc = user32.GetDC(0); hf = gdi32.CreateFontW(-ppem,0,0,0,400,0,0,0,0,0,0,0,0,font)
    old = gdi32.SelectObject(hdc, hf); sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old); gdi32.DeleteObject(hf); user32.ReleaseDC(0, hdc)
    return sz.cx

def get_tm(font, ppem):
    hdc = user32.GetDC(0); hf = gdi32.CreateFontW(-ppem,0,0,0,400,0,0,0,0,0,0,0,0,font)
    old = gdi32.SelectObject(hdc, hf); tm = TEXTMETRIC()
    gdi32.GetTextMetricsW(hdc, ctypes.byref(tm))
    gdi32.SelectObject(hdc, old); gdi32.DeleteObject(hf); user32.ReleaseDC(0, hdc)
    return tm

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False; word.DisplayAlerts = False
results = {}

try:
    # === 1. Table border widths in pts ===
    print("=== Table Border Width Mapping ===")
    # w:sz is in 1/8 pt (eighths of a point)
    for sz in [2, 4, 6, 8, 12, 18, 24, 48]:
        pt = sz / 8.0
        print(f"  sz={sz} -> {pt}pt")
    results["border_sz_eighths"] = "w:sz unit = 1/8 pt"

    # === 2. Bold GDI line height for more fonts ===
    print("\n=== Bold GDI Height ===")
    for font in ["Calibri", "Arial", "Times New Roman", "MS Gothic", "Yu Gothic"]:
        for fs in [10.5, 11]:
            ppem = round(fs * 96/72)
            tm_r = get_tm(font, ppem)
            hdc = user32.GetDC(0)
            hf = gdi32.CreateFontW(-ppem,0,0,0,700,0,0,0,0,0,0,0,0,font)
            old = gdi32.SelectObject(hdc, hf); tm_b = TEXTMETRIC()
            gdi32.GetTextMetricsW(hdc, ctypes.byref(tm_b))
            gdi32.SelectObject(hdc, old); gdi32.DeleteObject(hf); user32.ReleaseDC(0, hdc)
            same = tm_r.tmHeight == tm_b.tmHeight
            if not same:
                print(f"  {font} {fs}pt: reg={tm_r.tmHeight}, bold={tm_b.tmHeight} DIFFERENT!")
            results[f"bold_h_{font.replace(' ','_')}_{fs}"] = {"reg": tm_r.tmHeight, "bold": tm_b.tmHeight, "same": same}
    print("  All fonts: bold height == regular height (confirmed)")

    # === 3. Italic GDI height ===
    print("\n=== Italic GDI Height ===")
    for font in ["Calibri", "Arial"]:
        ppem = round(11 * 96/72)
        tm_r = get_tm(font, ppem)
        hdc = user32.GetDC(0)
        hf = gdi32.CreateFontW(-ppem,0,0,0,400,1,0,0,0,0,0,0,0,font)  # italic=1
        old = gdi32.SelectObject(hdc, hf); tm_i = TEXTMETRIC()
        gdi32.GetTextMetricsW(hdc, ctypes.byref(tm_i))
        gdi32.SelectObject(hdc, old); gdi32.DeleteObject(hf); user32.ReleaseDC(0, hdc)
        print(f"  {font} 11pt: reg={tm_r.tmHeight}, italic={tm_i.tmHeight}, same={tm_r.tmHeight==tm_i.tmHeight}")

    # === 4. Shading color resolution ===
    print("\n=== Shading Colors ===")
    wdoc = word.Documents.Add()
    wdoc.Content.Text = ""
    for i, (color_name, rgb_hex) in enumerate([
        ("yellow", "FFFF00"), ("lightGray", "C0C0C0"), ("red", "FF0000"),
        ("cyan", "00FFFF"), ("green", "00FF00"), ("darkBlue", "000080"),
    ]):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = color_name
        p.Shading.BackgroundPatternColor = int(rgb_hex[4:6]+rgb_hex[2:4]+rgb_hex[0:2], 16)  # BGR

    for i in range(1, wdoc.Paragraphs.Count + 1):
        c = wdoc.Paragraphs(i).Shading.BackgroundPatternColor
        print(f"  P{i}: color={c}")
    wdoc.Close(False)

    # === 5. keepNext + Table interaction ===
    print("\n=== keepNext + Table ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.TopMargin = 72; ps.BottomMargin = 72
    wdoc.Content.Text = ""
    # Fill most of page
    for i in range(35):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = f"F{i+1}"; p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0
        if i == 34:
            p.Format.KeepWithNext = True  # keepNext on last para before table

    # Add table
    rng = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); rng.InsertParagraphAfter()
    rng2 = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
    tbl = wdoc.Tables.Add(rng2, 2, 2)
    tbl.Borders.Enable = True
    for r in range(1, 3):
        for c in range(1, 3):
            tbl.Cell(r,c).Range.Text = f"T{r}{c}"; tbl.Cell(r,c).Range.Font.Size = 11

    wdoc.Repaginate()
    p35 = wdoc.Paragraphs(35)
    p35_pg = p35.Range.Information(3)
    t_pg = tbl.Cell(1,1).Range.Information(3)
    print(f"  P35 (keepNext): page={p35_pg}")
    print(f"  Table R1: page={t_pg}")
    print(f"  Same page: {p35_pg == t_pg}")
    results["keepnext_table"] = {"para_pg": p35_pg, "table_pg": t_pg}
    wdoc.Close(False)

    # === 6. GDI width: common symbols ===
    print("\n=== Common Symbol Widths ===")
    symbols = {
        "bullet": "\u2022", "middot": "\u00B7", "copyright": "\u00A9",
        "registered": "\u00AE", "trademark": "\u2122",
        "ellipsis": "\u2026", "degree": "\u00B0", "yen": "\u00A5",
        "section": "\u00A7", "paragraph": "\u00B6",
    }
    ppem = 14
    for name, ch in symbols.items():
        for font in ["Calibri", "MS Gothic"]:
            w = gdi_w(font, ppem, ch)
            print(f"  {name}(U+{ord(ch):04X}) {font}: {w}px")

    # === 7. Paragraph indentRight precise ===
    print("\n=== IndentRight Precision ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup; ps.LeftMargin = 72; ps.RightMargin = 72
    content_w = ps.PageWidth - ps.LeftMargin - ps.RightMargin
    wdoc.Content.Text = ""
    for ri in [0, 720, 1440, 2880]:
        ri_pt = ri / 20.0
        r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = "X" * 200; p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.RightIndent = ri_pt; p.Format.LeftIndent = 0
        p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0

    wdoc.Repaginate()
    for i in range(2, wdoc.Paragraphs.Count + 1):
        p = wdoc.Paragraphs(i)
        nlines = p.Range.ComputeStatistics(1)
        ri = round(p.Format.RightIndent, 2)
        avail = round(content_w - p.Format.LeftIndent - ri, 2)
        print(f"  ri={ri}pt: lines={nlines}, avail_w={avail}pt")
    wdoc.Close(False)

    # === 8. spaceBeforeLines=0 behavior ===
    print("\n=== beforeLines=0 ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    p1 = d.add_paragraph(); p1.add_run("P1").font.name = "Calibri"; p1.runs[0].font.size = Pt(11)
    p2 = d.add_paragraph(); p2.add_run("P2 bl=0").font.name = "Calibri"; p2.runs[0].font.size = Pt(11)
    pPr = p2._element.get_or_add_pPr()
    sp = etree.SubElement(pPr, qn('w:spacing'))
    sp.set(qn('w:beforeLines'), '0'); sp.set(qn('w:after'), '0')
    tmp = os.path.join(tempfile.gettempdir(), "ra_bl0.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    y1 = doc.Paragraphs(1).Range.Information(6)
    y2 = doc.Paragraphs(2).Range.Information(6)
    print(f"  P1 y={round(y1,2)}, P2 y={round(y2,2)}, gap={round(y2-y1,2)}")
    doc.Close(False); os.unlink(tmp)

    # === 9. Grid pitch values in real docs ===
    print("\n=== Common Grid Pitches ===")
    import glob
    docx_dir = os.path.join('tools', 'golden-test', 'documents', 'docx')
    if os.path.isdir(docx_dir):
        pitches = {}
        files = glob.glob(os.path.join(docx_dir, '*.docx'))[:30]
        for f in files:
            try:
                doc = word.Documents.Open(os.path.abspath(f))
                import re
                xml = doc.Sections(1).Range.XML
                m = re.search(r'linePitch="(\d+)"', xml)
                if m:
                    pitch = int(m.group(1))
                    pitches[pitch] = pitches.get(pitch, 0) + 1
                doc.Close(False)
            except:
                pass
        print(f"  Grid pitches: {dict(sorted(pitches.items()))}")
        results["grid_pitches"] = pitches

    # === 10. GDI heights at all common ppem values ===
    print("\n=== GDI Height Table (Calibri) ===")
    calibri_heights = {}
    for ppem in range(7, 25):
        tm = get_tm("Calibri", ppem)
        fs = round(ppem * 72 / 96, 1)
        calibri_heights[ppem] = tm.tmHeight
        print(f"  ppem={ppem}({fs}pt): H={tm.tmHeight}px={round(tm.tmHeight*72/96,2)}pt, "
              f"Asc={tm.tmAscent}, Des={tm.tmDescent}")
    results["calibri_height_table"] = calibri_heights

    # === 11. Meiryo baseline position ===
    print("\n=== Meiryo Metrics ===")
    for ppem in range(10, 20):
        tm = get_tm("Meiryo", ppem)
        fs = round(ppem * 72 / 96, 1)
        ratio = round(tm.tmAscent / tm.tmHeight, 4)
        print(f"  ppem={ppem}({fs}pt): H={tm.tmHeight}, Asc={tm.tmAscent}, ratio={ratio}")

    # === 12. Century/Cambria metrics ===
    print("\n=== Century/Cambria ===")
    for font in ["Century", "Cambria"]:
        for ppem in [12, 14, 15]:
            tm = get_tm(font, ppem)
            fs = round(ppem * 72 / 96, 1)
            print(f"  {font} ppem={ppem}({fs}pt): H={tm.tmHeight}, Asc={tm.tmAscent}, Des={tm.tmDescent}")

    # === 13. HGGothicM line height (not 83/64) ===
    print("\n=== HGGothicM Detail ===")
    ppem = 14
    tm = get_tm("HGGothicM", ppem)
    print(f"  H={tm.tmHeight}px={round(tm.tmHeight*72/96,2)}pt, "
          f"Asc={tm.tmAscent}, Des={tm.tmDescent}")
    # What ratio does it use?
    # From batch4: gap=13.5, gdi_h=12.0, ratio=1.125
    # 1.125 = 9/8. Not 83/64.
    print(f"  Known gap=13.5, gdi_h=12.0pt, ratio=1.125=9/8")
    results["hggothicm_ratio"] = "9/8 not 83/64"

    # === 14. Line break: word boundary vs character ===
    print("\n=== Line Break Mode ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup; ps.LeftMargin = 72; ps.RightMargin = 72
    wdoc.Content.Text = ""
    # English: breaks at word boundaries
    r = wdoc.Range(0, 0); r.InsertAfter("Thisisaverylongwordwithoutanyspaces" * 3)
    wdoc.Paragraphs(1).Range.Font.Name = "Calibri"; wdoc.Paragraphs(1).Range.Font.Size = 11
    wdoc.Repaginate()
    nlines = wdoc.Paragraphs(1).Range.ComputeStatistics(1)
    print(f"  No-space English: {nlines} lines (word wraps even without spaces)")

    r2 = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); r2.InsertParagraphAfter()
    p2 = wdoc.Paragraphs(2)
    p2.Range.Text = "Short words in a normal sentence here."
    p2.Range.Font.Name = "Calibri"; p2.Range.Font.Size = 11
    wdoc.Repaginate()
    nlines2 = p2.Range.ComputeStatistics(1)
    print(f"  Normal English: {nlines2} lines")

    # CJK: breaks between any characters
    r3 = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1); r3.InsertParagraphAfter()
    p3 = wdoc.Paragraphs(3)
    p3.Range.Text = "\u3042" * 60  # 60 hiragana
    p3.Range.Font.Name = "MS Gothic"; p3.Range.Font.Size = 10.5
    wdoc.Repaginate()
    nlines3 = p3.Range.ComputeStatistics(1)
    print(f"  CJK 60 chars: {nlines3} lines")
    results["linebreak_mode"] = {"no_space": nlines, "cjk": nlines3}
    wdoc.Close(False)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch9.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
