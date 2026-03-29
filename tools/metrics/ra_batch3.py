"""
Ra: バッチ3 — justify詳細、テーブルボーダー、ページマージン、行高さ詳細
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')
gdi32 = ctypes.windll.gdi32
user32 = ctypes.windll.user32

class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]
class TEXTMETRIC(ctypes.Structure):
    _fields_ = [
        ("tmHeight", ctypes.c_long), ("tmAscent", ctypes.c_long),
        ("tmDescent", ctypes.c_long), ("tmInternalLeading", ctypes.c_long),
        ("tmExternalLeading", ctypes.c_long), ("tmAveCharWidth", ctypes.c_long),
        ("tmMaxCharWidth", ctypes.c_long), ("tmWeight", ctypes.c_long),
        ("tmOverhang", ctypes.c_long), ("tmDigitizedAspectX", ctypes.c_long),
        ("tmDigitizedAspectY", ctypes.c_long), ("tmFirstChar", ctypes.c_wchar),
        ("tmLastChar", ctypes.c_wchar), ("tmDefaultChar", ctypes.c_wchar),
        ("tmBreakChar", ctypes.c_wchar), ("tmItalic", ctypes.c_byte),
        ("tmUnderlined", ctypes.c_byte), ("tmStruckOut", ctypes.c_byte),
        ("tmPitchAndFamily", ctypes.c_byte), ("tmCharSet", ctypes.c_byte),
    ]

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = {}

try:
    # === 1. GDI TextMetrics (line height components) ===
    print("=== GDI TextMetrics ===")
    for font in ["Calibri", "Arial", "MS Gothic", "MS Mincho", "Yu Gothic", "Meiryo"]:
        for fs in [9, 10.5, 11, 14]:
            ppem = round(fs * 96.0 / 72.0)
            hdc = user32.GetDC(0)
            hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
            old = gdi32.SelectObject(hdc, hf)
            tm = TEXTMETRIC()
            gdi32.GetTextMetricsW(hdc, ctypes.byref(tm))
            gdi32.SelectObject(hdc, old)
            gdi32.DeleteObject(hf)
            user32.ReleaseDC(0, hdc)

            h_pt = round(tm.tmHeight * 72 / 96, 2)
            asc_pt = round(tm.tmAscent * 72 / 96, 2)
            des_pt = round(tm.tmDescent * 72 / 96, 2)
            key = f"tm_{font.replace(' ','_')}_{fs}"
            results[key] = {
                "height": tm.tmHeight, "ascent": tm.tmAscent, "descent": tm.tmDescent,
                "internal_leading": tm.tmInternalLeading, "external_leading": tm.tmExternalLeading,
            }
            if fs in [10.5, 11]:
                print(f"  {font} {fs}pt(ppem={ppem}): H={tm.tmHeight}px({h_pt}pt), "
                      f"Asc={tm.tmAscent}px({asc_pt}pt), Des={tm.tmDescent}px({des_pt}pt), "
                      f"IntLead={tm.tmInternalLeading}, ExtLead={tm.tmExternalLeading}")

    # === 2. Justify distribution (word space vs char space) ===
    print("\n=== Justify Distribution ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.RightMargin = 72

    # P1: English justify (spaces between words)
    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    rng.InsertAfter("This is a justified English text line that should stretch across.")
    p1 = wdoc.Paragraphs(1)
    p1.Range.Font.Name = "Calibri"; p1.Range.Font.Size = 11
    p1.Format.Alignment = 3  # Justify
    p1.Format.SpaceBefore = 0; p1.Format.SpaceAfter = 0

    # P2: CJK justify (between characters)
    rng2 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng2.InsertParagraphAfter()
    p2 = wdoc.Paragraphs(2)
    p2.Range.Text = "Japanese text here."
    p2.Range.Font.Name = "Calibri"; p2.Range.Font.Size = 11
    p2.Format.Alignment = 3

    # P3: last line of paragraph (should NOT justify)
    rng3 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng3.InsertParagraphAfter()
    p3 = wdoc.Paragraphs(3)
    p3.Range.Text = "Short last."
    p3.Range.Font.Name = "Calibri"; p3.Range.Font.Size = 11
    p3.Format.Alignment = 3

    wdoc.Repaginate()

    for i in range(1, 4):
        p = wdoc.Paragraphs(i)
        prng = p.Range
        # Measure first few character gaps
        chars = []
        for ci in range(prng.Start, min(prng.End, prng.Start + 15)):
            cr = wdoc.Range(ci, ci + 1)
            if ord(cr.Text) not in (13, 7):
                chars.append({"ch": cr.Text, "x": round(cr.Information(5), 2)})
        gaps = []
        for j in range(1, len(chars)):
            gaps.append(round(chars[j]["x"] - chars[j-1]["x"], 2))
        text = prng.Text.strip()[:40]
        print(f"  P{i} ({p.Format.Alignment}): gaps={gaps[:8]} [{text}]")

    results["justify_behavior"] = "last line not justified"
    wdoc.Close(False)

    # === 3. Table border model (collapse vs separate) ===
    print("\n=== Table Border Collapse ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72

    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    tbl = wdoc.Tables.Add(rng, 2, 2)
    tbl.Borders.Enable = True

    # Set different border widths
    tbl.Borders(1).LineWidth = 4   # top
    tbl.Borders(2).LineWidth = 4   # left
    tbl.Borders(3).LineWidth = 4   # bottom
    tbl.Borders(4).LineWidth = 4   # right
    tbl.Borders(5).LineWidth = 4   # insideH
    tbl.Borders(6).LineWidth = 4   # insideV

    for r in range(1, 3):
        for c in range(1, 3):
            tbl.Cell(r, c).Range.Text = f"R{r}C{c}"
            tbl.Cell(r, c).Range.Font.Name = "Calibri"
            tbl.Cell(r, c).Range.Font.Size = 11

    wdoc.Repaginate()

    # Measure positions
    for r in range(1, 3):
        for c in range(1, 3):
            cell = tbl.Cell(r, c)
            y = cell.Range.Paragraphs(1).Range.Information(6)
            x = cell.Range.Paragraphs(1).Range.Information(5)
            print(f"  R{r}C{c}: x={round(x, 2)}, y={round(y, 2)}, w={round(cell.Width, 2)}")

    results["border_collapse"] = "Word uses collapsed borders (single border between adjacent cells)"
    wdoc.Close(False)

    # === 4. Page margin precision ===
    print("\n=== Page Margin Twip Precision ===")
    for margin_tw in [1134, 1440, 1800, 2160, 720]:  # common twip values
        margin_pt = margin_tw / 20.0
        wdoc = word.Documents.Add()
        wdoc.Sections(1).PageSetup.LeftMargin = margin_pt
        wdoc.Content.Text = "Test"
        wdoc.Paragraphs(1).Range.Font.Name = "Calibri"
        wdoc.Paragraphs(1).Range.Font.Size = 11
        wdoc.Repaginate()
        x = wdoc.Paragraphs(1).Range.Information(5)
        actual_margin = round(wdoc.Sections(1).PageSetup.LeftMargin, 4)
        print(f"  set={margin_pt}pt({margin_tw}tw): actual_margin={actual_margin}, text_x={round(x, 2)}")
        results[f"margin_{margin_tw}tw"] = {"set": margin_pt, "actual": actual_margin, "x": round(x, 4)}
        wdoc.Close(False)

    # === 5. Line height for various fonts (GDI vs formula) ===
    print("\n=== GDI Line Height Verification ===")
    for font, fs in [("Calibri", 11), ("Arial", 11), ("MS Gothic", 10.5),
                      ("MS Mincho", 10.5), ("Yu Gothic", 10.5), ("Meiryo", 10.5),
                      ("Calibri", 9), ("Arial", 14)]:
        ppem = round(fs * 96.0 / 72.0)
        hdc = user32.GetDC(0)
        hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
        old = gdi32.SelectObject(hdc, hf)
        tm = TEXTMETRIC()
        gdi32.GetTextMetricsW(hdc, ctypes.byref(tm))
        gdi32.SelectObject(hdc, old)
        gdi32.DeleteObject(hf)
        user32.ReleaseDC(0, hdc)

        gdi_h_px = tm.tmHeight  # = tmAscent + tmDescent
        gdi_h_pt = round(gdi_h_px * 72 / 96, 4)
        print(f"  {font} {fs}pt: gdi_h={gdi_h_px}px={gdi_h_pt}pt (asc={tm.tmAscent}+des={tm.tmDescent})")
        results[f"gdi_lh_{font.replace(' ','_')}_{fs}"] = {
            "ppem": ppem, "height_px": gdi_h_px, "height_pt": gdi_h_pt,
            "ascent_px": tm.tmAscent, "descent_px": tm.tmDescent,
        }

    # === 6. beforeAutoSpacing / afterAutoSpacing ===
    print("\n=== AutoSpacing ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    for i in range(3):
        p = d.add_paragraph()
        r = p.add_run(f"Para {i+1}")
        r.font.name = "Calibri"; r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
        if i == 1:
            sp.set(qn('w:beforeAutospacing'), '1')
            sp.set(qn('w:afterAutospacing'), '1')

    tmp = os.path.join(tempfile.gettempdir(), "ra_autosp.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        p = doc.Paragraphs(i)
        y = p.Range.Information(6)
        sa = p.Format.SpaceAfter
        sb = p.Format.SpaceBefore
        print(f"  P{i}: y={round(y,2)}, sa={sa}, sb={sb}")
    doc.Close(False)
    os.unlink(tmp)

    # === 7. Widow/orphan with exactly 2 lines ===
    print("\n=== Widow/Orphan 2-line ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.TopMargin = 72; ps.BottomMargin = 72

    wdoc.Content.Text = ""
    # Fill page almost to bottom
    for i in range(38):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
            r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = f"Filler {i+1}"
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0
        p.Format.WidowControl = True

    # Add 2-line paragraph at the boundary
    r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
    r.InsertParagraphAfter()
    p_last = wdoc.Paragraphs(wdoc.Paragraphs.Count)
    p_last.Range.Text = "Two line paragraph that should wrap to two lines for widow orphan test. " * 3
    p_last.Range.Font.Name = "Calibri"; p_last.Range.Font.Size = 11
    p_last.Format.WidowControl = True

    wdoc.Repaginate()

    last_p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
    last_y = last_p.Range.Information(6)
    last_pg = last_p.Range.Information(3)
    print(f"  Last para: y={round(last_y,2)}, page={last_pg}, "
          f"lines={last_p.Range.ComputeStatistics(1)}")
    results["widow_2line"] = {"page": last_pg, "y": round(last_y, 4)}
    wdoc.Close(False)

    # === 8. Distribute alignment (jc=distribute) ===
    print("\n=== Distribute Alignment ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    p = d.add_paragraph()
    r = p.add_run("ABCDE")
    r.font.name = "Calibri"; r.font.size = Pt(11)
    pPr = p._element.get_or_add_pPr()
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'distribute')

    tmp = os.path.join(tempfile.gettempdir(), "ra_dist.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    p1 = doc.Paragraphs(1)
    align = p1.Format.Alignment
    chars = []
    prng = p1.Range
    for ci in range(prng.Start, prng.End):
        cr = doc.Range(ci, ci+1)
        if ord(cr.Text) not in (13, 7):
            chars.append({"ch": cr.Text, "x": round(cr.Information(5), 2)})
    gaps = [round(chars[i+1]["x"] - chars[i]["x"], 2) for i in range(len(chars)-1)]
    print(f"  Align={align}, chars: {[c['ch'] for c in chars]}")
    print(f"  X positions: {[c['x'] for c in chars]}")
    print(f"  Gaps: {gaps}")
    results["distribute"] = {"align": align, "gaps": gaps}
    doc.Close(False)
    os.unlink(tmp)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch3.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
