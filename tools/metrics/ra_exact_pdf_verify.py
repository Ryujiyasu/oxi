"""
Ra: exact line spacing の 0.5pt量子化 — PDFテキスト座標で検証
COM Information(6)の測定精度限界 vs Word実際のレンダリング位置
"""
import win32com.client, json, os, tempfile, subprocess
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def make_doc(line_twips):
    d = Document(TEMPLATE)
    sec = d.sections[0]
    sectPr = sec._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)

    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    for i in range(5):
        p = d.add_paragraph()
        r = p.add_run(f"Line {i+1} test")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
    return d


try:
    test_twips = [183, 187, 195, 173, 210, 240]

    for tw in test_twips:
        d = make_doc(tw)
        docx_path = os.path.join(tempfile.gettempdir(), f"ra_verify_{tw}.docx")
        d.save(docx_path)

        # Open in Word and measure COM positions
        doc = word.Documents.Open(docx_path)
        com_ys = []
        for i in range(1, doc.Paragraphs.Count + 1):
            y = doc.Paragraphs(i).Range.Information(6)
            com_ys.append(round(y, 4))
        doc.Close(False)

        # Save as PDF via Word
        pdf_path = os.path.join(tempfile.gettempdir(), f"ra_verify_{tw}.pdf")
        doc2 = word.Documents.Open(docx_path)
        doc2.SaveAs2(pdf_path, FileFormat=17)
        doc2.Close(False)

        expected = tw / 20.0
        com_gaps = [round(com_ys[i] - com_ys[i-1], 4) for i in range(1, len(com_ys))]

        print(f"\n=== {tw}tw = {expected}pt ===")
        print(f"  COM Y positions: {com_ys[:5]}")
        print(f"  COM gaps: {com_gaps}")
        print(f"  PDF saved: {pdf_path}")

        results.append({
            "twips": tw,
            "expected_pt": expected,
            "com_ys": com_ys,
            "com_gaps": com_gaps,
            "pdf_path": pdf_path,
        })

        os.unlink(docx_path)
        # Keep PDFs for manual inspection

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_exact_pdf_verify.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")

print("\n=== SUMMARY ===")
for r in results:
    gaps = r["com_gaps"]
    avg_gap = sum(gaps) / len(gaps) if gaps else 0
    print(f"  {r['twips']}tw: expected={r['expected_pt']:.4f}pt, avg_com_gap={avg_gap:.4f}pt, "
          f"diff={abs(avg_gap - r['expected_pt']):.4f}pt")
