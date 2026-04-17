"""Minimal repro: ・ followed by chars of different classes.

Goal: derive the rule for ・ advance depending on next-char class.
Each paragraph starts with ・, followed by 5 chars of a specific class.
Measure x of each char and compute advances.
"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data",
                 "middle_dot_context_repro.docx")
)

# Each sample: (label, ・+text). Test ・ followed by various char classes.
SAMPLES = [
    # 1. ・CJK ideograph
    "・利用規約名を表記",
    # 2. ・hiragana
    "・あいうえお次に",
    # 3. ・katakana
    "・アイウエオ次に",
    # 4. ・Latin letter
    "・ABCDE next 次",
    # 5. ・digit
    "・12345 next 次",
    # 6. ・punct (。)
    "・。・、本文を見る",
    # 7. No bullet (baseline) — measure 利 vs 利
    "本文利用規約名です",
    # 8. Mid-text ・ (not at paragraph start)
    "本文の・中に位置する",
    # 9. ・ at start, followed by half-width space
    "・ 利用 次に",
    # 10. ・+hiragana with specific pattern d77a-like
    "・利用規約名を表記する",
]


def set_font(run, family="ＭＳ ゴシック", size_pt=12.0):
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)


def main():
    doc = Document()
    settings = doc.settings.element
    for existing in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(existing)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)

    for i, text in enumerate(SAMPLES, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"S{i}: {text}")
        set_font(r)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"[OK] {OUT}")


if __name__ == "__main__":
    main()
