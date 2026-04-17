"""Use python-docx to safely strip d77a to minimal variants and measure '（'.

Strategies tested (ordered):
1. keep_all — baseline (confirms current 10.5pt)
2. keep_only_idx10 — delete all body paragraphs except the one containing target
3. keep_idx10_and_title — keep idx=10 + first title only
"""
import os, sys, time, shutil
import win32com.client
from docx import Document

TMP = os.path.abspath("pipeline_data/_pydocx_tmp")
os.makedirs(TMP, exist_ok=True)
SRC = os.path.abspath(r"tools\golden-test\documents\docx\d77a58485f16_20240705_resources_data_outline_08.docx")
TARGET = "公共データ利用規約（第1.0版"

def make_keep_only(out_path, predicate):
    """Copy SRC to out_path, delete all body paragraphs where NOT predicate(text)."""
    shutil.copy(SRC, out_path)
    doc = Document(out_path)
    # Collect paragraphs in body (not in tables)
    # python-docx's doc.paragraphs gives top-level body paragraphs
    keep_count = 0
    del_count = 0
    to_delete = []
    for p in list(doc.paragraphs):
        if not predicate(p.text):
            to_delete.append(p)
            del_count += 1
        else:
            keep_count += 1
    # Remove by element manipulation
    for p in to_delete:
        el = p._element
        el.getparent().remove(el)
    doc.save(out_path)
    return keep_count, del_count

def measure(path):
    """Return dict of '（' advance keyed by font_size."""
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(path, ReadOnly=True); time.sleep(0.3)
        results = {}
        for p in doc.Paragraphs:
            text = p.Range.Text
            if '（' not in text: continue
            rng = p.Range
            for ci in range(1, rng.Characters.Count + 1):
                c = rng.Characters(ci)
                if c.Text == '（':
                    try:
                        x1 = c.Information(5); y1 = c.Information(6)
                        nxt = rng.Characters(ci + 1)
                        x2 = nxt.Information(5); y2 = nxt.Information(6)
                        if abs(y1 - y2) > 2: continue
                        fs = round(c.Font.Size, 1)
                        if fs not in results:
                            results[fs] = round(x2 - x1, 2)
                    except: pass
            if len(results) >= 3: break
        doc.Close(False)
        return results
    finally:
        word.Quit()

TESTS = [
    ("keep_any_paren", lambda t: '（' in t),
    ("keep_only_idx10_v2", lambda t: '公共データ利用規約（第1.0版）' in t),
    ("keep_idx10_strict", lambda t: t.startswith('「公共データ利用規約（第1.0版）」')),
    ("keep_empty_and_idx10", lambda t: t == '' or '公共データ利用規約（第1.0版）」の前身' in t),
    ("keep_first_5_paren", lambda t, c=[0]: ('（' in t and (c.__setitem__(0, c[0]+1) or c[0] <= 5))),
]

print(f"{'variant':<25} {'kept':>5} {'del':>5}  {'fs=14':>7}  {'fs=12':>7}  {'fs=10.5':>7}")
print('-' * 75)
for label, pred in TESTS:
    out = os.path.join(TMP, f"{label}.docx")
    kept, deleted = make_keep_only(out, pred)
    try:
        r = measure(out)
        fs14 = r.get(14.0, '-')
        fs12 = r.get(12.0, '-')
        fs105 = r.get(10.5, '-')
        marker = '' if fs12 == '-' else (' **compressed**' if (isinstance(fs12, float) and fs12 < 11.5) else ' (no compress)')
        print(f"{label:<25} {kept:>5} {deleted:>5}  {fs14:>7}  {fs12:>7}  {fs105:>7}{marker}")
    except Exception as e:
        print(f"{label:<25} {kept:>5} {deleted:>5} ERROR {e}")
