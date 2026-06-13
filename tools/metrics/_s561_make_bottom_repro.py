# -*- coding: utf-8 -*-
"""S561: derive Word's page-bottom rule for TRAILING EMPTY paragraphs (the
roudoujoken -1). Page = roudoujoken geometry (A4 portrait, top454/bottom233/
LR1701 tw), font MS明朝 10.5pt. K single-line filler CJK paras, then a tail:
 empty-tail: [empty][empty][MARKER]
 text-tail : [あ][あ][MARKER]   (control: is the leniency empty-specific?)
Measure MARKER's page in Word vs Oxi; the K where they DIFFER is the bug."""
import os,sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Emu
sys.stdout.reconfigure(encoding='utf-8')
TW=635  # EMU per twip
def setpg(sec):
    sec.page_width=Emu(11906*TW); sec.page_height=Emu(16838*TW)
    sec.top_margin=Emu(454*TW); sec.bottom_margin=Emu(233*TW)
    sec.left_margin=Emu(1701*TW); sec.right_margin=Emu(1701*TW)
    sec.header_distance=Emu(851*TW); sec.footer_distance=Emu(992*TW)
    # S561: match roudoujoken docGrid (type=lines linePitch=329tw=16.45pt)
    sectPr=sec._sectPr
    for g in sectPr.findall(qn('w:docGrid')): sectPr.remove(g)
    g=OxmlElement('w:docGrid'); g.set(qn('w:type'),'lines'); g.set(qn('w:linePitch'),'329')
    sectPr.append(g)
def mkpar(doc,text):
    p=doc.add_paragraph()
    r=p.add_run(text)
    r.font.size=Pt(10.5)
    rpr=r._r.get_or_add_rPr(); rf=rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'),'ＭＳ 明朝'); rf.set(qn('w:ascii'),'Century')
    # kill spacing for clean line stacking
    pf=p.paragraph_format; pf.space_before=Pt(0); pf.space_after=Pt(0); pf.line_spacing=1.0
    return p
def build(path,K,tail):
    doc=Document(); setpg(doc.sections[0])
    for i in range(K): mkpar(doc,'行%02dあいうえおかきくけこ'%i)
    if tail=='empty':
        mkpar(doc,''); mkpar(doc,'')
    else:
        mkpar(doc,'あ'); mkpar(doc,'あ')
    mkpar(doc,'MARKER終端')
    doc.save(path)
out='tools/golden-test/repros/s561_bottom_grid'; os.makedirs(out,exist_ok=True)
for K in range(50,58):
    build('%s/s561e_K%02d_repro.docx'%(out,K),K,'empty')
for K in (53,54,55):
    build('%s/s561t_K%02d_repro.docx'%(out,K),K,'text')
print('wrote', len(os.listdir(out)),'repros to',out)
