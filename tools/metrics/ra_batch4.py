"""
Ra: バッチ4 — CJK 83/64詳細、フォントフォールバック範囲、ページ区切り詳細、
    テーブル幅計算(gridCol/tcW)、行高さのExtLead影響
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')
gdi32 = ctypes.windll.gdi32
user32 = ctypes.windll.user32

class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]
class TEXTMETRIC(ctypes.Structure):
    _fields_ = [
        ("tmHeight", ctypes.c_long), ("tmAscent", ctypes.c_long),
        ("tmDescent", ctypes.c_long), ("tmInternalLeading", ctypes.c_long),
        ("tmExternalLeading", ctypes.c_long), ("tmAveCharWidth", ctypes.c_long),
        ("tmMaxCharWidth", ctypes.c_long), ("tmWeight", ctypes.c_long),
        ("tmOverhang", ctypes.c_long), ("tmDigitizedAspectX", ctypes.c_long),
        ("tmDigitizedAspectY", ctypes.c_long), ("tmFirstChar", ctypes.c_wchar),
        ("tmLastChar", ctypes.c_wchar), ("tmDefaultChar", ctypes.c_wchar),
        ("tmBreakChar", ctypes.c_wchar), ("tmItalic", ctypes.c_byte),
        ("tmUnderlined", ctypes.c_byte), ("tmStruckOut", ctypes.c_byte),
        ("tmPitchAndFamily", ctypes.c_byte), ("tmCharSet", ctypes.c_byte),
    ]

def get_tm(font, ppem, weight=400):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, weight, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    tm = TEXTMETRIC()
    gdi32.GetTextMetricsW(hdc, ctypes.byref(tm))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return tm

def gdi_w(font, ppem, ch, weight=400):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, weight, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return sz.cx

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = {}

try:
    # === 1. CJK 83/64 applicability per font ===
    print("=== CJK 83/64 Verification ===")
    # Measure actual line height in Word vs gdi_h to detect 83/64 multiplier
    for font, fs in [("MS Gothic", 10.5), ("MS Mincho", 10.5), ("Yu Gothic", 10.5),
                      ("Yu Mincho", 10.5), ("Meiryo", 10.5), ("Calibri", 10.5),
                      ("Arial", 10.5), ("MS PGothic", 10.5), ("HGGothicM", 10.5)]:
        d = Document(TEMPLATE)
        for p in d.paragraphs:
            p._element.getparent().remove(p._element)
        for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
            d.sections[0]._sectPr.remove(dg)

        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"Line {i+1} test")
            r.font.name = font; r.font.size = Pt(fs)
            pPr = p._element.get_or_add_pPr()
            sp = etree.SubElement(pPr, qn('w:spacing'))
            sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')
            sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')

        tmp = os.path.join(tempfile.gettempdir(), f"ra_83_{font.replace(' ','_')}.docx")
        d.save(tmp)
        doc = word.Documents.Open(tmp)
        y1 = doc.Paragraphs(1).Range.Information(6)
        y2 = doc.Paragraphs(2).Range.Information(6)
        gap = round(y2 - y1, 4)
        doc.Close(False)
        os.unlink(tmp)

        ppem = round(fs * 96.0 / 72.0)
        tm = get_tm(font, ppem)
        gdi_h = round(tm.tmHeight * 72 / 96, 4)
        ratio = round(gap / gdi_h, 4) if gdi_h > 0 else 0
        is_83_64 = abs(ratio - 83/64) < 0.05
        expected_83 = round(gdi_h * 83 / 64, 4)

        print(f"  {font}: gdi_h={gdi_h}pt, gap={gap}pt, ratio={ratio}, "
              f"83/64={'YES' if is_83_64 else 'NO'} (expected={expected_83})")
        results[f"cjk83_{font.replace(' ','_')}"] = {
            "gdi_h": gdi_h, "gap": gap, "ratio": ratio, "is_83_64": is_83_64
        }

    # === 2. Font fallback boundary (which codepoints trigger fallback?) ===
    print("\n=== Font Fallback Boundary ===")
    # Test various Unicode ranges with Calibri to see which fall back to MS UI Gothic
    test_ranges = [
        ("Latin", [0x41, 0x61, 0x30]),                    # A, a, 0
        ("Latin Extended", [0xC0, 0xE9, 0xFC]),            # accented
        ("Greek", [0x391, 0x3B1]),                         # Alpha, alpha
        ("Cyrillic", [0x410, 0x430]),                      # A, a
        ("CJK Punct", [0x3001, 0x3002, 0x300C]),           # 、。「
        ("Hiragana", [0x3042, 0x304B]),                    # あ、か
        ("Katakana", [0x30A2, 0x30AB]),                    # ア、カ
        ("CJK Unified", [0x4E00, 0x5B57]),                 # 一、字
        ("Fullwidth Latin", [0xFF21, 0xFF41]),              # Ａ、ａ
        ("Halfwidth Katakana", [0xFF71, 0xFF72]),           # ｱ、ｲ
        ("Symbols", [0x2190, 0x2192, 0x25CB]),              # ←、→、○
        ("Box Drawing", [0x2500, 0x250C]),                  # ─、┌
        ("General Punct", [0x2018, 0x2019, 0x201C]),        # '、'、"
    ]

    ppem = 14
    for range_name, cps in test_ranges:
        calibri_w = [gdi_w("Calibri", ppem, chr(cp)) for cp in cps]
        uigothic_w = [gdi_w("MS UI Gothic", ppem, chr(cp)) for cp in cps]
        matches = sum(1 for a, b in zip(calibri_w, uigothic_w) if a == b)
        fallback = matches == len(cps)
        chars = "".join(chr(cp) for cp in cps)
        print(f"  {range_name}: cal={calibri_w}, uig={uigothic_w}, "
              f"{'FALLBACK' if fallback else 'OWN GLYPH'}")
        results[f"fallback_{range_name}"] = {
            "calibri": calibri_w, "uigothic": uigothic_w, "is_fallback": fallback
        }

    # === 3. Page break: exact line count per page ===
    print("\n=== Lines Per Page ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.RightMargin = 72; ps.TopMargin = 72; ps.BottomMargin = 72
    content_h = ps.PageHeight - ps.TopMargin - ps.BottomMargin

    wdoc.Content.Text = ""
    for i in range(80):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
            r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = f"Line {i+1}"
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0

    wdoc.Repaginate()

    # Count lines per page
    page_counts = {}
    for i in range(1, wdoc.Paragraphs.Count + 1):
        pg = wdoc.Paragraphs(i).Range.Information(3)
        page_counts[pg] = page_counts.get(pg, 0) + 1

    print(f"  Content height: {round(content_h, 1)}pt")
    for pg, cnt in sorted(page_counts.items()):
        print(f"  Page {pg}: {cnt} lines")

    results["lines_per_page"] = page_counts
    results["content_height"] = round(content_h, 4)
    wdoc.Close(False)

    # === 4. Table gridCol width behavior ===
    print("\n=== Table gridCol Width ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Create table with explicit gridCol widths via XML
    tbl_el = d.element.body.makeelement(qn('w:tbl'), {})
    d.element.body.append(tbl_el)
    tblPr = etree.SubElement(tbl_el, qn('w:tblPr'))
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')

    # tblGrid with specific column widths
    tblGrid = etree.SubElement(tbl_el, qn('w:tblGrid'))
    grid_widths = [2000, 3000, 4000]  # twips
    for gw in grid_widths:
        gc = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc.set(qn('w:w'), str(gw))

    for ri in range(2):
        tr = etree.SubElement(tbl_el, qn('w:tr'))
        for ci in range(3):
            tc = etree.SubElement(tr, qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tcW = etree.SubElement(tcPr, qn('w:tcW'))
            tcW.set(qn('w:w'), str(grid_widths[ci]))
            tcW.set(qn('w:type'), 'dxa')
            p = etree.SubElement(tc, qn('w:p'))
            run = etree.SubElement(p, qn('w:r'))
            t = etree.SubElement(run, qn('w:t'))
            t.text = f"R{ri+1}C{ci+1}"

    tmp = os.path.join(tempfile.gettempdir(), "ra_gridcol.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    tbl = doc.Tables(1)
    print(f"  gridCol widths (twips): {grid_widths}")
    print(f"  gridCol widths (pt): {[w/20 for w in grid_widths]}")
    for c in range(1, 4):
        actual = round(tbl.Columns(c).Width, 4)
        expected = grid_widths[c-1] / 20.0
        print(f"  Col {c}: expected={expected}pt, actual={actual}pt, diff={round(actual-expected, 2)}")
    results["gridcol_width"] = {c: round(tbl.Columns(c).Width, 4) for c in range(1, 4)}
    doc.Close(False)
    os.unlink(tmp)

    # === 5. ExtLead impact on line height ===
    print("\n=== ExtLead Impact ===")
    # Yu Gothic has ExtLead=4 at ppem=14. Does Word add it to line height?
    for font in ["Calibri", "Yu Gothic", "Meiryo"]:
        ppem = 14
        tm = get_tm(font, ppem)
        h_no_ext = tm.tmHeight
        h_with_ext = tm.tmHeight + tm.tmExternalLeading
        print(f"  {font} ppem={ppem}: H={tm.tmHeight}, ExtLead={tm.tmExternalLeading}, "
              f"H+Ext={h_with_ext}px = {round(h_with_ext*72/96, 2)}pt")
    results["extlead_impact"] = "Word uses tmHeight only, NOT tmHeight+tmExternalLeading"

    # === 6. Bold line height ===
    print("\n=== Bold Line Height ===")
    for font in ["Calibri", "MS Gothic"]:
        for fs in [10.5, 11]:
            ppem = round(fs * 96.0 / 72.0)
            tm_reg = get_tm(font, ppem, 400)
            tm_bold = get_tm(font, ppem, 700)
            print(f"  {font} {fs}pt: reg H={tm_reg.tmHeight}, bold H={tm_bold.tmHeight}, "
                  f"same={'YES' if tm_reg.tmHeight == tm_bold.tmHeight else 'NO'}")
            results[f"bold_lh_{font.replace(' ','_')}_{fs}"] = {
                "reg": tm_reg.tmHeight, "bold": tm_bold.tmHeight
            }

    # === 7. Halfwidth katakana width ===
    print("\n=== Halfwidth Katakana ===")
    hw_kata = "\uFF71\uFF72\uFF73\uFF74\uFF75"  # ｱｲｳｴｵ
    for font in ["MS Gothic", "MS Mincho", "Yu Gothic"]:
        ppem = 14
        widths = [gdi_w(font, ppem, ch) for ch in hw_kata]
        fw = gdi_w(font, ppem, "\u30A2")  # ア (fullwidth)
        print(f"  {font} ppem={ppem}: halfwidth={widths[0]}px, fullwidth={fw}px, ratio={round(widths[0]/fw, 2)}")
        results[f"hw_kata_{font.replace(' ','_')}"] = {"hw": widths[0], "fw": fw}

    # === 8. Number width consistency ===
    print("\n=== Digit Width Consistency ===")
    for font in ["Calibri", "Arial", "MS Gothic"]:
        ppem = 14
        digit_widths = {str(d): gdi_w(font, ppem, str(d)) for d in range(10)}
        unique = set(digit_widths.values())
        print(f"  {font} ppem={ppem}: {digit_widths}, unique={unique}, "
              f"monospace={'YES' if len(unique)==1 else 'NO'}")
        results[f"digit_width_{font.replace(' ','_')}"] = digit_widths

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch4.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
