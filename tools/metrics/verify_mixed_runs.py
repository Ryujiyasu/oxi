"""
Mixed runs line height verification via Word COM.
Tests: When a line has multiple runs with different fonts/sizes,
does Word use max(line_height) across all runs?
Or max(ascent) + max(descent)?
"""
import win32com.client
import os, time, tempfile

def create_test_docx(path):
    """Create a .docx with mixed font runs on single lines."""
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.oxml.ns import qn

    doc = Document()

    # Set document grid to known pitch (disable grid snap for clean measurement)
    section = doc.sections[0]
    sectPr = section._sectPr
    # Remove grid snap by setting docGrid type to "default"
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)
    grid_el = sectPr.makeelement(qn('w:docGrid'), {qn('w:type'): 'default', qn('w:linePitch'): '360'})
    sectPr.append(grid_el)

    tests = [
        # (description, [(text, font, size_pt), ...])
        # Test 1: Same font, different sizes
        ("Calibri 10.5+14", [("Hello ", "Calibri", 10.5), ("World", "Calibri", 14)]),
        ("Calibri 10.5 only", [("Hello World", "Calibri", 10.5)]),
        ("Calibri 14 only", [("Hello World", "Calibri", 14)]),

        # Test 2: CJK + Western mixed
        ("YuGothic14+Calibri10.5", [("Test", "Yu Gothic", 14), ("abc", "Calibri", 10.5)]),
        ("YuGothic14 only", [("Test text", "Yu Gothic", 14)]),
        ("Calibri10.5 only", [("Test text", "Calibri", 10.5)]),

        # Test 3: Different CJK fonts
        ("YuGothic10.5+MSGothic10.5", [("AB", "Yu Gothic", 10.5), ("CD", "MS Gothic", 10.5)]),
        ("YuGothic10.5 only", [("ABCD", "Yu Gothic", 10.5)]),
        ("MSGothic10.5 only", [("ABCD", "MS Gothic", 10.5)]),

        # Test 4: Small + Very Large
        ("Calibri8+Calibri28", [("small", "Calibri", 8), ("BIG", "Calibri", 28)]),
        ("Calibri8 only", [("smallBIG", "Calibri", 8)]),
        ("Calibri28 only", [("smallBIG", "Calibri", 28)]),

        # Test 5: CJK small + Western large
        ("MSGothic10.5+Calibri24", [("AB", "MS Gothic", 10.5), ("XY", "Calibri", 24)]),
        ("MSGothic10.5 only", [("ABXY", "MS Gothic", 10.5)]),
        ("Calibri24 only", [("ABXY", "Calibri", 24)]),

        # Test 6: 3 different sizes in one line
        ("Cal8+Cal14+Cal24", [("a", "Calibri", 8), ("b", "Calibri", 14), ("c", "Calibri", 24)]),
        ("Calibri24 only (ref)", [("abc", "Calibri", 24)]),
    ]

    for desc, runs in tests:
        p = doc.add_paragraph()
        # Disable snap to grid
        pPr = p._element.get_or_add_pPr()
        snap = pPr.makeelement(qn('w:snapToGrid'), {qn('w:val'): '0'})
        pPr.append(snap)
        # Set single spacing explicitly
        spacing = pPr.makeelement(qn('w:spacing'), {
            qn('w:line'): '240', qn('w:lineRule'): 'auto',
            qn('w:before'): '0', qn('w:after'): '0'
        })
        pPr.append(spacing)

        for text, font, size in runs:
            run = p.add_run(text)
            run.font.name = font
            run.font.size = Pt(size)
            # Set eastAsia font for CJK
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rPr.makeelement(qn('w:rFonts'), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), font)

    doc.save(path)
    return [t[0] for t in tests]

def measure_line_heights(word, doc_path, descriptions):
    doc = word.Documents.Open(doc_path)
    time.sleep(1)
    doc.Repaginate()
    time.sleep(0.5)

    results = []
    for i, desc in enumerate(descriptions):
        para = doc.Paragraphs(i + 1)
        rng = para.Range
        # Move to start of paragraph
        rng.Collapse(1)  # wdCollapseStart
        y_start = rng.Information(6)  # wdVerticalPositionRelativeToPage

        # Get line height by measuring to next paragraph
        if i + 1 < len(descriptions):
            next_para = doc.Paragraphs(i + 2)
            next_rng = next_para.Range
            next_rng.Collapse(1)
            y_next = next_rng.Information(6)
            line_h = y_next - y_start
        else:
            line_h = None

        results.append((desc, y_start, line_h))

    doc.Close(0)
    return results

def main():
    tmp = os.path.join(tempfile.gettempdir(), "test_mixed_runs.docx")
    print("Creating test document...")
    descriptions = create_test_docx(tmp)

    print("Opening Word COM...")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        results = measure_line_heights(word, tmp, descriptions)

        print("\n=== Mixed Runs Line Height Results ===")
        print(f"{'Description':<35} {'Y pos':>8} {'Line H':>8}")
        print("-" * 55)

        for desc, y, h in results:
            h_str = f"{h:.2f}" if h is not None else "N/A"
            print(f"{desc:<35} {y:8.2f} {h_str:>8}")

        # Analysis
        print("\n=== Analysis ===")
        # Group by test sets
        groups = [
            ("Test1: Calibri sizes", 0, 3),
            ("Test2: YuGothic+Calibri", 3, 3),
            ("Test3: YuGothic+MSGothic", 6, 3),
            ("Test4: Calibri 8+28", 9, 3),
            ("Test5: MSGothic+Calibri24", 12, 3),
            ("Test6: 3 sizes", 15, 2),
        ]
        for name, start, count in groups:
            print(f"\n{name}:")
            for i in range(start, start + count):
                desc, y, h = results[i]
                h_str = f"{h:.2f}pt" if h is not None else "N/A"
                print(f"  {desc}: {h_str}")

            # Check if mixed == max of individuals
            mixed_h = results[start][2]
            individual_hs = [results[start + j][2] for j in range(1, count) if results[start + j][2] is not None]
            if mixed_h is not None and individual_hs:
                max_individual = max(individual_hs)
                diff = abs(mixed_h - max_individual)
                match = "MATCH" if diff < 0.05 else f"DIFF={diff:.2f}pt"
                print(f"  -> Mixed={mixed_h:.2f}, max(individuals)={max_individual:.2f} => {match}")

    finally:
        word.Quit()

    os.unlink(tmp)

if __name__ == "__main__":
    main()
