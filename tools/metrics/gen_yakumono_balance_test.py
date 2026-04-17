"""Test: does <w:balanceSingleByteDoubleByteWidth/> trigger ・ compression?"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data")
)
TEST_TEXT = "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"


def make(label, compat_flags, docgrid_type=None):
    doc = Document()
    settings = doc.settings.element
    for e in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(e)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)

    # Add compat flags (inside w:compat element)
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat"); settings.append(compat)
    for flag in compat_flags:
        el = OxmlElement(f"w:{flag}")
        compat.append(el)

    p = doc.add_paragraph()
    r = p.add_run(TEST_TEXT)
    rel = r._element
    rPr = rel.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); rel.insert(0, rPr)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "ＭＳ ゴシック")
    rFonts.set(qn("w:eastAsia"), "ＭＳ ゴシック")
    rFonts.set(qn("w:hAnsi"), "ＭＳ ゴシック")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)

    # docGrid
    if docgrid_type:
        sectPr = doc.element.body.find(qn("w:sectPr"))
        old = sectPr.find(qn("w:docGrid"))
        if old is not None: sectPr.remove(old)
        dg = OxmlElement("w:docGrid")
        dg.set(qn("w:type"), docgrid_type)
        dg.set(qn("w:linePitch"), "360")
        sectPr.append(dg)

    out = os.path.join(OUT_DIR, f"yakumono_bal_{label}.docx")
    doc.save(out)
    return out


def main():
    # B1: balanceSingleByteDoubleByteWidth only
    print(make("B1_balance", ["balanceSingleByteDoubleByteWidth"]))
    # B2: all d77a compat flags
    d77a_compat = ["spaceForUL", "balanceSingleByteDoubleByteWidth",
                   "doNotLeaveBackslashAlone", "ulTrailSpace",
                   "doNotExpandShiftReturn", "adjustLineHeightInTable",
                   "useFELayout"]
    print(make("B2_full_compat", d77a_compat))
    # B3: full compat + docGrid lines
    print(make("B3_full_and_grid", d77a_compat, docgrid_type="lines"))


if __name__ == "__main__":
    main()
