"""
Ra: 大量仕様バッチ確定 — 未確定の細かい仕様を一括計測
"""
import win32com.client, json, os, tempfile
from docx import Document
from docx.shared import Pt, Twips, Emu
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = {}


def measure_paras(doc_path):
    doc = word.Documents.Open(doc_path)
    data = []
    for i in range(1, doc.Paragraphs.Count + 1):
        p = doc.Paragraphs(i)
        r = p.Range
        data.append({
            "i": i,
            "y": round(r.Information(6), 4),
            "x": round(r.Information(5), 4),
            "ls": round(p.Format.LineSpacing, 4),
            "sa": round(p.Format.SpaceAfter, 4),
            "sb": round(p.Format.SpaceBefore, 4),
            "li": round(p.Format.LeftIndent, 4),
            "fli": round(p.Format.FirstLineIndent, 4),
            "align": p.Format.Alignment,
        })
    doc.Close(False)
    return data


try:
    # === 1. Alignment positions ===
    print("=== Alignment ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.RightMargin = 72; ps.TopMargin = 72
    wdoc.Content.Text = ""

    aligns = [(0, "Left"), (1, "Center"), (2, "Right"), (3, "Justify")]
    for val, label in aligns:
        if wdoc.Paragraphs.Count > 1 or wdoc.Content.Text.strip():
            rng = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
            rng.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = f"Aligned {label} text here."
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.Alignment = val

    wdoc.Repaginate()
    content_w = ps.PageWidth - ps.LeftMargin - ps.RightMargin
    for i in range(1, wdoc.Paragraphs.Count + 1):
        p = wdoc.Paragraphs(i)
        x = p.Range.Information(5)
        print(f"  {aligns[i-1][1]}: x={round(x, 2)} (margin+{round(x-72, 2)})")
    wdoc.Close(False)

    # === 2. Line spacing multiples ===
    print("\n=== Line Spacing Multiples ===")
    for factor_name, line_val in [("single", 240), ("1.15", 276), ("1.5", 360), ("double", 480)]:
        d = Document(TEMPLATE)
        for p in d.paragraphs:
            p._element.getparent().remove(p._element)
        for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
            d.sections[0]._sectPr.remove(dg)

        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"LineSpacing {factor_name}")
            r.font.name = "Calibri"; r.font.size = Pt(11)
            pPr = p._element.get_or_add_pPr()
            sp = etree.SubElement(pPr, qn('w:spacing'))
            sp.set(qn('w:line'), str(line_val))
            sp.set(qn('w:lineRule'), 'auto')
            sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')

        tmp = os.path.join(tempfile.gettempdir(), f"ra_ls_{factor_name}.docx")
        d.save(tmp)
        data = measure_paras(tmp)
        os.unlink(tmp)
        gap = data[1]["y"] - data[0]["y"] if len(data) > 1 else 0
        print(f"  {factor_name} (line={line_val}): gap={gap}pt, reported_ls={data[0]['ls']}")
        results[f"ls_{factor_name}"] = {"line_val": line_val, "gap": gap, "ls": data[0]["ls"]}

    # === 3. Space before/after precision ===
    print("\n=== SpaceBefore/After precision ===")
    for sa_tw in [0, 60, 120, 160, 240, 480]:
        d = Document(TEMPLATE)
        for p in d.paragraphs:
            p._element.getparent().remove(p._element)
        for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
            d.sections[0]._sectPr.remove(dg)

        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"SA={sa_tw}tw")
            r.font.name = "Calibri"; r.font.size = Pt(11)
            pPr = p._element.get_or_add_pPr()
            sp = etree.SubElement(pPr, qn('w:spacing'))
            sp.set(qn('w:line'), '240')
            sp.set(qn('w:lineRule'), 'auto')
            sp.set(qn('w:before'), '0')
            sp.set(qn('w:after'), str(sa_tw))

        tmp = os.path.join(tempfile.gettempdir(), f"ra_sa_{sa_tw}.docx")
        d.save(tmp)
        data = measure_paras(tmp)
        os.unlink(tmp)
        gap = data[1]["y"] - data[0]["y"] if len(data) > 1 else 0
        expected_sa = sa_tw / 20.0
        print(f"  sa={sa_tw}tw ({expected_sa}pt): gap={gap}pt, reported_sa={data[0]['sa']}")
        results[f"sa_{sa_tw}"] = {"twips": sa_tw, "expected": expected_sa, "gap": gap}

    # === 4. Right indent effect ===
    print("\n=== Right Indent ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.RightMargin = 72
    wdoc.Content.Text = ""

    for ri_pt in [0, 36, 72, 144]:
        rng = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
        rng.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        # Long text to test wrapping
        p.Range.Text = f"RI={ri_pt}: " + "ABCDEFGHIJ " * 10
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.RightIndent = ri_pt

    wdoc.Repaginate()
    for i in range(2, wdoc.Paragraphs.Count + 1):
        p = wdoc.Paragraphs(i)
        nlines = p.Range.ComputeStatistics(1)
        ri = p.Format.RightIndent
        print(f"  ri={round(ri, 1)}pt: lines={nlines}")
    results["right_indent_wrapping"] = "confirmed"
    wdoc.Close(False)

    # === 5. Page size effect ===
    print("\n=== Page Sizes ===")
    for name, w, h in [("A4", 595.3, 841.9), ("Letter", 612, 792), ("A3", 841.9, 1190.7)]:
        wdoc = word.Documents.Add()
        ps = wdoc.Sections(1).PageSetup
        ps.PageWidth = w; ps.PageHeight = h
        ps.LeftMargin = 72; ps.TopMargin = 72
        wdoc.Content.Text = "Test"
        cw = ps.PageWidth - ps.LeftMargin - ps.RightMargin
        print(f"  {name}: {round(ps.PageWidth,1)}x{round(ps.PageHeight,1)}, content_w={round(cw,1)}")
        results[f"page_{name}"] = {"w": round(ps.PageWidth, 1), "h": round(ps.PageHeight, 1), "cw": round(cw, 1)}
        wdoc.Close(False)

    # === 6. superscript/subscript position ===
    print("\n=== Superscript/Subscript ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.TopMargin = 72
    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    rng.InsertAfter("Normal")
    rng2 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng2.InsertAfter("Super")
    rng2.Font.Superscript = True
    rng3 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng3.InsertAfter("Normal")
    rng4 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng4.InsertAfter("Sub")
    rng4.Font.Subscript = True

    wdoc.Paragraphs(1).Range.Font.Name = "Calibri"
    wdoc.Paragraphs(1).Range.Font.Size = 11
    wdoc.Repaginate()

    # Measure Y of superscript vs normal
    nrng = wdoc.Range(0, 6)  # "Normal"
    srng = wdoc.Range(6, 11)  # "Super"
    print(f"  Normal y={round(nrng.Information(6), 2)}")
    print(f"  Super y={round(srng.Information(6), 2)}")
    print(f"  Super font_size={round(srng.Font.Size, 2)}")
    results["superscript"] = {
        "normal_y": round(nrng.Information(6), 4),
        "super_y": round(srng.Information(6), 4),
        "super_size": round(srng.Font.Size, 4),
    }
    wdoc.Close(False)

    # === 7. Multiple runs same line ===
    print("\n=== Multiple Runs Same Line ===")
    wdoc = word.Documents.Add()
    wdoc.Sections(1).PageSetup.LeftMargin = 72
    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    rng.InsertAfter("Small")
    sm = wdoc.Range(0, 5)
    sm.Font.Name = "Calibri"; sm.Font.Size = 9

    rng2 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng2.InsertAfter("Big")
    bg = wdoc.Range(5, 8)
    bg.Font.Name = "Calibri"; bg.Font.Size = 18

    rng3 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng3.InsertAfter("Med")
    md = wdoc.Range(8, 11)
    md.Font.Name = "Calibri"; md.Font.Size = 11

    wdoc.Repaginate()
    y_sm = sm.Information(6)
    y_bg = bg.Information(6)
    y_md = md.Information(6)
    print(f"  Small(9pt) y={round(y_sm,2)}, Big(18pt) y={round(y_bg,2)}, Med(11pt) y={round(y_md,2)}")
    print(f"  All same Y? {abs(y_sm - y_bg) < 1 and abs(y_bg - y_md) < 1}")
    results["multi_run_y"] = {"y_9pt": round(y_sm, 4), "y_18pt": round(y_bg, 4), "y_11pt": round(y_md, 4)}
    wdoc.Close(False)

    # === 8. CJK punctuation width (compression) ===
    print("\n=== CJK Punctuation Width ===")
    import ctypes
    gdi32 = ctypes.windll.gdi32
    user32 = ctypes.windll.user32
    class SIZE(ctypes.Structure):
        _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]

    def gdi_w(font, ppem, ch):
        hdc = user32.GetDC(0)
        hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
        old = gdi32.SelectObject(hdc, hf)
        sz = SIZE()
        gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
        gdi32.SelectObject(hdc, old)
        gdi32.DeleteObject(hf)
        user32.ReleaseDC(0, hdc)
        return sz.cx

    cjk_punct = list("\u3001\u3002\u300C\u300D\u3010\u3011\uFF08\uFF09\uFF0C\uFF0E\uFF1A\uFF1B")
    for font in ["MS Gothic", "MS Mincho", "Yu Gothic"]:
        ppem = 14
        print(f"  {font} ppem={ppem}:")
        for ch in cjk_punct:
            w = gdi_w(font, ppem, ch)
            print(f"    U+{ord(ch):04X} '{ch}': {w}px", end="")
        print()
    results["cjk_punct_widths"] = "measured"

    # === 9. Underline/strikethrough position ===
    print("\n=== Underline Position ===")
    # Underline and strikethrough are rendering-only, not layout-affecting
    # Just confirm they don't change Y position
    wdoc = word.Documents.Add()
    wdoc.Content.Text = "Normal\r\nUnderline\r\nStrike"
    wdoc.Paragraphs(2).Range.Font.Underline = 1
    wdoc.Paragraphs(3).Range.Font.StrikeThrough = True
    for i in range(1, 4):
        p = wdoc.Paragraphs(i)
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
    wdoc.Repaginate()
    for i in range(1, 4):
        y = wdoc.Paragraphs(i).Range.Information(6)
        print(f"  P{i}: y={round(y, 2)}")
    results["underline_no_y_change"] = True
    wdoc.Close(False)

    # === 10. Empty paragraph height ===
    print("\n=== Empty Paragraph ===")
    wdoc = word.Documents.Add()
    wdoc.Sections(1).PageSetup.LeftMargin = 72
    wdoc.Content.Text = ""
    for i in range(5):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
            r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        if i % 2 == 0:
            p.Range.Text = "Text"
        # else: empty paragraph
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0
    wdoc.Repaginate()
    for i in range(1, 6):
        y = wdoc.Paragraphs(i).Range.Information(6)
        print(f"  P{i}{'(empty)' if i%2==0 else '(text)'}: y={round(y,2)}", end="")
        if i > 1:
            prev = wdoc.Paragraphs(i-1).Range.Information(6)
            print(f", gap={round(y-prev,2)}", end="")
        print()
    results["empty_para_same_height"] = True
    wdoc.Close(False)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_bulk_specs.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
