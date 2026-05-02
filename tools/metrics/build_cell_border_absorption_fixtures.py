"""
Build cell-border-absorption fixtures for §13.3 refinement.

Spec §13.3 has only 4 data points (border ∈ {0, 4, 12, 24} half-pt). The
relationship is non-linear and non-trivial:
  border=0 (=0pt): text_x=77.0
  border=4 (=2pt): text_x=77.0 (no shift! absorbed)
  border=12 (=6pt): text_x=77.5 (+0.5pt)
  border=24 (=12pt): text_x=78.5 (+1.5pt)

Sweep border widths {1, 2, 3, 4, 6, 8, 10, 12, 16, 20, 24, 32, 40} half-pt
(= {0.5, 1, 1.5, 2, 3, 4, 5, 6, 8, 10, 12, 16, 20} pt) with default cell
padding (4.95pt) and Calibri 11pt body. Vary border on LEFT side only (so
text_x reflects left border + left padding).

Output: output/cell_border_absorption_fixtures/
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output",
                       "cell_border_absorption_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


def set_cell_border(cell, side="left", val="single", sz_halfpt=4, color="000000"):
    """Set a single border on one side of a cell via XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), str(sz_halfpt))
    border.set(qn("w:color"), color)
    border.set(qn("w:space"), "0")
    # Replace if existing
    existing = tcBorders.find(qn(f"w:{side}"))
    if existing is not None:
        tcBorders.remove(existing)
    tcBorders.append(border)


def set_table_cell_margins(table, top_tw=0, left_tw=99, bottom_tw=0, right_tw=99):
    """Set table-level cell margins (tblCellMar) — 99tw ≈ 4.95pt default."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblCellMar = tblPr.find(qn("w:tblCellMar"))
    if tblCellMar is None:
        tblCellMar = OxmlElement("w:tblCellMar")
        tblPr.append(tblCellMar)
    for side, w in (("top", top_tw), ("left", left_tw),
                    ("bottom", bottom_tw), ("right", right_tw)):
        elem = tblCellMar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tblCellMar.append(elem)
        elem.set(qn("w:w"), str(w))
        elem.set(qn("w:type"), "dxa")


def build_fixture(path, *, sz_halfpt):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # 2-row 1-col table; first row has the border, second is reference
    table = doc.add_table(rows=2, cols=1)
    # Set explicit Word default table cell padding (4.95pt = 99tw)
    set_table_cell_margins(table, left_tw=99, right_tw=99)
    # Set table alignment + width
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.width = Pt(400)
            if row_idx == 0:
                # Apply LEFT border with target sz
                set_cell_border(cell, side="left",
                                val="single", sz_halfpt=sz_halfpt)
            for p in cell.paragraphs:
                p.text = f"R{row_idx+1}: text"
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

    doc.save(path)


def main():
    sweep = [0, 1, 2, 3, 4, 6, 8, 10, 12, 16, 20, 24, 32, 40]  # half-pt units
    for sz in sweep:
        sz_pt = sz / 2.0
        name = f"CBA_sz{sz:02d}hp_({sz_pt}pt).docx"
        path = os.path.join(OUT_DIR, name)
        try:
            build_fixture(path, sz_halfpt=sz)
            print(f"  built {name}")
        except Exception as e:
            print(f"  ERR {name}: {e}")
    print(f"\nBuilt {len(sweep)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
