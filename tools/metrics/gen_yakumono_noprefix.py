"""Generate variants WITHOUT prefix — ・ is first char.

Test hypothesis: Word's ・ compression only fires when ・ is the absolute
first character of the paragraph (no Latin prefix before it).
"""
import os
from docx import Document
from docx.shared import Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data")
)

TEST_TEXT = "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"


def set_font(run, family="ＭＳ ゴシック", size_pt=12.0):
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)


def set_compress_punct(doc):
    settings = doc.settings.element
    for e in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(e)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)


def main():
    # Variant A: bare ・ at paragraph start, no prefix
    doc = Document()
    set_compress_punct(doc)
    p = doc.add_paragraph()
    r = p.add_run(TEST_TEXT)
    set_font(r)
    out = os.path.join(OUT_DIR, "yakumono_noprefix_A.docx")
    doc.save(out)
    print(f"[A] {out}")

    # Variant B: paragraph with small number prefix like d77a but not "V: "
    # (d77a's paragraphs have no label prefix; try without any prefix)
    doc = Document()
    set_compress_punct(doc)
    # Add several paragraphs all starting with ・
    for i in range(3):
        p = doc.add_paragraph()
        r = p.add_run(TEST_TEXT)
        set_font(r)
    out = os.path.join(OUT_DIR, "yakumono_noprefix_B.docx")
    doc.save(out)
    print(f"[B] {out}")


if __name__ == "__main__":
    main()
