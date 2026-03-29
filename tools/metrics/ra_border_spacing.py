"""
Ra: 段落ボーダーのspace値がレイアウトにどう影響するか計測
- pBdr/bottom space=N がcursor_yに加算されるか？
- spacing after との相互作用（collapse? 加算?）
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False


def make_doc(border_space=None, border_sz=8, sa_twips=300, sb_twips=480, no_grid=False):
    """Create test doc: P1(Calibri 26pt, optional border) -> P2(Calibri 14pt)"""
    d = Document(TEMPLATE)
    sec = d.sections[0]

    if no_grid:
        sectPr = sec._sectPr
        for dg in sectPr.findall(qn('w:docGrid')):
            sectPr.remove(dg)
        dg = etree.SubElement(sectPr, qn('w:docGrid'))
        dg.set(qn('w:linePitch'), '360')
    # else: keep template grid (lines, pitch=360)

    # Remove default paragraphs
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # P1: Calibri 26pt single, optional border
    p1 = d.add_paragraph()
    r1 = p1.add_run("AAAA")
    r1.font.name = "Calibri"
    r1.font.size = Pt(26)
    pf1 = p1.paragraph_format
    pf1.space_before = Pt(0)
    pf1.space_after = Twips(sa_twips)
    pf1.line_spacing_rule = 0  # single

    if border_space is not None:
        pPr = p1._element.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
        bottom = etree.SubElement(pBdr, qn('w:bottom'))
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(border_sz))
        bottom.set(qn('w:space'), str(border_space))
        bottom.set(qn('w:color'), '4F81BD')

    # P2: Calibri 14pt
    p2 = d.add_paragraph()
    r2 = p2.add_run("BBBB")
    r2.font.name = "Calibri"
    r2.font.size = Pt(14)
    pf2 = p2.paragraph_format
    pf2.space_before = Twips(sb_twips)
    pf2.space_after = Pt(0)
    pf2.line_spacing_rule = 0

    # P3: reference
    p3 = d.add_paragraph()
    r3 = p3.add_run("CCCC")
    r3.font.name = "Calibri"
    r3.font.size = Pt(11)
    pf3 = p3.paragraph_format
    pf3.space_before = Pt(0)
    pf3.space_after = Pt(0)
    pf3.line_spacing_rule = 0

    tmp = os.path.join(tempfile.gettempdir(), 'ra_border_sp.docx')
    d.save(tmp)
    return tmp


def measure_gap(path):
    """Open in Word, measure P1->P2 gap"""
    try:
        wdoc = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
        y1 = wdoc.Paragraphs(1).Range.Information(6)
        y2 = wdoc.Paragraphs(2).Range.Information(6)
        y3 = wdoc.Paragraphs(3).Range.Information(6)
        wdoc.Close(False)
        return {
            "P1_y": round(y1, 4),
            "P2_y": round(y2, 4),
            "P3_y": round(y3, 4),
            "gap12": round(y2 - y1, 4),
            "gap23": round(y3 - y2, 4),
        }
    except Exception as e:
        print(f"  ERROR: {e}", file=sys.stderr)
        try:
            wdoc.Close(False)
        except:
            pass
        return None


results = {}

# === Test 1: Baseline (no border) ===
print("=== Baseline: no border, sa=15pt, sb=24pt ===")
r = measure_gap(make_doc(sa_twips=300, sb_twips=480))
results["no_border_sa15_sb24"] = r
print(f"  gap12={r['gap12']}pt")

# === Test 2: Border space variations ===
print("\n=== Border space variations (sa=15pt, sb=24pt) ===")
for sp in [0, 1, 2, 3, 4, 5, 6, 8, 10, 12, 15, 20, 25, 30]:
    r = measure_gap(make_doc(border_space=sp, sa_twips=300, sb_twips=480))
    results[f"border_sp{sp}"] = r
    print(f"  space={sp:2d}pt  gap12={r['gap12']}pt")

# === Test 3: sa variations with border_space=4 ===
print("\n=== sa variations (border_space=4pt, sb=24pt) ===")
for sa_tw in [0, 40, 80, 100, 160, 200, 240, 300, 400, 480, 600, 800, 1000]:
    sa_pt = sa_tw / 20
    r = measure_gap(make_doc(border_space=4, sa_twips=sa_tw, sb_twips=480))
    results[f"sp4_sa{sa_tw}tw"] = r
    print(f"  sa={sa_pt:5.1f}pt  gap12={r['gap12']}pt")

# === Test 4: sb variations with border_space=4 ===
print("\n=== sb variations (border_space=4pt, sa=15pt) ===")
for sb_tw in [0, 40, 80, 100, 160, 200, 300, 400, 480, 600, 800, 1000]:
    sb_pt = sb_tw / 20
    r = measure_gap(make_doc(border_space=4, sa_twips=300, sb_twips=sb_tw))
    results[f"sp4_sb{sb_tw}tw"] = r
    print(f"  sb={sb_pt:5.1f}pt  gap12={r['gap12']}pt")

# === Test 5: No grid ===
print("\n=== No grid ===")
for sp in [None, 0, 4, 10]:
    label = f"nogrid_sp{sp}" if sp is not None else "nogrid_noborder"
    r = measure_gap(make_doc(border_space=sp, sa_twips=300, sb_twips=480, no_grid=True))
    results[label] = r
    print(f"  border_space={sp}  gap12={r['gap12']}pt")

# === Test 6: Border size (thickness) variations ===
print("\n=== Border size variations (border_space=4, sa=15, sb=24) ===")
for sz in [2, 4, 6, 8, 12, 18, 24, 36, 48]:
    r = measure_gap(make_doc(border_space=4, border_sz=sz, sa_twips=300, sb_twips=480))
    results[f"sp4_sz{sz}"] = r
    print(f"  sz={sz:2d} ({sz/8:.2f}pt)  gap12={r['gap12']}pt")

# === Summary ===
print("\n" + "=" * 70)
print("FULL SUMMARY")
print("=" * 70)
for name, r in results.items():
    if r:
        print(f"  {name:35s}  gap12={r['gap12']:7.2f}  gap23={r['gap23']:7.2f}")

# Save
out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_border_spacing.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nSaved: {out_path}")

word.Quit()
