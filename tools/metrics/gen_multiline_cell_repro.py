"""Generate a minimal repro docx with 8 single-cell tables.

Each table contains 1 paragraph that wraps to N=1..8 lines (MS Mincho 10.5pt).
Test condition: LM0 (no docGrid), same as b35.

Then COM-measure each table's rendered row height to build the formula.
"""
import os
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DOCX = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data",
                 "multiline_cell_repro.docx")
)

# Content width to fit exactly N "full-width MS Mincho 10.5pt chars" per line
# Page width 595pt, margins 70.9pt each side → content ≈ 453pt
# MS Mincho 10.5pt fullwidth char ≈ 10.5pt
# So per line ≈ 43 chars

LINES_TEXT = [
    "あ" * 10 + "。",  # short: wraps to 1 line
    "あ" * 43 + "。" + "い" * 43 + "、",  # 2 lines
    "う" * 43 + "。" + "え" * 43 + "、" + "お" * 30 + "。",  # 3 lines
    "か" * 43 + "。" + "き" * 43 + "、" + "く" * 43 + "。" + "け" * 30 + "、",  # 4
    "さ" * 43 + "。" + "し" * 43 + "、" + "す" * 43 + "。" + "せ" * 43 + "、" + "そ" * 30,  # 5
    "た" * 43 + "。" + "ち" * 43 + "、" + "つ" * 43 + "。" + "て" * 43 + "、" + "と" * 43 + "。" + "な" * 30,  # 6
    "は" * 43 + "。" + "ひ" * 43 + "、" + "ふ" * 43 + "。" + "へ" * 43 + "、" + "ほ" * 43 + "。" + "ま" * 43 + "、" + "み" * 30,  # 7
    "や" * 43 + "。" + "ゆ" * 43 + "、" + "よ" * 43 + "。" + "ら" * 43 + "、" + "り" * 43 + "。" + "る" * 43 + "、" + "れ" * 43 + "。" + "ろ" * 30,  # 8
]


def set_font(run, family="ＭＳ 明朝", size_pt=10.5):
    run.font.name = family
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szCs)


def main():
    doc = Document()
    # LM0 ensures no docGrid
    # python-docx default is fine; we don't add docGrid

    # Header intro
    h = doc.add_paragraph()
    r = h.add_run(f"multiline_cell_repro — {len(LINES_TEXT)} tables, MS Mincho 10.5pt")
    set_font(r, size_pt=10.5)

    for i, text in enumerate(LINES_TEXT, 1):
        doc.add_paragraph(f"--- Table {i} (expected ~{i} lines) ---")
        t = doc.add_table(rows=1, cols=1)
        t.autofit = False
        t.columns[0].width = Twips(9072)  # same as b35 tcW
        cell = t.rows[0].cells[0]
        # Clear cell default paragraph
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, size_pt=10.5)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"[OK] wrote {OUT_DOCX}")
    print(f"     {len(LINES_TEXT)} tables, 1 cell each, MS Mincho 10.5pt text")


if __name__ == "__main__":
    main()
