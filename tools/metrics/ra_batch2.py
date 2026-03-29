"""
Ra: バッチ2 — テキスト描画・フォント解決・特殊文字の仕様確定
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')
gdi32 = ctypes.windll.gdi32
user32 = ctypes.windll.user32

class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]

def gdi_w(font, ppem, ch):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return sz.cx

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = {}

try:
    # === 1. Space character width per font ===
    print("=== Space Width ===")
    for font in ["Calibri", "Arial", "Times New Roman", "MS Gothic", "MS Mincho", "MS UI Gothic"]:
        for fs in [9, 10.5, 11, 12]:
            ppem = round(fs * 96.0 / 72.0)
            w = gdi_w(font, ppem, " ")
            print(f"  {font} {fs}pt(ppem={ppem}): space={w}px = {round(w*72/96, 2)}pt")
            results[f"space_{font}_{fs}"] = w

    # === 2. Bold font width difference ===
    print("\n=== Bold Width Diff ===")
    for font in ["Calibri", "Arial", "MS Gothic"]:
        ppem = 14
        for ch in "AaBb0あ一":
            reg_w = gdi_w(font, ppem, ch)
            hdc = user32.GetDC(0)
            hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 700, 0, 0, 0, 0, 0, 0, 0, 0, font)
            old = gdi32.SelectObject(hdc, hf)
            sz = SIZE()
            gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
            bold_w = sz.cx
            gdi32.SelectObject(hdc, old)
            gdi32.DeleteObject(hf)
            user32.ReleaseDC(0, hdc)
            diff = bold_w - reg_w
            if diff != 0:
                print(f"  {font} '{ch}': reg={reg_w}, bold={bold_w}, diff={diff}px")
    results["bold_diff"] = "measured"

    # === 3. Italic font width ===
    print("\n=== Italic Width ===")
    for font in ["Calibri", "Arial"]:
        ppem = 14
        for ch in "AaBb":
            reg_w = gdi_w(font, ppem, ch)
            hdc = user32.GetDC(0)
            hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 1, 0, 0, 0, 0, 0, 0, 0, font)
            old = gdi32.SelectObject(hdc, hf)
            sz = SIZE()
            gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
            ital_w = sz.cx
            gdi32.SelectObject(hdc, old)
            gdi32.DeleteObject(hf)
            user32.ReleaseDC(0, hdc)
            diff = ital_w - reg_w
            if diff != 0:
                print(f"  {font} '{ch}': reg={reg_w}, italic={ital_w}, diff={diff}px")
            else:
                print(f"  {font} '{ch}': reg=italic={reg_w}px (same)")
    results["italic_width"] = "measured"

    # === 4. East Asian font resolution (rFonts) ===
    print("\n=== East Asian Font Resolution ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # P1: ascii=Calibri, eastAsia=MS Gothic
    p1 = d.add_paragraph()
    r1 = p1.add_run("Hello")
    rPr = r1._element.get_or_add_rPr()
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:eastAsia'), 'MS Gothic')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    r1.font.size = Pt(11)

    r1b = p1.add_run("日本語")
    rPr2 = r1b._element.get_or_add_rPr()
    rFonts2 = etree.SubElement(rPr2, qn('w:rFonts'))
    rFonts2.set(qn('w:ascii'), 'Calibri')
    rFonts2.set(qn('w:eastAsia'), 'MS Gothic')
    rFonts2.set(qn('w:hAnsi'), 'Calibri')
    r1b.font.size = Pt(11)

    tmp = os.path.join(tempfile.gettempdir(), "ra_eastasia.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)

    for i in range(doc.Paragraphs(1).Range.Start, min(doc.Paragraphs(1).Range.End, doc.Paragraphs(1).Range.Start + 10)):
        cr = doc.Range(i, i+1)
        ch = cr.Text
        font = cr.Font.Name
        if ord(ch) not in (13, 7):
            is_cjk = ord(ch) >= 0x3000
            print(f"  '{ch}'(U+{ord(ch):04X}): font={font} {'[CJK->eastAsia]' if is_cjk else '[Latin->ascii]'}")

    doc.Close(False)
    os.unlink(tmp)
    results["eastasia_font_resolution"] = "confirmed"

    # === 5. Theme font resolution ===
    print("\n=== Theme Font ===")
    wdoc = word.Documents.Add()
    p1 = wdoc.Paragraphs(1)
    p1.Range.Text = "Theme test"
    font_name = p1.Range.Font.Name
    font_size = p1.Range.Font.Size
    print(f"  Default font: {font_name}, size={font_size}")
    results["theme_default_font"] = font_name
    results["theme_default_size"] = round(font_size, 2)
    wdoc.Close(False)

    # === 6. Tab character width ===
    print("\n=== Tab Character ===")
    wdoc = word.Documents.Add()
    wdoc.Sections(1).PageSetup.LeftMargin = 72
    wdoc.Content.Text = "A\tB"
    wdoc.Paragraphs(1).Range.Font.Name = "Calibri"
    wdoc.Paragraphs(1).Range.Font.Size = 11
    wdoc.Repaginate()

    a_rng = wdoc.Range(0, 1)
    b_rng = wdoc.Range(2, 3)
    a_x = a_rng.Information(5)
    b_x = b_rng.Information(5)
    tab_advance = round(b_x - a_x, 2)
    print(f"  'A' x={round(a_x,2)}, 'B' x={round(b_x,2)}, tab_advance={tab_advance}pt")
    results["tab_advance"] = tab_advance
    wdoc.Close(False)

    # === 7. Soft hyphen / non-breaking space width ===
    print("\n=== Special Characters ===")
    special = {
        "NBSP": "\u00A0",
        "EN SPACE": "\u2002",
        "EM SPACE": "\u2003",
        "THIN SPACE": "\u2009",
        "ZERO WIDTH SPACE": "\u200B",
        "IDEOGRAPHIC SPACE": "\u3000",
    }
    ppem = 14
    for name, ch in special.items():
        for font in ["Calibri", "MS Gothic"]:
            w = gdi_w(font, ppem, ch)
            print(f"  {name} ({font} ppem={ppem}): {w}px = {round(w*72/96, 2)}pt")
    results["special_chars"] = "measured"

    # === 8. Line height with different fonts on same line ===
    print("\n=== Mixed Font Line Height Detail ===")
    wdoc = word.Documents.Add()
    wdoc.Sections(1).PageSetup.LeftMargin = 72
    # Remove grid
    wdoc.Content.Text = ""
    # P1: Calibri only
    rng = wdoc.Range(0, 0)
    rng.InsertAfter("Calibri only line")
    wdoc.Paragraphs(1).Range.Font.Name = "Calibri"
    wdoc.Paragraphs(1).Range.Font.Size = 11
    wdoc.Paragraphs(1).Format.SpaceBefore = 0
    wdoc.Paragraphs(1).Format.SpaceAfter = 0

    # P2: Arial 14pt only
    r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
    r.InsertParagraphAfter()
    p2 = wdoc.Paragraphs(2)
    p2.Range.Text = "Arial 14pt line"
    p2.Range.Font.Name = "Arial"; p2.Range.Font.Size = 14
    p2.Format.SpaceBefore = 0; p2.Format.SpaceAfter = 0

    # P3: Mixed Calibri 11 + Arial 14
    r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
    r.InsertParagraphAfter()
    p3 = wdoc.Paragraphs(3)
    p3.Range.Text = ""
    r3a = wdoc.Range(p3.Range.Start, p3.Range.Start)
    r3a.InsertAfter("CalSmall ")
    wdoc.Range(p3.Range.Start, p3.Range.Start + 9).Font.Name = "Calibri"
    wdoc.Range(p3.Range.Start, p3.Range.Start + 9).Font.Size = 11
    r3b = wdoc.Range(p3.Range.End - 1, p3.Range.End - 1)
    r3b.InsertAfter("ArialBig")
    wdoc.Range(p3.Range.End - 9, p3.Range.End - 1).Font.Name = "Arial"
    wdoc.Range(p3.Range.End - 9, p3.Range.End - 1).Font.Size = 14
    p3.Format.SpaceBefore = 0; p3.Format.SpaceAfter = 0

    # P4: baseline
    r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
    r.InsertParagraphAfter()
    p4 = wdoc.Paragraphs(4)
    p4.Range.Text = "Baseline"
    p4.Range.Font.Name = "Calibri"; p4.Range.Font.Size = 11
    p4.Format.SpaceBefore = 0; p4.Format.SpaceAfter = 0

    wdoc.Repaginate()
    for i in range(1, 5):
        y = wdoc.Paragraphs(i).Range.Information(6)
        print(f"  P{i}: y={round(y, 2)}", end="")
        if i > 1:
            prev = wdoc.Paragraphs(i-1).Range.Information(6)
            print(f", gap={round(y-prev, 2)}", end="")
        print()
    wdoc.Close(False)

    # === 9. Paragraph alignment values ===
    print("\n=== Alignment Values ===")
    print("  0=Left, 1=Center, 2=Right, 3=Justify, 4=Distribute")
    results["alignment_values"] = {0: "Left", 1: "Center", 2: "Right", 3: "Justify", 4: "Distribute"}

    # === 10. Default paragraph spacing for fresh doc ===
    print("\n=== Fresh Doc Defaults ===")
    wdoc = word.Documents.Add()
    p = wdoc.Paragraphs(1)
    p.Range.Text = "Default paragraph"
    print(f"  Font: {p.Range.Font.Name}, Size: {p.Range.Font.Size}")
    print(f"  SA: {p.Format.SpaceAfter}, SB: {p.Format.SpaceBefore}")
    print(f"  LS: {p.Format.LineSpacing}, Rule: {p.Format.LineSpacingRule}")
    print(f"  LI: {p.Format.LeftIndent}, FLI: {p.Format.FirstLineIndent}")
    print(f"  Align: {p.Format.Alignment}")
    results["fresh_doc_defaults"] = {
        "font": p.Range.Font.Name,
        "size": round(p.Range.Font.Size, 2),
        "sa": round(p.Format.SpaceAfter, 2),
        "sb": round(p.Format.SpaceBefore, 2),
        "ls": round(p.Format.LineSpacing, 2),
        "ls_rule": p.Format.LineSpacingRule,
        "li": round(p.Format.LeftIndent, 2),
        "align": p.Format.Alignment,
    }
    wdoc.Close(False)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch2.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
