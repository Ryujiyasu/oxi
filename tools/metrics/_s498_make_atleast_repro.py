# -*- coding: utf-8 -*-
"""S498 minimal repro: single-cell tables with atLeast row heights > content, to isolate
where Word places the cell content vertically (top? centered? +X?). Several variants:
atLeast in {none, 20, 30, 40, 50}pt, content = one CJK line (MS Mincho 10.5pt), vAlign top.
cp932-safe: CJK via unicode escapes."""
import os
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn

OUT = os.path.join(os.path.dirname(__file__), '..', 'golden-test', 'repros', 'atleast_cellY')
os.makedirs(OUT, exist_ok=True)
CJK = 'あいうえお'  # あいうえお

def set_cell_font(cell, name, sz):
    p = cell.paragraphs[0]
    r = p.add_run(CJK)
    r.font.size = Pt(sz)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): name, qn('w:ascii'): name, qn('w:hAnsi'): name})
    rpr.insert(0, rf)
    return p

def build(atleast_pt):
    doc = Document()
    # a marker paragraph before the table (reference Y)
    doc.add_paragraph('REF')
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    set_cell_font(cell, 'ＭＳ 明朝', 10.5)  # ＭＳ 明朝
    if atleast_pt:
        row = t.rows[0]
        row.height = Pt(atleast_pt)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    doc.add_paragraph('AFTER')
    name = 'atleast_%s.docx' % (atleast_pt if atleast_pt else 'none')
    path = os.path.join(OUT, name)
    doc.save(path)
    return path

for a in [None, 20, 30, 40, 50]:
    print('built', build(a))
print('DONE')
