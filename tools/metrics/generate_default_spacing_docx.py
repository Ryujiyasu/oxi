"""
Generate test .docx files with DEFAULT line spacing (no explicit setting).
This captures Word's natural auto-spacing behavior per font.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import json

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "docx_tests_default")

FONTS = [
    ("游明朝", "yumin"),
    ("游ゴシック", "yugothic"),
    ("Century", "century"),
    ("Times New Roman", "tnr"),
    ("Calibri", "calibri"),
    ("Arial", "arial"),
    ("ＭＳ 明朝", "msmincho"),
    ("ＭＳ ゴシック", "msgothic"),
]

SIZES = [10.5, 11, 12, 14]

# Long enough text to force multiple lines
TEXT_JA = "吾輩は猫である。名前はまだ無い。どこで生れたかとんと見当がつかぬ。何でも薄暗いじめじめした所でニャーニャー泣いていた事だけは記憶している。"
TEXT_EN = "The quick brown fox jumps over the lazy dog. Pack my box with five dozen liquor jugs. How vexingly quick daft zebras jump. The five boxing wizards jump quickly."


def create_docx(font: str, size: float, output_path: str):
    """Create a docx with NO explicit line spacing — pure Word defaults."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Pt(595)
    section.page_height = Pt(842)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    def add_styled_para(text, font_name, font_size):
        para = doc.add_paragraph()
        # NO line_spacing setting — let Word use defaults
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        # Set East Asian font
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            from lxml import etree
            r_fonts = etree.SubElement(r_pr, qn("w:rFonts"))
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        return para

    # Japanese paragraph (multiple lines)
    add_styled_para(TEXT_JA, font, size)

    # English paragraph (multiple lines)
    add_styled_para(TEXT_EN, font, size)

    doc.save(output_path)


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    manifest = []

    for font_name, font_id in FONTS:
        for size in SIZES:
            filename = f"{font_id}_{size}pt.docx"
            path = os.path.join(OUTPUT_DIR, filename)
            create_docx(font_name, size, path)
            manifest.append({
                "filename": filename,
                "font": font_name,
                "font_id": font_id,
                "size_pt": size,
            })
            print(f"  {filename}")

    with open(os.path.join(OUTPUT_DIR, "manifest.json"), "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    print(f"\n{len(manifest)} files generated in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
