"""Targeted repros to find when Word applies spacing collapse vs when it doesn't.

Motivation: S1-S6 (simple Table Grid, no pStyle) showed collapse applies.
But real 29dc6e row (pStyle='ac' on paragraphs) measured 37.5pt which matches
NON-collapse prediction (37.05pt), not collapse (33pt).

Isolate which pPr attribute of pStyle 'ac' disables collapse:
- widowControl=0
- wordWrap=0
- autoSpaceDE=0
- autoSpaceDN=0
- adjustRightInd=0
- jc=both

C1: baseline like S1 (no pStyle, no extra attrs) — expect collapse (33pt)
C2: + widowControl=0
C3: + wordWrap=0
C4: + autoSpaceDE=0
C5: + autoSpaceDN=0
C6: + jc=both
C7: ALL of above (mimics pStyle='ac') — hypothesis: no-collapse (37pt)
"""
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.abspath("tools/metrics/collapse_cond_repro")
os.makedirs(OUT_DIR, exist_ok=True)


def add_para_with_spacing(cell, text, sa_tw, sb_tw, extras):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for name in ["widowControl", "wordWrap", "autoSpaceDE", "autoSpaceDN",
                 "adjustRightInd", "jc", "spacing"]:
        existing = pPr.find(qn(f"w:{name}"))
        if existing is not None:
            pPr.remove(existing)
    # Add configured extras in order
    if extras.get("widowControl") == "0":
        e = OxmlElement("w:widowControl"); e.set(qn("w:val"), "0"); pPr.append(e)
    if extras.get("wordWrap") == "0":
        e = OxmlElement("w:wordWrap"); e.set(qn("w:val"), "0"); pPr.append(e)
    if extras.get("autoSpaceDE") == "0":
        e = OxmlElement("w:autoSpaceDE"); e.set(qn("w:val"), "0"); pPr.append(e)
    if extras.get("autoSpaceDN") == "0":
        e = OxmlElement("w:autoSpaceDN"); e.set(qn("w:val"), "0"); pPr.append(e)
    if extras.get("adjustRightInd") == "0":
        e = OxmlElement("w:adjustRightInd"); e.set(qn("w:val"), "0"); pPr.append(e)
    if extras.get("jc"):
        e = OxmlElement("w:jc"); e.set(qn("w:val"), extras["jc"]); pPr.append(e)
    # Spacing always
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(sb_tw))
    sp.set(qn("w:after"), str(sa_tw))
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "ＭＳ 明朝"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hAnsi"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    return p


def make(name, extras):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    # Delete default paragraph
    default_p = cell.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    add_para_with_spacing(cell, "代表者又は", 87, 87, extras)
    add_para_with_spacing(cell, "管理人の氏名（フリガナ）", 87, 87, extras)
    for i in range(8):
        r = tbl.add_row()
        r.cells[0].text = f"ref row {i+1}"
    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    make("C1_baseline", {})
    make("C2_widowCtrl0", {"widowControl": "0"})
    make("C3_wordWrap0", {"wordWrap": "0"})
    make("C4_autoSpaceDE0", {"autoSpaceDE": "0"})
    make("C5_autoSpaceDN0", {"autoSpaceDN": "0"})
    make("C6_jcBoth", {"jc": "both"})
    make("C7_all_acstyle", {
        "widowControl": "0", "wordWrap": "0",
        "autoSpaceDE": "0", "autoSpaceDN": "0",
        "adjustRightInd": "0", "jc": "both",
    })
