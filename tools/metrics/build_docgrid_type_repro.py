"""Test if docGrid type (lines vs linesAndChars) affects first-para sb suppression.

Hypothesis: type="linesAndChars" disables first-sb suppression in cells.

Data so far:
- BL1 (type=lines, before=87 only): 33pt = collapse + first-sb supp
- BL2 (type=lines, before+beforeLines=30): 35pt = collapse + first-sb supp
- 29dc6e (type=linesAndChars, linePitch=292, before=87+beforeLines=30): 37.5pt = collapse only (NO first-sb supp)

Test with type=linesAndChars:
- T1: linesAndChars, linePitch=360, before=87 only
- T2: linesAndChars, linePitch=292, before=87 only
- T3: linesAndChars, linePitch=292, before+beforeLines=30 (mimics 29dc6e)
- T4: lines, linePitch=292, before+beforeLines=30 (control: same as 29dc6e but type=lines)
"""
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.abspath("tools/metrics/docgrid_type_repro")
os.makedirs(OUT_DIR, exist_ok=True)


def mk_para(cell, text, sb_tw, sa_tw, sbLines=0, saLines=0):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for name in ["spacing"]:
        ex = pPr.find(qn(f"w:{name}"))
        if ex is not None: pPr.remove(ex)
    sp = OxmlElement("w:spacing")
    if sbLines: sp.set(qn("w:beforeLines"), str(sbLines))
    if sb_tw: sp.set(qn("w:before"), str(sb_tw))
    if saLines: sp.set(qn("w:afterLines"), str(saLines))
    if sa_tw: sp.set(qn("w:after"), str(sa_tw))
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "ＭＳ 明朝"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    return p


def make(name, grid_type, line_pitch, para_kwargs):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"

    sectPr = sec._sectPr
    for ex in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(ex)
    dg = OxmlElement("w:docGrid")
    dg.set(qn("w:type"), grid_type)
    dg.set(qn("w:linePitch"), str(line_pitch))
    if grid_type == "linesAndChars":
        dg.set(qn("w:charSpace"), "1453")
    sectPr.append(dg)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    default_p = cell.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    mk_para(cell, "para1", **para_kwargs)
    mk_para(cell, "para2", **para_kwargs)
    for i in range(6):
        r = tbl.add_row()
        r.cells[0].text = f"ref{i+1}"
    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    make("T1_LAC_linePitch360_before87only",
         "linesAndChars", 360, dict(sb_tw=87, sa_tw=87))
    make("T2_LAC_linePitch292_before87only",
         "linesAndChars", 292, dict(sb_tw=87, sa_tw=87))
    make("T3_LAC_linePitch292_fullmimic",
         "linesAndChars", 292, dict(sb_tw=87, sa_tw=87, sbLines=30, saLines=30))
    make("T4_LINES_linePitch292_fullmimic",
         "lines", 292, dict(sb_tw=87, sa_tw=87, sbLines=30, saLines=30))
