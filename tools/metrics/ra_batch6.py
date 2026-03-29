"""
Ra: バッチ6 — フォントメトリクス詳細、CJK文字分類境界、テーブルセル幅分配、
    段落ボーダー幅、spaceBefore先頭抑制詳細、テキスト色解決
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

gdi32 = ctypes.windll.gdi32
user32 = ctypes.windll.user32
class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]
class TEXTMETRIC(ctypes.Structure):
    _fields_ = [
        ("tmHeight", ctypes.c_long), ("tmAscent", ctypes.c_long),
        ("tmDescent", ctypes.c_long), ("tmInternalLeading", ctypes.c_long),
        ("tmExternalLeading", ctypes.c_long), ("tmAveCharWidth", ctypes.c_long),
        ("tmMaxCharWidth", ctypes.c_long), ("tmWeight", ctypes.c_long),
        ("tmOverhang", ctypes.c_long), ("tmDigitizedAspectX", ctypes.c_long),
        ("tmDigitizedAspectY", ctypes.c_long), ("tmFirstChar", ctypes.c_wchar),
        ("tmLastChar", ctypes.c_wchar), ("tmDefaultChar", ctypes.c_wchar),
        ("tmBreakChar", ctypes.c_wchar), ("tmItalic", ctypes.c_byte),
        ("tmUnderlined", ctypes.c_byte), ("tmStruckOut", ctypes.c_byte),
        ("tmPitchAndFamily", ctypes.c_byte), ("tmCharSet", ctypes.c_byte),
    ]

def get_tm(font, ppem):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    tm = TEXTMETRIC()
    gdi32.GetTextMetricsW(hdc, ctypes.byref(tm))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return tm

def gdi_w(font, ppem, ch):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return sz.cx

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = {}

try:
    # === 1. CJK character classification boundary ===
    print("=== CJK Char Boundary ===")
    # Which exact Unicode ranges are treated as CJK for font resolution?
    ppem = 14
    boundary_tests = [
        (0x2F00, 0x2FFF, "Kangxi Radicals"),
        (0x3000, 0x303F, "CJK Symbols"),
        (0x3040, 0x309F, "Hiragana"),
        (0x30A0, 0x30FF, "Katakana"),
        (0x3100, 0x312F, "Bopomofo"),
        (0x3200, 0x32FF, "Enclosed CJK"),
        (0x3400, 0x4DBF, "CJK Ext A"),
        (0x4E00, 0x9FFF, "CJK Unified"),
        (0xF900, 0xFAFF, "CJK Compat"),
        (0xFE30, 0xFE4F, "CJK Compat Forms"),
        (0xFF00, 0xFF5E, "Fullwidth ASCII"),
        (0xFF5F, 0xFFEF, "Halfwidth/Fullwidth"),
        (0x2000, 0x206F, "General Punct"),
        (0x2100, 0x214F, "Letterlike Symbols"),
        (0x2190, 0x21FF, "Arrows"),
        (0x2500, 0x257F, "Box Drawing"),
        (0x2580, 0x259F, "Block Elements"),
        (0x25A0, 0x25FF, "Geometric Shapes"),
        (0x2600, 0x26FF, "Misc Symbols"),
    ]

    for start, end, name in boundary_tests:
        # Test first char of range
        cp = start
        cal_w = gdi_w("Calibri", ppem, chr(cp))
        uig_w = gdi_w("MS UI Gothic", ppem, chr(cp))
        is_fb = cal_w == uig_w
        print(f"  U+{cp:04X} ({name}): Cal={cal_w}, MSUIG={uig_w}, {'FB' if is_fb else 'OWN'}")
        results[f"boundary_{name}"] = {"cp": cp, "fallback": is_fb}

    # === 2. GDI ascent for baseline calculation ===
    print("\n=== Ascent for Baseline ===")
    for font in ["Calibri", "Arial", "MS Gothic", "Yu Gothic", "Meiryo"]:
        for fs in [9, 10.5, 11]:
            ppem = round(fs * 96.0 / 72.0)
            tm = get_tm(font, ppem)
            asc_pt = round(tm.tmAscent * 72 / 96, 4)
            des_pt = round(tm.tmDescent * 72 / 96, 4)
            h_pt = round(tm.tmHeight * 72 / 96, 4)
            print(f"  {font} {fs}pt: asc={asc_pt}pt, des={des_pt}pt, h={h_pt}pt")

    # === 3. Column break character in OOXML ===
    print("\n=== Column Break Char ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    p = d.add_paragraph()
    r = p.add_run("Before break")
    r.font.name = "Calibri"; r.font.size = Pt(11)
    # Add column break via XML
    br = etree.SubElement(r._element, qn('w:br'))
    br.set(qn('w:type'), 'column')
    r2 = p.add_run("After break")
    r2.font.name = "Calibri"; r2.font.size = Pt(11)

    tmp = os.path.join(tempfile.gettempdir(), "ra_colbreak.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        p_doc = doc.Paragraphs(i)
        text = p_doc.Range.Text.strip()[:40]
        y = p_doc.Range.Information(6)
        print(f"  P{i}: y={round(y,2)} [{text}]")
    doc.Close(False); os.unlink(tmp)
    results["column_break_xml"] = "w:br type=column"

    # === 4. Text color resolution (theme vs explicit) ===
    print("\n=== Text Color ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Default color
    p1 = d.add_paragraph()
    r1 = p1.add_run("Default color")
    r1.font.name = "Calibri"; r1.font.size = Pt(11)

    # Explicit red
    p2 = d.add_paragraph()
    r2 = p2.add_run("Red text")
    r2.font.name = "Calibri"; r2.font.size = Pt(11)
    rPr = r2._element.get_or_add_rPr()
    color = etree.SubElement(rPr, qn('w:color'))
    color.set(qn('w:val'), 'FF0000')

    # Theme color
    p3 = d.add_paragraph()
    r3 = p3.add_run("Theme color")
    r3.font.name = "Calibri"; r3.font.size = Pt(11)
    rPr3 = r3._element.get_or_add_rPr()
    color3 = etree.SubElement(rPr3, qn('w:color'))
    color3.set(qn('w:val'), '000000')
    color3.set(qn('w:themeColor'), 'text1')

    tmp = os.path.join(tempfile.gettempdir(), "ra_color.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        p_doc = doc.Paragraphs(i)
        c = p_doc.Range.Font.Color
        print(f"  P{i}: color={c} (hex={c:06X} if positive)")
    doc.Close(False); os.unlink(tmp)

    # === 5. Table cell width distribution (auto) ===
    print("\n=== Table Auto Width Distribution ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.RightMargin = 72
    content_w = ps.PageWidth - ps.LeftMargin - ps.RightMargin

    wdoc.Content.Text = ""
    rng = wdoc.Range(0, 0)
    tbl = wdoc.Tables.Add(rng, 2, 4)
    tbl.Borders.Enable = True

    # Put varying content
    tbl.Cell(1,1).Range.Text = "Short"
    tbl.Cell(1,2).Range.Text = "A much longer text in this cell"
    tbl.Cell(1,3).Range.Text = "Med"
    tbl.Cell(1,4).Range.Text = "X"

    for r in range(1, 3):
        for c in range(1, 5):
            tbl.Cell(r, c).Range.Font.Name = "Calibri"
            tbl.Cell(r, c).Range.Font.Size = 11

    tbl.AutoFitBehavior(1)  # AutoFit to content
    wdoc.Repaginate()

    print(f"  Content width: {round(content_w, 1)}pt")
    total = 0
    for c in range(1, 5):
        w = round(tbl.Columns(c).Width, 2)
        total += w
        print(f"  Col {c}: {w}pt")
    print(f"  Total: {round(total, 2)}pt")
    results["autofit_content_dist"] = {c: round(tbl.Columns(c).Width, 4) for c in range(1, 5)}
    wdoc.Close(False)

    # === 6. SpaceBefore suppression details ===
    print("\n=== SpaceBefore Suppression ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.TopMargin = 72; ps.BottomMargin = 72
    wdoc.Content.Text = ""

    # Fill page 1
    for i in range(38):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
            r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = f"F{i+1}"
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0

    # Page 2 first para with large spaceBefore
    r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
    r.InsertParagraphAfter()
    p_top = wdoc.Paragraphs(wdoc.Paragraphs.Count)
    p_top.Range.Text = "Page2 First (sb=24)"
    p_top.Range.Font.Name = "Calibri"; p_top.Range.Font.Size = 11
    p_top.Format.SpaceBefore = 24

    # Page 2 second para with spaceBefore
    r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
    r.InsertParagraphAfter()
    p_second = wdoc.Paragraphs(wdoc.Paragraphs.Count)
    p_second.Range.Text = "Page2 Second (sb=24)"
    p_second.Range.Font.Name = "Calibri"; p_second.Range.Font.Size = 11
    p_second.Format.SpaceBefore = 24

    wdoc.Repaginate()

    p2_first = wdoc.Paragraphs(39)
    p2_second = wdoc.Paragraphs(40)
    y_first = p2_first.Range.Information(6)
    y_second = p2_second.Range.Information(6)
    pg_first = p2_first.Range.Information(3)
    pg_second = p2_second.Range.Information(3)
    print(f"  Page2 P1 (sb=24): y={round(y_first,2)}, page={pg_first}")
    print(f"  Page2 P2 (sb=24): y={round(y_second,2)}, page={pg_second}")
    print(f"  Gap: {round(y_second - y_first, 2)}")
    print(f"  P1 y matches topMargin? {abs(y_first - 74.5) < 1}")
    results["sb_suppression"] = {
        "p2_first_y": round(y_first, 4),
        "p2_second_y": round(y_second, 4),
        "gap": round(y_second - y_first, 4),
    }
    wdoc.Close(False)

    # === 7. Paragraph border width measurement ===
    print("\n=== Paragraph Border Width ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    for sz_val in [2, 4, 6, 12, 18, 24]:
        p = d.add_paragraph()
        r = p.add_run(f"Border sz={sz_val}")
        r.font.name = "Calibri"; r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
        bottom = etree.SubElement(pBdr, qn('w:bottom'))
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(sz_val))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')

    tmp = os.path.join(tempfile.gettempdir(), "ra_pbdr_sz.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        y = doc.Paragraphs(i).Range.Information(6)
        if i > 1:
            prev_y = doc.Paragraphs(i-1).Range.Information(6)
            gap = round(y - prev_y, 2)
            print(f"  P{i}: gap={gap}")
    doc.Close(False); os.unlink(tmp)

    # === 8. Ascent ratio for baseline rendering ===
    print("\n=== Baseline Ratio ===")
    for font in ["Calibri", "Arial", "MS Gothic", "Yu Gothic"]:
        ppem = 14
        tm = get_tm(font, ppem)
        ratio = round(tm.tmAscent / tm.tmHeight, 4)
        print(f"  {font}: asc/height = {tm.tmAscent}/{tm.tmHeight} = {ratio}")
        results[f"baseline_ratio_{font.replace(' ','_')}"] = ratio

    # === 9. En-dash / Em-dash width ===
    print("\n=== Dash Widths ===")
    dashes = {
        "hyphen": 0x002D,
        "en-dash": 0x2013,
        "em-dash": 0x2014,
        "minus": 0x2212,
        "figure-dash": 0x2012,
    }
    for font in ["Calibri", "Arial", "MS Gothic"]:
        ppem = 14
        for name, cp in dashes.items():
            w = gdi_w(font, ppem, chr(cp))
            print(f"  {font} {name}(U+{cp:04X}): {w}px")

    # === 10. Word spacing in justify ===
    print("\n=== Word Spacing in Justify ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.RightMargin = 72
    wdoc.Content.Text = ""

    # Two-word line that spans full width
    rng = wdoc.Range(0, 0)
    rng.InsertAfter("Word spacing test with several words to fill the entire line width completely now.")
    p = wdoc.Paragraphs(1)
    p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
    p.Format.Alignment = 3  # Justify

    # Add second line so first line gets justified
    rng2 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
    rng2.InsertParagraphAfter()
    p2 = wdoc.Paragraphs(2)
    p2.Range.Text = "Second line."
    p2.Range.Font.Name = "Calibri"; p2.Range.Font.Size = 11

    wdoc.Repaginate()

    # Measure space widths in justified line
    prng = wdoc.Paragraphs(1).Range
    space_positions = []
    for ci in range(prng.Start, min(prng.End, prng.Start + 100)):
        cr = wdoc.Range(ci, ci + 1)
        if cr.Text == " ":
            x = cr.Information(5)
            space_positions.append(round(x, 2))

    # Natural space width
    natural_sw = gdi_w("Calibri", 15, " ")
    natural_sw_pt = round(natural_sw * 72 / 96, 2)

    if len(space_positions) >= 2:
        # Measure gap around first space
        before_space = wdoc.Range(prng.Start + 3, prng.Start + 4)  # 'd' of "Word"
        after_space = wdoc.Range(prng.Start + 5, prng.Start + 6)  # 's' of "spacing"
        before_x = before_space.Information(5)
        after_x = after_space.Information(5)
        justified_space = round(after_x - before_x, 2)
        # Get char width of 'd'
        d_w = gdi_w("Calibri", 15, "d")
        actual_space_w = round(justified_space - d_w * 72 / 96, 2)
        print(f"  Natural space: {natural_sw}px = {natural_sw_pt}pt")
        print(f"  Justified 'd'->'s' gap: {justified_space}pt, minus 'd'({round(d_w*72/96,2)}pt) = space={actual_space_w}pt")

    wdoc.Close(False)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch6.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
