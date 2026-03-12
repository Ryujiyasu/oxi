"""
Generate test .docx files with isolated line spacing settings to derive Word's formula.
Three sets:
  A) w:line="240" (single) + snapToGrid=false  → base font line height
  B) w:line="276" (1.15x)  + snapToGrid=false  → 1.15x multiplier effect
  C) w:line="240" (single) + snapToGrid=true    → grid snapping effect
"""
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree
import os
import json

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "docx_tests_isolated")

FONTS = [
    ("游明朝", "yumin"),
    ("游ゴシック", "yugothic"),
    ("Century", "century"),
    ("Times New Roman", "tnr"),
    ("Calibri", "calibri"),
    ("Arial", "arial"),
    ("ＭＳ 明朝", "msmincho"),
    ("ＭＳ ゴシック", "msgothic"),
]

SIZES = [10.5, 11, 12, 14]

TEXT_JA = "吾輩は猫である。名前はまだ無い。どこで生れたかとんと見当がつかぬ。何でも薄暗いじめじめした所でニャーニャー泣いていた事だけは記憶している。"
TEXT_EN = "The quick brown fox jumps over the lazy dog. Pack my box with five dozen liquor jugs. How vexingly quick daft zebras jump. The five boxing wizards jump quickly."

MODES = [
    ("single_nogrid", 240, False, False),   # Single spacing, no grid
    ("115_nogrid",    276, False, False),    # 1.15x spacing, no grid
    ("single_grid",  240, True,  True),     # Single spacing, with grid
    ("default",      None, True, True),     # No explicit line, inherit from style (with grid)
]


def create_docx(font: str, size: float, line_val, snap_to_grid: bool, use_grid: bool, output_path: str):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Pt(595)
    section.page_height = Pt(842)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    # Control docGrid
    sect_pr = section._sectPr
    doc_grid = sect_pr.find(qn("w:docGrid"))
    if doc_grid is None:
        doc_grid = etree.SubElement(sect_pr, qn("w:docGrid"))

    if use_grid:
        doc_grid.set(qn("w:type"), "lines")
        doc_grid.set(qn("w:linePitch"), "360")
    else:
        doc_grid.set(qn("w:type"), "default")
        # Remove linePitch
        if qn("w:linePitch") in doc_grid.attrib:
            del doc_grid.attrib[qn("w:linePitch")]

    def add_para(text, font_name, font_size):
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # Set explicit line spacing
        if line_val is not None:
            # Use OxmlElement to set exact w:line and w:lineRule
            p_pr = para._element.get_or_add_pPr()
            spacing = p_pr.find(qn("w:spacing"))
            if spacing is None:
                spacing = etree.SubElement(p_pr, qn("w:spacing"))
            spacing.set(qn("w:line"), str(line_val))
            spacing.set(qn("w:lineRule"), "auto")

        # Control snapToGrid
        p_pr = para._element.get_or_add_pPr()
        snap = p_pr.find(qn("w:snapToGrid"))
        if snap is None:
            snap = etree.SubElement(p_pr, qn("w:snapToGrid"))
        snap.set(qn("w:val"), "1" if snap_to_grid else "0")

        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)

        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = etree.SubElement(r_pr, qn("w:rFonts"))
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)

    add_para(TEXT_JA, font, size)
    add_para(TEXT_EN, font, size)

    doc.save(output_path)


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    manifest = []

    for mode_name, line_val, snap, grid in MODES:
        for font_name, font_id in FONTS:
            for size in SIZES:
                filename = f"{mode_name}_{font_id}_{size}pt.docx"
                path = os.path.join(OUTPUT_DIR, filename)
                create_docx(font_name, size, line_val, snap, grid, path)
                manifest.append({
                    "filename": filename,
                    "font": font_name,
                    "font_id": font_id,
                    "size_pt": size,
                    "mode": mode_name,
                    "line_val": line_val,
                    "snap_to_grid": snap,
                    "doc_grid": grid,
                })
                print(f"  {filename}")

    with open(os.path.join(OUTPUT_DIR, "manifest.json"), "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    print(f"\n{len(manifest)} files generated in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
