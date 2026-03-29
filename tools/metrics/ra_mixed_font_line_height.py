"""
Ra: Mixed font run (CJK + Latin 混在行) の行高さをCOM計測で確定
- CJK(MS Gothic) + Latin(Calibri) 混在行の行高さ
- max(各runの行高さ) で合っているか
- grid snap の適用タイミング (snap前にmax? max後にsnap?)
- フォントサイズが異なる混在の挙動
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def make_doc(scenario, no_grid=False, grid_pitch=360):
    d = Document(TEMPLATE)
    sec = d.sections[0]
    sectPr = sec._sectPr

    if no_grid:
        for dg in sectPr.findall(qn('w:docGrid')):
            sectPr.remove(dg)
    else:
        for dg in sectPr.findall(qn('w:docGrid')):
            sectPr.remove(dg)
        dg = etree.SubElement(sectPr, qn('w:docGrid'))
        dg.set(qn('w:type'), 'lines')
        dg.set(qn('w:linePitch'), str(grid_pitch))

    # Remove default paragraphs
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    if scenario == "latin_only":
        # Calibri 10.5pt x 3 paragraphs
        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"Latin only line {i+1} ABCDEFG")
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    elif scenario == "cjk_only":
        # MS Gothic 10.5pt x 3 paragraphs
        for i in range(3):
            p = d.add_paragraph()
            r = p.add_run(f"日本語のみ行{i+1}あいうえお")
            r.font.name = "MS Gothic"
            r.font.size = Pt(10.5)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    elif scenario == "mixed_same_size":
        # Mixed: Calibri + MS Gothic, both 10.5pt x 3 paragraphs
        for i in range(3):
            p = d.add_paragraph()
            r1 = p.add_run("Hello ")
            r1.font.name = "Calibri"
            r1.font.size = Pt(10.5)
            r2 = p.add_run("日本語")
            r2.font.name = "MS Gothic"
            r2.font.size = Pt(10.5)
            r3 = p.add_run(f" World{i+1}")
            r3.font.name = "Calibri"
            r3.font.size = Pt(10.5)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    elif scenario == "mixed_diff_size":
        # Mixed: Calibri 10.5pt + MS Gothic 14pt
        for i in range(3):
            p = d.add_paragraph()
            r1 = p.add_run("Hello ")
            r1.font.name = "Calibri"
            r1.font.size = Pt(10.5)
            r2 = p.add_run("日本語")
            r2.font.name = "MS Gothic"
            r2.font.size = Pt(14)
            r3 = p.add_run(f" World{i+1}")
            r3.font.name = "Calibri"
            r3.font.size = Pt(10.5)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    elif scenario == "mixed_latin_larger":
        # Mixed: Calibri 14pt + MS Gothic 10.5pt (Latin larger)
        for i in range(3):
            p = d.add_paragraph()
            r1 = p.add_run("Hello ")
            r1.font.name = "Calibri"
            r1.font.size = Pt(14)
            r2 = p.add_run("日本語")
            r2.font.name = "MS Gothic"
            r2.font.size = Pt(10.5)
            r3 = p.add_run(f" World{i+1}")
            r3.font.name = "Calibri"
            r3.font.size = Pt(14)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    elif scenario == "mixed_meiryo":
        # Mixed: Calibri 10.5pt + Meiryo 10.5pt (Meiryo has large ascender)
        for i in range(3):
            p = d.add_paragraph()
            r1 = p.add_run("Hello ")
            r1.font.name = "Calibri"
            r1.font.size = Pt(10.5)
            r2 = p.add_run("メイリオ")
            r2.font.name = "Meiryo"
            r2.font.size = Pt(10.5)
            r3 = p.add_run(f" World{i+1}")
            r3.font.name = "Calibri"
            r3.font.size = Pt(10.5)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    return d


def measure(doc_path, scenario, grid_info=""):
    doc = word.Documents.Open(doc_path)
    try:
        data = {"scenario": scenario, "grid": grid_info, "paragraphs": []}
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            rng = para.Range
            y = rng.Information(6)  # wdVerticalPositionRelativeToPage
            ls = para.Format.LineSpacing
            data["paragraphs"].append({
                "index": i,
                "y_pt": round(y, 4),
                "line_spacing_pt": round(ls, 4),
                "text": rng.Text.strip()[:40]
            })
        # Compute gaps (= effective line height since sa=sb=0)
        for i in range(1, len(data["paragraphs"])):
            gap = data["paragraphs"][i]["y_pt"] - data["paragraphs"][i-1]["y_pt"]
            data["paragraphs"][i]["line_height"] = round(gap, 4)
        return data
    finally:
        doc.Close(False)


scenarios = ["latin_only", "cjk_only", "mixed_same_size", "mixed_diff_size", "mixed_latin_larger", "mixed_meiryo"]

try:
    # Test with grid
    print("=" * 60)
    print("WITH GRID (pitch=360 twips = 18pt)")
    print("=" * 60)
    for sc in scenarios:
        d = make_doc(sc, no_grid=False, grid_pitch=360)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_mixed_{sc}.docx")
        d.save(tmp)
        data = measure(tmp, sc, "lines/360")
        results.append(data)
        print(f"\n--- {sc} ---")
        for p in data["paragraphs"]:
            lh_str = f"  line_height={p.get('line_height', '-')}" if 'line_height' in p else ""
            print(f"  P{p['index']}: y={p['y_pt']}pt, ls={p['line_spacing_pt']}pt{lh_str}")
        os.unlink(tmp)

    # Test without grid
    print("\n" + "=" * 60)
    print("WITHOUT GRID (no docGrid)")
    print("=" * 60)
    for sc in scenarios:
        d = make_doc(sc, no_grid=True)
        tmp = os.path.join(tempfile.gettempdir(), f"ra_mixed_ng_{sc}.docx")
        d.save(tmp)
        data = measure(tmp, sc, "no_grid")
        results.append(data)
        print(f"\n--- {sc} (no grid) ---")
        for p in data["paragraphs"]:
            lh_str = f"  line_height={p.get('line_height', '-')}" if 'line_height' in p else ""
            print(f"  P{p['index']}: y={p['y_pt']}pt, ls={p['line_spacing_pt']}pt{lh_str}")
        os.unlink(tmp)

finally:
    word.Quit()

# Save results
out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_mixed_font_line_height.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)

print(f"\nResults saved to {out_path}")

# Analysis
print("\n=== ANALYSIS ===")
grid_results = [r for r in results if r["grid"] == "lines/360"]
no_grid_results = [r for r in results if r["grid"] == "no_grid"]

for label, res_list in [("GRID", grid_results), ("NO GRID", no_grid_results)]:
    print(f"\n{label}:")
    for r in res_list:
        if len(r["paragraphs"]) >= 2:
            lh = r["paragraphs"][1].get("line_height", "N/A")
            print(f"  {r['scenario']}: line_height = {lh}pt")

# Check if mixed = max(latin, cjk)
for label, res_list in [("GRID", grid_results), ("NO GRID", no_grid_results)]:
    latin = next((r for r in res_list if r["scenario"] == "latin_only"), None)
    cjk = next((r for r in res_list if r["scenario"] == "cjk_only"), None)
    mixed = next((r for r in res_list if r["scenario"] == "mixed_same_size"), None)
    if latin and cjk and mixed:
        lh_latin = latin["paragraphs"][1].get("line_height", 0)
        lh_cjk = cjk["paragraphs"][1].get("line_height", 0)
        lh_mixed = mixed["paragraphs"][1].get("line_height", 0)
        expected = max(lh_latin, lh_cjk)
        print(f"\n{label}: latin={lh_latin}, cjk={lh_cjk}, mixed={lh_mixed}, max(lat,cjk)={expected}")
        if abs(lh_mixed - expected) < 0.1:
            print(f"  => CONFIRMED: mixed line height = max(run heights)")
        else:
            print(f"  => DIFFERENT: mixed={lh_mixed} vs expected max={expected}")
