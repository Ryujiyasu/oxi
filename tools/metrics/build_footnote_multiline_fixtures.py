"""
Build footnote multi-line decomposition fixtures via python-docx.

Each fixture has exactly ONE footnote with controlled text length. The
.docx files are saved to:
  output/footnote_multiline_fixtures/

When Word COM becomes free, run a separate `measure_*.py` to open each
fixture and record block_h.

Sweep design:
  Visual line count 1, 2, 3, 4, 5 × footnote font size {8, 10.5, 14, 18}.
  Variable: text length tuned to land each visual-line target.

python-docx footnote support: limited (footnotes are not directly exposed in
python-docx 1.2). We construct footnotes by inserting raw XML into the
footnotes.xml part. The doc has one body paragraph with a footnoteReference,
linked via relationship to footnotes.xml.

Note: python-docx doesn't make this trivial. Easier path: subclass with
direct XML manipulation, or use a "with footnote" template.

For this stage we generate fixtures as plain docx files with footnoteReference
inserted via XML, and footnote content placed in footnotes.xml.
"""
import os
import zipfile
import shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

OUT_DIR = os.path.join(os.path.dirname(__file__), "output",
                       "footnote_multiline_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


FOOTNOTES_XML_TMPL = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:id="1">
    <w:p>
      <w:pPr>
        <w:pStyle w:val="FootnoteText"/>
      </w:pPr>
      <w:r>
        <w:rPr>
          <w:rStyle w:val="FootnoteReference"/>
        </w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}"/>
          <w:sz w:val="{sz_half}"/>
          <w:szCs w:val="{sz_half}"/>
        </w:rPr>
        <w:t xml:space="preserve">{text}</w:t>
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>"""


def build_fixture(path, *, fn_text, fn_font_size_pt, fn_font="Calibri"):
    """Build a docx with one footnote on the first body paragraph."""
    # Author body via python-docx, then post-process to add footnote XML
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Pt(36)
    sec.footer_distance = Pt(36)

    # Body content
    for i in range(60):
        p = doc.add_paragraph(f"B{i+1}")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Insert footnote reference into B1's paragraph, after its run
    body = doc.element.body
    p1 = doc.paragraphs[0]
    # Build a w:r with footnoteReference id=1
    r_ref = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "FootnoteReference")
    rpr.append(rstyle)
    r_ref.append(rpr)
    fn_ref = OxmlElement("w:footnoteReference")
    fn_ref.set(qn("w:id"), "1")
    r_ref.append(fn_ref)
    p1._p.append(r_ref)

    # Save first
    doc.save(path)

    # Now post-process: add footnotes.xml part and the relationship.
    # python-docx stores its package as a zip; we re-zip with extras.
    tmp_path = path + ".tmp"
    sz_half = int(round(fn_font_size_pt * 2))  # half-pt units for w:sz
    fn_xml = FOOTNOTES_XML_TMPL.format(
        font=fn_font, sz_half=sz_half, text=fn_text.replace("&", "&amp;").replace("<", "&lt;")
    )

    # Read the existing zip, add footnotes.xml + content_types + relationships
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                data = zin.read(name)
                if name == "[Content_Types].xml":
                    # Add footnotes content type
                    if "Override PartName=\"/word/footnotes.xml\"" not in data.decode("utf-8"):
                        ct = data.decode("utf-8")
                        insert = (
                            '<Override PartName="/word/footnotes.xml" '
                            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                        )
                        ct = ct.replace("</Types>", insert + "</Types>")
                        data = ct.encode("utf-8")
                elif name == "word/_rels/document.xml.rels":
                    # Add footnotes relationship
                    rels = data.decode("utf-8")
                    if "footnotes.xml" not in rels:
                        insert = (
                            '<Relationship Id="rIdFn1" '
                            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
                            'Target="footnotes.xml"/>'
                        )
                        rels = rels.replace("</Relationships>", insert + "</Relationships>")
                        data = rels.encode("utf-8")
                zout.writestr(name, data)
            # Write footnotes.xml
            zout.writestr("word/footnotes.xml", fn_xml.encode("utf-8"))

    shutil.move(tmp_path, path)


def main():
    cases = []
    # Visual line count via text length (for 10.5pt Calibri default)
    line_targets = [
        (1, 50),    # ~1 line at 10.5pt body width
        (2, 130),
        (3, 210),
        (4, 290),
        (5, 370),
    ]
    for vl, n_chars in line_targets:
        text = ("alpha bravo " * (n_chars // 12 + 1))[:n_chars]
        for sz in [8, 10.5, 14, 18]:
            name = f"FN_vl{vl}_sz{str(sz).replace('.', 'p')}.docx"
            path = os.path.join(OUT_DIR, name)
            try:
                build_fixture(path, fn_text=text, fn_font_size_pt=sz)
                cases.append({"path": name, "vl_target": vl, "fn_size": sz, "text_chars": len(text)})
                print(f"  built {name} ({len(text)} chars, target={vl} lines, sz={sz}pt)")
            except Exception as e:
                print(f"  ERR {name}: {e}")

    print(f"\nBuilt {len(cases)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
