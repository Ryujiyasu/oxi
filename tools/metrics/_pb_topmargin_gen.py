"""Derive whether Word rounds the TOP page margin to a 10-twip multiple.

Oxi's parser applies `to_pt_round10` to the top margin only (the 2026-04-13
"Top margin rounds for content Y, bottom stays exact" rule, derived on the JP
corpus).  Two Latin documents contradict it:

    legal__0014c86f  top=2376tw (118.80pt)  Word body box top measured 118.80
    reference__0042471c top=1134tw (56.70pt) Word first baseline 68.04
                        - Cambria A 0.950195 x 12 = 56.638 (~56.70, not 56.50)

This probe pins it directly: one page per arm, a single Cambria-12 line at the
top of the body, with the top margin swept across values that round DOWN
(x4 -> -0.2pt), round UP (x6 -> +0.2pt) and are already exact (x0).  The first
baseline is read from the Word PDF; box_top = baseline - A*fs.

    round-to-10tw  =>  1134 and 1130 give the SAME baseline
    exact          =>  1134 sits 0.2pt below 1130

usage:  python tools/metrics/_pb_topmargin_gen.py gen
        python tools/metrics/_pb_topmargin_gen.py read <pdf>
"""

import sys
import os

OUT = os.path.join("pipeline_data", "pptx_probes", "topmargin")
# (label, top twips) - pairs that differ by <10tw so rounding collapses them
ARMS = [
    ("t1130", 1130),  # exact multiple (control)
    ("t1134", 1134),  # rounds DOWN to 1130 under the current rule
    ("t1140", 1140),  # exact multiple (control, +10)
    ("t2370", 2370),  # exact multiple (control)
    ("t2376", 2376),  # rounds UP to 2380 - the legal__0014c86f value
    ("t2380", 2380),  # exact multiple (control)
]

FS = 12.0
FAMILY = "Cambria"
# A = ascent part of the taller metric set (S1047); Cambria hhea == win -> win
A = 1946.0 / 2048.0


def gen():
    import docx
    from docx.shared import Twips, Pt
    from docx.enum.section import WD_SECTION

    os.makedirs(OUT, exist_ok=True)
    d = docx.Document()

    for i, (label, top) in enumerate(ARMS):
        if i > 0:
            d.add_section(WD_SECTION.NEW_PAGE)
        sec = d.sections[-1]
        sec.page_width = Twips(11906)
        sec.page_height = Twips(16838)
        sec.top_margin = Twips(top)
        sec.bottom_margin = Twips(1440)
        sec.left_margin = Twips(1440)
        sec.right_margin = Twips(1440)
        sec.header_distance = Twips(720)
        sec.footer_distance = Twips(720)

        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("TOPMARK_%s" % label)
        r.font.name = FAMILY
        r.font.size = Pt(FS)

    path = os.path.join(OUT, "topmargin.docx")
    d.save(path)
    print("wrote", path, "arms:", len(ARMS))


def read(pdf):
    import fitz

    doc = fitz.open(pdf)
    print("pages", doc.page_count)
    rows = []
    for pi in range(doc.page_count):
        raw = doc[pi].get_text("rawdict")
        best = None
        for b in raw["blocks"]:
            if b["type"] != 0:
                continue
            for l in b["lines"]:
                for s in l["spans"]:
                    txt = "".join(c["c"] for c in s["chars"])
                    if "TOPMARK_" in txt:
                        best = (s["origin"][1], txt.strip(), s["size"], s["font"])
        if best:
            rows.append((pi, best))

    print("%-8s %-10s %10s %10s %10s" % ("arm", "declared", "baseline", "box_top", "delta"))
    for (pi, (bl, txt, size, font)) in rows:
        label = txt.replace("TOPMARK_", "")
        tw = dict((a, t) for a, t in ARMS).get(label)
        declared = tw / 20.0 if tw else float("nan")
        box_top = bl - A * size
        print(
            "%-8s %8.2fpt %10.2f %10.3f %10.3f  (%s %.2f)"
            % (label, declared, bl, box_top, box_top - declared, font, size)
        )

    # pairwise verdict
    by = {}
    for (pi, (bl, txt, size, font)) in rows:
        by[txt.replace("TOPMARK_", "")] = bl - A * size
    for lo, mid, hi in (("t1130", "t1134", "t1140"), ("t2370", "t2376", "t2380")):
        if lo in by and mid in by:
            print(
                "  %s vs %s: %+0.3f   (0.00 => rounds down / +0.20 => exact)"
                % (mid, lo, by[mid] - by[lo])
            )
        if hi in by and mid in by:
            print(
                "  %s vs %s: %+0.3f   (0.00 => rounds up   / -0.30 => exact)"
                % (mid, hi, by[mid] - by[hi])
            )


def bake():
    import win32com.client

    docx = os.path.abspath(os.path.join(OUT, "topmargin.docx"))
    pdf = os.path.abspath(os.path.join(OUT, "topmargin.pdf"))
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(docx, ReadOnly=True)
        doc.ExportAsFixedFormat(pdf, 17)
        doc.Close(False)
    finally:
        word.Quit()
    print("baked", pdf)
    return pdf


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "gen":
        gen()
    elif len(sys.argv) > 1 and sys.argv[1] == "bake":
        read(bake())
    elif len(sys.argv) > 2 and sys.argv[1] == "read":
        read(sys.argv[2])
    else:
        print(__doc__)
