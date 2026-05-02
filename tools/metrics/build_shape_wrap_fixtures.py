"""
Build shape-wrap Y-effect fixtures for §17 refinement.

Tests how a floating shape with various wrap types affects body paragraph Y
positions. Variables:
  - wrap type: square, topAndBottom, tight, none
  - shape position: anchored at known x/y (e.g., x=72, y=200pt from page top)
  - shape size: 100×60pt

Each fixture has 30 body paragraphs (Calibri 11pt) so we can see paragraph
positions before, beside, and after the shape.

For pure spec investigation, the python-docx XML approach for floating shapes
is non-trivial. Use a direct OOXML write via ElementTree.
"""
import os
import zipfile
import shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output", "shape_wrap_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


# Inline drawing template for a floating rectangle anchored at absolute position.
# We use a simple <wp:anchor> with relativeFrom="page" for both H/V.
DRAWING_XML_TMPL = """
<mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <mc:Choice Requires="wps">
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                 distT="0" distB="0" distL="114300" distR="114300"
                 simplePos="0" relativeHeight="251659264" behindDoc="0" locked="0"
                 layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
        <wp:extent cx="{w_emu}" cy="{h_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        {wrap_xml}
        <wp:docPr id="1" name="ShapeForWrapTest"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:cNvSpPr/>
              <wps:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:solidFill><a:srgbClr val="DDDDDD"/></a:solidFill>
              </wps:spPr>
              <wps:bodyPr wrap="square"/>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </mc:Choice>
</mc:AlternateContent>
""".strip()


WRAP_XMLS = {
    "none":         '<wp:wrapNone/>',
    "square":       '<wp:wrapSquare wrapText="bothSides"/>',
    "tight":        '<wp:wrapTight wrapText="bothSides"><wp:wrapPolygon edited="0"><wp:start x="0" y="0"/><wp:lineTo x="0" y="21600"/><wp:lineTo x="21600" y="21600"/><wp:lineTo x="21600" y="0"/><wp:lineTo x="0" y="0"/></wp:wrapPolygon></wp:wrapTight>',
    "topAndBottom": '<wp:wrapTopAndBottom/>',
}


def pt_to_emu(pt):
    return int(round(pt * 12700))


def build_fixture(path, *, wrap_type, shape_x_pt=72, shape_y_pt=200,
                   shape_w_pt=100, shape_h_pt=60):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # Body content
    for i in range(30):
        p = doc.add_paragraph(f"B{i+1}")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Add the shape to the FIRST paragraph as a floating drawing
    first_p = doc.paragraphs[0]
    drawing_xml = DRAWING_XML_TMPL.format(
        x_emu=pt_to_emu(shape_x_pt),
        y_emu=pt_to_emu(shape_y_pt),
        w_emu=pt_to_emu(shape_w_pt),
        h_emu=pt_to_emu(shape_h_pt),
        wrap_xml=WRAP_XMLS[wrap_type],
    )
    # Wrap drawing in a w:r and append as the first run
    from lxml import etree
    drawing_elem = etree.fromstring(drawing_xml)
    r = OxmlElement("w:r")
    r.append(drawing_elem)
    # Insert at start of first paragraph (before existing runs)
    first_p._p.insert(0, r)
    # Move the run into a fresh paragraph above (so it doesn't disturb body)
    # Actually just keep it in p1; floating doesn't disturb inline flow.

    doc.save(path)


def main():
    wrap_types = ["none", "square", "topAndBottom", "tight"]
    for wt in wrap_types:
        name = f"SW_{wt}.docx"
        path = os.path.join(OUT_DIR, name)
        try:
            build_fixture(path, wrap_type=wt)
            print(f"  built {name}")
        except Exception as e:
            print(f"  ERR {name}: {e}")
    print(f"\nBuilt {len(wrap_types)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
