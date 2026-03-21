"""Re-measure COM line height table using Japanese government template.
Uses ja_gov_template.docx (rPrDefault=MS明朝 10.5pt) instead of python-docx defaults.
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

def make_test(font_name, font_size, grid_pitch_twips):
    doc = Document(TEMPLATE)
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    sectPr = section._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)
    if grid_pitch_twips > 0:
        dg = etree.SubElement(sectPr, qn('w:docGrid'))
        dg.set(qn('w:type'), 'lines')
        dg.set(qn('w:linePitch'), str(grid_pitch_twips))
    else:
        # No type = no grid snap
        dg = etree.SubElement(sectPr, qn('w:docGrid'))
        dg.set(qn('w:linePitch'), '360')

    for text in ['AAAA', 'BBBB']:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, qn('w:rFonts'))
        rfonts.set(qn('w:eastAsia'), font_name)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    tmp = os.path.join(tempfile.gettempdir(), 'ra_real_env.docx')
    doc.save(tmp)
    return tmp

fonts = [
    ('ＭＳ 明朝', 'MS Mincho'), ('ＭＳ ゴシック', 'MS Gothic'),
    ('メイリオ', 'Meiryo'), ('游ゴシック', 'Yu Gothic'), ('游明朝', 'Yu Mincho'),
    ('Calibri', 'Calibri'), ('Century', 'Century'),
    ('Times New Roman', 'Times New Roman'), ('Arial', 'Arial'),
    ('Cambria', 'Cambria'),
]

sizes = [s / 2.0 for s in range(16, 57)]  # 8.0 to 28.0
results = {}
count = 0
total = len(fonts) * len(sizes)

# Phase 1: no_grid
for font_com, font_label in fonts:
    results[font_label] = {}
    for size in sizes:
        tmp = make_test(font_com, size, 0)
        try:
            wdoc = word.Documents.Open(os.path.abspath(tmp), ReadOnly=True)
            y1 = wdoc.Paragraphs(1).Range.Information(6)
            y2 = wdoc.Paragraphs(2).Range.Information(6)
            gap = round(y2 - y1, 4)
            wdoc.Close(False)
            if gap > 0:
                sk = str(int(size)) if size == int(size) else str(size)
                results[font_label][sk] = {'no_grid': round(gap * 20) / 20}
        except:
            pass
        count += 1
        if count % 50 == 0:
            print(f'  no_grid: {count}/{total}', file=sys.stderr)

print(f'Phase 1 done: {count} measurements', file=sys.stderr)

# Phase 2: grid pitches
for font_com, font_label in fonts:
    for pitch in [300, 336, 350, 357, 360]:
        for size in [8, 9, 10, 10.5, 11, 12, 13, 14, 16, 18, 20, 24, 26, 28]:
            tmp = make_test(font_com, size, pitch)
            try:
                wdoc = word.Documents.Open(os.path.abspath(tmp), ReadOnly=True)
                y1 = wdoc.Paragraphs(1).Range.Information(6)
                y2 = wdoc.Paragraphs(2).Range.Information(6)
                gap = round(y2 - y1, 4)
                wdoc.Close(False)
                if gap > 0:
                    sk = str(int(size)) if size == int(size) else str(size)
                    if sk not in results[font_label]:
                        results[font_label][sk] = {}
                    gk = f'grid{pitch}' if pitch != 360 else 'default_grid'
                    results[font_label][sk][gk] = round(gap * 20) / 20
            except:
                pass

word.Quit()

# Save
out = os.path.join(os.path.dirname(__file__), '..', '..', 'crates', 'oxidocs-core', 'src', 'font', 'data', 'com_line_height_table.json')
with open(out, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)

print(f'\nSaved to {out}')
print(f'Sample (real-env values):')
for font in ['MS Mincho', 'MS Gothic', 'Calibri', 'Cambria', 'Meiryo']:
    for size in ['10.5', '11', '12', '14']:
        d = results.get(font, {}).get(size, {})
        if d:
            print(f'  {font:>12s} {size:>5s}pt: {d}')
