"""Test if beforeLines/afterLines changes cell spacing behavior.

Hypothesis: real 29dc6e (37.5pt row) uses beforeLines=30 before=87 afterLines=30 after=87
while S1 (33pt row) uses before=87 after=87 only. Maybe beforeLines presence
flips from "collapse + first-sb supp" to "collapse only".

docGrid linePitch determines beforeLines resolution:
- beforeLines=30 = 30/100 × linePitch = 0.30 × linePitch pt
- If linePitch=360tw=18pt (default): 5.4pt
- If linePitch=240tw=12pt: 3.6pt
"""
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.abspath("tools/metrics/beforelines_repro")
os.makedirs(OUT_DIR, exist_ok=True)


def mk_para(cell, text, *, sb_tw=0, sa_tw=0, sbLines=0, saLines=0):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for name in ["spacing"]:
        ex = pPr.find(qn(f"w:{name}"))
        if ex is not None: pPr.remove(ex)
    sp = OxmlElement("w:spacing")
    if sbLines: sp.set(qn("w:beforeLines"), str(sbLines))
    if sb_tw: sp.set(qn("w:before"), str(sb_tw))
    if saLines: sp.set(qn("w:afterLines"), str(saLines))
    if sa_tw: sp.set(qn("w:after"), str(sa_tw))
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "ＭＳ 明朝"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    return p


def make(name, p1_kwargs, p2_kwargs):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"
    # Add docGrid by injecting sectPr manipulation
    sectPr = sec._sectPr
    for ex in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(ex)
    dg = OxmlElement("w:docGrid")
    dg.set(qn("w:type"), "lines")
    dg.set(qn("w:linePitch"), "360")  # Default Word 18pt pitch
    sectPr.append(dg)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    default_p = cell.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    mk_para(cell, "para1", **p1_kwargs)
    mk_para(cell, "para2", **p2_kwargs)
    for i in range(6):
        r = tbl.add_row()
        r.cells[0].text = f"ref{i+1}"
    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    kw = dict(sb_tw=87, sa_tw=87)
    # BL1: baseline (before=87, after=87 only) — already S1, expect 33pt
    make("BL1_before_only", kw, kw)
    # BL2: add beforeLines/afterLines=30 (both)
    kw_bl = dict(sb_tw=87, sa_tw=87, sbLines=30, saLines=30)
    make("BL2_beforeLines30", kw_bl, kw_bl)
    # BL3: only beforeLines (no afterLines)
    kw_bl_only = dict(sb_tw=87, sa_tw=87, sbLines=30)
    make("BL3_beforeLines_only", kw_bl_only, kw_bl_only)
    # BL4: high beforeLines (50%)
    kw_bl50 = dict(sb_tw=87, sa_tw=87, sbLines=50, saLines=50)
    make("BL4_beforeLines50", kw_bl50, kw_bl50)
    # BL5: beforeLines only no before twips
    kw_bl_no_tw = dict(sbLines=30, saLines=30)
    make("BL5_beforeLines_notwip", kw_bl_no_tw, kw_bl_no_tw)
