"""Author minimal .docx repros to probe Word's fn reserve algorithm.

Scenarios:
  R1: 1 page, 10 body paras no fns, then 1 para with 5 fns at end
      → Does Word pre-reserve all 5 fns before last para, forcing it onto page 2?
  R2: 1 page, 10 body paras no fns, then 1 para with 3 fns
      → Compare reserve pre-allocation vs actual fit
  R3: 1 page, 10 body paras no fns, then 5 paras each with 1 fn
      → Streaming: each para's fn is committed as it lands
  R4: 1 page, 10 body paras no fns, then 3 paras each with 2 fns, widow/orphan off
      → Test per-line reserve without widow shifting

We use python-docx to author. Measure with Word COM separately.
"""
import os
from docx import Document
from docx.shared import Pt, Cm

OUT_DIR = r"tools\metrics\fn_reserve_repro"
os.makedirs(OUT_DIR, exist_ok=True)

FILLER = "これは本文の段落です。"
LONG_BODY = FILLER * 5  # ~5 lines of content per paragraph at 11pt in narrow column

def add_body(doc, n=10):
    for i in range(n):
        p = doc.add_paragraph(f"{i+1}. " + LONG_BODY)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def add_para_with_fns(doc, text, fn_texts):
    """Add a paragraph with N footnote references trailing it."""
    p = doc.add_paragraph(text)
    for t in fn_texts:
        # python-docx lacks a footnote API, so we add via XML
        # We'll defer this — skip for now and just build body
        pass
    # TODO: footnote XML injection
    return p

def build_r1():
    """Body-heavy then 1 para with 5 fns at end."""
    doc = Document()
    add_body(doc, n=10)
    # 1 para with 5 fns would go here — for now just plain para
    doc.add_paragraph("ラスト段落 ①②③④⑤ を参照")  # placeholder
    path = os.path.join(OUT_DIR, "R1_body10_thenpara5fn.docx")
    doc.save(path)
    print(f"Saved {path}")
    return path

def build_r2():
    doc = Document()
    add_body(doc, n=10)
    doc.add_paragraph("ラスト段落 ①②③ を参照")
    path = os.path.join(OUT_DIR, "R2_body10_thenpara3fn.docx")
    doc.save(path)
    print(f"Saved {path}")
    return path

def build_r3():
    doc = Document()
    add_body(doc, n=10)
    for i in range(5):
        doc.add_paragraph(f"fn-para {i+1} ① を参照")
    path = os.path.join(OUT_DIR, "R3_body10_then5para1fn.docx")
    doc.save(path)
    print(f"Saved {path}")
    return path

def build_all():
    r1 = build_r1()
    r2 = build_r2()
    r3 = build_r3()
    print("\nNOTE: python-docx lacks footnote API; these scaffolds are body-only.")
    print("For true repro, inject <w:footnoteReference> into document.xml manually,")
    print("or open in Word and add footnotes via References > Insert Footnote.")
    return [r1, r2, r3]

if __name__ == "__main__":
    build_all()
