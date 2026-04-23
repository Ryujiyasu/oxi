"""Test whether the pStyle reference itself (vs inline attrs) changes collapse behavior.

Hypothesis: 29dc6e's cell uses <w:pStyle w:val="ac"/> which references a custom
style with its own pPr (line=210 lineRule=exact, widowControl=0, etc.). My C7
repro applied those attrs inline but measured 33pt (same as S1/simple).

Test: add a real custom style "ac" to styles.xml and reference it via pStyle.
"""
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.abspath("tools/metrics/pstyle_ac_repro")
os.makedirs(OUT_DIR, exist_ok=True)


def add_ac_style(doc):
    """Inject custom style 'ac' into styles.xml mimicking 29dc6e's style."""
    styles = doc.styles.element
    # Check if "ac" exists
    for s in styles.findall(qn("w:style")):
        if s.get(qn("w:styleId")) == "ac":
            return
    ac = OxmlElement("w:style")
    ac.set(qn("w:type"), "paragraph")
    ac.set(qn("w:customStyle"), "1")
    ac.set(qn("w:styleId"), "ac")

    name = OxmlElement("w:name")
    name.set(qn("w:val"), "一太郎")
    ac.append(name)

    pPr = OxmlElement("w:pPr")
    for tag in ["widowControl", "wordWrap", "autoSpaceDE", "autoSpaceDN", "adjustRightInd"]:
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:val"), "0"); pPr.append(e)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "210")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "both"); pPr.append(jc)
    ac.append(pPr)

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), "ＭＳ 明朝")
    rPr.append(rFonts)
    spr = OxmlElement("w:spacing"); spr.set(qn("w:val"), "-1"); rPr.append(spr)
    for tag in ("sz", "szCs"):
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:val"), "21"); rPr.append(e)
    ac.append(rPr)

    styles.append(ac)


def mk_para(cell, text, with_ac=True):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for name in ["pStyle", "spacing", "rPr"]:
        ex = pPr.find(qn(f"w:{name}"))
        if ex is not None: pPr.remove(ex)
    if with_ac:
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), "ac")
        pPr.append(ps)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:beforeLines"), "30")
    sp.set(qn("w:before"), "87")
    sp.set(qn("w:afterLines"), "30")
    sp.set(qn("w:after"), "87")
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    # rPr on paragraph mark with character spacing 0
    pmark_rpr = OxmlElement("w:rPr")
    csp = OxmlElement("w:spacing"); csp.set(qn("w:val"), "0"); pmark_rpr.append(csp)
    pPr.append(pmark_rpr)
    # Run with text
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "ＭＳ 明朝"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    return p


def make(name, with_ac, p_count=2):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"

    sectPr = sec._sectPr
    for ex in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(ex)
    dg = OxmlElement("w:docGrid")
    dg.set(qn("w:type"), "linesAndChars")
    dg.set(qn("w:linePitch"), "292")
    dg.set(qn("w:charSpace"), "1453")
    sectPr.append(dg)

    if with_ac:
        add_ac_style(doc)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    default_p = cell.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    for i in range(p_count):
        mk_para(cell, f"para{i+1}" if i < 2 else "", with_ac=with_ac)
    for i in range(6):
        r = tbl.add_row()
        r.cells[0].text = f"ref{i+1}"
    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    make("P1_without_ac_2p", with_ac=False)
    make("P2_with_ac_2p", with_ac=True)
    make("P3_without_ac_1p", with_ac=False, p_count=1)
    make("P4_with_ac_1p", with_ac=True, p_count=1)
