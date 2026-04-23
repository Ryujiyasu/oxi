"""Test whether <w:adjustLineHeightInTable/> compat flag changes cell spacing behavior.

29dc6e settings.xml has this compat flag; my minimal repros don't. This flag
makes Word adjust line heights within tables.
"""
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.abspath("tools/metrics/adjustlh_repro")
os.makedirs(OUT_DIR, exist_ok=True)


def add_compat_flag(doc, flag_name):
    """Add a <w:{flag_name}/> element to <w:compat> in settings.xml."""
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    # Check not already present
    if compat.find(qn(f"w:{flag_name}")) is None:
        el = OxmlElement(f"w:{flag_name}")
        compat.insert(0, el)


def mk_para(cell, text, sb_tw=87, sa_tw=87, sbLines=30, saLines=30):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for name in ["spacing"]:
        ex = pPr.find(qn(f"w:{name}"))
        if ex is not None: pPr.remove(ex)
    sp = OxmlElement("w:spacing")
    if sbLines: sp.set(qn("w:beforeLines"), str(sbLines))
    if sb_tw: sp.set(qn("w:before"), str(sb_tw))
    if saLines: sp.set(qn("w:afterLines"), str(saLines))
    if sa_tw: sp.set(qn("w:after"), str(sa_tw))
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "ＭＳ 明朝"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    return p


def make(name, compat_flags=[]):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"

    for flag in compat_flags:
        add_compat_flag(doc, flag)

    sectPr = sec._sectPr
    for ex in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(ex)
    dg = OxmlElement("w:docGrid")
    dg.set(qn("w:type"), "linesAndChars")
    dg.set(qn("w:linePitch"), "292")
    dg.set(qn("w:charSpace"), "1453")
    sectPr.append(dg)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    default_p = cell.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    mk_para(cell, "para1")
    mk_para(cell, "para2")
    for i in range(6):
        r = tbl.add_row()
        r.cells[0].text = f"ref{i+1}"
    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    make("A1_no_compat")
    make("A2_adjustLH", ["adjustLineHeightInTable"])
    make("A3_adjustLH_useFE", ["adjustLineHeightInTable", "useFELayout"])
    make("A4_useFE", ["useFELayout"])
