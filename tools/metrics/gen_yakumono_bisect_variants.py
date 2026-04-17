"""Generate 5 minimal repro variants to isolate d77a's yakumono trigger.

Each variant adds ONE d77a-like property to the baseline repro.
Same test paragraph: "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"
(d77a para 28's exact content, MS Gothic 12pt)

Variants:
  V1: baseline (no extras)
  V2: +firstLineChars=100 + firstLine=210
  V3: +ind left=210 right=210
  V4: +docGrid type=lines linePitch=360 (explicit type)
  V5: +all d77a compat settings
  V6: V2+V3+V4+V5 combined
"""
import os
from docx import Document
from docx.shared import Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data")
)

# Exact d77a para 28 text
TEST_TEXT = "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"


def set_font(run, family="ＭＳ ゴシック", size_pt=12.0):
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)


def set_ind(p, left=None, right=None, first_line=None, first_line_chars=None):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr"); p._element.insert(0, pPr)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind"); pPr.append(ind)
    if left is not None: ind.set(qn("w:left"), str(left))
    if right is not None: ind.set(qn("w:right"), str(right))
    if first_line is not None: ind.set(qn("w:firstLine"), str(first_line))
    if first_line_chars is not None: ind.set(qn("w:firstLineChars"), str(first_line_chars))


def set_compress_punct(doc):
    settings = doc.settings.element
    for e in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(e)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)


def set_docgrid(doc, type_val=None, line_pitch=360):
    sectPr = doc.element.body.find(qn("w:sectPr"))
    if sectPr is None: return
    old = sectPr.find(qn("w:docGrid"))
    if old is not None: sectPr.remove(old)
    dg = OxmlElement("w:docGrid")
    if type_val is not None: dg.set(qn("w:type"), type_val)
    dg.set(qn("w:linePitch"), str(line_pitch))
    sectPr.append(dg)


def set_compat_flags(doc, flags):
    """flags: list of (name, val) tuples or element names (no value = self-closing)."""
    settings = doc.settings.element
    for fname, fval in flags:
        if fval is None:
            # self-closing flag like <w:adjustLineHeightInTable/>
            existing = settings.find(qn(f"w:{fname}"))
            if existing is not None: settings.remove(existing)
            el = OxmlElement(f"w:{fname}")
            settings.append(el)
        else:
            # compatSetting with name/uri/val
            el = OxmlElement("w:compatSetting")
            el.set(qn("w:name"), fname)
            el.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
            el.set(qn("w:val"), str(fval))
            settings.append(el)


def build_variant(label, mods):
    doc = Document()
    set_compress_punct(doc)
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: {TEST_TEXT}")
    set_font(r)
    # Apply paragraph mods
    if "ind" in mods:
        set_ind(p, **mods["ind"])
    # Apply section mods
    if "docgrid" in mods:
        set_docgrid(doc, **mods["docgrid"])
    # Apply compat mods
    if "compat" in mods:
        set_compat_flags(doc, mods["compat"])

    out = os.path.join(OUT_DIR, f"yakumono_bisect_{label}.docx")
    doc.save(out)
    return out


def main():
    variants = [
        # V1: baseline
        ("V1", {}),
        # V2: firstLineChars indent
        ("V2", {"ind": {"left": 210, "right": 210, "first_line": 210, "first_line_chars": 100}}),
        # V3: paragraph indent only (no firstLine)
        ("V3", {"ind": {"left": 210, "right": 210}}),
        # V4: docGrid type=lines (explicit)
        ("V4", {"docgrid": {"type_val": "lines", "line_pitch": 360}}),
        # V5: d77a compat settings
        ("V5", {"compat": [
            ("adjustLineHeightInTable", None),
            ("compatibilityMode", 15),
            ("overrideTableStyleFontSizeAndJustification", 1),
        ]}),
        # V6: combined
        ("V6", {
            "ind": {"left": 210, "right": 210, "first_line": 210, "first_line_chars": 100},
            "docgrid": {"type_val": "lines", "line_pitch": 360},
            "compat": [
                ("adjustLineHeightInTable", None),
                ("compatibilityMode", 15),
            ],
        }),
    ]
    for label, mods in variants:
        path = build_variant(label, mods)
        print(f"[{label}] {path}")
        # Verify mods applied
        import zipfile, re
        with zipfile.ZipFile(path) as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
            settings_xml = zf.read("word/settings.xml").decode("utf-8")
        csc = re.search(r'<w:characterSpacingControl[^/]*/>', settings_xml)
        print(f"  csc: {csc.group(0) if csc else 'MISSING'}")
        dg = re.search(r'<w:docGrid[^/]*/>', doc_xml)
        print(f"  docGrid: {dg.group(0) if dg else 'MISSING'}")
        ind = re.search(r'<w:ind[^/]*/>', doc_xml)
        print(f"  ind: {ind.group(0) if ind else '(none)'}")


if __name__ == "__main__":
    main()
