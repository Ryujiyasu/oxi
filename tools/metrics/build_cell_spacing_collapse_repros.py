"""Minimal repros for paragraph spacing collapse inside a table cell.

Hypothesis:
  When a single cell contains 2+ consecutive paragraphs, Word collapses
  the space_after of paragraph N with the space_before of paragraph N+1
  using max(sa, sb) instead of adding them.

Test cases (all 10.5pt MS Mincho, 1 table with 1 row 1 cell):
  S1: 2 paragraphs, sa=sb=87tw (4.35pt), line=240tw exact
  S2: 2 paragraphs, sa=60tw, sb=120tw, line=240tw exact  (max=6pt, add=9pt → Δ3pt)
  S3: 2 paragraphs, sa=0, sb=120tw, line=240tw exact     (max=6pt, add=6pt → Δ0)
  S4: 3 paragraphs, sa=sb=100tw each, line=240tw exact   (2 collapses → Δ10pt)
  S5: 2 paragraphs, sa=200tw, sb=100tw, line=240tw exact (max=10pt, add=15pt → Δ5pt)
  S6: 1 paragraph alone (baseline, no collapse possible)
"""
import os
from docx import Document
from docx.shared import Twips, Pt

OUT_DIR = os.path.abspath("tools/metrics/cell_spacing_repro")
os.makedirs(OUT_DIR, exist_ok=True)


def add_para_with_spacing(cell, text, sa_tw, sb_tw, run_sz_pt=10.5):
    p = cell.add_paragraph()
    # Remove default empty run
    pPr = p._p.get_or_add_pPr()
    # Add w:spacing
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(sb_tw))
    sp.set(qn("w:after"), str(sa_tw))
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(run_sz_pt)
    run.font.name = "ＭＳ 明朝"
    # East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hAnsi"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    return p


def remove_default_first_para(cell):
    """Cells come with one default empty paragraph; we want to clear it."""
    # Keep only our added paragraphs (skip the default one)
    pass  # We'll instead delete the default first paragraph later


def make(name, paras, extra_rows=8):
    """paras = list of (text, sa_tw, sb_tw). extra_rows fills page below."""
    doc = Document()
    # Page setup: A4, margins
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)

    # Set default style font size to 10.5pt MS Mincho
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"

    # Create a 1-cell table row with our paragraphs
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    # Delete the default paragraph inserted by python-docx
    default_p = cell.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    for (txt, sa, sb) in paras:
        add_para_with_spacing(cell, txt, sa, sb)
    # Add some follow-up rows so we can see subsequent row position
    for i in range(extra_rows):
        r = tbl.add_row()
        r.cells[0].text = f"ref row {i+1}"

    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    # S1: sa=sb=87tw (matches 29dc6e pattern)
    make("S1_sa87_sb87_2p", [("para1", 87, 87), ("para2", 87, 87)])
    # S2: sa=60, sb=120
    make("S2_sa60_sb120_2p", [("para1", 60, 120), ("para2", 60, 120)])
    # S3: sa=0, sb=120
    make("S3_sa0_sb120_2p", [("para1", 0, 120), ("para2", 0, 120)])
    # S4: 3 paragraphs
    make("S4_sa100_sb100_3p", [
        ("para1", 100, 100),
        ("para2", 100, 100),
        ("para3", 100, 100),
    ])
    # S5: asymmetric
    make("S5_sa200_sb100_2p", [("para1", 200, 100), ("para2", 200, 100)])
    # S6: baseline 1 paragraph
    make("S6_single_para", [("soloPara", 100, 100)])
