"""
Ra: バッチ5 — フォント解決詳細、テーブルセル垂直マージン、ページ区切り条件、
    spacing collapse詳細、firstLine indent + justify、CJK kinsoku
"""
import win32com.client, json, os, tempfile, ctypes
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

gdi32 = ctypes.windll.gdi32
user32 = ctypes.windll.user32
class SIZE(ctypes.Structure):
    _fields_ = [("cx", ctypes.c_long), ("cy", ctypes.c_long)]

def gdi_w(font, ppem, ch):
    hdc = user32.GetDC(0)
    hf = gdi32.CreateFontW(-ppem, 0, 0, 0, 400, 0, 0, 0, 0, 0, 0, 0, 0, font)
    old = gdi32.SelectObject(hdc, hf)
    sz = SIZE()
    gdi32.GetTextExtentPoint32W(hdc, ch, 1, ctypes.byref(sz))
    gdi32.SelectObject(hdc, old)
    gdi32.DeleteObject(hf)
    user32.ReleaseDC(0, hdc)
    return sz.cx

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = {}

try:
    # === 1. Kinsoku (line-start/end prohibited chars) ===
    print("=== Kinsoku Line Break ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)
    # Narrow column to force line breaks
    d.sections[0].page_width = Pt(300)
    d.sections[0].left_margin = Pt(36)
    d.sections[0].right_margin = Pt(36)

    # Text with kinsoku chars
    texts = [
        "あいうえおかきくけこ。さしすせそ",  # 。at end
        "あいうえおかきくけこ「さしすせそ」",  # 「 at start prohibited
        "テストですよね、テストです。",  # 、。 prohibited at line start
    ]
    for text in texts:
        p = d.add_paragraph()
        r = p.add_run(text)
        r.font.name = "MS Gothic"; r.font.size = Pt(10.5)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
        sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')

    tmp = os.path.join(tempfile.gettempdir(), "ra_kinsoku.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        p = doc.Paragraphs(i)
        nlines = p.Range.ComputeStatistics(1)
        text = p.Range.Text.strip()[:30]
        print(f"  P{i}: lines={nlines} [{text}]")
        # Find line break positions
        prng = p.Range
        prev_y = None
        line_ends = []
        for ci in range(prng.Start, min(prng.End, prng.Start + 50)):
            cr = doc.Range(ci, ci + 1)
            y = cr.Information(6)
            ch = cr.Text
            if prev_y is not None and abs(y - prev_y) > 1:
                line_ends.append(prev_ch)
            prev_y = y
            prev_ch = ch
        if line_ends:
            print(f"    Line-end chars: {[f'U+{ord(c):04X}({c})' for c in line_ends]}")
    doc.Close(False)
    os.unlink(tmp)

    # === 2. firstLineIndent + justify interaction ===
    print("\n=== FirstLine + Justify ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    p = d.add_paragraph()
    r = p.add_run("First line indented and justified text that wraps to multiple lines for testing purposes here. Second line continues.")
    r.font.name = "Calibri"; r.font.size = Pt(11)
    pPr = p._element.get_or_add_pPr()
    ind = etree.SubElement(pPr, qn('w:ind'))
    ind.set(qn('w:firstLine'), '720')  # 36pt
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'both')

    tmp = os.path.join(tempfile.gettempdir(), "ra_fli_just.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    p1 = doc.Paragraphs(1)
    print(f"  FLI={p1.Format.FirstLineIndent}, Align={p1.Format.Alignment}")
    print(f"  Lines: {p1.Range.ComputeStatistics(1)}")
    # Check first char of line 1 vs line 2
    prng = p1.Range
    prev_y = None; line_starts = []
    for ci in range(prng.Start, min(prng.End, prng.Start + 120)):
        cr = doc.Range(ci, ci + 1)
        y = cr.Information(6); x = cr.Information(5)
        ch = cr.Text
        if ord(ch) in (13, 7): continue
        if prev_y is None or abs(y - prev_y) > 1:
            line_starts.append({"x": round(x, 2), "y": round(y, 2), "ch": ch})
        prev_y = y
    for ls in line_starts:
        print(f"    Line start: x={ls['x']}, y={ls['y']}, ch='{ls['ch']}'")
    results["fli_justify"] = line_starts
    doc.Close(False); os.unlink(tmp)

    # === 3. Spacing collapse: sa+sb with different styles ===
    print("\n=== Spacing Collapse Detail ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.TopMargin = 72
    wdoc.Content.Text = ""

    configs = [
        (0, 0, "sa=0,sb=0"),
        (10, 0, "sa=10,sb=0"),
        (0, 10, "sa=0,sb=10"),
        (10, 10, "sa=10,sb=10"),
        (20, 5, "sa=20,sb=5"),
        (5, 20, "sa=5,sb=20"),
    ]
    for sa, sb, label in configs:
        if wdoc.Paragraphs.Count > 1 or wdoc.Content.Text.strip():
            r = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
            r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = label
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.SpaceAfter = sa
        p.Format.SpaceBefore = sb

    wdoc.Repaginate()
    print("  (noGrid, Calibri 11pt)")
    for i in range(1, wdoc.Paragraphs.Count + 1):
        p = wdoc.Paragraphs(i)
        y = p.Range.Information(6)
        sa = p.Format.SpaceAfter; sb = p.Format.SpaceBefore
        gap = ""
        if i > 1:
            prev_y = wdoc.Paragraphs(i-1).Range.Information(6)
            gap = f", gap={round(y - prev_y, 2)}"
        print(f"  P{i}: y={round(y,2)}, sa={sa}, sb={sb}{gap}")
    wdoc.Close(False)

    # === 4. Page break behavior: exactly at page bottom ===
    print("\n=== Page Break Boundary ===")
    wdoc = word.Documents.Add()
    ps = wdoc.Sections(1).PageSetup
    ps.LeftMargin = 72; ps.TopMargin = 72; ps.BottomMargin = 72
    wdoc.Content.Text = ""

    for i in range(42):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
            r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = f"L{i+1}"
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11
        p.Format.SpaceBefore = 0; p.Format.SpaceAfter = 0

    wdoc.Repaginate()
    # Find exact boundary
    for i in range(36, 42):
        p = wdoc.Paragraphs(i)
        y = round(p.Range.Information(6), 2)
        pg = p.Range.Information(3)
        print(f"  P{i}: y={y}, page={pg}")
    results["page_break_boundary"] = "measured"
    wdoc.Close(False)

    # === 5. Table cell top margin (tblCellMar top) ===
    print("\n=== Table Cell Top Margin ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Create table with explicit tblCellMar
    tbl_el = d.element.body.makeelement(qn('w:tbl'), {})
    d.element.body.append(tbl_el)
    tblPr = etree.SubElement(tbl_el, qn('w:tblPr'))
    # Set table-level cell margins
    tblCellMar = etree.SubElement(tblPr, qn('w:tblCellMar'))
    top_mar = etree.SubElement(tblCellMar, qn('w:top'))
    top_mar.set(qn('w:w'), '100'); top_mar.set(qn('w:type'), 'dxa')  # 5pt
    bottom_mar = etree.SubElement(tblCellMar, qn('w:bottom'))
    bottom_mar.set(qn('w:w'), '100'); bottom_mar.set(qn('w:type'), 'dxa')
    left_mar = etree.SubElement(tblCellMar, qn('w:left'))
    left_mar.set(qn('w:w'), '108'); left_mar.set(qn('w:type'), 'dxa')  # default 5.4pt

    tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(tblBorders, qn(f'w:{side}'))
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')

    for ri in range(2):
        tr = etree.SubElement(tbl_el, qn('w:tr'))
        for ci in range(2):
            tc = etree.SubElement(tr, qn('w:tc'))
            p = etree.SubElement(tc, qn('w:p'))
            run = etree.SubElement(p, qn('w:r'))
            rPr = etree.SubElement(run, qn('w:rPr'))
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:ascii'), 'Calibri')
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), '22')
            t = etree.SubElement(run, qn('w:t'))
            t.text = f"R{ri+1}C{ci+1}"

    tmp = os.path.join(tempfile.gettempdir(), "ra_cellmar.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    tbl = doc.Tables(1)
    for r in range(1, 3):
        cell = tbl.Cell(r, 1)
        y = cell.Range.Paragraphs(1).Range.Information(6)
        x = cell.Range.Paragraphs(1).Range.Information(5)
        tp = round(cell.TopPadding, 2)
        lp = round(cell.LeftPadding, 2)
        print(f"  R{r}: x={round(x,2)}, y={round(y,2)}, topPad={tp}, leftPad={lp}")
    results["tblCellMar_top"] = "measured"
    doc.Close(False); os.unlink(tmp)

    # === 6. Multiple font fallback targets ===
    print("\n=== Fallback: Arial CJK ===")
    ppem = 14
    # Does Arial use same fallback as Calibri?
    for ch in ["\u3042", "\u4E00", "\u30A2"]:
        cal_w = gdi_w("Calibri", ppem, ch)
        ari_w = gdi_w("Arial", ppem, ch)
        uig_w = gdi_w("MS UI Gothic", ppem, ch)
        print(f"  '{ch}': Calibri={cal_w}, Arial={ari_w}, MSUIG={uig_w}")
    results["arial_fallback_same"] = "measured"

    # === 7. Bold CJK font name resolution ===
    print("\n=== Bold CJK Font ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    p = d.add_paragraph()
    r = p.add_run("Bold CJK")
    r.font.name = "MS Gothic"; r.font.size = Pt(10.5); r.font.bold = True

    tmp = os.path.join(tempfile.gettempdir(), "ra_boldcjk.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    rng = doc.Paragraphs(1).Range
    font_name = rng.Font.Name
    font_bold = rng.Font.Bold
    print(f"  Font: {font_name}, Bold: {font_bold}")
    results["bold_cjk_font"] = font_name
    doc.Close(False); os.unlink(tmp)

    # === 8. Line height with lineSpacing exact < font height ===
    print("\n=== Exact LS < Font Height ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    for dg in d.sections[0]._sectPr.findall(qn('w:docGrid')):
        d.sections[0]._sectPr.remove(dg)

    for tw in [120, 160, 200]:  # 6pt, 8pt, 10pt (all < Calibri 11pt natural 13.5pt)
        p = d.add_paragraph()
        r = p.add_run(f"Exact {tw/20}pt < natural")
        r.font.name = "Calibri"; r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:line'), str(tw)); sp.set(qn('w:lineRule'), 'exact')
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')

    tmp = os.path.join(tempfile.gettempdir(), "ra_exact_small.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        p = doc.Paragraphs(i)
        y = p.Range.Information(6)
        ls = p.Format.LineSpacing
        if i > 1:
            prev = doc.Paragraphs(i-1).Range.Information(6)
            gap = round(y - prev, 2)
            print(f"  P{i}: y={round(y,2)}, ls={ls}, gap={gap}")
    doc.Close(False); os.unlink(tmp)

    # === 9. beforeLines/afterLines with specific linePitch ===
    print("\n=== beforeLines/afterLines ===")
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)
    # Template has linePitch=360 (18pt)
    for bl_val in [50, 100, 200]:  # 50/100 * 18 = 9, 18, 36pt
        p = d.add_paragraph()
        r = p.add_run(f"beforeLines={bl_val}")
        r.font.name = "Calibri"; r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:beforeLines'), str(bl_val))
        sp.set(qn('w:after'), '0')
        sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')

    tmp = os.path.join(tempfile.gettempdir(), "ra_beforelines.docx")
    d.save(tmp)
    doc = word.Documents.Open(tmp)
    for i in range(1, doc.Paragraphs.Count + 1):
        p = doc.Paragraphs(i)
        y = p.Range.Information(6)
        sb = p.Format.SpaceBefore
        print(f"  P{i}: y={round(y,2)}, sb={sb}")
        if i > 1:
            prev = doc.Paragraphs(i-1).Range.Information(6)
            print(f"    gap={round(y-prev, 2)}")
    doc.Close(False); os.unlink(tmp)

    # === 10. Text box anchor: paragraph index behavior ===
    print("\n=== TextBox Anchor Index ===")
    wdoc = word.Documents.Add()
    wdoc.Sections(1).PageSetup.LeftMargin = 72
    wdoc.Content.Text = ""
    for i in range(5):
        if i > 0:
            r = wdoc.Range(wdoc.Content.End-1, wdoc.Content.End-1)
            r.InsertParagraphAfter()
        p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
        p.Range.Text = f"Anchor para {i+1}"
        p.Range.Font.Name = "Calibri"; p.Range.Font.Size = 11

    # Add textbox anchored to P3
    tb = wdoc.Shapes.AddTextbox(1, 300, 0, 100, 50, wdoc.Paragraphs(3).Range)
    tb.RelativeVerticalPosition = 2  # paragraph
    tb.Top = 0
    tf = tb.TextFrame
    tf.TextRange.Text = "TB anchored to P3"

    wdoc.Repaginate()
    print(f"  TB anchor range start: {tb.Anchor.Start}")
    print(f"  P3 range start: {wdoc.Paragraphs(3).Range.Start}")
    print(f"  TB top: {round(tb.Top, 2)}")
    p3_y = wdoc.Paragraphs(3).Range.Information(6)
    print(f"  P3 y: {round(p3_y, 2)}")
    results["textbox_anchor"] = {
        "anchor_start": tb.Anchor.Start,
        "p3_start": wdoc.Paragraphs(3).Range.Start,
        "tb_top": round(tb.Top, 4),
        "p3_y": round(p3_y, 4),
    }
    wdoc.Close(False)

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_batch5.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
