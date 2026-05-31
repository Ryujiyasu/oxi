"""S467 minimal repro: isolate the list-group-transition vertical drift.
Build ListBullet x2 -> ListNumber x2 -> ListBullet x1 -> Heading2 -> body x2
(mirrors gen2_067) and measure Word(COM) vs Oxi(GDI dump) paragraph tops to
pinpoint why Oxi's list-transition gap is ~1pt short of Word's."""
import os, io, json, subprocess, statistics
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import win32com.client as win32

REPO = r"C:\Users\ryuji\oxi-main"
OUT = os.path.join(REPO, "tools", "golden-test", "repros", "grid_snap")
RENDERER = os.path.join(REPO, "tools", "oxi-gdi-renderer", "target", "release", "oxi-gdi-renderer.exe")
VPOS, PAGE = 6, 3


def build():
    d = docx.Document()
    sec = d.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.25)
    st = d.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    pf = st.paragraph_format
    pf.line_spacing = 1.15; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_after = Pt(10); pf.space_before = Pt(0)
    seq = [
        ("List Bullet", "Bullet one regular reviews"),
        ("List Bullet", "Bullet two this policy effective"),
        ("List Number", "Number one please do not"),
        ("List Number", "Number two please review"),
        ("List Bullet", "Bullet three please contact"),
        ("Heading 2", "Purpose"),
        ("Normal", "We are pleased to present"),
        ("Normal", "Based on our assessment"),
    ]
    for style, text in seq:
        try:
            d.add_paragraph(text, style=style)
        except KeyError:
            d.add_paragraph(text)
    p = os.path.join(OUT, "list_transition.docx")
    d.save(p)
    return p


def oxi_tops(path):
    dump = r"C:/Users/ryuji/AppData/Local/Temp/s467_list.json"
    subprocess.run([RENDERER, path, r"C:/Users/ryuji/AppData/Local/Temp/s467list", "150",
                    "--dump-layout=" + dump], capture_output=True, text=True)
    d = json.load(io.open(dump, encoding="utf-8"))
    P = {}
    for e in d["pages"][0]["elements"]:
        if e.get("type") != "text" or e.get("para_idx") is None:
            continue
        # skip marker glyphs (x near left gutter) — keep body text
        if e.get("x", 0) < 100:  # marker at ~90; body text at ~108+
            continue
        P[e["para_idx"]] = min(P.get(e["para_idx"], 1e9), e["y"])
    return [P[k] for k in sorted(P)]


def main():
    path = build()
    word = win32.gencache.EnsureDispatch("Word.Application"); word.Visible = False
    doc = word.Documents.Open(path, ReadOnly=True)
    wrows = []
    for para in doc.Paragraphs:
        st = doc.Range(para.Range.Start, para.Range.Start)
        if st.Information(PAGE) != 1:
            continue
        t = para.Range.Text.strip()
        if not t:
            continue
        wrows.append((round(st.Information(VPOS), 3), para.Style.NameLocal, t[:18]))
    doc.Close(False); word.Quit()
    oxi = oxi_tops(path)
    lines = ["idx  style          word_y   oxi_y   drift   word_gap  oxi_gap"]
    pw = po = None
    for i, (wy, sty, t) in enumerate(wrows):
        oy = oxi[i] if i < len(oxi) else None
        dr = (oy - wy) if oy is not None else None
        wg = (wy - pw) if pw is not None else 0.0
        og = (oy - po) if (oy is not None and po is not None) else 0.0
        lines.append("%2d  %-13s %7.2f %7s %7s  %+7.2f  %+7.2f  %s" % (
            i, sty[:13], wy, ("%.2f" % oy) if oy else "-",
            ("%+.2f" % dr) if dr is not None else "-", wg, og, t))
        pw = wy; po = oy
    txt = "\n".join(lines)
    io.open(os.path.join(REPO, "tools", "metrics", "_s467_repro_list.out"), "w", encoding="utf-8").write(txt)
    print(txt.encode("ascii", "replace").decode())


if __name__ == "__main__":
    main()
