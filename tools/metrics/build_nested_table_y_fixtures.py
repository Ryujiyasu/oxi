"""
Build §16 nested-table Y position fixtures.

Each fixture has a 2-level nested table:
  - Outer table: 1 row × 1 cell, with varying cell paddings
  - Inner table: 1 row × 1 cell inside the outer cell

We measure:
  - Outer cell top Y (= row top y)
  - Inner table first paragraph Y
  - Difference = inner table offset within outer cell

Variables:
  - outer cell topPadding: 0, 4.95 (default), 10, 20 pt
  - outer cell bottomPadding: 0, 5
  - whether outer cell has a leading paragraph BEFORE the nested table
  - inner table first row content (1 line)

Output: output/nested_table_y_fixtures/
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output", "nested_table_y_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


def set_cell_margins_xml(tc, top_tw=0, left_tw=99, bottom_tw=0, right_tw=99):
    """Set <w:tcMar/> on the cell element."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, w in (("top", top_tw), ("left", left_tw),
                    ("bottom", bottom_tw), ("right", right_tw)):
        elem = tcMar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tcMar.append(elem)
        elem.set(qn("w:w"), str(w))
        elem.set(qn("w:type"), "dxa")


def build_fixture(path, *, outer_top_pad_tw, outer_bot_pad_tw,
                   include_leading_para=False):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # Body P1 reference
    p = doc.add_paragraph("P1 reference")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Outer table
    outer = doc.add_table(rows=1, cols=1)
    outer_cell = outer.cell(0, 0)
    outer_cell.width = Pt(400)
    set_cell_margins_xml(outer_cell._tc,
                         top_tw=outer_top_pad_tw, left_tw=99,
                         bottom_tw=outer_bot_pad_tw, right_tw=99)

    # Clear outer cell's default paragraph
    for p in list(outer_cell.paragraphs):
        p._element.getparent().remove(p._element)

    # Optionally add a leading paragraph before the inner table
    if include_leading_para:
        from docx.oxml import OxmlElement
        new_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        new_p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22")
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = "OuterLead"
        r.append(t)
        new_p.append(r)
        outer_cell._tc.append(new_p)

    # Inner table
    inner = outer_cell.add_table(rows=1, cols=1)
    inner.autofit = False
    inner_cell = inner.cell(0, 0)
    inner_cell.width = Pt(380)
    inner_cell.text = "InnerCell"
    for p in inner_cell.paragraphs:
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Body P-end reference (Word requires a non-table paragraph after a table)
    p = doc.add_paragraph("P-end reference")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    doc.save(path)


def main():
    cases = [
        # (filename, outer_top_pad_tw, outer_bot_pad_tw, leading_para)
        ("NT_top0_bot0_noLead.docx",     0,    0,   False),
        ("NT_top0_bot100_noLead.docx",   0,    100, False),
        ("NT_top99_bot99_noLead.docx",   99,   99,  False),
        ("NT_top100_bot100_noLead.docx", 100,  100, False),
        ("NT_top200_bot100_noLead.docx", 200,  100, False),
        ("NT_top400_bot100_noLead.docx", 400,  100, False),
        ("NT_top99_bot99_lead.docx",     99,   99,  True),
        ("NT_top200_bot100_lead.docx",   200,  100, True),
    ]
    for name, top_tw, bot_tw, lead in cases:
        path = os.path.join(OUT_DIR, name)
        try:
            build_fixture(path, outer_top_pad_tw=top_tw,
                          outer_bot_pad_tw=bot_tw,
                          include_leading_para=lead)
            print(f"  built {name}")
        except Exception as e:
            print(f"  ERR {name}: {e}")
    print(f"\nBuilt {len(cases)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
