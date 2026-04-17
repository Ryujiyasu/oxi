"""LM1 variant: docGrid type=lines linePitch=350 matching b35.

Same 8 tables with n=1..9 lines, but emit explicit w:docGrid type=lines.
Tests whether b35's formula row_h(n=1)=18, row_h(n=3)=42.25 etc replicates.
"""
import os
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DOCX = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data",
                 "multiline_cell_repro_lm1.docx")
)

LINES_TEXT = [
    "あ" * 10 + "。",
    "あ" * 43 + "。" + "い" * 43 + "、",
    "う" * 43 + "。" + "え" * 43 + "、" + "お" * 30 + "。",
    "か" * 43 + "。" + "き" * 43 + "、" + "く" * 43 + "。" + "け" * 30 + "、",
    "さ" * 43 + "。" + "し" * 43 + "、" + "す" * 43 + "。" + "せ" * 43 + "、" + "そ" * 30,
    "た" * 43 + "。" + "ち" * 43 + "、" + "つ" * 43 + "。" + "て" * 43 + "、" + "と" * 43 + "。" + "な" * 30,
    "は" * 43 + "。" + "ひ" * 43 + "、" + "ふ" * 43 + "。" + "へ" * 43 + "、" + "ほ" * 43 + "。" + "ま" * 43 + "、" + "み" * 30,
    "や" * 43 + "。" + "ゆ" * 43 + "、" + "よ" * 43 + "。" + "ら" * 43 + "、" + "り" * 43 + "。" + "る" * 43 + "、" + "れ" * 43 + "。" + "ろ" * 30,
]


def set_font(run, family="ＭＳ 明朝", size_pt=10.5):
    run.font.name = family
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)


def main():
    doc = Document()
    h = doc.add_paragraph()
    r = h.add_run(f"multiline_cell_repro_lm1 — {len(LINES_TEXT)} tables, MS Mincho 10.5pt, docGrid=lines 350tw")
    set_font(r)

    for i, text in enumerate(LINES_TEXT, 1):
        doc.add_paragraph(f"--- Table {i} ---")
        t = doc.add_table(rows=1, cols=1)
        t.autofit = False
        t.columns[0].width = Twips(9072)
        cell = t.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r)

    # Find sectPr and inject explicit docGrid type=lines linePitch=350
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        # Remove existing docGrid (if any)
        old = sectPr.find(qn("w:docGrid"))
        if old is not None:
            sectPr.remove(old)
        dg = OxmlElement("w:docGrid")
        dg.set(qn("w:type"), "lines")
        dg.set(qn("w:linePitch"), "350")
        sectPr.append(dg)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"[OK] {OUT_DOCX}")


if __name__ == "__main__":
    main()
