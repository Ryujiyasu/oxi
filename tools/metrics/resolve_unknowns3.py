#!/usr/bin/env python3
"""
Additional targeted tests to resolve remaining precision questions.
"""
import os
import sys
import tempfile
import time
import pythoncom
pythoncom.CoInitialize()
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import win32com.client


TMPDIR = tempfile.gettempdir()


def docx_to_pdf_word(docx_path, pdf_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    try:
        abs_docx = os.path.abspath(docx_path).replace("/", "\\")
        abs_pdf = os.path.abspath(pdf_path).replace("/", "\\")
        doc = word.Documents.Open(abs_docx)
        time.sleep(1)
        doc.SaveAs(abs_pdf, FileFormat=17)
        time.sleep(0.5)
        doc.Close(False)
    finally:
        word.Quit()


def get_text_positions(pdf_path, page_idx=0):
    doc = fitz.open(pdf_path)
    if len(doc) <= page_idx:
        doc.close()
        return []
    page = doc[page_idx]
    blocks = page.get_text("dict")["blocks"]
    positions = []
    for b in blocks:
        if "lines" in b:
            for line in b["lines"]:
                text = "".join(span["text"] for span in line["spans"])
                y = line["bbox"][1]
                y_bottom = line["bbox"][3]
                font_size = line["spans"][0]["size"] if line["spans"] else 0
                positions.append({"y": y, "y_bottom": y_bottom, "text": text.strip(), "font_size": font_size})
    doc.close()
    return positions


def get_page_count(pdf_path):
    doc = fitz.open(pdf_path)
    n = len(doc)
    doc.close()
    return n


def set_grid(doc, line_pitch_twips, char_space=None):
    sectPr = doc.sections[0]._sectPr
    existing = sectPr.find(qn('w:docGrid'))
    if existing is not None:
        sectPr.remove(existing)
    attrs = f'w:type="linesAndChars" w:linePitch="{line_pitch_twips}"'
    if char_space is not None:
        attrs += f' w:charSpace="{char_space}"'
    grid = parse_xml(f'<w:docGrid {nsdecls("w")} {attrs}/>')
    sectPr.append(grid)


def set_snap_false(para):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    snap = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
    pPr.append(snap)


def set_no_spacing(doc):
    style_el = doc.styles['Normal'].element
    pPr = style_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        style_el.append(pPr)
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:line="240" w:lineRule="auto"/>')
    pPr.append(sp)


def set_line_spacing(para, line_val, rule="auto"):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:line="{line_val}" w:lineRule="{rule}"/>')
    pPr.append(sp)


def run_test(name, doc):
    docx_path = os.path.join(TMPDIR, f"test_{name}.docx")
    pdf_path = os.path.join(TMPDIR, f"test_{name}.pdf")
    doc.save(docx_path)
    docx_to_pdf_word(docx_path, pdf_path)
    positions = get_text_positions(pdf_path)
    pages = get_page_count(pdf_path)
    os.unlink(docx_path)
    os.unlink(pdf_path)
    return positions, pages


def print_positions(positions):
    prev_y = None
    for p in positions:
        gap = f" gap={p['y']-prev_y:.2f}" if prev_y is not None else ""
        h = p['y_bottom'] - p['y']
        print(f"  y={p['y']:.2f} h={h:.2f}{gap}  sz={p['font_size']:.1f} \"{p['text'][:50]}\"")
        prev_y = p['y']


# ============================================================
# TEST A: Grid snap formula verification
# Different grid pitches with same font size
# ============================================================
def testA_grid_formula():
    print("\n=== TEST A: Grid snap formula - various pitches ===")
    print("  Font: 11pt (default), single spacing (w:line=240)")

    for pitch_twips in [240, 272, 300, 320, 360, 400, 480]:
        pitch_pt = pitch_twips / 20.0
        doc = Document()
        set_no_spacing(doc)
        set_grid(doc, pitch_twips)

        for i in range(5):
            doc.add_paragraph(f"Line {i+1}")

        positions, _ = run_test(f"grid_p{pitch_twips}", doc)
        if len(positions) >= 2:
            gap = positions[1]['y'] - positions[0]['y']
            print(f"  pitch={pitch_pt:.1f}pt ({pitch_twips}tw): gap={gap:.2f}pt")


# ============================================================
# TEST B: snap=false line heights for various fonts
# ============================================================
def testB_snap_false_fonts():
    print("\n=== TEST B: snap=false line heights for various fonts ===")

    fonts = [
        ("Calibri", 11), ("Calibri", 10.5),
        ("Times New Roman", 11), ("Times New Roman", 10.5),
        ("Arial", 11),
        ("Century", 10.5),
        ("Yu Mincho", 10.5), ("Yu Gothic", 10.5),
        ("MS Mincho", 10.5), ("MS Gothic", 10.5),
    ]

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 360)  # 18pt grid

    # Reference line
    doc.add_paragraph("REFERENCE (snap=true, default)")

    for font_name, font_size in fonts:
        p = doc.add_paragraph()
        r = p.add_run(f"{font_name} {font_size}pt")
        r.font.name = font_name
        r.font.size = Pt(font_size)
        set_snap_false(p)

    # Another reference
    doc.add_paragraph("REFERENCE END (snap=true, default)")

    positions, _ = run_test("snap_false_fonts", doc)
    print_positions(positions)


# ============================================================
# TEST C: snap=false - does default font minimum apply?
# ============================================================
def testC_default_min():
    print("\n=== TEST C: snap=false - default font minimum ===")
    print("  Use tiny font (6pt) with snap=false to see if minimum applies")

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 360)  # grid, but snap=false

    doc.add_paragraph("Normal 11pt (snap=true)")

    p = doc.add_paragraph()
    r = p.add_run("Tiny 6pt snap=false")
    r.font.size = Pt(6)
    set_snap_false(p)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Tiny 6pt snap=false #2")
    r2.font.size = Pt(6)
    set_snap_false(p2)

    doc.add_paragraph("Normal 11pt (snap=true)")

    # Same test without grid
    doc2 = Document()
    set_no_spacing(doc2)

    doc2.add_paragraph("Normal 11pt (no grid)")

    p3 = doc2.add_paragraph()
    r3 = p3.add_run("Tiny 6pt no grid")
    r3.font.size = Pt(6)

    p4 = doc2.add_paragraph()
    r4 = p4.add_run("Tiny 6pt no grid #2")
    r4.font.size = Pt(6)

    doc2.add_paragraph("Normal 11pt (no grid)")

    pos1, _ = run_test("default_min_grid", doc)
    pos2, _ = run_test("default_min_nogrid", doc2)

    print("\n  With grid (snap=false for tiny):")
    print_positions(pos1)
    print("\n  No grid:")
    print_positions(pos2)


# ============================================================
# TEST D: afterLines / beforeLines exact values
# ============================================================
def testD_before_after_lines():
    print("\n=== TEST D: beforeLines/afterLines exact spacing ===")

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 360)  # 18pt pitch

    doc.add_paragraph("Line 1 (reference)")

    p = doc.add_paragraph("beforeLines=50")
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="50" w:before="180"/>')
    pPr.append(sp)

    p2 = doc.add_paragraph("beforeLines=100")
    pPr2 = p2._element.find(qn('w:pPr'))
    if pPr2 is None:
        pPr2 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p2._element.insert(0, pPr2)
    sp2 = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="100" w:before="360"/>')
    pPr2.append(sp2)

    p3 = doc.add_paragraph("beforeLines=200")
    pPr3 = p3._element.find(qn('w:pPr'))
    if pPr3 is None:
        pPr3 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p3._element.insert(0, pPr3)
    sp3 = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="200" w:before="720"/>')
    pPr3.append(sp3)

    doc.add_paragraph("After (reference)")

    positions, _ = run_test("beforelines_exact", doc)
    print_positions(positions)

    print(f"\n  Grid pitch = 18pt (360tw)")
    print(f"  Expected: beforeLines=50 -> 9pt, 100 -> 18pt, 200 -> 36pt")


# ============================================================
# TEST E: Space before/after in twips vs actual
# ============================================================
def testE_space_twips():
    print("\n=== TEST E: Space before in twips precision ===")

    doc = Document()
    set_no_spacing(doc)

    doc.add_paragraph("Ref line 1")

    # space before = 120 twips = 6pt
    p1 = doc.add_paragraph("before=120tw (6pt)")
    pPr1 = p1._element.find(qn('w:pPr'))
    if pPr1 is None:
        pPr1 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p1._element.insert(0, pPr1)
    sp1 = parse_xml(f'<w:spacing {nsdecls("w")} w:before="120"/>')
    pPr1.append(sp1)

    # space before = 240 twips = 12pt
    p2 = doc.add_paragraph("before=240tw (12pt)")
    pPr2 = p2._element.find(qn('w:pPr'))
    if pPr2 is None:
        pPr2 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p2._element.insert(0, pPr2)
    sp2 = parse_xml(f'<w:spacing {nsdecls("w")} w:before="240"/>')
    pPr2.append(sp2)

    doc.add_paragraph("After")

    positions, _ = run_test("space_twips", doc)
    print_positions(positions)


if __name__ == "__main__":
    testA_grid_formula()
    testB_snap_false_fonts()
    testC_default_min()
    testD_before_after_lines()
    testE_space_twips()
    print("\n=== ALL TESTS COMPLETE ===")
