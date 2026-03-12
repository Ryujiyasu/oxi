#!/usr/bin/env python3
"""
Resolve unknowns by measuring ACTUAL layout positions via Word COM.
Uses PDF export + PyMuPDF to measure exact text Y positions.
"""
import os
import sys
import tempfile
import time
import pythoncom
pythoncom.CoInitialize()
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import win32com.client


TMPDIR = tempfile.gettempdir()


def docx_to_pdf_word(docx_path, pdf_path):
    """Convert docx to PDF using Word COM."""
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    try:
        abs_docx = os.path.abspath(docx_path).replace("/", "\\")
        abs_pdf = os.path.abspath(pdf_path).replace("/", "\\")
        doc = word.Documents.Open(abs_docx)
        time.sleep(1)
        doc.SaveAs(abs_pdf, FileFormat=17)  # wdFormatPDF
        time.sleep(0.5)
        doc.Close(False)
    finally:
        word.Quit()


def get_text_positions(pdf_path, page_idx=0):
    """Extract text and Y positions from a PDF page."""
    doc = fitz.open(pdf_path)
    if len(doc) <= page_idx:
        doc.close()
        return []
    page = doc[page_idx]
    blocks = page.get_text("dict")["blocks"]
    positions = []
    for b in blocks:
        if "lines" in b:
            for line in b["lines"]:
                text = "".join(span["text"] for span in line["spans"])
                y = line["bbox"][1]  # top Y
                font_size = line["spans"][0]["size"] if line["spans"] else 0
                positions.append({"y": y, "text": text.strip(), "font_size": font_size})
    doc.close()
    return positions


def get_page_count(pdf_path):
    doc = fitz.open(pdf_path)
    n = len(doc)
    doc.close()
    return n


def set_grid(doc, line_pitch_twips, char_space=None):
    """Set docGrid linesAndChars on the document."""
    sectPr = doc.sections[0]._sectPr
    existing = sectPr.find(qn('w:docGrid'))
    if existing is not None:
        sectPr.remove(existing)
    attrs = f'w:type="linesAndChars" w:linePitch="{line_pitch_twips}"'
    if char_space is not None:
        attrs += f' w:charSpace="{char_space}"'
    grid = parse_xml(f'<w:docGrid {nsdecls("w")} {attrs}/>')
    sectPr.append(grid)


def set_snap_false(para):
    """Set snapToGrid=false on a paragraph."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    snap = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
    pPr.append(snap)


def set_before_lines(para, before_lines, before_twips):
    """Set beforeLines spacing on a paragraph."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="{before_lines}" w:before="{before_twips}"/>')
    pPr.append(sp)


def set_line_spacing(para, line_val, rule="auto"):
    """Set explicit line spacing."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:line="{line_val}" w:lineRule="{rule}"/>')
    pPr.append(sp)


def set_no_spacing(doc):
    """Remove default space_after and set single line spacing on Normal style."""
    style_el = doc.styles['Normal'].element
    pPr = style_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        style_el.append(pPr)
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:line="240" w:lineRule="auto"/>')
    pPr.append(sp)


def run_test(name, doc):
    """Save doc, convert to PDF via Word, extract positions."""
    docx_path = os.path.join(TMPDIR, f"test_{name}.docx")
    pdf_path = os.path.join(TMPDIR, f"test_{name}.pdf")
    doc.save(docx_path)
    docx_to_pdf_word(docx_path, pdf_path)
    positions = get_text_positions(pdf_path)
    pages = get_page_count(pdf_path)
    os.unlink(docx_path)
    os.unlink(pdf_path)
    return positions, pages


def print_positions(positions):
    prev_y = None
    for p in positions:
        gap = f" gap={p['y']-prev_y:.2f}" if prev_y is not None else ""
        print(f"  y={p['y']:.2f}{gap}  \"{p['text'][:50]}\"")
        prev_y = p['y']


# ============================================================
# TEST 1: snap_to_grid=false actual line heights
# ============================================================
def test1_snap_false_positions():
    print("\n=== TEST 1: snap_to_grid=false - actual Y positions ===")

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 272)  # 13.6pt pitch

    for i in range(5):
        doc.add_paragraph(f"Line {i+1} snap=true")

    p = doc.add_paragraph("Line 6 snap=FALSE")
    set_snap_false(p)

    p = doc.add_paragraph("Line 7 snap=FALSE")
    set_snap_false(p)

    for i in range(3):
        doc.add_paragraph(f"Line {i+8} snap=true")

    positions, pages = run_test("snap_false", doc)
    print(f"  Pages: {pages}")
    print_positions(positions)


# ============================================================
# TEST 2: Empty paragraph heights
# ============================================================
def test2_empty_heights():
    print("\n=== TEST 2: Empty paragraph heights ===")

    doc = Document()
    set_no_spacing(doc)

    doc.add_paragraph("Text line 1")
    doc.add_paragraph("Text line 2")
    doc.add_paragraph("")  # empty
    doc.add_paragraph("Text after empty")
    doc.add_paragraph("")  # empty
    doc.add_paragraph("")  # empty
    doc.add_paragraph("Text after 2 empties")

    positions, pages = run_test("empty_heights", doc)
    print(f"  Pages: {pages}")
    print_positions(positions)
    print("  NOTE: Empty paragraphs won't appear in PDF text extraction.")
    print("  Gap between 'Text line 2' and 'Text after empty' shows empty para height.")


# ============================================================
# TEST 3: Table cell grid snapping
# ============================================================
def test3_table_grid():
    print("\n=== TEST 3: Table cell grid snapping ===")

    # Version A: with grid
    doc_a = Document()
    set_no_spacing(doc_a)
    set_grid(doc_a, 360)  # 18pt pitch

    table_a = doc_a.add_table(rows=5, cols=1)
    for i in range(5):
        table_a.cell(i, 0).text = f"Row {i+1}"
    doc_a.add_paragraph("After table A (grid)")

    # Version B: no grid
    doc_b = Document()
    set_no_spacing(doc_b)

    table_b = doc_b.add_table(rows=5, cols=1)
    for i in range(5):
        table_b.cell(i, 0).text = f"Row {i+1}"
    doc_b.add_paragraph("After table B (no grid)")

    pos_a, pages_a = run_test("table_grid", doc_a)
    pos_b, pages_b = run_test("table_nogrid", doc_b)

    print(f"\n  Version A (grid, linePitch=360): pages={pages_a}")
    print_positions(pos_a)
    print(f"\n  Version B (no grid): pages={pages_b}")
    print_positions(pos_b)


# ============================================================
# TEST 4: beforeLines with snap=false
# ============================================================
def test4_beforelines():
    print("\n=== TEST 4: beforeLines + snap_to_grid ===")

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 272)  # 13.6pt pitch

    doc.add_paragraph("Reference line 1")
    doc.add_paragraph("Reference line 2")

    # beforeLines=100, snap=true
    p1 = doc.add_paragraph("beforeLines=100, snap=true")
    set_before_lines(p1, 100, 272)

    doc.add_paragraph("After snap=true beforeLines")

    # beforeLines=100, snap=false
    p2 = doc.add_paragraph("beforeLines=100, snap=false")
    set_before_lines(p2, 100, 272)
    set_snap_false(p2)

    doc.add_paragraph("After snap=false beforeLines")

    positions, pages = run_test("beforelines", doc)
    print(f"  Pages: {pages}")
    print_positions(positions)


# ============================================================
# TEST 5: Page overflow behavior
# ============================================================
def test5_overflow():
    print("\n=== TEST 5: Page overflow behavior ===")

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 272)

    # Set tight margins
    ps = doc.sections[0]
    ps.top_margin = Twips(851)    # 42.55pt
    ps.bottom_margin = Twips(142) # 7.1pt

    for i in range(57):
        doc.add_paragraph(f"Grid line {i+1}")

    p = doc.add_paragraph("Final snap=false with beforeLines")
    set_snap_false(p)
    set_before_lines(p, 100, 272)

    positions, pages = run_test("overflow57", doc)
    print(f"  57 lines + beforeLines: pages={pages}")
    if positions:
        print(f"  Last 3:")
        for p in positions[-3:]:
            print(f"    y={p['y']:.2f}  \"{p['text'][:50]}\"")

    # Now with 58 lines
    doc2 = Document()
    set_no_spacing(doc2)
    set_grid(doc2, 272)
    ps2 = doc2.sections[0]
    ps2.top_margin = Twips(851)
    ps2.bottom_margin = Twips(142)

    for i in range(58):
        doc2.add_paragraph(f"Grid line {i+1}")

    p2 = doc2.add_paragraph("Final snap=false with beforeLines (58 lines)")
    set_snap_false(p2)
    set_before_lines(p2, 100, 272)

    positions2, pages2 = run_test("overflow58", doc2)
    print(f"\n  58 lines + beforeLines: pages={pages2}")
    if positions2:
        # Show page 1 last 3 and page 2
        print(f"  Page 1 last 3:")
        for p in positions2[-3:]:
            print(f"    y={p['y']:.2f}  \"{p['text'][:50]}\"")

        # Also check page 2
        pdf_path = os.path.join(TMPDIR, "test_overflow58.pdf")
        # Already cleaned up, re-run for page 2
        doc2_2 = Document()
        set_no_spacing(doc2_2)
        set_grid(doc2_2, 272)
        ps2_2 = doc2_2.sections[0]
        ps2_2.top_margin = Twips(851)
        ps2_2.bottom_margin = Twips(142)
        for i in range(58):
            doc2_2.add_paragraph(f"Grid line {i+1}")
        p2_2 = doc2_2.add_paragraph("Final snap=false with beforeLines (58 lines)")
        set_snap_false(p2_2)
        set_before_lines(p2_2, 100, 272)
        docx_path = os.path.join(TMPDIR, "test_overflow58b.docx")
        pdf_path = os.path.join(TMPDIR, "test_overflow58b.pdf")
        doc2_2.save(docx_path)
        docx_to_pdf_word(docx_path, pdf_path)
        pos_p2 = get_text_positions(pdf_path, page_idx=1)
        if pos_p2:
            print(f"  Page 2:")
            print_positions(pos_p2)
        else:
            print(f"  Page 2: empty (all content fits on page 1)")
        os.unlink(docx_path)
        os.unlink(pdf_path)


# ============================================================
# TEST 6: Different font sizes with grid
# ============================================================
def test6_font_sizes_grid():
    print("\n=== TEST 6: Different font sizes with grid snap ===")

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 360)  # 18pt pitch

    for size in [10, 11, 12, 14, 16, 18, 20, 24]:
        p = doc.add_paragraph()
        r = p.add_run(f"Font size {size}pt")
        r.font.size = Pt(size)

    positions, pages = run_test("fontsizes_grid", doc)
    print(f"  Pages: {pages}")
    print_positions(positions)


# ============================================================
# TEST 7: Line spacing 1.15x (w:line=276) with grid
# ============================================================
def test7_line_spacing_115():
    print("\n=== TEST 7: Line spacing 1.15x (w:line=276) with grid ===")

    doc = Document()
    set_no_spacing(doc)
    set_grid(doc, 360)  # 18pt pitch

    for i in range(5):
        p = doc.add_paragraph(f"Line {i+1} at 1.15x")
        set_line_spacing(p, 276)  # 276/240 = 1.15

    positions, pages = run_test("line115_grid", doc)
    print(f"  Pages: {pages}")
    print_positions(positions)

    # Same without grid
    doc2 = Document()
    set_no_spacing(doc2)

    for i in range(5):
        p = doc2.add_paragraph(f"Line {i+1} at 1.15x")
        set_line_spacing(p, 276)

    positions2, pages2 = run_test("line115_nogrid", doc2)
    print(f"\n  Without grid:")
    print_positions(positions2)


if __name__ == "__main__":
    test1_snap_false_positions()
    test2_empty_heights()
    test3_table_grid()
    test4_beforelines()
    test5_overflow()
    test6_font_sizes_grid()
    test7_line_spacing_115()
    print("\n=== ALL TESTS COMPLETE ===")
