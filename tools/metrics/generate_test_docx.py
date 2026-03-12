"""
Generate test .docx files for Word text shaping analysis.

36 test cases: 4 fonts × 3 sizes × 3 language patterns.
Each docx contains a single paragraph with known text,
making it easy to extract and compare glyph positions.

Requirements: pip install python-docx
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json

# --- Configuration ---

FONTS = [
    "游明朝",          # Yu Mincho - standard Japanese serif
    "游ゴシック",       # Yu Gothic - standard Japanese sans-serif
    "Century",          # Common English serif in Japanese Word
    "Times New Roman",  # Universal serif baseline
]

SIZES = [10.5, 11, 12]  # pt - common document sizes

TEXT_PATTERNS = {
    "ja": "吾輩は猫である。名前はまだ無い。どこで生れたかとんと見当がつかぬ。",
    "en": "The quick brown fox jumps over the lazy dog. Pack my box with five dozen liquor jugs.",
    "mixed": "Word文書のレイアウトエンジンは、Latin textと日本語の混在を正しく処理する必要がある。",
}

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "docx_tests")


def create_test_docx(font: str, size: float, lang: str, text: str, output_path: str):
    """Create a minimal docx with a single styled paragraph."""
    doc = Document()

    # Remove default paragraph (empty body)
    for p in doc.paragraphs:
        p.clear()

    # Set page size to A4 and fixed margins to eliminate variables
    section = doc.sections[0]
    section.page_width = Pt(595)
    section.page_height = Pt(842)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    # Add metadata paragraph (not rendered, for identification)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"[TEST] font={font} size={size} lang={lang}")
    meta_run.font.size = Pt(8)
    meta_run.font.name = "Arial"

    # Add test paragraph
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = Pt(size * 1.5)  # Fixed line spacing
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)

    # For Japanese fonts, also set East Asian font
    from docx.oxml.ns import qn
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        from lxml import etree
        r_fonts = etree.SubElement(r_pr, qn("w:rFonts"))
    r_fonts.set(qn("w:eastAsia"), font)

    # Add single-character reference lines for per-glyph measurement
    doc.add_paragraph()  # spacer
    ref_header = doc.add_paragraph()
    ref_header_run = ref_header.add_run("[REFERENCE GLYPHS]")
    ref_header_run.font.size = Pt(8)
    ref_header_run.font.name = "Arial"

    # Output individual characters for precise width measurement
    for char in text[:20]:  # First 20 chars as reference
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.space_before = Pt(0)
        ref_para.paragraph_format.space_after = Pt(0)
        ref_run = ref_para.add_run(char)
        ref_run.font.name = font
        ref_run.font.size = Pt(size)
        r_pr2 = ref_run._element.get_or_add_rPr()
        r_fonts2 = r_pr2.find(qn("w:rFonts"))
        if r_fonts2 is None:
            from lxml import etree
            r_fonts2 = etree.SubElement(r_pr2, qn("w:rFonts"))
        r_fonts2.set(qn("w:eastAsia"), font)

    doc.save(output_path)


def generate_all():
    """Generate all 36 test docx files."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    manifest = []

    for font in FONTS:
        for size in SIZES:
            for lang, text in TEXT_PATTERNS.items():
                # Sanitize font name for filename
                safe_font = font.replace(" ", "_").replace("ゴシック", "gothic").replace("游明朝", "yu_mincho").replace("游", "yu_")
                filename = f"{safe_font}_{size}pt_{lang}.docx"
                output_path = os.path.join(OUTPUT_DIR, filename)

                create_test_docx(font, size, lang, text, output_path)

                manifest.append({
                    "filename": filename,
                    "font": font,
                    "size_pt": size,
                    "lang": lang,
                    "text": text,
                })

                print(f"  Generated: {filename}")

    # Write manifest for downstream tools
    manifest_path = os.path.join(OUTPUT_DIR, "manifest.json")
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    print(f"\n{len(manifest)} test files generated in {OUTPUT_DIR}/")
    print(f"Manifest written to {manifest_path}")


if __name__ == "__main__":
    generate_all()
