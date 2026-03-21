"""Generate complete line height lookup table via COM measurement.
Measures actual rendering line height (Y coordinate diff) for all font×size×grid combinations.
"""
import win32com.client, json, sys, os, time

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = {}

fonts = [
    ('ＭＳ ゴシック', 'MS Gothic'),
    ('ＭＳ 明朝', 'MS Mincho'),
    ('メイリオ', 'Meiryo'),
    ('游ゴシック', 'Yu Gothic'),
    ('游明朝', 'Yu Mincho'),
    ('Calibri', 'Calibri'),
    ('Century', 'Century'),
    ('Times New Roman', 'Times New Roman'),
    ('Arial', 'Arial'),
]

# Half-point sizes from 8 to 28
sizes = [s / 2.0 for s in range(16, 57)]  # 8.0, 8.5, ..., 28.0

# Grid conditions
grid_configs = [
    ('noGrid', False, 0),       # No docGrid type
    ('grid360', True, 360),     # type=lines, linePitch=360 (default)
    ('grid300', True, 300),     # type=lines, linePitch=300
    ('grid336', True, 336),     # type=lines, linePitch=336 (common in gov docs)
    ('grid350', True, 350),     # type=lines, linePitch=350
    ('grid357', True, 357),     # type=lines, linePitch=357
]

def measure_line_height(font_com_name, font_size, has_grid, grid_pitch_twips):
    """Measure actual line height by creating 2 paragraphs and taking Y diff."""
    try:
        doc = word.Documents.Add()
        sec = doc.Sections(1)
        sec.PageSetup.TopMargin = 72
        sec.PageSetup.BottomMargin = 72
        sec.PageSetup.PageHeight = 841.9
        sec.PageSetup.PageWidth = 595.3

        # Set docGrid via XML manipulation is complex, so use default (type=lines)
        # For noGrid, we'll use a different approach

        p1 = doc.Paragraphs(1)
        p1.Range.Text = 'AAAA'
        p1.Range.Font.Name = font_com_name
        p1.Range.Font.Size = font_size
        p1.Format.SpaceBefore = 0
        p1.Format.SpaceAfter = 0
        p1.Format.LineSpacingRule = 0  # Single

        p2 = doc.Paragraphs.Add()
        p2.Range.InsertBefore('BBBB')
        p2.Range.Font.Name = font_com_name
        p2.Range.Font.Size = font_size
        p2.Format.SpaceBefore = 0
        p2.Format.SpaceAfter = 0
        p2.Format.LineSpacingRule = 0

        y1 = doc.Paragraphs(1).Range.Information(6)
        y2 = doc.Paragraphs(2).Range.Information(6)
        gap = round(y2 - y1, 4)

        doc.Close(False)
        return gap
    except Exception as e:
        try:
            doc.Close(False)
        except:
            pass
        return None

def measure_with_docx(font_com_name, font_size, grid_pitch_twips):
    """Measure using python-docx to control docGrid precisely."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        import tempfile

        doc = Document()

        # Set page size
        section = doc.sections[0]
        section.page_height = 841.9 * 12700  # EMU
        section.page_width = 595.3 * 12700

        # Set docGrid
        sectPr = section._sectPr
        # Remove existing docGrid
        for dg in sectPr.findall(qn('w:docGrid')):
            sectPr.remove(dg)

        if grid_pitch_twips > 0:
            from lxml import etree
            dg = etree.SubElement(sectPr, qn('w:docGrid'))
            dg.set(qn('w:type'), 'lines')
            dg.set(qn('w:linePitch'), str(grid_pitch_twips))
        # else: no docGrid = no grid snap

        # Add 2 paragraphs
        p1 = doc.paragraphs[0]
        p1.text = 'AAAA'
        run1 = p1.runs[0]
        run1.font.name = font_com_name
        run1.font.size = int(font_size * 12700)  # EMU... no, Pt class

        from docx.shared import Pt
        run1.font.size = Pt(font_size)

        pf1 = p1.paragraph_format
        pf1.space_before = Pt(0)
        pf1.space_after = Pt(0)
        pf1.line_spacing_rule = 0  # SINGLE

        p2 = doc.add_paragraph('BBBB')
        run2 = p2.runs[0]
        run2.font.name = font_com_name
        run2.font.size = Pt(font_size)
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after = Pt(0)
        pf2.line_spacing_rule = 0

        # Save temp file
        tmp = os.path.join(tempfile.gettempdir(), 'ra_lh_test.docx')
        doc.save(tmp)

        # Open with Word COM and measure
        wdoc = word.Documents.Open(os.path.abspath(tmp), ReadOnly=True)
        y1 = wdoc.Paragraphs(1).Range.Information(6)
        y2 = wdoc.Paragraphs(2).Range.Information(6)
        gap = round(y2 - y1, 4)
        wdoc.Close(False)

        return gap
    except Exception as e:
        print(f'  Error: {e}', file=sys.stderr)
        return None

# === Main measurement loop ===
total = len(fonts) * len(sizes)
done = 0

for font_com, font_label in fonts:
    results[font_label] = {}

    for size in sizes:
        size_key = str(size)

        # Measure with default grid (type=lines, pitch=360)
        gap = measure_line_height(font_com, size, True, 360)
        if gap and gap > 0:
            results[font_label][size_key] = {
                'default_grid': round(gap * 20) / 20,  # Round to 0.05pt
            }

        done += 1
        if done % 20 == 0:
            print(f'  Progress: {done}/{total}', file=sys.stderr)

print(f'Default grid measurements done', file=sys.stderr)

# Now measure noGrid using python-docx
try:
    from docx import Document
    has_docx = True
except ImportError:
    has_docx = False
    print('python-docx not available, skipping noGrid measurements', file=sys.stderr)

if has_docx:
    for font_com, font_label in fonts:
        for size in [8, 9, 10, 10.5, 11, 12, 14, 16, 18, 20, 24, 26, 28]:
            size_key = str(size)
            gap = measure_with_docx(font_com, size, 0)  # No grid
            if gap and gap > 0:
                if size_key not in results[font_label]:
                    results[font_label][size_key] = {}
                results[font_label][size_key]['no_grid'] = round(gap * 20) / 20

        # Also measure with common grid pitches
        for pitch in [300, 336, 350, 357]:
            for size in [10.5, 11, 12]:
                size_key = str(size)
                gap = measure_with_docx(font_com, size, pitch)
                if gap and gap > 0:
                    if size_key not in results[font_label]:
                        results[font_label][size_key] = {}
                    results[font_label][size_key][f'grid{pitch}'] = round(gap * 20) / 20

    print(f'NoGrid + varied pitch measurements done', file=sys.stderr)

word.Quit()

# Save
out_path = 'pipeline_data/com_line_height_table.json'
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)

# Summary
total_entries = sum(len(v) for v in results.values())
print(f'\nSaved {out_path}: {len(results)} fonts, {total_entries} size entries')
for font, data in results.items():
    print(f'  {font}: {len(data)} sizes')
    # Show a few
    for size in ['10.5', '12', '14']:
        if size in data:
            print(f'    {size}pt: {data[size]}')
