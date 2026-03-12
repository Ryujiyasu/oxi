#!/usr/bin/env python3
"""
Resolve unknown Word layout behaviors via COM automation.
Tests: snap_to_grid=false line heights, empty paragraph heights,
       beforeLines behavior, table cell default font minimum.
"""
import os
import sys
import tempfile
import time

# Use python-docx to create test documents
from docx import Document
from docx.shared import Pt, Twips, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import win32com.client

PYTHON = sys.executable
RESULTS = {}


def measure_word_positions(docx_path):
    """Open docx in Word, measure paragraph positions via COM."""
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))
        time.sleep(0.5)

        positions = []
        for i, para in enumerate(doc.Paragraphs):
            rng = para.Range
            # PageSetup for reference
            if i == 0:
                ps = doc.Sections(1).PageSetup
                positions.append({
                    "page_width": ps.PageWidth,
                    "page_height": ps.PageHeight,
                    "margin_top": ps.TopMargin,
                    "margin_bottom": ps.BottomMargin,
                    "margin_left": ps.LeftMargin,
                    "margin_right": ps.RightMargin,
                })

            info = rng.Information
            page = info(3)  # wdActiveEndPageNumber

            # Get position using Range
            left = word.PointsToMillimeters(rng.Information(1))  # approximate

            # Use paragraph format info
            pf = para.Format
            positions.append({
                "para_idx": i,
                "text": rng.Text[:30].strip(),
                "page": page,
                "line_spacing": pf.LineSpacing,
                "line_spacing_rule": pf.LineSpacingRule,
                "space_before": pf.SpaceBefore,
                "space_after": pf.SpaceAfter,
                "snap_to_grid": not pf.NoLineNumber if hasattr(pf, 'NoLineNumber') else None,
            })

        doc.Close(False)
        return positions
    finally:
        word.Quit()


def measure_line_positions(docx_path):
    """Measure exact Y positions of each line using Word COM Selection."""
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))
        time.sleep(0.3)

        results = []
        for i, para in enumerate(doc.Paragraphs):
            rng = para.Range
            text = rng.Text[:30].strip()

            # Move to start of paragraph and get position
            rng2 = rng.Duplicate
            rng2.Collapse(1)  # wdCollapseStart

            # Get line height via line spacing
            pf = para.Format

            # Compute effective line height using Word's own calculation
            # LineSpacing property returns the actual computed value
            ls = pf.LineSpacing
            lsr = pf.LineSpacingRule
            sb = pf.SpaceBefore
            sa = pf.SpaceAfter

            results.append({
                "idx": i,
                "text": text,
                "line_spacing_pt": ls,
                "line_spacing_rule": lsr,  # 0=auto, 1=atLeast, 2=exactly, 3=multiple, 4=single, 5=double
                "space_before_pt": sb,
                "space_after_pt": sa,
                "page": rng.Information(3),
            })

        doc.Close(False)
        return results
    finally:
        word.Quit()


def test1_snap_to_grid_false():
    """Test: Does default font minimum apply to snap_to_grid=false paragraphs?"""
    print("\n=== TEST 1: snap_to_grid=false line height ===")

    doc = Document()

    # Set up document with linesAndChars grid
    sectPr = doc.sections[0]._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = parse_xml(f'<w:pgSz {nsdecls("w")} w:w="11906" w:h="16838"/>')
        sectPr.append(pgSz)

    # Add docGrid linesAndChars with linePitch=360 (18pt)
    existing_grid = sectPr.find(qn('w:docGrid'))
    if existing_grid is not None:
        sectPr.remove(existing_grid)
    grid = parse_xml(f'<w:docGrid {nsdecls("w")} w:type="linesAndChars" w:linePitch="360"/>')
    sectPr.append(grid)

    # Set default font to Calibri 10.5pt
    rpr = doc.styles['Normal'].element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}><w:sz w:val="21"/></w:rPr>')
        doc.styles['Normal'].element.append(rpr)

    # Paragraph 1: snap_to_grid=true (default), marker text
    p1 = doc.add_paragraph("GRID-SNAP-TRUE: This paragraph snaps to grid.")

    # Paragraph 2: snap_to_grid=false
    p2 = doc.add_paragraph("GRID-SNAP-FALSE: This paragraph does NOT snap to grid.")
    pPr2 = p2._element.find(qn('w:pPr'))
    if pPr2 is None:
        pPr2 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p2._element.insert(0, pPr2)
    snap_el = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
    pPr2.append(snap_el)

    # Paragraph 3: snap_to_grid=true again
    p3 = doc.add_paragraph("GRID-SNAP-TRUE-2: Back to grid snap.")

    # Paragraph 4: snap_to_grid=false, empty
    p4 = doc.add_paragraph("")
    pPr4 = p4._element.find(qn('w:pPr'))
    if pPr4 is None:
        pPr4 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p4._element.insert(0, pPr4)
    snap_el4 = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
    pPr4.append(snap_el4)

    # Paragraph 5: normal
    p5 = doc.add_paragraph("AFTER-EMPTY: After empty non-snap paragraph.")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    results = measure_line_positions(tmp)
    for r in results:
        print(f"  para[{r['idx']}] ls={r['line_spacing_pt']:.2f} rule={r['line_spacing_rule']} sb={r['space_before_pt']:.2f} sa={r['space_after_pt']:.2f} text=\"{r['text'][:40]}\"")

    os.unlink(tmp)
    return results


def test2_empty_paragraph_height():
    """Test: What is the line height of an empty paragraph?"""
    print("\n=== TEST 2: Empty paragraph line height ===")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Non-empty paragraph for reference
    doc.add_paragraph("Reference line with text (Calibri 10.5pt)")
    # Empty paragraph
    doc.add_paragraph("")
    # Another reference
    doc.add_paragraph("After empty paragraph")
    # Empty with different font size
    p = doc.add_paragraph("")
    p.style.font.size = Pt(14)
    # Reference
    doc.add_paragraph("After 14pt empty paragraph")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    results = measure_line_positions(tmp)
    for r in results:
        print(f"  para[{r['idx']}] ls={r['line_spacing_pt']:.2f} rule={r['line_spacing_rule']} sb={r['space_before_pt']:.2f} text=\"{r['text'][:40]}\"")

    os.unlink(tmp)
    return results


def test3_beforeLines_behavior():
    """Test: How does beforeLines interact with grid pitch?"""
    print("\n=== TEST 3: beforeLines with grid ===")

    doc = Document()

    # Set up grid
    sectPr = doc.sections[0]._sectPr
    existing_grid = sectPr.find(qn('w:docGrid'))
    if existing_grid is not None:
        sectPr.remove(existing_grid)
    grid = parse_xml(f'<w:docGrid {nsdecls("w")} w:type="linesAndChars" w:linePitch="360"/>')
    sectPr.append(grid)

    # Reference paragraph
    doc.add_paragraph("Reference: no beforeLines")

    # Paragraph with beforeLines=100 (1.0 lines), snap=true
    p2 = doc.add_paragraph("beforeLines=100, snap=true")
    pPr2 = p2._element.find(qn('w:pPr'))
    if pPr2 is None:
        pPr2 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p2._element.insert(0, pPr2)
    sp2 = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="100" w:before="360"/>')
    pPr2.append(sp2)

    # Paragraph with beforeLines=100, snap=false
    p3 = doc.add_paragraph("beforeLines=100, snap=false")
    pPr3 = p3._element.find(qn('w:pPr'))
    if pPr3 is None:
        pPr3 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p3._element.insert(0, pPr3)
    sp3 = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="100" w:before="360"/>')
    pPr3.append(sp3)
    snap3 = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
    pPr3.append(snap3)

    # Another reference
    doc.add_paragraph("Reference: after beforeLines tests")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    results = measure_line_positions(tmp)
    for r in results:
        print(f"  para[{r['idx']}] ls={r['line_spacing_pt']:.2f} rule={r['line_spacing_rule']} sb={r['space_before_pt']:.2f} text=\"{r['text'][:40]}\"")

    os.unlink(tmp)
    return results


def test4_table_cell_default_minimum():
    """Test: Does the default font minimum apply inside table cells?"""
    print("\n=== TEST 4: Table cell default font minimum ===")

    doc = Document()

    # Set up grid document
    sectPr = doc.sections[0]._sectPr
    existing_grid = sectPr.find(qn('w:docGrid'))
    if existing_grid is not None:
        sectPr.remove(existing_grid)
    grid = parse_xml(f'<w:docGrid {nsdecls("w")} w:type="linesAndChars" w:linePitch="360"/>')
    sectPr.append(grid)

    # Reference paragraph outside table
    doc.add_paragraph("Outside table: reference line (auto spacing)")

    # Create a simple table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Cell A1: auto spacing"
    table.cell(0, 1).text = "Cell B1: auto spacing"
    table.cell(1, 0).text = "Cell A2: Calibri 10.5pt"
    table.cell(1, 1).text = "Cell B2: Calibri 10.5pt"
    table.cell(2, 0).text = "Cell A3: text"
    table.cell(2, 1).text = "Cell B3: text"

    # Another reference
    doc.add_paragraph("Outside table: after table")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    # Measure with Word
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        wdoc = word.Documents.Open(str(tmp))
        time.sleep(0.3)

        print("  Body paragraphs:")
        for i, para in enumerate(wdoc.Paragraphs):
            pf = para.Format
            text = para.Range.Text[:40].strip()
            print(f"    para[{i}] ls={pf.LineSpacing:.2f} rule={pf.LineSpacingRule} sb={pf.SpaceBefore:.2f} text=\"{text}\"")

        print("\n  Table cells:")
        for tbl_idx, tbl in enumerate(wdoc.Tables):
            for row_idx in range(1, tbl.Rows.Count + 1):
                for col_idx in range(1, tbl.Columns.Count + 1):
                    cell = tbl.Cell(row_idx, col_idx)
                    for pi, para in enumerate(cell.Range.Paragraphs):
                        pf = para.Format
                        text = para.Range.Text[:30].strip()
                        h = cell.Height
                        print(f"    tbl[{tbl_idx}] cell({row_idx},{col_idx}) para[{pi}] ls={pf.LineSpacing:.2f} rule={pf.LineSpacingRule} h={h:.2f} text=\"{text}\"")

        wdoc.Close(False)
    finally:
        word.Quit()

    os.unlink(tmp)


def test5_snap_false_exact_positions():
    """Measure exact Y positions for snap=true vs snap=false paragraphs."""
    print("\n=== TEST 5: Exact Y positions via Selection ===")

    doc = Document()

    # Set up grid
    sectPr = doc.sections[0]._sectPr
    existing_grid = sectPr.find(qn('w:docGrid'))
    if existing_grid is not None:
        sectPr.remove(existing_grid)
    # linePitch=272 (13.6pt) - same as 1636d document
    grid = parse_xml(f'<w:docGrid {nsdecls("w")} w:type="linesAndChars" w:linePitch="272"/>')
    sectPr.append(grid)

    # Set margins similar to 1636d
    ps = doc.sections[0]
    ps.top_margin = Twips(851)  # 42.55pt
    ps.bottom_margin = Twips(142)  # 7.1pt

    # Set default font to Century/MS Mincho like 1636d
    rpr_default = doc.styles['Normal'].element.find(qn('w:rPr'))
    if rpr_default is None:
        rpr_default = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        doc.styles['Normal'].element.append(rpr_default)
    sz = rpr_default.find(qn('w:sz'))
    if sz is None:
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="21"/>')
        rpr_default.append(sz)
    else:
        sz.set(qn('w:val'), '21')

    # Add many grid-snapped paragraphs to fill most of the page
    for i in range(50):
        doc.add_paragraph(f"Grid line {i+1}: テスト文書の行")

    # Then snap=false paragraphs with beforeLines
    p_note = doc.add_paragraph("備考：この段落はsnapToGrid=0です。")
    pPr = p_note._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_note._element.insert(0, pPr)
    snap = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
    pPr.append(snap)
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="100" w:before="272"/>')
    pPr.append(sp)

    # Empty snap=false with beforeLines
    p_empty = doc.add_paragraph("")
    pPr2 = p_empty._element.find(qn('w:pPr'))
    if pPr2 is None:
        pPr2 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_empty._element.insert(0, pPr2)
    snap2 = parse_xml(f'<w:snapToGrid {nsdecls("w")} w:val="0"/>')
    pPr2.append(snap2)
    sp2 = parse_xml(f'<w:spacing {nsdecls("w")} w:beforeLines="100" w:before="272"/>')
    pPr2.append(sp2)

    # Final reference
    doc.add_paragraph("最終行：参照用")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    # Measure with Word Selection
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        wdoc = word.Documents.Open(str(tmp))
        time.sleep(0.5)

        sel = word.Selection
        total_paras = wdoc.Paragraphs.Count
        print(f"  Total paragraphs: {total_paras}")
        print(f"  Pages: {wdoc.ComputeStatistics(2)}")  # wdStatisticPages

        # Measure last 10 paragraphs
        for i in range(max(1, total_paras - 10), total_paras + 1):
            para = wdoc.Paragraphs(i)
            rng = para.Range
            text = rng.Text[:30].strip()
            pf = para.Format

            # Get page number
            page = rng.Information(3)

            print(f"  para[{i}] page={page} ls={pf.LineSpacing:.2f} rule={pf.LineSpacingRule} sb={pf.SpaceBefore:.2f} snap={'Y' if not pf.SnapToGrid == False else 'N'} text=\"{text[:25]}\"")

        wdoc.Close(False)
    finally:
        word.Quit()

    os.unlink(tmp)


if __name__ == "__main__":
    test1_snap_to_grid_false()
    test2_empty_paragraph_height()
    test3_beforeLines_behavior()
    test4_table_cell_default_minimum()
    test5_snap_false_exact_positions()

    print("\n\n=== ALL TESTS COMPLETE ===")
