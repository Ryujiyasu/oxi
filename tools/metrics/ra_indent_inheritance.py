"""
Ra: インデント継承チェーンの仕様確定
- docDefaults → style → paragraph の継承順序
- Chars単位 (indLeftChars) vs twip単位 (ind left) の優先順位
- mirrorIndents の挙動
- スタイル間の継承 (Heading1 basedOn Normal)
"""
import win32com.client, json, os, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def test_indent_chars_vs_twips():
    """indLeftChars vs ind left priority."""
    d = Document(TEMPLATE)
    sec = d.sections[0]
    sectPr = sec._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)

    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # P1: ind left=720tw(36pt) only
    p1 = d.add_paragraph()
    p1.add_run("P1: left=720tw only").font.name = "Calibri"
    p1.runs[0].font.size = Pt(11)
    pPr1 = p1._element.get_or_add_pPr()
    ind1 = etree.SubElement(pPr1, qn('w:ind'))
    ind1.set(qn('w:left'), '720')

    # P2: indLeftChars=200 only (200/100 * charWidth)
    p2 = d.add_paragraph()
    p2.add_run("P2: leftChars=200 only").font.name = "Calibri"
    p2.runs[0].font.size = Pt(11)
    pPr2 = p2._element.get_or_add_pPr()
    ind2 = etree.SubElement(pPr2, qn('w:ind'))
    ind2.set(qn('w:leftChars'), '200')

    # P3: both ind left=720 AND leftChars=400
    p3 = d.add_paragraph()
    p3.add_run("P3: left=720 + leftChars=400").font.name = "Calibri"
    p3.runs[0].font.size = Pt(11)
    pPr3 = p3._element.get_or_add_pPr()
    ind3 = etree.SubElement(pPr3, qn('w:ind'))
    ind3.set(qn('w:left'), '720')
    ind3.set(qn('w:leftChars'), '400')

    # P4: no indent (baseline)
    p4 = d.add_paragraph()
    p4.add_run("P4: no indent baseline").font.name = "Calibri"
    p4.runs[0].font.size = Pt(11)

    return d


def test_style_inheritance():
    """Style inheritance chain for indents."""
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Set Normal style indent
    normal = d.styles['Normal']
    normal.paragraph_format.left_indent = Pt(18)

    # P1: Normal style (should inherit 18pt indent)
    p1 = d.add_paragraph()
    p1.add_run("P1: Normal style, inherited indent").font.name = "Calibri"
    p1.runs[0].font.size = Pt(11)

    # P2: Normal style + explicit indent override
    p2 = d.add_paragraph()
    p2.add_run("P2: Normal + explicit 36pt").font.name = "Calibri"
    p2.runs[0].font.size = Pt(11)
    p2.paragraph_format.left_indent = Pt(36)

    # P3: Normal style + explicit 0 indent
    p3 = d.add_paragraph()
    p3.add_run("P3: Normal + explicit 0pt").font.name = "Calibri"
    p3.runs[0].font.size = Pt(11)
    p3.paragraph_format.left_indent = Pt(0)

    return d


def measure(doc_path, label):
    doc = word.Documents.Open(doc_path)
    try:
        data = {"label": label, "paragraphs": []}
        ml = doc.Sections(1).PageSetup.LeftMargin
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            rng = para.Range
            cr = doc.Range(rng.Start, rng.Start + 1)
            x = cr.Information(5)
            data["paragraphs"].append({
                "index": i,
                "x_pt": round(x, 4),
                "margin_rel": round(x - ml, 4),
                "left_indent": round(para.Format.LeftIndent, 4),
                "first_line_indent": round(para.Format.FirstLineIndent, 4),
                "text": rng.Text.strip()[:50],
            })
        data["margin_left"] = round(ml, 4)
        return data
    finally:
        doc.Close(False)


try:
    # Test 1: Chars vs Twips
    d1 = test_indent_chars_vs_twips()
    tmp1 = os.path.join(tempfile.gettempdir(), "ra_indent_chars.docx")
    d1.save(tmp1)
    data1 = measure(tmp1, "chars_vs_twips")
    results.append(data1)
    os.unlink(tmp1)

    print("=== chars_vs_twips ===")
    for p in data1["paragraphs"]:
        print(f"  P{p['index']}: x=margin+{p['margin_rel']}pt, li={p['left_indent']}pt  [{p['text']}]")

    # Test 2: Style inheritance
    d2 = test_style_inheritance()
    tmp2 = os.path.join(tempfile.gettempdir(), "ra_indent_inherit.docx")
    d2.save(tmp2)
    data2 = measure(tmp2, "style_inheritance")
    results.append(data2)
    os.unlink(tmp2)

    print(f"\n=== style_inheritance ===")
    for p in data2["paragraphs"]:
        print(f"  P{p['index']}: x=margin+{p['margin_rel']}pt, li={p['left_indent']}pt  [{p['text']}]")

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_indent_inheritance.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
