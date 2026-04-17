"""Generate a minimal docx with yakumono chars in both body AND cell contexts.

Uses the same text in both contexts so per-char x comparison is direct.
Tests whether Word compresses yakumono pairs identically or differently
between body vs cell.
"""
import os
from docx import Document
from docx.shared import Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data",
                 "yakumono_context_repro.docx")
)

# Text with yakumono pairs: ）。」」、。（（ etc.
# Include known-compressible pairs between CJK trigger chars.
TEXT_SAMPLES = [
    # 1. Short, no overflow (baseline)
    "今日は良い天気。明日は雨。",
    # 2. Short, no overflow
    "「本利用ルール」というもの。",
    # 3. Short, no overflow
    "以下（以下「本コンテンツ」）の利用。",
    # 4. Short, no overflow
    "第1条、「契約」はここに。第2条。",
    # 5. Previous d77a-like body paragraph
    "本ウェブサイトで公開している情報（以下「コンテンツ」といいます。）は、別の利用ルールが適用されるコンテンツを除き、",
    # 6. OVERFLOW TEST: line-end yakumono pair that JUST overflows natural width.
    # d77a's paragraph "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"
    # At 10.5pt × 40 chars = 420pt. Line width ~453pt. Natural fits. But with bullet and
    # indent, Word may be tight at line-end — test if 。 compresses at end.
    "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。",
    # 7. Explicit overflow: 42 chars of pure CJK + yakumono pair at end (.。)
    # 42 × 10.5 = 441pt. Line width ~453pt. Last 。 at or near overflow.
    "これは四十二文字の日本語テキスト例で行末圧縮のテスト用に作成しました、。",
    # 8. Heavy yakumono density: many brackets
    "（第一章）「契約」（第二章）「解除」（第三章）「賠償」について定める。",
]


def set_font(run, family="ＭＳ 明朝", size_pt=10.5):
    run.font.name = family
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)


def main():
    doc = Document()
    # Replace compressPunctuation setting (python-docx adds doNotCompress default)
    settings = doc.settings.element
    for existing in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(existing)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)

    # Header
    h = doc.add_paragraph()
    r = h.add_run("yakumono_context_repro — body + cell comparison (MS Mincho 10.5pt, compressPunctuation)")
    set_font(r)

    # BODY section
    doc.add_paragraph().add_run("=== BODY ===")
    for i, text in enumerate(TEXT_SAMPLES, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"B{i}: {text}")
        set_font(r)

    # CELL section — single-cell table, same texts
    doc.add_paragraph().add_run("=== CELL ===")
    for i, text in enumerate(TEXT_SAMPLES, 1):
        t = doc.add_table(rows=1, cols=1)
        t.autofit = False
        t.columns[0].width = Twips(9072)
        cell = t.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"C{i}: {text}")
        set_font(r)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"[OK] {OUT}")


if __name__ == "__main__":
    main()
