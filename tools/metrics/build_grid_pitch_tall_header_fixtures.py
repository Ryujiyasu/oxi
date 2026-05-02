"""
Build grid-pitch sweep fixtures for §8.2 tall-header pushdown follow-up.

Phase E in ra2_tall_header_pushdown.py confirmed grid offset at pitch=360tw
(+2.0pt exact) but had +0.5pt residual at pitch=480tw (1 data point).
Sweep pitches {300, 360, 420, 480, 540, 600 tw} and N_header_lines ∈ {3, 4}
to determine whether the residual is truly +0.5pt at pitch=480tw or just
quantization noise.

Each fixture is a Word doc with:
  - LayoutMode=LineGrid, custom linePitch (twips)
  - 3-line or 4-line Calibri 14pt header (forces tall-header overflow)
  - Calibri 11pt body with 2 paragraphs (so we can measure body line height)
  - tm=72, hdrDist=36, fdDist=36

Saved to output/grid_pitch_tall_header_fixtures/
"""
import os
import zipfile
import shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output",
                       "grid_pitch_tall_header_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


def add_doc_grid(section, line_pitch_tw, char_space_tw=0):
    """Add or update a w:docGrid element on a section's sectPr."""
    sectPr = section._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        docGrid = OxmlElement("w:docGrid")
        sectPr.append(docGrid)
    docGrid.set(qn("w:type"), "lines")
    docGrid.set(qn("w:linePitch"), str(line_pitch_tw))
    if char_space_tw:
        docGrid.set(qn("w:charSpace"), str(char_space_tw))


def build_fixture(path, *, line_pitch_tw, hdr_lines, hdr_size_pt=14):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Pt(36)
    sec.footer_distance = Pt(36)
    add_doc_grid(sec, line_pitch_tw)

    # Header content via python-docx's first-page header
    header = sec.header
    # Remove default empty paragraph
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    for i in range(hdr_lines):
        p = header.add_paragraph(f"H{i+1}")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(hdr_size_pt)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Body: 2 paragraphs to measure body line height
    p1 = doc.add_paragraph("B1")
    p2 = doc.add_paragraph("B2")
    for p in (p1, p2):
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.save(path)


def main():
    pitches = [300, 360, 420, 480, 540, 600]
    hdr_line_counts = [3, 4]
    cases = []
    for pitch in pitches:
        for n in hdr_line_counts:
            name = f"GP_pitch{pitch}tw_hdr{n}.docx"
            path = os.path.join(OUT_DIR, name)
            try:
                build_fixture(path, line_pitch_tw=pitch, hdr_lines=n)
                cases.append({"path": name, "pitch_tw": pitch, "hdr_lines": n})
                print(f"  built {name}")
            except Exception as e:
                print(f"  ERR {name}: {e}")

    print(f"\nBuilt {len(cases)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
