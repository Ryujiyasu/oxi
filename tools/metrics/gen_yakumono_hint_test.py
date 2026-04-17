"""Test: does w:hint='eastAsia' on rFonts trigger ・ compression?

d77a para 28 XML has <w:rFonts ... w:hint="eastAsia"/>. My previous repros
used python-docx default rFonts without w:hint attribute. Test if this
specific attribute is the compression trigger.

Also test variations with pPr/rPr rFonts (not just run rPr).
"""
import os
from docx import Document
from docx.shared import Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "pipeline_data")
)

TEST_TEXT = "・利用規約名を表記する際に表示スペースの制約により略称があると利便性が高まる。"


def set_compress_punct(doc):
    settings = doc.settings.element
    for e in settings.findall(qn("w:characterSpacingControl")):
        settings.remove(e)
    csc = OxmlElement("w:characterSpacingControl")
    csc.set(qn("w:val"), "compressPunctuation")
    settings.append(csc)


def make_variant_with_hint(label, with_hint, with_ppr_rpr):
    doc = Document()
    set_compress_punct(doc)
    p = doc.add_paragraph()
    # If with_ppr_rpr, add rFonts to pPr/rPr
    if with_ppr_rpr:
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr"); p._element.insert(0, pPr)
        pPr_rPr = pPr.find(qn("w:rPr"))
        if pPr_rPr is None:
            pPr_rPr = OxmlElement("w:rPr"); pPr.append(pPr_rPr)
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), "ＭＳ ゴシック")
        rf.set(qn("w:eastAsia"), "ＭＳ ゴシック")
        rf.set(qn("w:hAnsi"), "ＭＳ ゴシック")
        pPr_rPr.append(rf)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); pPr_rPr.append(sz)

    # Add run with specific rFonts
    r = p.add_run(TEST_TEXT)
    rel = r._element
    rPr = rel.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); rel.insert(0, rPr)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "ＭＳ ゴシック")
    rFonts.set(qn("w:eastAsia"), "ＭＳ ゴシック")
    rFonts.set(qn("w:hAnsi"), "ＭＳ ゴシック")
    if with_hint:
        rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)

    out = os.path.join(OUT_DIR, f"yakumono_hint_{label}.docx")
    doc.save(out)
    return out


def main():
    variants = [
        # H1: no hint, no pPr/rPr (baseline)
        ("H1_plain", False, False),
        # H2: with hint=eastAsia on run rFonts
        ("H2_hint", True, False),
        # H3: with pPr/rPr rFonts + hint on run
        ("H3_pprRpr_hint", True, True),
        # H4: with pPr/rPr rFonts, no hint
        ("H4_pprRpr_nohint", False, True),
    ]
    for label, with_hint, with_ppr_rpr in variants:
        path = make_variant_with_hint(label, with_hint, with_ppr_rpr)
        print(f"[{label}] {path}")


if __name__ == "__main__":
    main()
