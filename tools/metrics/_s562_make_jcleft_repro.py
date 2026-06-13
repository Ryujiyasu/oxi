# -*- coding: utf-8 -*-
"""S562: derive Word's compat15 jc=LEFT yakumono behavior (oikomi vs oidashi).
roudoujoken (compat15, jc=left) WRAPS the (5)裁量 line (overflow ~14pt > demand
cap). Does compat15 jc=left do ANY oikomi (small-overflow compress) or always
oidashi (wrap)? Sweep: a jc=left CJK line ending in a yakumono, line length
tuned so the yakumono sits at increasing overflow. Measure Word's wrap point.

Each doc: page A4 portrait, MS明朝 10.5pt, jc=left, a single paragraph of K CJK
chars + a comma 、 (compressible yakumono). Vary K so the 、 lands just at/over
the right margin. If Word fits 、 (compresses prev or hangs) at small overflow but
wraps at large -> oikomi cap exists; if it wraps at ANY overflow -> no oikomi.
"""
import os,sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Emu
sys.stdout.reconfigure(encoding='utf-8')
TW=635
def set_compat15(doc):
    # settings.xml compatibilityMode=15
    st=doc.settings.element
    cs=st.find(qn('w:compat'))
    if cs is None:
        cs=OxmlElement('w:compat'); st.append(cs)
    c=OxmlElement('w:compatSetting')
    c.set(qn('w:name'),'compatibilityMode'); c.set(qn('w:uri'),'http://schemas.microsoft.com/office/word'); c.set(qn('w:val'),'15')
    cs.append(c)
def build(path, K):
    doc=Document(); set_compat15(doc)
    sec=doc.sections[0]
    sec.page_width=Emu(11906*TW); sec.page_height=Emu(16838*TW)
    sec.top_margin=Emu(1000*TW); sec.bottom_margin=Emu(1000*TW)
    sec.left_margin=Emu(1701*TW); sec.right_margin=Emu(1701*TW)
    # content width = (11906-3402)/20 = 425.2pt ; fullwidth char 10.5 -> ~40 chars/line
    p=doc.add_paragraph()
    pf=p.paragraph_format; pf.alignment=0  # left
    pf.space_before=Pt(0); pf.space_after=Pt(0); pf.line_spacing=1.0
    # K CJK chars then a comma, designed to wrap near char K
    txt='あ'*K+'、つづく'
    r=p.add_run(txt); r.font.size=Pt(10.5)
    rpr=r._r.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'ＭＳ 明朝')
    doc.save(path)
out='tools/golden-test/repros/s562_jcleft_yakumono'; os.makedirs(out,exist_ok=True)
# content 425.2pt / 10.5 = 40.5 chars. Sweep K so the 、 (char K+1) is near 40.
for K in range(38,44):
    build('%s/s562_K%02d_repro.docx'%(out,K),K)
print('wrote', len([f for f in os.listdir(out)]),'repros')
