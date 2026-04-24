"""Author minimal repros for the 'split-box bottom padding' spec.

Pattern (from d77a p.7 observation):
- 1x1 table with borders, contains a long paragraph that wraps to N lines
- Table is placed near bottom of p.1 so M <= N lines spill to p.2
- AFTER the table: empty paragraph + content paragraph (the [解説])

Measure on p.2:
- last_new_page_y: last continuation line y
- first_body_after_y: first paragraph y after table
- gap = first_body_after_y - last_new_page_y

Hypothesis (from memory DERIVED formula):
    gap = line_height * (1 + has_trailing_empty_cell_para)

Variations:
  SB_A: no trailing empty cell para, no empty-para-after-table
  SB_B: no trailing empty cell para, WITH empty-para-after-table
  SB_C: WITH trailing empty cell para, no empty-para-after-table
  SB_D: WITH trailing empty cell para, WITH empty-para-after-table
  SB_E: smaller font (10pt instead of 10.5pt)
  SB_F: Meiryo 11pt (different metric)
"""
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import sys
    print("python-docx required: pip install python-docx")
    sys.exit(1)


OUT_DIR = Path(__file__).parent / "split_box_padding_repro"
OUT_DIR.mkdir(exist_ok=True)

MS_GOTHIC = "ＭＳ ゴシック"
MEIRYO = "Meiryo"


def set_table_borders_all(table, sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_run_font(run, font_name: str, pt: float):
    r = run._r
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    for child_tag in ('w:rFonts', 'w:sz', 'w:szCs'):
        old = rPr.find(qn(child_tag))
        if old is not None:
            rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(pt * 2)))
    rPr.append(szCs)


def set_section_docgrid(section, line_pitch_tw: int = 400):
    sectPr = section._sectPr
    old = sectPr.find(qn('w:docGrid'))
    if old is not None:
        sectPr.remove(old)
    dg = OxmlElement('w:docGrid')
    dg.set(qn('w:type'), 'linesAndChars')
    dg.set(qn('w:linePitch'), str(line_pitch_tw))
    dg.set(qn('w:charSpace'), '0')
    sectPr.append(dg)


def strip_pprdefault_spacing(doc):
    """Remove docDefaults pPr spacing (w:after=200 w:line=276) so repro
    matches d77a-style docDefaults (empty pPrDefault)."""
    styles = doc.styles.element
    for dd in styles.iter(qn('w:docDefaults')):
        for ppr_default in dd.iter(qn('w:pPrDefault')):
            # Remove all children (pPr and its spacing)
            for child in list(ppr_default):
                ppr_default.remove(child)


def build_base(filler: int = 18, font: str = MS_GOTHIC, pt: float = 10.5) -> Document:
    doc = Document()
    strip_pprdefault_spacing(doc)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    set_section_docgrid(section, line_pitch_tw=400)
    for i in range(filler):
        p = doc.add_paragraph()
        run = p.add_run(f"Filler line {i+1:02d}. " + "あ" * 30)
        set_run_font(run, font, pt)
    return doc


def _set_para_rpr_font(p, font_name: str, pt: float):
    """Set pPr/rPr font+size (affects empty paragraph line height)."""
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn('w:rPr'))
    if old is not None:
        pPr.remove(old)
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(pt * 2)))
    rPr.append(sz)
    pPr.append(rPr)


def add_cell_para(cell, text: str, font=MS_GOTHIC, pt=10.5, first_para=False):
    if first_para:
        p = cell.paragraphs[0]
        for r in list(p._p.findall(qn('w:r'))):
            p._p.remove(r)
    else:
        p = cell.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, font, pt)
    else:
        # Empty paragraph — set pPr/rPr font so line height is not inherited
        _set_para_rpr_font(p, font, pt)
    return p


def add_body_para(doc, text: str, font=MS_GOTHIC, pt=10.5):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, font, pt)
    else:
        _set_para_rpr_font(p, font, pt)
    return p


def build_repro(name: str, trailing_empty_cell: bool, empty_para_after_table: bool,
                font: str = MS_GOTHIC, pt: float = 10.5, filler: int = 18):
    """Build 1 repro."""
    doc = build_base(filler=filler, font=font, pt=pt)

    # 1x1 bordered table
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    cell = table.cell(0, 0)

    # Long content paragraph (wraps to ~8-10 lines)
    long_text = "".join(f"{i+1:02d}文字目アイウエオカキクケコサシスセソタチツテト" for i in range(18))
    add_cell_para(cell, long_text, font=font, pt=pt, first_para=True)

    # Optional trailing empty paragraph INSIDE cell
    if trailing_empty_cell:
        add_cell_para(cell, "", font=font, pt=pt, first_para=False)

    # After table
    if empty_para_after_table:
        add_body_para(doc, "", font=font, pt=pt)
    add_body_para(doc, "［解説］これは解説段落のテキストです。", font=font, pt=pt)
    add_body_para(doc, "ここは本文段落その 1 です。" + "本文内容。" * 15, font=font, pt=pt)

    path = OUT_DIR / f"{name}.docx"
    doc.save(path)
    return path


def main():
    # Filler count tuned so the table's first N lines fit on p.1 and remainder
    # spills to p.2. With no docDefault spacing, needs more filler than before.
    repros = [
        ("SB_A", False, False, MS_GOTHIC, 10.5, 28),
        ("SB_B", False, True,  MS_GOTHIC, 10.5, 28),
        ("SB_C", True,  False, MS_GOTHIC, 10.5, 28),
        ("SB_D", True,  True,  MS_GOTHIC, 10.5, 28),
        ("SB_E", True,  True,  MS_GOTHIC, 9.0,  32),
        ("SB_F", True,  True,  MEIRYO,   11.0, 26),
    ]
    for name, te, eat, font, pt, fil in repros:
        path = build_repro(name, te, eat, font, pt, fil)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
