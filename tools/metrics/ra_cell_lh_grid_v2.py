"""
Ra: #47の矛盾を解決 — adjustLineHeightInTable + grid snap の正確な関係
- COM Documents.Add() で作ったテーブル vs テンプレートのテーブル
- adjustLineHeightInTable の実際の値を確認
- compat mode による影響
- 実在する低SSIM文書のテーブルでのgrid snap挙動
"""
import win32com.client, json, os, tempfile
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.DispatchEx('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []


def test_com_created_table():
    """Table created via COM (not python-docx) — should match Word's default behavior."""
    wdoc = word.Documents.Add()
    try:
        ps = wdoc.Sections(1).PageSetup
        ps.LeftMargin = 72; ps.TopMargin = 72

        wdoc.Content.Text = ""

        # Add body paragraphs first (for grid comparison)
        for i in range(3):
            if i > 0:
                rng = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
                rng.InsertParagraphAfter()
            p = wdoc.Paragraphs(wdoc.Paragraphs.Count)
            p.Range.Text = f"Body para {i+1}"
            p.Range.Font.Name = "Calibri"
            p.Range.Font.Size = 11
            p.Format.SpaceBefore = 0
            p.Format.SpaceAfter = 0

        # Add table via COM
        rng = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
        rng.InsertParagraphAfter()
        rng2 = wdoc.Range(wdoc.Content.End - 1, wdoc.Content.End - 1)
        tbl = wdoc.Tables.Add(rng2, 3, 2)
        tbl.Borders.Enable = True

        for r in range(1, 4):
            for c in range(1, 3):
                tbl.Cell(r, c).Range.Text = f"R{r}C{c}"
                tbl.Cell(r, c).Range.Font.Name = "Calibri"
                tbl.Cell(r, c).Range.Font.Size = 11
                for pi in range(1, tbl.Cell(r, c).Range.Paragraphs.Count + 1):
                    tbl.Cell(r, c).Range.Paragraphs(pi).Format.SpaceBefore = 0
                    tbl.Cell(r, c).Range.Paragraphs(pi).Format.SpaceAfter = 0

        wdoc.Repaginate()

        data = {"scenario": "com_table"}

        # Check compatibility settings
        data["compat_mode"] = wdoc.CompatibilityMode

        # Body paragraph gaps
        data["body_gaps"] = []
        for i in range(1, 4):
            y = wdoc.Paragraphs(i).Range.Information(6)
            data["body_gaps"].append(round(y, 4))

        body_gap = data["body_gaps"][1] - data["body_gaps"][0]
        data["body_line_gap"] = round(body_gap, 4)

        # Table row gaps
        data["table_rows"] = []
        for r in range(1, 4):
            y = tbl.Cell(r, 1).Range.Paragraphs(1).Range.Information(6)
            data["table_rows"].append({"row": r, "y": round(y, 4)})

        for i in range(1, len(data["table_rows"])):
            data["table_rows"][i]["gap"] = round(
                data["table_rows"][i]["y"] - data["table_rows"][i - 1]["y"], 4)

        return data
    finally:
        wdoc.Close(False)


def test_template_table():
    """Table in ja_gov_template (which has grid settings)."""
    d = Document(TEMPLATE)
    for p in d.paragraphs:
        p._element.getparent().remove(p._element)

    # Check grid settings in template
    sec = d.sections[0]
    sectPr = sec._sectPr
    dg = sectPr.find(qn('w:docGrid'))
    grid_info = {}
    if dg is not None:
        grid_info["type"] = dg.get(qn('w:type'), 'none')
        grid_info["linePitch"] = dg.get(qn('w:linePitch'), 'none')

    # Add body paras
    for i in range(3):
        p = d.add_paragraph()
        r = p.add_run(f"Body para {i+1}")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pPr = p._element.get_or_add_pPr()
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')

    # Add table via XML
    tbl_el = d.element.body.makeelement(qn('w:tbl'), {})
    d.element.body.append(tbl_el)
    tblPr = etree.SubElement(tbl_el, qn('w:tblPr'))
    tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(tblBorders, qn(f'w:{side}'))
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')

    for row_i in range(3):
        tr = etree.SubElement(tbl_el, qn('w:tr'))
        for col_i in range(2):
            tc = etree.SubElement(tr, qn('w:tc'))
            p = etree.SubElement(tc, qn('w:p'))
            pPr = etree.SubElement(p, qn('w:pPr'))
            sp = etree.SubElement(pPr, qn('w:spacing'))
            sp.set(qn('w:before'), '0')
            sp.set(qn('w:after'), '0')
            run = etree.SubElement(p, qn('w:r'))
            rPr = etree.SubElement(run, qn('w:rPr'))
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:ascii'), 'Calibri')
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), '22')
            t = etree.SubElement(run, qn('w:t'))
            t.text = f"R{row_i+1}C{col_i+1}"

    tmp = os.path.join(tempfile.gettempdir(), "ra_tpl_table.docx")
    d.save(tmp)

    doc = word.Documents.Open(tmp)
    try:
        data = {"scenario": "template_table", "grid": grid_info}
        data["compat_mode"] = doc.CompatibilityMode

        # Body gaps
        data["body_gaps"] = []
        for i in range(1, 4):
            y = doc.Paragraphs(i).Range.Information(6)
            data["body_gaps"].append(round(y, 4))
        data["body_line_gap"] = round(data["body_gaps"][1] - data["body_gaps"][0], 4)

        # Table gaps
        tbl = doc.Tables(1)
        data["table_rows"] = []
        for r in range(1, 4):
            y = tbl.Cell(r, 1).Range.Paragraphs(1).Range.Information(6)
            data["table_rows"].append({"row": r, "y": round(y, 4)})

        for i in range(1, len(data["table_rows"])):
            data["table_rows"][i]["gap"] = round(
                data["table_rows"][i]["y"] - data["table_rows"][i - 1]["y"], 4)

        # Check adjustLineHeightInTable
        compat_val = doc.Compatibility(12)  # wdNoTabHangIndent=12?
        data["adjust_lh_compat"] = compat_val

        return data
    finally:
        doc.Close(False)
        os.unlink(tmp)


def test_compat_adjustlh():
    """Directly check Compatibility(65) = adjustLineHeightInTable."""
    wdoc = word.Documents.Add()
    try:
        data = {"scenario": "compat_check"}
        # Compatibility constants for adjustLineHeightInTable
        # The constant is wdAdjustLineHeightInTable = 65 (0-based from VBA docs)
        # But COM uses different numbering...
        # Let's check several values
        for i in [12, 36, 65, 66, 67, 68, 69, 70]:
            try:
                val = wdoc.Compatibility(i)
                data[f"compat_{i}"] = val
            except:
                data[f"compat_{i}"] = "error"

        data["compat_mode"] = wdoc.CompatibilityMode
        return data
    finally:
        wdoc.Close(False)


try:
    d1 = test_com_created_table()
    results.append(d1)
    print("=== com_table ===")
    print(f"  compat_mode={d1['compat_mode']}")
    print(f"  Body line gap: {d1['body_line_gap']}pt")
    for r in d1["table_rows"]:
        print(f"  R{r['row']}: y={r['y']}, gap={r.get('gap', '-')}")

    d2 = test_template_table()
    results.append(d2)
    print(f"\n=== template_table ===")
    print(f"  grid={d2['grid']}, compat_mode={d2['compat_mode']}")
    print(f"  Body line gap: {d2['body_line_gap']}pt")
    for r in d2["table_rows"]:
        print(f"  R{r['row']}: y={r['y']}, gap={r.get('gap', '-')}")
    print(f"  adjust_lh_compat={d2.get('adjust_lh_compat', '?')}")

    d3 = test_compat_adjustlh()
    results.append(d3)
    print(f"\n=== compat_check (compat_mode={d3['compat_mode']}) ===")
    for k, v in sorted(d3.items()):
        if k.startswith("compat_"):
            print(f"  {k} = {v}")

    # Compare body vs table gaps
    print(f"\n=== COMPARISON ===")
    print(f"  COM table: body_gap={d1['body_line_gap']}, table_gap={d1['table_rows'][1].get('gap', '?')}")
    print(f"  Template:  body_gap={d2['body_line_gap']}, table_gap={d2['table_rows'][1].get('gap', '?')}")

finally:
    word.Quit()

out_path = os.path.join(os.path.dirname(__file__), 'output', 'ra_cell_lh_grid_v2.json')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f"\nResults saved to {out_path}")
