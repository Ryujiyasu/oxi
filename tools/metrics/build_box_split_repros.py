"""Author minimal repro docx files: single-cell 1-row tables with content
spanning 2 pages. Vary paragraph count, line height, cell padding, and borders
so the close-border y formula can be derived.

Output dir: tools/metrics/box_split_repro/
"""
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Twips, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import sys
    print("python-docx required: pip install python-docx")
    sys.exit(1)


OUT_DIR = Path(__file__).parent / "box_split_repro"
OUT_DIR.mkdir(exist_ok=True)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            elem = tcBorders.find(qn(f'w:{edge}'))
            if elem is None:
                elem = OxmlElement(f'w:{edge}')
                tcBorders.append(elem)
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), str(kwargs[edge]))
            elem.set(qn('w:color'), 'auto')


def set_table_borders_all(table, sz=4):
    """Set all table borders (top/bottom/left/right/insideH/insideV)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    # Remove any existing tblBorders
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def set_cell_margins(cell, top=None, bottom=None, left=None, right=None):
    """Set cell margins (tcMar) in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
    # Remove any existing tcMar
    old = tcPr.find(qn('w:tcMar'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcMar)


def set_cell_valign(cell, val: str):
    """Set cell vertical alignment: top | center | bottom."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:vAlign'))
    if old is not None:
        tcPr.remove(old)
    el = OxmlElement('w:vAlign')
    el.set(qn('w:val'), val)
    tcPr.append(el)


def set_para_spacing(p, before_tw: int | None = None, after_tw: int | None = None):
    """Set paragraph spacing before/after in twips."""
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if before_tw is not None:
        spacing.set(qn('w:before'), str(before_tw))
    if after_tw is not None:
        spacing.set(qn('w:after'), str(after_tw))


def set_para_flag(p, tag: str, val: str = "0"):
    """Set a boolean pPr flag like widowControl, keepLines, keepNext."""
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn(f'w:{tag}'))
    if old is not None:
        pPr.remove(old)
    el = OxmlElement(f'w:{tag}')
    if val:
        el.set(qn('w:val'), val)
    pPr.append(el)


def set_table_cell_spacing(table, tw: int):
    """Set tblCellSpacing (tw = half-space around cells in twips)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblCellSpacing'))
    if old is not None:
        tblPr.remove(old)
    el = OxmlElement('w:tblCellSpacing')
    el.set(qn('w:w'), str(tw))
    el.set(qn('w:type'), 'dxa')
    tblPr.append(el)


def make_repro(name: str, n_paragraphs: int, line_count_per_para: int = 1,
               font_pt: int = 10.5, line_exact_pt: int | None = None,
               pad_bottom_twips: int = 0, pad_top_twips: int = 0,
               filler_paragraphs: int = 20,
               widow_control: bool | None = None,
               keep_lines_last: bool = False,
               keep_next_last: bool = False,
               space_before_tw: int | None = None,
               space_after_tw: int | None = None,
               v_align: str | None = None,
               cell_spacing_tw: int | None = None):
    """Build a docx with `filler_paragraphs` of body text, then a single-cell
    1-row table whose content is tall enough to split across page boundary.

    Each cell paragraph is `line_count_per_para` lines; total content
    is `n_paragraphs * line_count_per_para` lines.
    """
    doc = Document()
    # Set A4 page size, default margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Filler paragraphs to push table to just before page bottom
    for i in range(filler_paragraphs):
        p = doc.add_paragraph(f"Filler line {i+1}. " + "あ" * 30)
        p.runs[0].font.size = Pt(font_pt)

    # Single-cell 1-row table with all borders
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_all(table)
    if cell_spacing_tw is not None:
        set_table_cell_spacing(table, cell_spacing_tw)
    cell = table.cell(0, 0)
    # Set cell padding
    if pad_bottom_twips or pad_top_twips:
        set_cell_margins(cell, top=pad_top_twips, bottom=pad_bottom_twips)
    if v_align is not None:
        set_cell_valign(cell, v_align)
    # Replace default paragraph
    cell.text = ""
    for pi in range(n_paragraphs):
        if pi == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        # Make content long enough to wrap to `line_count_per_para` lines
        base = f"Row{pi+1}: " + "あ" * (40 * line_count_per_para)
        run = p.add_run(base)
        run.font.size = Pt(font_pt)
        if line_exact_pt is not None:
            pPr = p._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), str(line_exact_pt * 20))
            spacing.set(qn('w:lineRule'), 'exact')
            # Remove existing spacing
            old = pPr.find(qn('w:spacing'))
            if old is not None:
                pPr.remove(old)
            pPr.append(spacing)
        if space_before_tw is not None or space_after_tw is not None:
            set_para_spacing(p, before_tw=space_before_tw, after_tw=space_after_tw)
        if widow_control is False:
            set_para_flag(p, "widowControl", "0")
        is_last = (pi == n_paragraphs - 1)
        if is_last and keep_lines_last:
            set_para_flag(p, "keepLines", "")
        if is_last and keep_next_last:
            set_para_flag(p, "keepNext", "")

    path = OUT_DIR / f"{name}.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    # Repro A: default margins, 8 single-line paragraphs, default padding
    p1 = make_repro("repro_A_default", n_paragraphs=8, line_count_per_para=1,
                    filler_paragraphs=25)
    # Repro B: larger cell padding (200 twips = 10pt top+bottom)
    p2 = make_repro("repro_B_pad10", n_paragraphs=8, line_count_per_para=1,
                    pad_top_twips=200, pad_bottom_twips=200,
                    filler_paragraphs=25)
    # Repro C: exact line spacing 20pt, fewer paragraphs but still splits
    p3 = make_repro("repro_C_exact20", n_paragraphs=6, line_count_per_para=1,
                    line_exact_pt=20, filler_paragraphs=28)
    # Repro D: small padding, 2-line paragraphs
    p4 = make_repro("repro_D_2line", n_paragraphs=4, line_count_per_para=2,
                    pad_bottom_twips=100, filler_paragraphs=25)

    # Step-2 repro matrix: OAT variants from Repro A baseline (pad_b=0).
    # Each isolates one candidate variable for close_border_y formula.
    base = dict(n_paragraphs=8, line_count_per_para=1, filler_paragraphs=25)

    p5 = make_repro("repro_E_widowOff", **base, widow_control=False)
    p6 = make_repro("repro_F_keepLines", **base, keep_lines_last=True)
    p7 = make_repro("repro_G_keepNext", **base, keep_next_last=True)
    p8 = make_repro("repro_H_vAlignCenter", **base, v_align="center")
    p9 = make_repro("repro_I_vAlignBottom", **base, v_align="bottom")
    p10 = make_repro("repro_J_spaceAfter10", **base, space_after_tw=200)
    p11 = make_repro("repro_K_spaceBefore10", **base, space_before_tw=200)
    p12 = make_repro("repro_L_cellSpacing40", **base, cell_spacing_tw=40)

    print(f"Created 4 (A-D):\n  {p1}\n  {p2}\n  {p3}\n  {p4}")
    print(f"Created 8 (E-L):\n  {p5}\n  {p6}\n  {p7}\n  {p8}\n  {p9}\n  {p10}\n  {p11}\n  {p12}")
