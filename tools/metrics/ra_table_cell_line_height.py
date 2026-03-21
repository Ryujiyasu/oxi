"""Ra: テーブルセル内の行高さ完全解明
全条件をCOM計測して公式を導出する。
"""
import win32com.client, json, os, sys, tempfile
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = os.path.join(os.path.dirname(__file__), 'ja_gov_template.docx')

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False

results = []

def add(param, value, context):
    results.append({
        'parameter': param, 'value': round(value, 4),
        'unit': 'pt', 'domain': 'table_cell_line_height', 'context': context
    })
    print(f'  {param} = {value:.2f}pt', file=sys.stderr)

def make_table_doc(grid_type, grid_pitch, font_name, font_size, compat65=False):
    """Create docx with a 3-row table, specific grid, font settings."""
    doc = Document(TEMPLATE)
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    # Set docGrid
    sectPr = section._sectPr
    for dg in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(dg)
    if grid_type:
        dg = etree.SubElement(sectPr, qn('w:docGrid'))
        dg.set(qn('w:type'), grid_type)
        dg.set(qn('w:linePitch'), str(grid_pitch))
    elif grid_pitch:
        # type absent, linePitch only
        dg = etree.SubElement(sectPr, qn('w:docGrid'))
        dg.set(qn('w:linePitch'), str(grid_pitch))
    # else: no docGrid at all

    # Set compat65 (adjustLineHeightInTable)
    if compat65:
        settings_part = None
        for rel in doc.part.rels.values():
            if 'settings' in rel.reltype:
                settings_part = rel.target_part
                break
        if settings_part:
            settings_xml = settings_part._element
            compat = settings_xml.find(qn('w:compat'))
            if compat is None:
                compat = etree.SubElement(settings_xml, qn('w:compat'))
            cs = etree.SubElement(compat, qn('w:compatSetting'))
            cs.set(qn('w:name'), 'compatibilityMode')
            cs.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
            cs.set(qn('w:val'), '11')  # Word 2003 compat = adjustLineHeightInTable

    # Add table
    tbl = doc.add_table(rows=3, cols=2)
    for row in tbl.rows:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run('テスト')
            run.font.name = font_name
            run.font.size = Pt(font_size)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = etree.SubElement(rpr, qn('w:rFonts'))
            rfonts.set(qn('w:eastAsia'), font_name)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    tmp = os.path.join(tempfile.gettempdir(), 'ra_table_lh.docx')
    doc.save(tmp)
    return tmp

def measure_table_gap(docx_path):
    """Open docx, measure Row1->Row2 gap in first table."""
    try:
        wdoc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
        tbl = wdoc.Tables(1)
        y1 = tbl.Rows(1).Range.Information(6)
        y2 = tbl.Rows(2).Range.Information(6)
        gap = round(y2 - y1, 4)
        wdoc.Close(False)
        return gap
    except Exception as e:
        print(f'  ERROR: {e}', file=sys.stderr)
        return None

# === 1. Grid type × font × size ===
print('=== 1. Grid type effect on table cell line height ===', file=sys.stderr)

fonts = [
    ('ＭＳ 明朝', 'MS_Mincho'),
    ('ＭＳ ゴシック', 'MS_Gothic'),
    ('Calibri', 'Calibri'),
    ('Century', 'Century'),
    ('メイリオ', 'Meiryo'),
]

grid_configs = [
    ('lines', 360, 'lines_360'),
    ('lines', 272, 'lines_272'),
    ('linesAndChars', 360, 'lAC_360'),
    ('linesAndChars', 272, 'lAC_272'),
    (None, 360, 'noType_360'),      # linePitch only, no type
    (None, None, 'noGrid'),          # no docGrid element
]

for font_com, font_label in fonts:
    for size in [10.5, 11, 12]:
        for grid_type, grid_pitch, grid_label in grid_configs:
            tmp = make_table_doc(grid_type, grid_pitch, font_com, size)
            gap = measure_table_gap(tmp)
            if gap:
                add(f'cell_{font_label}_{size}_{grid_label}', gap,
                    f'font={font_label} size={size} grid={grid_label}')

print('=== 2. Normal paragraph comparison (same settings) ===', file=sys.stderr)

# For each config, also measure normal paragraph gap for comparison
for font_com, font_label in [('ＭＳ 明朝', 'MS_Mincho'), ('Calibri', 'Calibri')]:
    for size in [10.5, 11]:
        for grid_type, grid_pitch, grid_label in grid_configs:
            doc = Document(TEMPLATE)
            section = doc.sections[0]
            section.top_margin = Pt(72)
            sectPr = section._sectPr
            for dg in sectPr.findall(qn('w:docGrid')):
                sectPr.remove(dg)
            if grid_type:
                dg = etree.SubElement(sectPr, qn('w:docGrid'))
                dg.set(qn('w:type'), grid_type)
                dg.set(qn('w:linePitch'), str(grid_pitch))
            elif grid_pitch:
                dg = etree.SubElement(sectPr, qn('w:docGrid'))
                dg.set(qn('w:linePitch'), str(grid_pitch))

            for text in ['AAA', 'BBB']:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = font_com
                run.font.size = Pt(size)
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = etree.SubElement(rpr, qn('w:rFonts'))
                rfonts.set(qn('w:eastAsia'), font_com)
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)

            tmp = os.path.join(tempfile.gettempdir(), 'ra_normal_lh.docx')
            doc.save(tmp)
            try:
                wdoc = word.Documents.Open(os.path.abspath(tmp), ReadOnly=True)
                y1 = wdoc.Paragraphs(1).Range.Information(6)
                y2 = wdoc.Paragraphs(2).Range.Information(6)
                gap = round(y2 - y1, 4)
                wdoc.Close(False)
                add(f'normal_{font_label}_{size}_{grid_label}', gap,
                    f'NORMAL font={font_label} size={size} grid={grid_label}')
            except:
                pass

# === 3. adjustLineHeightInTable (compat65) ===
print('=== 3. adjustLineHeightInTable effect ===', file=sys.stderr)

for font_com, font_label in [('ＭＳ 明朝', 'MS_Mincho'), ('Calibri', 'Calibri')]:
    for size in [10.5, 11]:
        # compat65=True (adjustLineHeightInTable=true)
        tmp = make_table_doc('lines', 360, font_com, size, compat65=True)
        gap = measure_table_gap(tmp)
        if gap:
            add(f'cell_{font_label}_{size}_compat65', gap,
                f'font={font_label} size={size} adjustLH=true')

        # compat65=False (default)
        tmp = make_table_doc('lines', 360, font_com, size, compat65=False)
        gap = measure_table_gap(tmp)
        if gap:
            add(f'cell_{font_label}_{size}_default', gap,
                f'font={font_label} size={size} adjustLH=false(default)')

word.Quit()

# === Analysis ===
print('\n=== ANALYSIS ===', file=sys.stderr)
print(f'\nTotal measurements: {len(results)}')

# Compare cell vs normal
print('\nCell vs Normal comparison:')
cell_vals = {r['parameter']: r['value'] for r in results if r['parameter'].startswith('cell_')}
normal_vals = {r['parameter'].replace('normal_', 'cell_'): r['value'] for r in results if r['parameter'].startswith('normal_')}

print(f'{"Config":<40s} {"Cell":>8s} {"Normal":>8s} {"Same?":>6s}')
for key in sorted(cell_vals.keys()):
    cv = cell_vals[key]
    nv = normal_vals.get(key, None)
    if nv:
        same = 'YES' if abs(cv - nv) < 0.1 else 'NO'
        print(f'{key:<40s} {cv:8.2f} {nv:8.2f} {same:>6s}')

# Save
out_path = 'pipeline_data/ra_table_cell_lh.json'
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f'\nSaved: {out_path}')
