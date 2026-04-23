"""Exact mimic of 29dc6e 代表者又は管理人 row to find why real measures 37.5pt
while simple S1 measures 33pt. Adds:
- trHeight=371
- 2 cells (cell 2 empty, gridSpan-equivalent with wider width)
- vAlign=center on cell 2
- pStyle="ac" attributes inline
- rPr spacing=0 inside pPr (character spacing on mark)
"""
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.abspath("tools/metrics/d29_mimic_repro")
os.makedirs(OUT_DIR, exist_ok=True)


def mk_para_ac(cell, text, sa_tw, sb_tw):
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for name in ["widowControl", "wordWrap", "autoSpaceDE", "autoSpaceDN",
                 "adjustRightInd", "jc", "spacing", "rPr"]:
        ex = pPr.find(qn(f"w:{name}"))
        if ex is not None: pPr.remove(ex)
    # pStyle ac attrs
    for tag in ["widowControl", "wordWrap", "autoSpaceDE", "autoSpaceDN", "adjustRightInd"]:
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:val"), "0"); pPr.append(e)
    e = OxmlElement("w:jc"); e.set(qn("w:val"), "both"); pPr.append(e)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(sb_tw))
    sp.set(qn("w:after"), str(sa_tw))
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    # rPr for paragraph mark with character spacing 0
    pmark_rpr = OxmlElement("w:rPr")
    csp = OxmlElement("w:spacing"); csp.set(qn("w:val"), "0"); pmark_rpr.append(csp)
    pPr.append(pmark_rpr)
    # Add run
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "ＭＳ 明朝"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hAnsi"), "ＭＳ 明朝")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)
    csp2 = OxmlElement("w:spacing"); csp2.set(qn("w:val"), "0"); rPr.append(csp2)
    return p


def set_row_trheight(row, twips):
    trPr = row._tr.get_or_add_trPr()
    for ex in trPr.findall(qn("w:trHeight")):
        trPr.remove(ex)
    th = OxmlElement("w:trHeight")
    th.set(qn("w:val"), str(twips))
    trPr.append(th)


def set_cell_valign(cell, val):
    tcPr = cell._tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(ex)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), val)
    tcPr.append(va)


def make(name, has_cell2=True, trHeight=371, vAlign_center=True):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Twips(1134)
    sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1134)
    sec.right_margin = Twips(1134)
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "ＭＳ 明朝"

    cols = 2 if has_cell2 else 1
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.style = "Table Grid"
    if trHeight:
        set_row_trheight(tbl.rows[0], trHeight)

    # cell 1: 2 paragraphs
    cell1 = tbl.rows[0].cells[0]
    default_p = cell1.paragraphs[0]
    default_p._element.getparent().remove(default_p._element)
    mk_para_ac(cell1, "代表者又は", 87, 87)
    mk_para_ac(cell1, "管理人の氏名（フリガナ）", 87, 87)

    # cell 2: empty
    if has_cell2:
        cell2 = tbl.rows[0].cells[1]
        default_p2 = cell2.paragraphs[0]
        default_p2._element.getparent().remove(default_p2._element)
        mk_para_ac(cell2, "", 87, 87)
        if vAlign_center:
            set_cell_valign(cell2, "center")

    # ref rows
    for i in range(6):
        r = tbl.add_row()
        r.cells[0].text = f"ref{i+1}"

    out = os.path.join(OUT_DIR, name + ".docx")
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    make("D1_1cell_acstyle", has_cell2=False, trHeight=None, vAlign_center=False)
    make("D2_2cells_vAlign", has_cell2=True, trHeight=None, vAlign_center=True)
    make("D3_2cells_trH371", has_cell2=True, trHeight=371, vAlign_center=False)
    make("D4_2cells_vAlign_trH371", has_cell2=True, trHeight=371, vAlign_center=True)
