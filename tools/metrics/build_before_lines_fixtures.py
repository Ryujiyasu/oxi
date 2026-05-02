"""
Build fixtures for §2.4 beforeLines/afterLines formula validation.

Spec §2.4 formula:
  before_pt = beforeLines / 100 * line_pitch
With grid snap applied if snap_to_grid=true.

Sweep:
  - beforeLines ∈ {0, 50, 100, 150, 200, 250} (units: 1/100ths of a line)
  - linePitch ∈ {300, 360, 480} twips (= 15, 18, 24 pt grid)
  - body Calibri 11pt

Each fixture: 2 body paragraphs, P2 has the beforeLines applied.
Measure: P2.y - P1.y (the gap), compare to predicted gap.

Predicted gap = body_line_h + before_pt (for P2 with before set)
- body_line_h ≈ grid_pitch (when grid snap active)
- before_pt = beforeLines / 100 * line_pitch_pt
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.join(os.path.dirname(__file__), "output", "before_lines_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)


def add_doc_grid(section, line_pitch_tw):
    sectPr = section._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        docGrid = OxmlElement("w:docGrid")
        sectPr.append(docGrid)
    docGrid.set(qn("w:type"), "lines")
    docGrid.set(qn("w:linePitch"), str(line_pitch_tw))


def set_before_lines(p, before_lines):
    """Set <w:spacing w:beforeLines="N"/> on paragraph properties."""
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:beforeLines"), str(before_lines))
    spacing.set(qn("w:before"), "0")  # explicit 0pt before; beforeLines takes priority
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:afterLines"), "0")


def build_fixture(path, *, before_lines, line_pitch_tw):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    add_doc_grid(sec, line_pitch_tw)

    p1 = doc.add_paragraph("P1 baseline")
    p2 = doc.add_paragraph("P2 with beforeLines")
    p3 = doc.add_paragraph("P3 baseline")
    for p in (p1, p2, p3):
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    set_before_lines(p2, before_lines)

    doc.save(path)


def main():
    cases = []
    for pitch in [300, 360, 480]:
        for bl in [0, 50, 100, 150, 200, 250]:
            name = f"BL_pitch{pitch}tw_bl{bl:03d}.docx"
            path = os.path.join(OUT_DIR, name)
            try:
                build_fixture(path, before_lines=bl, line_pitch_tw=pitch)
                cases.append({"path": name, "pitch_tw": pitch, "before_lines": bl})
                print(f"  built {name}")
            except Exception as e:
                print(f"  ERR {name}: {e}")
    print(f"\nBuilt {len(cases)} fixtures in {OUT_DIR}")


if __name__ == "__main__":
    main()
