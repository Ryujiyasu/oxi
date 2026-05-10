"""Series PB_DRIFT_EMPTY — empty paragraph height behaviour.

Document: paragraph A (text) → paragraph B (empty) → paragraph C (text).
Empty paragraph height is determined by:
- Paragraph mark's font size (rPr inside pPr OR para's Run rPr)
- Paragraph's line spacing rule

Variants vary the empty paragraph's properties:
- DR_EMP_01: empty para no overrides (inherit body fs=10.5 lh=auto)
- DR_EMP_02: empty para fs=14 (via paragraph mark rPr)
- DR_EMP_03: empty para fs=8
- DR_EMP_04: empty para lh=Exactly 14
- DR_EMP_05: empty para lh=Exactly 21
- DR_EMP_06: empty para fs=14 + lh=Exactly 21
- DR_EMP_07: empty para mult 1.5
- DR_EMP_08: 5 consecutive empty paras (test cumulative)
- DR_EMP_09: empty para w:rPr placed at para_mark only (not run)
- DR_EMP_10: A=fs=10.5 / B=empty fs=11.5 lh=Exact 16 / C=fs=10.5 (bd90b00 pi=11 trigger)
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent


def make_run(para, text="あ", font_name="ＭＳ 明朝", size_pt=10.5):
    if not text: return None
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.insert(0, rFonts)
    return run


def set_para_mark_font(para, font_name="ＭＳ 明朝", size_pt=10.5):
    """Set paragraph mark rPr (inside pPr) — controls empty paragraph height."""
    pPr = para._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:rPr")): pPr.remove(old)
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
    rPr.append(sz)
    pPr.append(rPr)


def apply_lh(para, spec):
    pf = para.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    if spec == "auto" or spec is None: return
    if isinstance(spec, tuple):
        kind, val = spec
        if kind == "exact":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(val)
        elif kind == "multiple":
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = float(val)


def set_section(section):
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    sectPr = section._sectPr
    for ref in sectPr.findall(qn("w:headerReference")): sectPr.remove(ref)
    for ref in sectPr.findall(qn("w:footerReference")): sectPr.remove(ref)
    for old in sectPr.findall(qn("w:docGrid")): sectPr.remove(old)


def build(variant_id, empty_specs):
    """empty_specs is list of dicts: each dict for one empty paragraph
    {paramark_fs: float, lh: spec}."""
    doc = Document()
    set_section(doc.sections[0])
    # Paragraph A
    pA = doc.add_paragraph()
    pA.paragraph_format.space_before = Pt(0); pA.paragraph_format.space_after = Pt(0)
    make_run(pA, "A", size_pt=10.5)
    # Empty paragraphs
    for spec in empty_specs:
        pE = doc.add_paragraph()
        if spec.get('paramark_fs'):
            set_para_mark_font(pE, size_pt=spec['paramark_fs'])
        apply_lh(pE, spec.get('lh'))
    # Paragraph C
    pC = doc.add_paragraph()
    pC.paragraph_format.space_before = Pt(0); pC.paragraph_format.space_after = Pt(0)
    make_run(pC, "C", size_pt=10.5)
    out = HERE / f"{variant_id}.docx"
    doc.save(str(out))


VARIANTS = [
    ("DR_EMP_01", [dict()]),  # inherit body
    ("DR_EMP_02", [dict(paramark_fs=14.0)]),
    ("DR_EMP_03", [dict(paramark_fs=8.0)]),
    ("DR_EMP_04", [dict(lh=("exact", 14.0))]),
    ("DR_EMP_05", [dict(lh=("exact", 21.0))]),
    ("DR_EMP_06", [dict(paramark_fs=14.0, lh=("exact", 21.0))]),
    ("DR_EMP_07", [dict(lh=("multiple", 1.5))]),
    ("DR_EMP_08", [dict()] * 5),  # 5 consecutive empties
    ("DR_EMP_09", [dict(paramark_fs=10.5)]),  # explicit but matching body
    ("DR_EMP_10", [dict(paramark_fs=11.5, lh=("exact", 16.0))]),  # bd90b00 pi=11 ish
]


def main():
    print(f"Generating {len(VARIANTS)} PB_DRIFT_EMPTY variants in {HERE}")
    for vid, specs in VARIANTS:
        build(vid, specs)
        n = len(specs)
        print(f"  {vid}: {n} empty para(s), specs={specs}")
    print("Done.")


if __name__ == "__main__":
    main()
