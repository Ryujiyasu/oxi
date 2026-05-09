"""Series PB_DG — page-break tolerance with docGrid present.

Same baseline as PB_OF but adds <w:docGrid w:linePitch="400" w:type="lines"/>
or "linesAndChars" to the section.

Layout invariants (same as PB_OF):
- A4, top=bottom=72pt, no header/footer
- 47 fill paragraphs (MS Mincho 10.5pt, lh=Exactly 14pt)
- 1 test paragraph with space-before X = D + 25.7
- D values: -3, -1, 0, +1, +3, +5, +7, +10pt

Variants:
- PB_DG_A_NN: docGrid type="lines" linePitch=400 (20pt)
- PB_DG_B_NN: docGrid type="linesAndChars" linePitch=400 charPitch=210
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
N_FILL = 47
LH_PT = 14.0
D_VALUES = [-3, -1, 0, 1, 3, 5, 7, 10]


def make_run(para, text, font_name="ＭＳ 明朝", size_pt=10.5):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.insert(0, rFonts)
    return run


def set_para_lh_exact(para, pt_value):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pt_value)


def disable_widow_control(para):
    pPr = para._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:widowControl")):
        pPr.remove(old)
    el = OxmlElement("w:widowControl")
    el.set(qn("w:val"), "0")
    pPr.insert(0, el)


def set_section_geometry_and_grid(section, grid_type="lines", line_pitch=400, char_pitch=None):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    sectPr = section._sectPr
    for ref in sectPr.findall(qn("w:headerReference")):
        sectPr.remove(ref)
    for ref in sectPr.findall(qn("w:footerReference")):
        sectPr.remove(ref)
    # Remove existing docGrid
    for old in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(old)
    # Add new docGrid
    grid = OxmlElement("w:docGrid")
    grid.set(qn("w:type"), grid_type)
    grid.set(qn("w:linePitch"), str(line_pitch))
    if char_pitch is not None:
        grid.set(qn("w:charSpace"), str(char_pitch))
    sectPr.append(grid)


def build(variant_id: str, D: float, grid_type: str, line_pitch: int, char_pitch=None):
    X = D + 25.7
    doc = Document()
    set_section_geometry_and_grid(doc.sections[0], grid_type, line_pitch, char_pitch)

    for i in range(N_FILL):
        p = doc.add_paragraph()
        set_para_lh_exact(p, LH_PT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        disable_widow_control(p)
        make_run(p, "あ")

    test_p = doc.add_paragraph()
    set_para_lh_exact(test_p, LH_PT)
    test_p.paragraph_format.space_before = Pt(X)
    test_p.paragraph_format.space_after = Pt(0)
    disable_widow_control(test_p)
    make_run(test_p, "TEST")

    out = HERE / f"{variant_id}.docx"
    doc.save(str(out))


def main():
    print(f"Generating PB_DG variants in {HERE}")
    # Series A: docGrid type="lines"
    for i, D in enumerate(D_VALUES, 1):
        vid = f"PB_DG_A_{i:02d}"
        build(vid, D, "lines", 400)
        print(f"  {vid}: D={D:+3d}pt grid=lines pitch=400")
    # Series B: docGrid type="linesAndChars"
    for i, D in enumerate(D_VALUES, 1):
        vid = f"PB_DG_B_{i:02d}"
        build(vid, D, "linesAndChars", 400, 210)
        print(f"  {vid}: D={D:+3d}pt grid=linesAndChars pitch=400 chars=210")
    print("Done.")


if __name__ == "__main__":
    main()
