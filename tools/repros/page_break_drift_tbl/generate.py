"""Series PB_DRIFT_TBL — cursor advance through table entry/exit.

Document structure: body paragraph A → table (single cell, single para B)
→ body paragraph C. Measure y_A, y_B (paragraph in cell), y_C.

Word: y_B and y_C reflect actual layout positions. table entry/exit
advance includes cell padding, borders, row top/bottom overhead.

Variants vary table parameters:
- DR_TBL_01: 1×1 table, no border, no padding, single cell para fs=10.5
- DR_TBL_02: 1×1 table, default Word border (0.5pt), no padding
- DR_TBL_03: 1×1 table, default border + default cell margins (108 dxa = 5.4pt)
- DR_TBL_04: 1×1 table, no border, no padding, multi-line cell (3 paras)
- DR_TBL_05: 2-row table, no border, no padding
- DR_TBL_06: trHeight = exact 30pt, no border, no padding
- DR_TBL_07: trHeight = atLeast 30pt, no border, no padding
- DR_TBL_08: 1×1 table, fs=14 in cell (heading-like)
- DR_TBL_09: 1×1 table, lh=Exact 16 in cell para
- DR_TBL_10: 1×1 table, body para fs=10.5 with cell using small fs=8
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches, Twips
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent


def make_run(para, text="あ", font_name="ＭＳ 明朝", size_pt=10.5):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.insert(0, rFonts)
    return run


def set_section(section):
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    sectPr = section._sectPr
    for ref in sectPr.findall(qn("w:headerReference")): sectPr.remove(ref)
    for ref in sectPr.findall(qn("w:footerReference")): sectPr.remove(ref)
    for old in sectPr.findall(qn("w:docGrid")): sectPr.remove(old)


def remove_borders(table):
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); table._element.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}"); el.set(qn("w:val"), "nil")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_default_borders(table, sz=4):
    """Set thin Word default borders (sz in 1/8pt; 4 = 0.5pt)."""
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); table._element.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_no_padding(cell):
    tcPr = cell._element.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")): tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0"); m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def set_default_padding(cell, dxa=108):
    tcPr = cell._element.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")): tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(dxa)); m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def set_row_height(row, val_pt, rule):
    """rule: 'atLeast' or 'exact'."""
    trPr = row._element.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")): trPr.remove(old)
    el = OxmlElement("w:trHeight")
    el.set(qn("w:val"), str(int(val_pt * 20)))
    el.set(qn("w:hRule"), rule)
    trPr.append(el)


def apply_lh_exact(para, val):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(val)


def build_basic_table(doc, n_rows=1, fs=10.5, lh_spec=None,
                     borders="none", padding="none",
                     row_height=None, n_paras_per_cell=1):
    table = doc.add_table(rows=n_rows, cols=1)
    if borders == "none":
        remove_borders(table)
    elif borders == "default":
        set_default_borders(table)
    if padding == "none":
        for row in table.rows:
            for cell in row.cells:
                set_no_padding(cell)
    elif padding == "default":
        for row in table.rows:
            for cell in row.cells:
                set_default_padding(cell, 108)
    if row_height:
        for row in table.rows:
            set_row_height(row, row_height[0], row_height[1])
    # Cell paragraphs
    for r_idx, row in enumerate(table.rows):
        cell = row.cells[0]
        # Reuse first auto-created paragraph
        existing = cell.paragraphs[0]
        if lh_spec:
            apply_lh_exact(existing, lh_spec)
        existing.paragraph_format.space_before = Pt(0)
        existing.paragraph_format.space_after = Pt(0)
        make_run(existing, "あ", size_pt=fs)
        from docx.text.paragraph import Paragraph
        for _ in range(n_paras_per_cell - 1):
            new_p = OxmlElement("w:p")
            cell._element.append(new_p)
            np = Paragraph(new_p, cell)
            if lh_spec:
                apply_lh_exact(np, lh_spec)
            np.paragraph_format.space_before = Pt(0)
            np.paragraph_format.space_after = Pt(0)
            make_run(np, "い", size_pt=fs)
    return table


def build_doc(variant_id, body_fs=10.5, **table_kwargs):
    doc = Document()
    set_section(doc.sections[0])
    # Body paragraph A
    pA = doc.add_paragraph()
    pA.paragraph_format.space_before = Pt(0)
    pA.paragraph_format.space_after = Pt(0)
    make_run(pA, "A", size_pt=body_fs)
    # Table
    build_basic_table(doc, **table_kwargs)
    # Body paragraph C
    pC = doc.add_paragraph()
    pC.paragraph_format.space_before = Pt(0)
    pC.paragraph_format.space_after = Pt(0)
    make_run(pC, "C", size_pt=body_fs)
    out = HERE / f"{variant_id}.docx"
    doc.save(str(out))


VARIANTS = [
    ("DR_TBL_01", dict(fs=10.5, borders="none", padding="none")),
    ("DR_TBL_02", dict(fs=10.5, borders="default", padding="none")),
    ("DR_TBL_03", dict(fs=10.5, borders="default", padding="default")),
    ("DR_TBL_04", dict(fs=10.5, borders="none", padding="none", n_paras_per_cell=3)),
    ("DR_TBL_05", dict(fs=10.5, borders="none", padding="none", n_rows=2)),
    ("DR_TBL_06", dict(fs=10.5, borders="none", padding="none", row_height=(30.0, "exact"))),
    ("DR_TBL_07", dict(fs=10.5, borders="none", padding="none", row_height=(30.0, "atLeast"))),
    ("DR_TBL_08", dict(fs=14.0, borders="none", padding="none")),
    ("DR_TBL_09", dict(fs=10.5, borders="none", padding="none", lh_spec=16.0)),
    ("DR_TBL_10", dict(fs=8.0, borders="none", padding="none")),
]


def main():
    print(f"Generating {len(VARIANTS)} PB_DRIFT_TBL variants in {HERE}")
    for vid, kwargs in VARIANTS:
        build_doc(vid, **kwargs)
        print(f"  {vid}: {kwargs}")
    print("Done.")


if __name__ == "__main__":
    main()
