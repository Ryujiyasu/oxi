"""Series PB_OF — page-break overflow tolerance characterization.

Goal: characterize Word's tolerance for last-paragraph overflow at exact
page boundary. For each target overflow D (pt), generate a docx where
the test paragraph would overflow page 1 by D pt under Oxi's strict
`cursor_y + lh > page_bottom` rule. Measure with COM whether Word
places the test paragraph on page 1 (fits) or page 2 (breaks).

Layout invariants (held constant across variants):
- Page: A4 (8.27 in × 11.69 in = 595.3 × 841.7 pt)
- Margins: top=bottom=72pt (1in), left=right=72pt
- Content area height: 841.7 - 144 = 697.7 pt
- No header, no footer
- Body paragraphs: MS Mincho 10.5pt, lh = Exactly 14pt
- 47 fill paragraphs each "あ" → cursor advance 47 * 14 = 658pt
  (cursor lands at 72 + 658 = 730pt at end of fill)
- 1 test paragraph: spacing-before = X pt, lh = Exactly 14pt
  - test paragraph top = 730 + X
  - test paragraph bottom = 744 + X
  - Page bottom = 72 + 697.7 = 769.7 pt (using exact A4 height)

For target D = (test bottom) - (page bottom) = (744 + X) - 769.7 = X - 25.7
  → X = D + 25.7

For D values -15..+10pt: X values 10.7..35.7pt.

Variants:
  PB_OF_01 D=-15  X=10.7
  PB_OF_02 D=-10  X=15.7
  PB_OF_03 D= -5  X=20.7
  PB_OF_04 D= -2  X=23.7
  PB_OF_05 D= -1  X=24.7
  PB_OF_06 D=  0  X=25.7  ← exact boundary
  PB_OF_07 D=+ 1  X=26.7
  PB_OF_08 D=+ 2  X=27.7
  PB_OF_09 D=+ 3  X=28.7
  PB_OF_10 D=+ 5  X=30.7
  PB_OF_11 D=+ 7  X=32.7
  PB_OF_12 D=+10  X=35.7
"""
from __future__ import annotations
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Mm, Inches
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: pip install python-docx", file=sys.stderr); sys.exit(1)


HERE = Path(__file__).resolve().parent
N_FILL = 47
LH_PT = 14.0  # Exactly line height for both fill and test paragraphs

VARIANTS = [
    # (variant_id, D_overflow_pt)
    ("PB_OF_01", -15),
    ("PB_OF_02", -10),
    ("PB_OF_03",  -5),
    ("PB_OF_04",  -2),
    ("PB_OF_05",  -1),
    ("PB_OF_06",   0),
    ("PB_OF_07",   1),
    ("PB_OF_08",   2),
    ("PB_OF_09",   3),
    ("PB_OF_10",   5),
    ("PB_OF_11",   7),
    ("PB_OF_12",  10),
]


def make_run(para, text, font_name="ＭＳ 明朝", size_pt=10.5):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.insert(0, rFonts)
    return run


def set_para_lh_exact(para, pt_value):
    """Set paragraph line spacing to Exactly N pt."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pt_value)


def set_para_space_before(para, pt_value):
    """Set paragraph spacing-before in pt."""
    pf = para.paragraph_format
    pf.space_before = Pt(pt_value)


def set_section_geometry(section):
    # A4
    section.page_width = Mm(210)   # 595.28pt
    section.page_height = Mm(297)  # 841.89pt
    # 1-inch margins
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    # Disable headers/footers explicitly via sectPr
    sectPr = section._sectPr
    # Remove any existing headerReference / footerReference
    for ref in sectPr.findall(qn("w:headerReference")):
        sectPr.remove(ref)
    for ref in sectPr.findall(qn("w:footerReference")):
        sectPr.remove(ref)


def disable_widow_control(para):
    """Disable widowControl so the test paragraph isn't pushed to next page."""
    pPr = para._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:widowControl")):
        pPr.remove(old)
    el = OxmlElement("w:widowControl")
    el.set(qn("w:val"), "0")
    pPr.insert(0, el)


def build(variant_id: str, D_overflow_pt: float) -> None:
    X_space_before = D_overflow_pt + 25.7
    doc = Document()
    section = doc.sections[0]
    set_section_geometry(section)

    # Set default paragraph: no spacing-before/after, lh=Exactly 14
    # (We'll override per paragraph.)

    fill_paras = []
    for i in range(N_FILL):
        p = doc.add_paragraph()
        set_para_lh_exact(p, LH_PT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        disable_widow_control(p)
        make_run(p, "あ")
        fill_paras.append(p)

    # Test paragraph with spacing-before
    test_p = doc.add_paragraph()
    set_para_lh_exact(test_p, LH_PT)
    test_p.paragraph_format.space_before = Pt(X_space_before)
    test_p.paragraph_format.space_after = Pt(0)
    disable_widow_control(test_p)
    make_run(test_p, "TEST")

    out = HERE / f"{variant_id}.docx"
    doc.save(str(out))
    return X_space_before


def main():
    print(f"Generating {len(VARIANTS)} PB_OF variants in {HERE}")
    print(f"  N_fill={N_FILL}, lh=Exactly {LH_PT}pt")
    for vid, D in VARIANTS:
        X = build(vid, D)
        print(f"  {vid}: D={D:+4d}pt → space_before={X:.1f}pt")
    print("\nDone. Run measure.py next.")


if __name__ == "__main__":
    sys.exit(main() or 0)
