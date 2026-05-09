"""Series PB_DRIFT_LH — per-paragraph cursor advance under line-height rule transitions.

Each variant: 3 paragraphs A/B/C, all fs=10.5 MS Mincho, varying lh rule.
- auto: lh = font default (no override)
- Exactly N: lh_rule=exact, lh=N pt
- Multiple X: lh_rule=multiple, lh = X * font_default

bd90b00 has many lh=Exact 16pt paragraphs at fs=11.5/10.5 (Day 32 part 6
identified pi=11 as a transition trigger). This series isolates lh
transitions while holding fs constant.
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent

# (variant_id, A_spec, B_spec, C_spec) where spec is ("auto" | ("exact", N) | ("multiple", X))
VARIANTS = [
    ("DR_LH_01", "auto", "auto", "auto"),                    # baseline
    ("DR_LH_02", "auto", ("exact", 14.0), "auto"),
    ("DR_LH_03", "auto", ("exact", 16.0), "auto"),
    ("DR_LH_04", "auto", ("exact", 18.0), "auto"),
    ("DR_LH_05", "auto", ("exact", 21.0), "auto"),           # 1.5x of 10.5*1.33
    ("DR_LH_06", "auto", ("multiple", 1.15), "auto"),
    ("DR_LH_07", "auto", ("multiple", 1.5), "auto"),
    ("DR_LH_08", ("exact", 16.0), ("exact", 16.0), ("exact", 16.0)),  # all exact
    ("DR_LH_09", ("exact", 14.0), ("exact", 16.0), ("exact", 14.0)),  # exact transition
    ("DR_LH_10", "auto", ("exact", 14.0), ("exact", 14.0)),  # auto→exact, no return
]


def make_run(para, text="あ", font_name="ＭＳ 明朝", size_pt=10.5):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.insert(0, rFonts)
    return run


def apply_lh(para, spec):
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if spec == "auto":
        return
    if isinstance(spec, tuple):
        kind, val = spec
        if kind == "exact":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(val)
        elif kind == "multiple":
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = float(val)


def set_section_geometry(section):
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    sectPr = section._sectPr
    for ref in sectPr.findall(qn("w:headerReference")): sectPr.remove(ref)
    for ref in sectPr.findall(qn("w:footerReference")): sectPr.remove(ref)
    for old in sectPr.findall(qn("w:docGrid")): sectPr.remove(old)


def build(variant_id, A_spec, B_spec, C_spec):
    doc = Document()
    set_section_geometry(doc.sections[0])
    for spec in (A_spec, B_spec, C_spec):
        p = doc.add_paragraph()
        apply_lh(p, spec)
        make_run(p)
    out = HERE / f"{variant_id}.docx"
    doc.save(str(out))


def fmt(spec):
    if spec == "auto": return "auto"
    if isinstance(spec, tuple): return f"{spec[0][:3]}{spec[1]}"
    return str(spec)


def main():
    print(f"Generating {len(VARIANTS)} PB_DRIFT_LH variants in {HERE}")
    for vid, A, B, C in VARIANTS:
        build(vid, A, B, C)
        print(f"  {vid}: A={fmt(A):>10} / B={fmt(B):>10} / C={fmt(C):>10}")
    print("Done.")


if __name__ == "__main__":
    main()
