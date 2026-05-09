"""Series PB_DRIFT_CUMUL — long sequence of identical paragraphs, expose cumulative-round.

Single variant per (font, fs, lh-rule) combination. 30 paragraphs each
"あ", measure y of every paragraph. Word uses cumulative ceil for body
LM0; Oxi uses per-paragraph round → per-paragraph advance diverges.

Variants:
- DR_CU_01: 30× MS Mincho 10.5 lh=auto
- DR_CU_02: 30× MS Mincho 11   lh=auto
- DR_CU_03: 30× MS Mincho 11.5 lh=auto
- DR_CU_04: 30× MS Mincho 12   lh=auto
- DR_CU_05: 30× MS Mincho 14   lh=auto
- DR_CU_06: 30× MS Mincho 10.5 lh=Multiple 1.15
- DR_CU_07: 30× MS Mincho 10.5 lh=Multiple 1.5
- DR_CU_08: 30× MS Mincho 10.5 lh=Exactly 14   (no cumulative — exact is fixed)
- DR_CU_09: 30× MS Mincho 11.5 lh=Exactly 16   (Day 32 part 6 trigger context)
- DR_CU_10: 30× Yu Mincho 10.5 lh=auto         (different font metrics)
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
N = 30

VARIANTS = [
    ("DR_CU_01", "ＭＳ 明朝", 10.5, "auto"),
    ("DR_CU_02", "ＭＳ 明朝", 11.0, "auto"),
    ("DR_CU_03", "ＭＳ 明朝", 11.5, "auto"),
    ("DR_CU_04", "ＭＳ 明朝", 12.0, "auto"),
    ("DR_CU_05", "ＭＳ 明朝", 14.0, "auto"),
    ("DR_CU_06", "ＭＳ 明朝", 10.5, ("multiple", 1.15)),
    ("DR_CU_07", "ＭＳ 明朝", 10.5, ("multiple", 1.5)),
    ("DR_CU_08", "ＭＳ 明朝", 10.5, ("exact", 14.0)),
    ("DR_CU_09", "ＭＳ 明朝", 11.5, ("exact", 16.0)),
    ("DR_CU_10", "Yu Mincho", 10.5, "auto"),
]


def make_run(para, text, font_name, size_pt):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.insert(0, rFonts)
    return run


def apply_lh(para, spec):
    pf = para.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    if spec == "auto": return
    if isinstance(spec, tuple):
        kind, val = spec
        if kind == "exact":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(val)
        elif kind == "multiple":
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = float(val)


def set_section_geometry(section):
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    sectPr = section._sectPr
    for ref in sectPr.findall(qn("w:headerReference")): sectPr.remove(ref)
    for ref in sectPr.findall(qn("w:footerReference")): sectPr.remove(ref)
    for old in sectPr.findall(qn("w:docGrid")): sectPr.remove(old)


def build(variant_id, font, fs, lh_spec):
    doc = Document()
    set_section_geometry(doc.sections[0])
    for _ in range(N):
        p = doc.add_paragraph()
        apply_lh(p, lh_spec)
        make_run(p, "あ", font, fs)
    out = HERE / f"{variant_id}.docx"
    doc.save(str(out))


def main():
    print(f"Generating {len(VARIANTS)} PB_DRIFT_CUMUL variants in {HERE}")
    for vid, font, fs, lh in VARIANTS:
        build(vid, font, fs, lh)
        print(f"  {vid}: {font} {fs}pt lh={lh}  N={N}")
    print("Done.")


if __name__ == "__main__":
    main()
