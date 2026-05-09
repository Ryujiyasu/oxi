"""Series PB_DRIFT_FS — per-paragraph cursor advance under font-size transitions.

For each variant, generate a 3-paragraph docx:
  Para A (fs=A_fs, lh=auto)
  Para B (fs=B_fs, lh=auto)  ← test paragraph
  Para C (fs=C_fs, lh=auto)

Each paragraph contains "あ".

Measurements (separate measure.py):
- Word: COM Information(6) per paragraph → y_A, y_B, y_C
- Oxi: oxi-gdi-renderer --dump-layout → first-text-element y per pi
- advance_AB = y_B - y_A   (cursor advance after A = lh_A)
- advance_BC = y_C - y_B   (cursor advance after B = lh_B)

Comparison reveals: does Oxi compute lh_A and lh_B the same as Word for
each fs? If consistent, no fs-transition drift. If divergent, the
cursor advance formula is the drift source.

Layout: A4, 1in margins, no header/footer/grid (clean baseline).
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent

# (variant_id, A_fs, B_fs, C_fs)
VARIANTS = [
    ("DR_FS_01", 10.5, 10.5, 10.5),  # baseline constant
    ("DR_FS_02", 10.5, 11.0, 10.5),
    ("DR_FS_03", 10.5, 12.0, 10.5),
    ("DR_FS_04", 10.5, 14.0, 10.5),  # large jump up
    ("DR_FS_05", 14.0, 10.5, 14.0),  # large jump down then up
    ("DR_FS_06", 11.5, 10.5, 11.5),  # Day 32 part 6 fs change context
    ("DR_FS_07", 10.5,  8.0, 10.5),  # small
    ("DR_FS_08",  8.0, 10.5,  8.0),  # start small
    ("DR_FS_09", 10.5, 16.0, 10.5),  # very large
    ("DR_FS_10",  9.0, 10.5,  9.0),  # 9pt context (common in 0e7af)
]


def make_run(para, text, font_name="ＭＳ 明朝", size_pt=10.5):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    rPr.insert(0, rFonts)
    return run


def set_section_geometry(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    sectPr = section._sectPr
    for ref in sectPr.findall(qn("w:headerReference")):
        sectPr.remove(ref)
    for ref in sectPr.findall(qn("w:footerReference")):
        sectPr.remove(ref)
    for old in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(old)


def build(variant_id, A_fs, B_fs, C_fs):
    doc = Document()
    set_section_geometry(doc.sections[0])

    for fs in (A_fs, B_fs, C_fs):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        make_run(p, "あ", size_pt=fs)

    out = HERE / f"{variant_id}.docx"
    doc.save(str(out))


def main():
    print(f"Generating {len(VARIANTS)} PB_DRIFT_FS variants in {HERE}")
    for vid, A, B, C in VARIANTS:
        build(vid, A, B, C)
        print(f"  {vid}: A={A:.1f} / B={B:.1f} / C={C:.1f}")
    print("Done.")


if __name__ == "__main__":
    main()
