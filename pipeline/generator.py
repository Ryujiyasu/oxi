import anthropic
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from .config import (
    ANTHROPIC_API_KEY, CLAUDE_MODEL, DOCX_DIR, BATCH_SIZE,
)

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

SYSTEM_PROMPT = """
あなたはWordのレイアウトエンジンのエキスパートQAエンジニアです。
Wordのレンダリングエンジンのバグを引き出しやすい.docxの構造を設計してください。

## 方針
- ターゲット: Word 365 Windows 日本語ロケール
- difficulty: HIGH のケースのみ生成する
- 単純な段落1つではなく、レイアウトが複雑になる組み合わせを重視
- 日本語文字列を積極的に使う
- 長い段落（3行以上折り返すもの）を含める
- 複数段落・複数要素の組み合わせを含める

## 重点カテゴリ（難しいものを優先）
1. 表の中の禁則処理 — セル内での行折り返し＋禁則
2. 複数フォントサイズ混在行 — 10.5pt + 24pt + 8pt を同一行に
3. ページをまたぐ表 — 行の途中でページが切れる
4. 段落前後のスペーシング — beforeLines/afterLines + グリッドスナップ
5. CJK + Latin混在の行折り返し — 英単語途中での改行禁止
6. ネストした表 — 表の中に表
7. テキストボックス + 本文の回り込み — wrapSquare/wrapTight
8. 箇条書き + インデント階層 — 3段以上のネスト
9. セル結合 + 自動幅計算 — 横結合＋縦結合の組み合わせ
10. 行間 exact/atLeast + CJKフォント — グリッドスナップとの干渉

## 出力形式（JSON配列のみ）
[
  {
    "id": "一意のID（英数字とアンダースコアのみ）",
    "category": "カテゴリ名",
    "description": "何を検証するか",
    "difficulty": "LOW | MEDIUM | HIGH",
    "elements": [
      {
        "type": "paragraph | table | image | header | footer",
        "content": "要素の詳細仕様（日本語テキストを含める）",
        "style": "スタイル名（省略可）",
        "font": "フォント名（省略可）",
        "font_size": 10.5
      }
    ]
  }
]
"""


def generate_test_documents(count: int = BATCH_SIZE) -> list[str]:
    """テスト文書を生成して.docxファイルのパスリストを返す"""

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=8192,
        system=SYSTEM_PROMPT,
        messages=[{
            "role": "user",
            "content": (
                f"Wordレンダリングのエッジケースを {count} 件設計してください。"
                "日本語文書で発生しやすいレイアウトのズレに注目してください。"
            )
        }]
    )

    raw = response.content[0].text.strip()
    if "```json" in raw:
        raw = raw.split("```json")[1].split("```")[0].strip()
    elif "```" in raw:
        raw = raw.split("```")[1].split("```")[0].strip()

    specs = json.loads(raw)

    Path(DOCX_DIR).mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    generated_paths = []

    for i, spec in enumerate(specs):
        doc_id = spec.get("id", f"{timestamp}_{i:04d}")
        docx_path = str(Path(DOCX_DIR) / f"{doc_id}.docx")
        _build_docx(spec, docx_path)
        generated_paths.append(docx_path)
        print(f"  生成: {doc_id}.docx ({spec.get('category', '?')})")

    print(f"[OK] test documents: {len(generated_paths)}")
    return generated_paths


def _build_docx(spec: dict, output_path: str):
    """仕様からpython-docxで.docxファイルを構築する"""
    doc = Document()
    doc.styles['Normal'].font.name = '游ゴシック'
    doc.styles['Normal'].font.size = Pt(10.5)

    for element in spec.get("elements", []):
        el_type   = element.get("type", "paragraph")
        content   = element.get("content", "")
        font_name = element.get("font", "游ゴシック")
        try:
            font_size = float(element.get("font_size", 10.5))
        except (ValueError, TypeError):
            font_size = 10.5

        if el_type == "paragraph":
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.name = font_name
            run.font.size = Pt(font_size)

        elif el_type == "table":
            try:
                table_spec = json.loads(content) if isinstance(content, str) else content
                rows = table_spec.get("rows", 2)
                cols = table_spec.get("cols", 2)
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                for r, row_data in enumerate(table_spec.get("cells", [])[:rows]):
                    for c, cell_text in enumerate(row_data[:cols]):
                        table.cell(r, c).text = str(cell_text)
            except Exception:
                doc.add_paragraph(f"[table placeholder: {str(content)[:50]}]")

        elif el_type == "header":
            section = doc.sections[0]
            section.header.paragraphs[0].text = content

        elif el_type == "footer":
            section = doc.sections[0]
            section.footer.paragraphs[0].text = content

        else:
            doc.add_paragraph(content)

    doc.save(output_path)
