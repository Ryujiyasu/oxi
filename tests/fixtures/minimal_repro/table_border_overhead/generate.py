"""Generate minimal repro .docx variants for table border overhead spec.

Each variant has identical content (single "あ" 10.5pt MS Mincho cell),
varying ONLY in border configuration. This isolates border-overhead from
content/padding/font effects.

Run on Windows with: pip install python-docx
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


HERE = Path(__file__).resolve().parent

# (variant_name, num_rows, top_sz, bot_sz, left_sz, right_sz, insideH_sz)
# sz units: 1/8 pt. sz=4 → 0.5pt, sz=8 → 1.0pt, sz=16 → 2.0pt. None → no border.
VARIANTS: list[tuple[str, int, int | None, int | None, int | None, int | None, int | None]] = [
    # --- Original 10 (commit d4dd7d1) ---
    ("1row_none",          1, None, None, None, None, None),
    ("1row_outer4",        1, 4,    4,    4,    4,    None),
    ("1row_outer8",        1, 8,    8,    8,    8,    None),
    ("1row_outer16",       1, 16,   16,   16,   16,   None),
    ("2row_outer4",        2, 4,    4,    4,    4,    None),
    ("2row_outer4_ih4",    2, 4,    4,    4,    4,    4),
    ("3row_outer4_ih4",    3, 4,    4,    4,    4,    4),
    ("1row_top4_only",     1, 4,    None, None, None, None),
    ("1row_bot4_only",     1, None, 4,    None, None, None),
    ("1row_topbot8",       1, 8,    8,    None, None, None),
    # --- Round 2: disambiguate the +0.5pt multi-row+ih residual ---
    # (a) Does residual scale with N? (constant or per-row?)
    ("5row_outer4_ih4",    5, 4,    4,    4,    4,    4),
    ("10row_outer4_ih4",  10, 4,    4,    4,    4,    4),
    ("20row_outer4_ih4",  20, 4,    4,    4,    4,    4),
    # (b) Does residual scale with insideH width?
    ("2row_outer4_ih8",    2, 4,    4,    4,    4,    8),
    ("2row_outer4_ih16",   2, 4,    4,    4,    4,    16),
    ("3row_outer4_ih16",   3, 4,    4,    4,    4,    16),
    # (c) Are left/right involved at all in multi-row+ih?
    ("3row_topbot4_ih4",   3, 4,    4,    None, None, 4),
    # (d) Insulate content height — bigger font for residual sensitivity
    # (variant generated separately because font size differs from baseline)
]


def _border_elem(name: str, sz: int | None) -> OxmlElement:
    """Build a single <w:top|bottom|left|right|insideH> element."""
    el = OxmlElement(f"w:{name}")
    if sz is None:
        el.set(qn("w:val"), "nil")
    else:
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
    return el


def _set_borders(table, top, bot, left, right, ih):
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    # Drop any existing tblBorders to avoid double declaration
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    tblBorders.append(_border_elem("top",     top))
    tblBorders.append(_border_elem("left",    left))
    tblBorders.append(_border_elem("bottom",  bot))
    tblBorders.append(_border_elem("right",   right))
    tblBorders.append(_border_elem("insideH", ih))
    # insideV doesn't matter (single column) but include for completeness
    tblBorders.append(_border_elem("insideV", None))
    tblPr.append(tblBorders)


def _set_cell_no_padding(cell):
    """Force cell margins to 0 so padding doesn't muddy the measurement."""
    tcPr = cell._element.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _set_run_font(run, name: str = "ＭＳ 明朝", size_pt: float = 10.5):
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def build(variant: tuple) -> None:
    name, rows, top, bot, left, right, ih = variant
    doc = Document()
    # Page setup intentionally default — we measure relative Y.
    table = doc.add_table(rows=rows, cols=1)
    _set_borders(table, top, bot, left, right, ih)
    for r_idx in range(rows):
        cell = table.rows[r_idx].cells[0]
        _set_cell_no_padding(cell)
        # Replace the auto-created paragraph's run
        para = cell.paragraphs[0]
        run = para.add_run("あ")
        _set_run_font(run)
    # Anchor paragraph BEFORE the table so we have a known Y reference for COM.
    # python-docx adds tables to the end; insert a marker paragraph above.
    body = doc.element.body
    table_el = table._element
    marker = OxmlElement("w:p")
    body.insert(list(body).index(table_el), marker)
    out = HERE / f"{name}.docx"
    doc.save(str(out))
    print(f"  wrote {out.name}")


def main() -> int:
    print(f"Generating {len(VARIANTS)} variants in {HERE}")
    for v in VARIANTS:
        build(v)
    print("Done. Next: python measure.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
