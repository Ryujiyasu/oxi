"""Generate minimal repro .docx variants for table cell row-height drift.

Each variant: 1-row 1-col table whose single cell holds N PARAGRAPHS (each
just "あ"), no padding, no borders. Varies N, font family, font size.

We deliberately use separate <w:p> elements (not <w:br/> soft breaks) because:
1. kyodoken10 (the real-world doc that motivated this minimal repro) uses
   paragraph breaks
2. Oxi's parser does not currently handle <w:br/> as a forced line break
   inside a paragraph (separate bug, see br_gap.md)

Run on Windows with: pip install python-docx
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


HERE = Path(__file__).resolve().parent

# (variant_name, n_lines, font_name, font_size_pt)
VARIANTS: list[tuple[str, int, str, float]] = [
    # N-scaling for the baseline font
    ("1cell_1line_mincho10p5",    1, "ＭＳ 明朝", 10.5),
    ("1cell_2line_mincho10p5",    2, "ＭＳ 明朝", 10.5),
    ("1cell_3line_mincho10p5",    3, "ＭＳ 明朝", 10.5),
    ("1cell_5line_mincho10p5",    5, "ＭＳ 明朝", 10.5),
    ("1cell_10line_mincho10p5",  10, "ＭＳ 明朝", 10.5),
    ("1cell_20line_mincho10p5",  20, "ＭＳ 明朝", 10.5),
    ("1cell_50line_mincho10p5",  50, "ＭＳ 明朝", 10.5),
    # Font-size sweep at fixed N=10
    ("1cell_10line_mincho10",    10, "ＭＳ 明朝", 10.0),
    ("1cell_10line_mincho11",    10, "ＭＳ 明朝", 11.0),
    ("1cell_10line_mincho12",    10, "ＭＳ 明朝", 12.0),
    # Font-family sweep at fixed N=10
    ("1cell_10line_calibri11",   10, "Calibri",  11.0),
    ("1cell_10line_yumin10p5",   10, "Yu Mincho", 10.5),
]


def _set_cell_no_padding(cell):
    tcPr = cell._element.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _no_borders(table):
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "nil")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _make_run(para, text: str, font_name: str, size_pt: float):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.insert(0, rFonts)
    return run


def build(variant: tuple[str, int, str, float]) -> None:
    name, n_lines, font, size = variant
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _no_borders(table)
    cell = table.rows[0].cells[0]
    _set_cell_no_padding(cell)

    # The auto-created first paragraph
    para = cell.paragraphs[0]
    _make_run(para, "あ", font, size)

    # Add (n_lines - 1) more paragraphs to the cell
    tc = cell._element
    for _ in range(n_lines - 1):
        new_p = OxmlElement("w:p")
        # Insert before the cell's tcPr (if any) by appending — tcPr stays first
        tc.append(new_p)
        # Build a python-docx Paragraph wrapper around the new w:p
        from docx.text.paragraph import Paragraph  # local import to avoid top-level docx noise
        np = Paragraph(new_p, cell)
        _make_run(np, "あ", font, size)

    # Anchor paragraph BEFORE the table for Y reference
    body = doc.element.body
    table_el = table._element
    marker = OxmlElement("w:p")
    body.insert(list(body).index(table_el), marker)

    out = HERE / f"{name}.docx"
    doc.save(str(out))
    print(f"  wrote {out.name}")


def main() -> int:
    print(f"Generating {len(VARIANTS)} variants in {HERE}")
    for v in VARIANTS:
        build(v)
    print("Done. Next: python measure.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
